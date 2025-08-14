import ezdxf
import math
import csv
import os
import re
import glob
import locale
from docx import Document
from docx.shared import Inches
from datetime import datetime
from decimal import Decimal, getcontext
import pandas as pd
import locale
import openpyxl
from openpyxl.styles import Alignment, Font
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import logging
EPS_BULGE = 1e-9  # pode ficar aqui mesmo, no topo dos helpers
import sys

# Diretório para logs
BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
LOG_DIR = os.path.join(BASE_DIR, 'static', 'logs')
os.makedirs(LOG_DIR, exist_ok=True)

# Arquivo de log específico para poligonal_fechada
log_file = os.path.join(LOG_DIR, f'poligonal_fechada_{datetime.now().strftime("%Y%m%d_%H%M%S")}.log')

# Configuração básica do logger
logger = logging.getLogger(__name__)
logger.setLevel(logging.DEBUG)

fmt = logging.Formatter('%(asctime)s [%(levelname)s] %(message)s')

# FileHandler (UTF-8)
file_handler = logging.FileHandler(log_file, encoding='utf-8')
file_handler.setFormatter(fmt)

# StreamHandler → stdout (aparece no LOG RENDER)
stream_handler = logging.StreamHandler(sys.stdout)
stream_handler.setFormatter(fmt)

# Evitar handlers duplicados (em caso de reload/import)
if not logger.handlers:
    logger.addHandler(file_handler)
    logger.addHandler(stream_handler)

# (opcional) deixar propagar para o root também
logger.propagate = True

stream_handler = logging.StreamHandler(sys.stdout)
stream_handler.setFormatter(logging.Formatter('%(asctime)s [%(levelname)s] %(message)s'))
logger.addHandler(stream_handler)
logger.propagate = True  # repassa para o root também

logger.info("[AZ] Log de poligonal_fechada gravando em: %s", log_file)

getcontext().prec = 28  # Define a precisão para 28 casas decimais

MESES_PT_BR = {
    'January': 'janeiro',
    'February': 'fevereiro',
    'March': 'março',
    'April': 'abril',
    'May': 'maio',
    'June': 'junho',
    'July': 'julho',
    'August': 'agosto',
    'September': 'setembro',
    'October': 'outubro',
    'November': 'novembro',
    'December': 'dezembro'
}

def calcular_area_poligonal(pontos):
    """Calcula a área da poligonal fechada usando a fórmula shoelace."""
    n = len(pontos)
    area = 0.0
    for i in range(n):
        x1, y1 = pontos[i]
        x2, y2 = pontos[(i + 1) % n]
        area += (x1 * y2) - (x2 * y1)
    return abs(area) / 2


def limpar_dxf_e_converter_r2010(original_path, saida_path):
    """
    Lê um DXF original e regrava o arquivo com a versão R2010,
    preservando LWPOLYLINE (com bulge). Não cria geometria nova.
    """
    try:
        doc_antigo = ezdxf.readfile(original_path)
        msp_antigo = doc_antigo.modelspace()
        doc_novo = ezdxf.new(dxfversion='R2010')
        msp_novo = doc_novo.modelspace()

        encontrou_polilinha = False

        for entity in msp_antigo.query('LWPOLYLINE'):
            if entity.closed:
                # Pega (x, y, bulge) como tuplas
                pontos_xyb = entity.get_points('xyb')  # [(x, y, bulge), ...]
                # Regrava preservando o bulge
                msp_novo.add_lwpolyline(
                    pontos_xyb,
                    format='xyb',
                    close=True,
                    dxfattribs={'layer': entity.dxf.layer}
                )
                encontrou_polilinha = True
                break

        if not encontrou_polilinha:
            raise ValueError("Nenhuma polilinha fechada encontrada no DXF original.")

        doc_novo.saveas(saida_path)
        logger.info(f"✅ DXF convertido e salvo como R2010 em: {saida_path}")
        return saida_path

    except Exception as e:
        logger.error(f"❌ Erro ao converter DXF para R2010: {e}")
        return original_path


def get_document_info_from_dxf(dxf_file_path):
    try:
        doc = ezdxf.readfile(dxf_file_path)
        msp = doc.modelspace()

        lines = []
        perimeter_dxf = 0.0
        area_dxf = 0.0
        ponto_az = None
        ordered_points = []
        ordered_points_with_bulge = []

        # Lê a PRIMEIRA LWPOLYLINE fechada
        for entity in msp.query('LWPOLYLINE'):
            if entity.closed:
                pts_xyb = entity.get_points('xyb')  # lista de tuplas (x, y, bulge)
                # Se houver último = primeiro, remove duplicata
                if len(pts_xyb) >= 2 and (pts_xyb[0][0], pts_xyb[0][1]) == (pts_xyb[-1][0], pts_xyb[-1][1]):
                    pts_xyb = pts_xyb[:-1]

                # Monta listas
                for (x, y, b) in pts_xyb:
                    ordered_points.append((x, y))
                    ordered_points_with_bulge.append({'x': x, 'y': y, 'bulge_next': float(b or 0.0)})

                logger.info(
                    f"pts_bulge: n={len(ordered_points_with_bulge)} | exemplo={ordered_points_with_bulge[:2]}"
                )

                # Linhas + perímetro
                n = len(ordered_points)
                for i in range(n):
                    p1 = ordered_points[i]
                    p2 = ordered_points[(i + 1) % n]
                    lines.append((p1, p2))
                    perimeter_dxf += ((p2[0] - p1[0]) ** 2 + (p2[1] - p1[1]) ** 2) ** 0.5

                # Área por shoelace
                x = [p[0] for p in ordered_points]
                y = [p[1] for p in ordered_points]
                area_dxf = abs(sum(x[i] * y[(i + 1) % n] - x[(i + 1) % n] * y[i] for i in range(n)) / 2.0)
                break

        if not lines:
            logger.info("Nenhuma polilinha fechada encontrada no arquivo DXF.")
            return None, [], 0.0, 0.0, None, None, []

        # Ponto Az (TEXT → INSERT → POINT)
        for entity in msp.query('TEXT'):
            if "Az" in (entity.dxf.text or ""):
                ponto_az = (entity.dxf.insert.x, entity.dxf.insert.y, 0.0)
                break

        if ponto_az is None:
            for entity in msp.query('INSERT'):
                if "Az" in (entity.dxf.name or ""):
                    ponto_az = (entity.dxf.insert.x, entity.dxf.insert.y, 0.0)
                    break

        if ponto_az is None:
            for entity in msp.query('POINT'):
                ponto_az = (entity.dxf.location.x, entity.dxf.location.y, 0.0)
                break

        if ponto_az is None:
            # Fallback: primeiro vértice
            ponto_az = (ordered_points[0][0], ordered_points[0][1], 0.0)
            logger.warning("⚠️ Ponto Az não encontrado no DXF. Usando fallback (primeiro ponto).")

        logger.info(f"Linhas processadas: {len(lines)}")
        logger.info(f"Perímetro do DXF: {perimeter_dxf:.2f} m")
        logger.info(f"Área do DXF: {area_dxf:.2f} m²")

        # >>> RETORNA 7 ITENS <<<
        return doc, lines, perimeter_dxf, area_dxf, ponto_az, msp, ordered_points_with_bulge

    except Exception as e:
        logger.error(f"Erro ao obter informações do documento: {e}")
        # >>> TAMBÉM retorna 7 itens no erro <<<
        return None, [], 0.0, 0.0, None, None, []


# # Função que processa as linhas da poligonal
# def get_document_info_from_dxf(dxf_file_path):
#     try:
#         doc = ezdxf.readfile(dxf_file_path)  
#         msp = doc.modelspace()  

#         lines = []
#         perimeter_dxf = 0
#         area_dxf = 0
#         ponto_az = None
#         area_poligonal = None

#         for entity in msp.query('LWPOLYLINE'):
#             if entity.closed:
#                 points = entity.get_points('xy')
                
#                 # Verifica e remove vértice repetido no final, se houver
#                 if points[0] == points[-1]:
#                     points.pop()
                
#                 num_points = len(points)

#                 for i in range(num_points):
#                     start_point = (points[i][0], points[i][1])
#                     end_point = (points[(i + 1) % num_points][0], points[(i + 1) % num_points][1])
#                     lines.append((start_point, end_point))

#                     segment_length = ((end_point[0] - start_point[0]) ** 2 + 
#                                       (end_point[1] - start_point[1]) ** 2) ** 0.5
#                     perimeter_dxf += segment_length

#                 x = [point[0] for point in points]
#                 y = [point[1] for point in points]
#                 area_dxf = abs(sum(x[i] * y[(i + 1) % num_points] - x[(i + 1) % num_points] * y[i] for i in range(num_points)) / 2)

#                 break  

#         if not lines:
#             print("Nenhuma polilinha encontrada no arquivo DXF.")
#             return None, [], 0, 0, None, None

#         for entity in msp.query('TEXT'):
#             if "Az" in entity.dxf.text:
#                 ponto_az = (entity.dxf.insert.x, entity.dxf.insert.y, 0)
#                 print(f"Ponto Az encontrado em texto: {ponto_az}")

#         for entity in msp.query('INSERT'):
#             if "Az" in entity.dxf.name:
#                 ponto_az = (entity.dxf.insert.x, entity.dxf.insert.y, 0)
#                 print(f"Ponto Az encontrado no bloco: {ponto_az}")

#         for entity in msp.query('POINT'):
#             ponto_az = (entity.dxf.location.x, entity.dxf.location.y, 0)
#             print(f"Ponto Az encontrado como ponto: {ponto_az}")
            
#         if not ponto_az:
#             print("Ponto Az não encontrado no arquivo DXF.")
#             return None, lines, 0, 0, None, None

#         print(f"Linhas processadas: {len(lines)}")
#         print(f"Perímetro do DXF: {perimeter_dxf:.2f} metros")
#         print(f"Área do DXF: {area_dxf:.2f} metros quadrados")

#         return doc, lines, perimeter_dxf, area_dxf, ponto_az, area_poligonal

#     except Exception as e:
#         print(f"Erro ao obter informações do documento: {e}")
#         return None, [], 0, 0, None, None


def is_clockwise(points):
    """
    Verifica se a poligonal está no sentido horário.
    Retorna True se for horário, False se anti-horário.
    """
    area = 0.0
    for i in range(len(points)):
        j = (i + 1) % len(points)
        area += points[i][0] * points[j][1]
        area -= points[j][0] * points[i][1]
    return area < 0

def ensure_counterclockwise(points):
    """
    Garante que a lista de pontos esteja no sentido anti-horário.
    Se estiver no sentido horário, inverte a ordem dos pontos.
    """
    if is_clockwise(points):
        points.reverse()
    return points

    
# 🔹 Função para definir a fonte padrão
def set_default_font(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    
def calculate_point_on_line(start, end, distance):
    """
    Calcula um ponto a uma determinada distância sobre uma linha.
    """
    dx, dy = end[0] - start[0], end[1] - start[1]
    length = math.hypot(dx, dy)
    return (
        start[0] + dx / length * distance,
        start[1] + dy / length * distance
    )

def add_azimuth_arc_to_dxf(msp, ponto_az, v1, azimute):
    """
    Adiciona o arco do azimute ao DXF usando ezdxf.
    """
    try:
        print(f"Iniciando a adição do arco de azimute. Azimute: {azimute}°")

        # Criar camada 'Azimute', se não existir
        if 'Azimute' not in msp.doc.layers:
            msp.doc.layers.new(name='Azimute', dxfattribs={'color': 1})
            print("Camada 'Azimute' criada com sucesso.")

        # Traçar segmento entre Az e V1
        msp.add_line(start=ponto_az, end=v1, dxfattribs={'layer': 'Azimute'})
        print(f"Segmento entre Az e V1 desenhado de {ponto_az} para {v1}")

        # Traçar segmento para o norte
        north_point = (ponto_az[0], ponto_az[1] + 2)
        msp.add_line(start=ponto_az, end=north_point, dxfattribs={'layer': 'Azimute'})
        print(f"Linha para o norte desenhada com sucesso de {ponto_az} para {north_point}")

        # Calcular o ponto inicial (1 metro de Az para V1)
        start_arc = calculate_point_on_line(ponto_az, v1, 1)

        # Calcular o ponto final (1 metro de Az para o Norte)
        end_arc = calculate_point_on_line(ponto_az, north_point, 1)

        # Traçar o arco do azimute no sentido horário
        msp.add_arc(
            center=ponto_az,
            radius=1,
            start_angle=math.degrees(math.atan2(start_arc[1] - ponto_az[1], start_arc[0] - ponto_az[0])),
            end_angle=math.degrees(math.atan2(end_arc[1] - ponto_az[1], end_arc[0] - ponto_az[0])),
            dxfattribs={'layer': 'Azimute'}
        )
        print(f"Arco do azimute desenhado com sucesso com valor de {azimute}° no ponto {ponto_az}")

       # Adicionar rótulo do azimute diretamente com o texto "Azimute:"
        azimuth_label = f"Azimute: {convert_to_dms(azimute)}"  # Incluir o prefixo "Azimute:"

        # Calcular a posição do rótulo
        label_position = (
            ponto_az[0] + 1.0 * math.cos(math.radians(azimute / 2)),
            ponto_az[1] + 1.0 * math.sin(math.radians(azimute / 2))
        )

        # Adicionar o texto ao desenho
        msp.add_text(
            azimuth_label,
            dxfattribs={
                'height': 0.25,
                'layer': 'Azimute',
                'insert': label_position  # Define a posição diretamente
            }
        )

        print(f"Rótulo do azimute adicionado com sucesso: '{azimuth_label}' em {label_position}")


    except Exception as e:
        print(f"Erro na função `add_azimuth_arc_to_dxf`: {e}")

def calculate_polygon_area(points):
    """
    Calcula a área de uma poligonal fechada utilizando o método do produto cruzado com alta precisão.
    :param points: Lista de coordenadas [(x1, y1), (x2, y2), ...].
    :return: Área da poligonal.
    """
    n = len(points)
    if n < 3:
        return Decimal(0)  # Não é uma poligonal válida

    area = Decimal(0)
    for i in range(n):
        x1, y1 = Decimal(points[i][0]), Decimal(points[i][1])
        x2, y2 = Decimal(points[(i + 1) % n][0]), Decimal(points[(i + 1) % n][1])
        area += x1 * y2 - y1 * x2

    return abs(area) / 2
    
def convert_to_dms(decimal_degrees):
    """
    Converte graus decimais para o formato graus, minutos e segundos (DMS).
    """
    try:
        # Verificar se o valor é NaN
        if math.isnan(decimal_degrees):
            raise ValueError("Valor de entrada é NaN")

        degrees = int(decimal_degrees)
        minutes = int((abs(decimal_degrees) - abs(degrees)) * 60)
        seconds = round((abs(decimal_degrees) - abs(degrees) - minutes / 60) * 3600, 2)
        return f"{degrees}°{minutes}'{seconds}\""
    except Exception as e:
        print(f"Erro na conversão para DMS: {e}")
        return "0°0'0.00\""  # Valor padrão em caso de erro

def calculate_distance(p1, p2):
    return math.sqrt((p2[0] - p1[0])**2 + (p2[1] - p1[1])**2)


def degrees_to_dms(angle):
    """
    Converte um ângulo em graus decimais para o formato graus, minutos e segundos (° ' ").
    """
    degrees = int(angle)
    minutes = int((angle - degrees) * 60)
    seconds = round((angle - degrees - minutes / 60) * 3600, 2)
    return f"{degrees}°{minutes}'{seconds}\""
import math

def calculate_azimuth(p1, p2):
    """
    Calcula o azimute entre dois pontos p1 e p2.
    
    :param p1: Coordenadas do ponto 1 (E, N)
    :param p2: Coordenadas do ponto 2 (E, N)
    :return: Azimute em graus decimais
    """
    delta_x = p2[0] - p1[0]
    delta_y = p2[1] - p1[1]
    
    azimuth = math.degrees(math.atan2(delta_x, delta_y)) % 360  # Garantir valor entre 0° e 360°
    
    return azimuth


def create_arrow_block(doc, block_name="ARROW"):
    """
    Cria um bloco no DXF representando uma seta sólida como um triângulo.
    """
    if block_name in doc.blocks:
        return  # O bloco já existe

    block = doc.blocks.new(name=block_name)

    # Definir o triângulo da seta
    length = 0.5  # Comprimento da seta
    base_half_length = length / 2

    tip = (0, 0)  # Ponta da seta no eixo de coordenadas
    base1 = (-base_half_length, -length)
    base2 = (base_half_length, -length)

    block.add_solid([base1, base2, tip])
import math

# def add_giro_angular_arc_to_dxf(doc_dxf, v1, az, v2, radius=1.0):
#     """
#     Adiciona um arco representando o giro angular horário no espaço de modelo do DXF já aberto.
#     """
#     try:
#         msp = doc_dxf if hasattr(doc_dxf, "add_line") else doc_dxf.modelspace()
#         doc = msp.doc  # referência ao documento para camadas/estilos

#         # Traçar a reta entre V1 e Az
#         msp.add_line(start=v1[:2], end=az[:2])
#         print(f"Linha entre V1 e Az traçada com sucesso.")

#         # Definir os pontos de apoio
#         def calculate_displacement(point1, point2, distance):
#             dx = point2[0] - point1[0]
#             dy = point2[1] - point1[1]
#             magnitude = math.hypot(dx, dy)
#             return (
#                 point1[0] + (dx / magnitude) * distance,
#                 point1[1] + (dy / magnitude) * distance,
#             )

#         # Calcular os pontos de apoio
#         ponto_inicial = calculate_displacement(v1, v2, radius)  # 2m na reta V1-V2
#         ponto_final = calculate_displacement(v1, az, radius)   # 2m na reta V1-Az

#         # Calcular os ângulos dos vetores
#         angle_v2 = math.degrees(math.atan2(ponto_inicial[1] - v1[1], ponto_inicial[0] - v1[0]))
#         angle_az = math.degrees(math.atan2(ponto_final[1] - v1[1], ponto_final[0] - v1[0]))

#         # Calcular o giro angular no sentido horário
#         giro_angular = (angle_az - angle_v2) % 360  # Garantir que o ângulo esteja no intervalo [0, 360)
#         if giro_angular < 0:  # Caso negativo, ajustar para o sentido horário
#             giro_angular += 360

#         print(f"Giro angular calculado corretamente: {giro_angular:.2f}°")

#         # Traçar o arco
#         msp.add_arc(center=v1[:2], radius=radius, start_angle=angle_v2, end_angle=angle_az)
#         print(f"Arco do giro angular traçado com sucesso.")

#         # Adicionar rótulo ao arco
#         label_offset = 3.0
#         deslocamento_x=3
#         deslocamento_y=-3
#         angle_middle = math.radians((angle_v2 + angle_az) / 2)
#         label_position = (
#             v1[0] + (label_offset+deslocamento_x) * math.cos(angle_middle),
#             v1[1] + (label_offset+deslocamento_y) * math.sin(angle_middle),
#         )
#         # Converter o ângulo para DMS e exibir no rótulo
#         giro_angular_dms = f"Giro Angular:{convert_to_dms(giro_angular)}"
#         msp.add_text(
#             giro_angular_dms,
#             dxfattribs={
#                 'height': 0.3,
#                 'layer': 'Labels',
#                 'insert': label_position  # Define a posição do texto
#             }
#         )
#         print(f"Rótulo do giro angular ({giro_angular_dms}) adicionado com sucesso.")

#     except Exception as e:
#         print(f"Erro ao adicionar o arco do giro angular ao DXF: {e}") 




def add_giro_angular_arc_to_dxf(doc_dxf, v1, az, v2, radius=2.0):
    """
    Adiciona um arco representando o giro angular (setor entre V1→V2 e V1→Az)
    no modelspace do DXF já aberto. Aceita tanto 'doc' quanto 'msp' como 1º parâmetro.
    """
    try:
        # ── 1) normaliza msp/doc ─────────────────────────────────────────────
        msp = doc_dxf if hasattr(doc_dxf, "add_line") else doc_dxf.modelspace()
        doc = msp.doc

        # ── 2) garante camadas (opcional) ────────────────────────────────────
        try:
            if "GiroAZ" not in doc.layers:
                doc.layers.new("GiroAZ")
            if "Labels" not in doc.layers:
                doc.layers.new("Labels")
        except Exception:
            pass

        # ── 3) reta V1–Az (debug/apoio) ──────────────────────────────────────
        msp.add_line(start=v1[:2], end=az[:2], dxfattribs={"layer": "GiroAZ"})
        print("Linha entre V1 e Az traçada com sucesso.")

        # ── 4) helper de deslocamento com proteção ───────────────────────────
        def calculate_displacement(point1, point2, distance):
            dx = point2[0] - point1[0]
            dy = point2[1] - point1[1]
            magnitude = math.hypot(dx, dy)
            if magnitude == 0:
                return (point1[0], point1[1])
            return (
                point1[0] + (dx / magnitude) * distance,
                point1[1] + (dy / magnitude) * distance,
            )

        # ── 5) pontos de apoio a partir de V1 ────────────────────────────────
        ponto_inicial = calculate_displacement(v1, v2, radius)  # na direção V1→V2
        ponto_final   = calculate_displacement(v1, az, radius)  # na direção V1→Az

        # ── 6) ângulos (em graus) ────────────────────────────────────────────
        angle_v2 = math.degrees(math.atan2(ponto_inicial[1] - v1[1], ponto_inicial[0] - v1[0]))
        angle_az = math.degrees(math.atan2(ponto_final[1] - v1[1],   ponto_final[0] - v1[0]))

        # giro CCW entre as direções (mantém compatibilidade visual atual)
        giro_angular = (angle_az - angle_v2) % 360
        print(f"Giro angular calculado corretamente: {giro_angular:.2f}°")

        # ── 7) arco CCW de angle_v2 → angle_az ───────────────────────────────
        msp.add_arc(center=v1[:2], radius=radius, start_angle=angle_v2, end_angle=angle_az,
                    dxfattribs={"layer": "GiroAZ"})
        print("Arco do giro angular traçado com sucesso.")

        # ── 8) rótulo no meio do setor (tratando wrap-around) ────────────────
        label_offset   = 3.0
        desloc_x, desloc_y = 3.0, -3.0
        sweep_ccw = (angle_az - angle_v2) % 360
        angle_middle = math.radians((angle_v2 + sweep_ccw / 2.0) % 360)

        label_position = (
            v1[0] + (label_offset + desloc_x) * math.cos(angle_middle),
            v1[1] + (label_offset + desloc_y) * math.sin(angle_middle),
        )

        giro_angular_dms = f"Giro Angular: {convert_to_dms(giro_angular)}"

        txt = msp.add_text(
            giro_angular_dms,
            dxfattribs={'height': 0.3, 'layer': 'Labels'}
        
        ).set_dxf_attrib('insert', label_position)

        print(f"Rótulo do giro angular ({giro_angular_dms}) adicionado com sucesso.")

    except Exception as e:
        print(f"Erro ao adicionar o arco do giro angular ao DXF: {e}")



def calculate_arc_angles(p1, p2, p3):
    """
    Calcula os ângulos de início e fim do arco para representar o ângulo interno da poligonal.
    O arco deve sempre ser desenhado dentro da poligonal.
    """
    try:
        # Vetores a partir de p2
        dx1, dy1 = p1[0] - p2[0], p1[1] - p2[1]  # Vetor de p2 para p1
        dx2, dy2 = p3[0] - p2[0], p3[1] - p2[1]  # Vetor de p2 para p3

        # Ângulos dos vetores em relação ao eixo X
        angle1 = math.degrees(math.atan2(dy1, dx1)) % 360
        angle2 = math.degrees(math.atan2(dy2, dx2)) % 360

        # Determinar o ângulo interno
        internal_angle = (angle2 - angle1) % 360
#         if internal_angle > 180:
#             internal_angle = 360 - internal_angle  # Complementar para 360°

        # Ajustar os ângulos de início e fim para sempre formar o ângulo interno dentro da poligonal
        start_angle = angle1
        end_angle = angle1 + internal_angle  # Sempre desenha o arco no sentido correto

        # Se o ângulo interno for maior que 180°, inverter os ângulos para garantir o lado interno
        if internal_angle > 180:
            start_angle, end_angle = end_angle, start_angle

        return start_angle % 360, end_angle % 360

    except Exception as e:
        print(f"Erro ao calcular ângulos do arco: {e}")
        return 0, 0  # Retorno seguro em caso de erro




def insert_and_rotate_arrow(msp, center, radius, angle, rotation_adjustment, block_name="ARROW"):
    """
    Insere um bloco da seta alinhado ao arco e rotacionado conforme especificado.
    """
    arrow_tip_x = center[0] + radius * math.cos(angle)
    arrow_tip_y = center[1] + radius * math.sin(angle)
    arrow_tip = (arrow_tip_x, arrow_tip_y)

    rotation_angle = math.degrees(angle) + rotation_adjustment

    msp.add_blockref(
        block_name,
        insert=arrow_tip,
        dxfattribs={"rotation": rotation_angle}
    )
import math





def internal_angle_decimal(msp, ordered_points, angulos_excel):
    try:
        total_points = len(ordered_points)
        
        for i, p2 in enumerate(ordered_points):
            if i == 0:
                print("⏩ Ignorando arco e rótulo para V1")
                continue
            p1 = ordered_points[i - 1] if i > 0 else ordered_points[-1]
            p3 = ordered_points[(i + 1) % total_points]

            def calculate_displacement(base_point, direction_point, distance):
                dx = direction_point[0] - base_point[0]
                dy = direction_point[1] - base_point[1]
                magnitude = math.hypot(dx, dy)
                return (
                    base_point[0] + (dx / magnitude) * distance,
                    base_point[1] + (dy / magnitude) * distance,
                )

            lado_antes = math.hypot(p2[0] - p1[0], p2[1] - p1[1])
            lado_depois = math.hypot(p3[0] - p2[0], p3[1] - p2[1])
            lado_menor = min(lado_antes, lado_depois)

            radius = lado_menor * 0.8 if lado_menor <= 0.5 else 1.0

            ponto_inicial = calculate_displacement(p2, p3, radius)
            ponto_final = calculate_displacement(p2, p1, radius)

            start_angle = math.degrees(math.atan2(ponto_inicial[1] - p2[1], ponto_inicial[0] - p2[0]))
            end_angle = math.degrees(math.atan2(ponto_final[1] - p2[1], ponto_final[0] - p2[0]))

            if end_angle < start_angle:
                end_angle += 360

            internal_angle_dms = angulos_excel[i]

            msp.add_arc(
                center=p2,
                radius=radius,
                start_angle=start_angle,
                end_angle=end_angle,
                dxfattribs={'layer': 'Internal_Arcs'}
            )

            label_offset = 1
            label_position = (
                p2[0] + label_offset * math.cos(math.radians((start_angle + end_angle) / 2)),
                p2[1] + label_offset * math.sin(math.radians((start_angle + end_angle) / 2))
            )

            msp.add_text(
                internal_angle_dms,
                dxfattribs={
                    'height': 0.3,
                    'layer': 'Labels',
                    'insert': label_position
                }
            )

            print(f"Vértice V{i+1}: Ângulo interno {internal_angle_dms}")

    except Exception as e:
        print(f"Erro ao adicionar ângulos internos ao DXF: {e}")




def calculate_internal_angle(p1, p2, p3):
    try:
        # Vetores a partir do ponto central (p2)
        dx1, dy1 = p1[0] - p2[0], p1[1] - p2[1]
        dx2, dy2 = p3[0] - p2[0], p3[1] - p2[1]

        # Ângulos dos vetores em relação ao eixo X
        angle1 = math.atan2(dy1, dx1)
        angle2 = math.atan2(dy2, dx2)

        # Ângulo formado entre os vetores
        internal_angle = (angle2 - angle1) % (2 * math.pi)

        # Se o ângulo for maior que 180°, use o suplementar
        if internal_angle > math.pi:
            internal_angle = (2 * math.pi) - internal_angle

        # Retornar o ângulo em graus
        return math.degrees(internal_angle)

    except Exception as e:
        print(f"Erro inesperado ao calcular o ângulo interno: {e}")
        return 0


def calculate_label_position(p2, start_angle, end_angle, radius=1.8):
    """
    Calcula a posição do rótulo do ângulo interno, deslocando-o para uma posição central no arco.
    
    :param p2: Ponto central do arco (vértice da poligonal)
    :param start_angle: Ângulo de início do arco
    :param end_angle: Ângulo de fim do arco
    :param radius: Raio do deslocamento para evitar sobreposição
    :return: Posição (x, y) onde o rótulo do ângulo interno será colocado
    """
    try:
        # Calcular o ângulo médio entre os dois ângulos do arco
        mid_angle = math.radians((start_angle + end_angle) / 2)

        # Calcular a posição do rótulo deslocado no ângulo médio
        label_x = p2[0] + radius * math.cos(mid_angle)
        label_y = p2[1] + radius * math.sin(mid_angle)

        return (label_x, label_y)

    except Exception as e:
        print(f"Erro ao calcular posição do rótulo do ângulo: {e}")
        return p2  # Retorna o próprio ponto central caso ocorra erro


import math

def add_label_and_distance(msp, start_point, end_point, label, distance):
    """
    Adiciona rótulos e distâncias no espaço de modelo usando ezdxf.
    """
    try:
        # Calcular ponto médio
        mid_point = (
            (start_point[0] + end_point[0]) / 2,
            (start_point[1] + end_point[1]) / 2
        )

        # Vetor da linha
        dx = end_point[0] - start_point[0]
        dy = end_point[1] - start_point[1]
        length = math.hypot(dx, dy)

        # Ângulo da linha
        angle = math.degrees(math.atan2(dy, dx))

        # Corrigir para manter a leitura sempre da esquerda para a direita
        if angle < -90 or angle > 90:
            angle += 180  

        # Afastar o rótulo da linha
        offset = -0.5  # Ajuste o valor para mais afastamento
        perp_x = -dy / length * offset
        perp_y = dx / length * offset
        displaced_mid_point = (mid_point[0] + perp_x, mid_point[1] + perp_y)

        # Criar layer se não existir
        if "Distance_Labels" not in msp.doc.layers:
            msp.doc.layers.new(name="Distance_Labels", dxfattribs={"color": 2})  # Define cor para melhor visualização

        # Adicionar o rótulo da distância no LAYER correto
        msp.add_text(
            f"{distance:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".") + "",
            dxfattribs={
                "height": 0.3,
                "layer": "Distance_Labels",  # Novo LAYER
                "rotation": angle,
                "insert": displaced_mid_point
            }
        )
        print(f"Distância {distance:.2f} m adicionada corretamente em {displaced_mid_point}")

    except Exception as e:
        print(f"Erro ao adicionar rótulo de distância: {e}")

def calculate_angular_turn(p1, p2, p3):
    """
    Calcula o giro angular no ponto `p2` entre os segmentos `p1-p2` e `p2-p3` no sentido horário.
    Retorna o ângulo em graus.
    """
    import math
    
    dx1, dy1 = p1[0] - p2[0], p1[1] - p2[1]  # Vetor do segmento p1-p2
    dx2, dy2 = p3[0] - p2[0], p3[1] - p2[1]  # Vetor do segmento p2-p3

    angle1 = math.atan2(dy1, dx1)
    angle2 = math.atan2(dy2, dx2)

    # Calcula o ângulo horário
    angular_turn = (angle2 - angle1) % (2 * math.pi)
    angular_turn_degrees = math.degrees(angular_turn)

    return angular_turn_degrees

def add_angle_visualization_to_dwg(msp, ordered_points, angulos_decimais, sentido_poligonal):
    """
    Desenha os arcos de ângulos internos SEMPRE para dentro.
    Premissa: 'ordered_points' já está no sentido desejado (horário ou anti_horário).
    'sentido_poligonal' deve ser 'horario' ou 'anti_horario'.
    """
    try:
        n = len(ordered_points)
        if n < 3:
            return  # nada a fazer

        for i in range(n):
            # vizinhos (p1 <- p2 -> p3)
            p2 = _xy(ordered_points[i])
            p1 = _xy(ordered_points[i - 1])
            p3 = _xy(ordered_points[(i + 1) % n])

            # ângulos das DIREÇÕES adjacentes no vértice
            ang_in  = _ang(p2, p1)  # direção de entrada (segmento p1->p2, visto a partir de p2)
            ang_out = _ang(p2, p3)  # direção de saída   (segmento p2->p3, visto a partir de p2)

            interno = float(angulos_decimais[i])

            # raio: 10% do menor lado adjacente (com piso de 1 cm)
            lado1 = math.hypot(p1[0] - p2[0], p1[1] - p2[1])
            lado2 = math.hypot(p3[0] - p2[0], p3[1] - p2[1])
            raio = max(0.01, 0.10 * min(lado1, lado2))

            # ===== REGRA FUNDAMENTAL (ezdxf varre SEMPRE CCW de start->end) =====
            if sentido_poligonal == 'anti_horario':
                # para ficar "por dentro", varra CCW de SAÍDA para ENTRADA
                start_angle = ang_out
                end_angle   = (start_angle + interno) % 360.0
                mid_angle   = (start_angle + interno / 2.0) % 360.0
            else:
                # 'horario' → varra CCW de ENTRADA para SAÍDA (minor arc é o interno)
                start_angle = ang_in
                end_angle   = (start_angle + interno) % 360.0
                mid_angle   = (start_angle + interno / 2.0) % 360.0

            # desenha arco
            msp.add_arc(
                center=(p2[0], p2[1]),
                radius=raio,
                start_angle=start_angle,
                end_angle=end_angle,
                dxfattribs={'layer': 'Internal_Arcs'}
            )

            # rótulo no meio do arco
            distancia_label = raio + 1.0
            pos_label = (
                p2[0] + distancia_label * math.cos(math.radians(mid_angle)),
                p2[1] + distancia_label * math.sin(math.radians(mid_angle)),
            )

            # se já tem sua função utilitária, mantenha
            ang_texto = convert_to_dms(interno) if 'convert_to_dms' in globals() else f"{interno:.4f}°"

            msp.add_text(
                ang_texto,
                dxfattribs={
                    'height': 0.7,
                    'layer': 'Labels',
                    'insert': pos_label,
                    # rotação amigável
                    'rotation': mid_angle if mid_angle <= 180 else mid_angle - 180
                }
            )

    except Exception as e:
        print(f"Erro ao adicionar ângulos internos ao DXF: {e}")

#HELPERS ESPECIFICOS SOMENTE PARA O ANGULO_AZ

# nomes de layers “típicas” onde você pode marcar o AZ
_AZ_LAYER_HINTS = {"AZ", "PONTO_AZ", "AMARRACAO", "AMARRAÇÃO", "AZIMUTE", "AZ_POINT", "AZ_PONTO"}

def _clean_mtext(text: str) -> str:
    """Remove marcações comuns de MTEXT (\\P, \\~ etc.) e espaços."""
    if text is None:
        return ""
    s = str(text)
    # remove sequências de formatação básicas do MTEXT
    s = re.sub(r"\\[A-Za-z]+", "", s)       # \P, \~ etc.
    s = re.sub(r"{|}", "", s)               # chaves
    s = s.replace("\n", " ").replace("\r", " ")
    return s.strip()

def _is_az_label(text: str) -> bool:
    """Retorna True se o texto parecer um rótulo 'AZ'."""
    if text is None:
        return False
    s = _clean_mtext(text).strip()
    # aceita “AZ”, “Az”, “az” exatamente; e também com pequenos adornos (ex: 'AZ:', '(AZ)')
    return bool(re.fullmatch(r"(?i)\s*\(?\s*AZ\s*[:]?\s*\)?\s*", s))

def robust_find_ponto_az(msp, fallback=None):
    """
    Tenta descobrir o Ponto_AZ no DXF, nesta ordem:
      1) TEXT/MTEXT cujo conteúdo seja 'AZ' (case-insensitive).
      2) POINT/CIRCLE em layers sugestivas (AZ, PONTO_AZ, AMARRAÇÃO...).
      3) INSERT de um bloco com nome sugestivo ('AZ', 'PONTO_AZ', etc.).
      4) Fallback informado (ex.: V1) ou None.
    Retorna (x, y) ou None.
    """
    doc = msp.doc
    # 1) TEXT/MTEXT
    try:
        for e in msp.query("TEXT"):
            try:
                if _is_az_label(e.dxf.text):
                    p = e.dxf.insert  # (x, y, z)
                    return (float(p[0]), float(p[1]))
            except Exception:
                pass
        for e in msp.query("MTEXT"):
            try:
                raw = e.text  # mtext content
                if _is_az_label(raw):
                    p = e.dxf.insert
                    return (float(p[0]), float(p[1]))
            except Exception:
                pass
    except Exception:
        # se deu qualquer erro lendo textos, seguimos
        pass

    # 2) POINT/CIRCLE em layers sugestivas
    try:
        for e in msp.query("POINT"):
            try:
                layer = (e.dxf.layer or "").upper()
                if layer in _AZ_LAYER_HINTS:
                    p = e.dxf.location
                    return (float(p[0]), float(p[1]))
            except Exception:
                pass
        for e in msp.query("CIRCLE"):
            try:
                layer = (e.dxf.layer or "").upper()
                if layer in _AZ_LAYER_HINTS:
                    p = e.dxf.center
                    return (float(p[0]), float(p[1]))
            except Exception:
                pass
    except Exception:
        pass

    # 3) INSERT de bloco com nome sugestivo
    try:
        for e in msp.query("INSERT"):
            try:
                name = (e.dxf.name or "").upper()
                layer = (e.dxf.layer or "").upper()
                if (name in _AZ_LAYER_HINTS) or (layer in _AZ_LAYER_HINTS):
                    p = e.dxf.insert
                    return (float(p[0]), float(p[1]))
            except Exception:
                pass
    except Exception:
        pass

    # 4) fallback (ex.: V1)
    return fallback


def azimute_graus_de_norte(dx: float, dy: float) -> float:
    """
    Azimute geodésico em graus 0–360, medido a partir do NORTE (eixo +Y) no sentido horário.
    Para coordenadas projetadas (UTM): dx = x_dest - x_ori, dy = y_dest - y_ori.
    """
    import math
    # atan2 retorna ângulo relativo ao Eixo X; para azimute a partir do Norte, invertemos:
    ang = math.degrees(math.atan2(dx, dy))
    return (ang + 360.0) % 360.0

def distancia_euclidea(p1: tuple[float, float], p2: tuple[float, float]) -> float:
    import math
    return math.hypot(p2[0] - p1[0], p2[1] - p1[1])

def giro_angular_sentido_horario(az_origem: float, az_dest: float) -> float:
    """
    Gera o giro (0–360) em sentido HORÁRIO necessário para ir do az_origem para o az_dest.
    Por convenção do desenho catastral/engenharia, giro horário = (az_dest - az_origem) mod 360.
    """
    return (az_dest - az_origem) % 360.0

def ensure_text_style_STANDARD(msp) -> None:
    """Garante que exista um estilo de texto STANDARD, para compatibilidade DWG/DXF diversos."""
    if "STANDARD" not in msp.doc.styles:
        msp.doc.styles.new("STANDARD")

def ensure_layer(doc, name: str, color: int | None = None):
    """Cria a layer se não existir (opcional: definir cor)."""
    if name not in doc.layers:
        if color is None:
            doc.layers.new(name)
        else:
            doc.layers.new(name, dxfattribs={"color": color})

def add_rotulo(msp, texto: str, xy: tuple[float, float], altura: float = 0.3, layer: str = "Rotulos"):
    # garante estilo/layer
    try:
        ensure_text_style_STANDARD(msp)
    except Exception:
        pass
    try:
        ensure_layer(msp.doc, layer)
    except Exception:
        pass

    # cria o TEXT e posiciona via dxf.insert (compatível com versões sem set_pos)
    txt = msp.add_text(
        texto,
        dxfattribs={
            "height": altura,
            "layer": layer,
            "style": "STANDARD",
            "insert": xy,   # <- POSIÇÃO AQUI
        },
    )

    # Se sua versão tiver set_pos, ótimo; se não, ignora
    try:
        txt.set_pos(xy)  # algumas versões possuem; se não tiver, cai no except
    except Exception:
        pass

    return txt


# ==== HELPERS_ANGULOS_DXF_BEGIN ====
def _log_info(msg):
    try:
        logger.info(msg)
    except Exception:
        print(msg)

def _log_error(msg):
    try:
        logger.error(msg)
    except Exception:
        print(msg)

def _convert_to_dms_safe(graus):
    if "convert_to_dms" in globals():
        return convert_to_dms(graus)
    g = float(graus)
    d = int(g)
    m_f = abs((g - d) * 60.0)
    m = int(m_f)
    s = (m_f - m) * 60.0
    return f"{abs(d):02d}°{m:02d}'{s:06.3f}\"{'S' if g<0 else ''}"

def _angle_deg(dx, dy):
    return (math.degrees(math.atan2(dy, dx)) + 360.0) % 360.0

def _chord_angle(pA, pB):
    return _angle_deg(pB['x'] - pA['x'], pB['y'] - pA['y'])

def _angle_diff_abs(a, b):
    d = (b - a + 540.0) % 360.0 - 180.0
    return abs(d)

def _bulge_tangents_deg(pA, pB, bulge):
    alpha = _chord_angle(pA, pB)
    if abs(bulge) < 1e-12:
        return alpha, alpha
    theta = 4.0 * math.degrees(math.atan(bulge))
    s = 1.0 if bulge > 0 else -1.0
    offset = 90.0 - (abs(theta) / 2.0)
    tan_start = (alpha + s * offset) % 360.0
    tan_end   = (alpha - s * offset) % 360.0
    return tan_start, tan_end

def _polygon_orientation(pts_xyb) -> int:
    # +1 CCW, -1 CW
    area2 = 0.0
    n = len(pts_xyb)
    for i in range(n):
        x1, y1 = pts_xyb[i]['x'], pts_xyb[i]['y']
        x2, y2 = pts_xyb[(i+1) % n]['x'], pts_xyb[(i+1) % n]['y']
        area2 += x1*y2 - x2*y1
    return +1 if area2 > 0 else -1

def _ensure_orientation(points, sentido_desejado):
    pts = list(points)
    orient = _polygon_orientation(pts)
    target = +1 if sentido_desejado == 'anti_horario' else -1
    if orient != target:
        pts.reverse()
        for p in pts:
            p['bulge_next'] = -p.get('bulge_next', 0.0)
    return pts

def _extract_poly_points_with_bulge(doc_dxf):
    msp = doc_dxf.modelspace()
    for e in msp.query("LWPOLYLINE"):
        if not e.closed:
            continue
        pts = []
        for v in list(e):
            x, y = v.dxf.x, v.dxf.y
            bulge_next = float(v.dxf.bulge or 0.0)
            pts.append({'x': x, 'y': y, 'bulge_next': bulge_next})
        return pts
    raise ValueError("Nenhuma LWPOLYLINE fechada encontrada.")

def _chain_lines_closed(msp, tol=1e-6):
    lines = list(msp.query("LINE"))
    if not lines:
        raise ValueError("Não há LWPOLYLINE e nem LINEs no DXF.")
    segs = [((l.dxf.start.x, l.dxf.start.y), (l.dxf.end.x, l.dxf.end.y)) for l in lines]

    def _close(a, b):
        return (abs(a[0]-b[0]) <= tol) and (abs(a[1]-b[1]) <= tol)

    used = [False]*len(segs)
    path = [segs[0][0], segs[0][1]]
    used[0] = True

    changed = True
    while changed:
        changed = False
        for i, (s, e) in enumerate(segs):
            if used[i]:
                continue
            if _close(path[-1], s):
                path.append(e); used[i] = True; changed = True
            elif _close(path[-1], e):
                path.append(s); used[i] = True; changed = True
            elif _close(path[0], e):
                path.insert(0, s); used[i] = True; changed = True
            elif _close(path[0], s):
                path.insert(0, e); used[i] = True; changed = True

    if not _close(path[0], path[-1]):
        raise ValueError("LINEs não formam anel fechado (ou tolerância insuficiente).")

    path = path[:-1]  # remove duplicado final
    pts = [{'x': x, 'y': y, 'bulge_next': 0.0} for (x, y) in path]
    return pts

def _ensure_poly_from_dxf(doc_dxf):
    msp = doc_dxf.modelspace()
    try:
        pts = _extract_poly_points_with_bulge(doc_dxf)
        return pts, None
    except Exception:
        _log_info("Nenhuma LWPOLYLINE fechada; tentando unir LINEs...")
        pts = _chain_lines_closed(msp)
        xy = [(p['x'], p['y']) for p in pts]
        e = msp.add_lwpolyline(xy, format="xy", dxfattribs={"closed": True})
        _log_info("LWPOLYLINE criada a partir de LINEs.")
        return pts, e

def _internal_angles_and_concavity(pts_xyb, sentido_poligonal):
    import math
    n = len(pts_xyb)
    if n < 3:
        return [], []

    s = _polygon_orientation(pts_xyb)  # +1 CCW, -1 CW
    internos_deg = []
    concavo = []

    EPS = 1e-12
    for i in range(n):
        p_prev = pts_xyb[(i-1) % n]
        p      = pts_xyb[i]
        p_next = pts_xyb[(i+1) % n]

        ux = p['x'] - p_prev['x']
        uy = p['y'] - p_prev['y']
        vx = p_next['x'] - p['x']
        vy = p_next['y'] - p['y']

        # evita degenerações
        nu = math.hypot(ux, uy)
        nv = math.hypot(vx, vy)
        if nu < EPS or nv < EPS:
            internos_deg.append(0.0)
            concavo.append(False)
            continue
        ux /= nu; uy /= nu
        vx /= nv; vy /= nv

        cross = ux*vy - uy*vx
        dot   = ux*vx + uy*vy
        theta = math.atan2(cross, dot)  # (-pi, pi]

        internal = math.pi - s*theta
        # normaliza para [0, 2pi)
        while internal < 0:
            internal += 2*math.pi
        while internal >= 2*math.pi:
            internal -= 2*math.pi

        internos_deg.append(math.degrees(internal))

        # côncavo se o giro tiver sinal oposto à orientação
        concavo.append((s*theta) < 0)

    return internos_deg, concavo

def _draw_internal_angles(msp, points, internos_deg, sentido_poligonal, raio_frac=0.10):
    n = len(points)
    for i in range(n):
        p1 = points[i - 1]
        p2 = points[i]
        p3 = points[(i + 1) % n]

        ang_in  = _chord_angle(p2, p1)
        ang_out = _chord_angle(p2, p3)
        interno = internos_deg[i]

        lado1 = math.hypot(p1['x'] - p2['x'], p1['y'] - p2['y'])
        lado2 = math.hypot(p3['x'] - p2['x'], p3['y'] - p2['y'])
        raio = max(0.01, raio_frac * min(lado1, lado2))

        if sentido_poligonal == 'anti_horario':
            start = ang_out
        else:
            start = ang_in
        end = (start + interno) % 360.0
        mid = (start + interno / 2.0) % 360.0

        try:
            if "Internal_Arcs" not in msp.doc.layers:
                msp.doc.layers.add("Internal_Arcs")
        except Exception:
            pass

        msp.add_arc(
            center=(p2['x'], p2['y']),
            radius=raio,
            start_angle=start,
            end_angle=end,
            dxfattribs={'layer': 'Internal_Arcs'}
        )

        pos = (
            p2['x'] + (raio + 1.0) * math.cos(math.radians(mid)),
            p2['y'] + (raio + 1.0) * math.sin(math.radians(mid)),
        )
        texto = _convert_to_dms_safe(interno)
        try:
            if "Labels" not in msp.doc.layers:
                msp.doc.layers.add("Labels")
        except Exception:
            pass

        msp.add_text(
            texto,
            dxfattribs={
                'height': 0.7,
                'layer': 'Labels',
                'insert': pos,
                'rotation': mid if mid <= 180 else mid - 180
            }
        )

def _dist2(a, b):
    return (a[0]-b[0])**2 + (a[1]-b[1])**2

def escolher_ponto_az_externo(v1_xy, ponto_az_dxf, pontos_aberta, ponto_amarracao):
    """
    Decide o 'ponto AZ externo' a usar nos cálculos/desenho:
    - Prioriza ponto Az do DXF (se existir e ≠ V1),
    - senão pega o ponto externo mais próximo de V1 vindo da poligonal ABERTA,
    - senão usa o ponto de amarração,
    - senão retorna None.
    Todos como tuplas (x,y,0.0).
    """
    EPS2 = 1e-12
    vx, vy = float(v1_xy[0]), float(v1_xy[1])

    # 1) Az do DXF
    if ponto_az_dxf is not None:
        ax, ay = float(ponto_az_dxf[0]), float(ponto_az_dxf[1])
        if _dist2((ax, ay), (vx, vy)) > EPS2:
            return (ax, ay, 0.0)

    # 2) ponto externo mais próximo (ABERTA)
    if pontos_aberta:
        # cada item pode ser (x,y) ou (x,y,...) → normalize
        candidatos = []
        for p in pontos_aberta:
            px, py = float(p[0]), float(p[1])
            if _dist2((px, py), (vx, vy)) > EPS2:
                candidatos.append((px, py))
        if candidatos:
            px, py = min(candidatos, key=lambda q: _dist2(q, (vx, vy)))
            return (px, py, 0.0)

    # 3) ponto de amarração
    if ponto_amarracao is not None:
        px, py = float(ponto_amarracao[0]), float(ponto_amarracao[1])
        if _dist2((px, py), (vx, vy)) > EPS2:
            return (px, py, 0.0)

    return None

def add_distance_label(msp, p1, p2, distancia):
    try:
        mid = ((p1[0]+p2[0])/2.0, (p1[1]+p2[1])/2.0)
        msp.add_text(
            f"{distancia:,.2f}".replace(",", "").replace(".", ","),
            dxfattribs={"height": 0.25, "layer": "Azimute", "insert": (mid[0], mid[1])}
        )
        logger.info("Rótulo de distância adicionado com sucesso.")
    except Exception as e:
        logger.error(f"Erro ao adicionar rótulo de distância: {e}")


# ==== HELPERS_ANGULOS_DXF_END ====
#DAQUI PARA BAIXO HELPERS RELATIVOS A EXISTENCIA DE BULGE NA POLIGONAL




def _deg(a_rad): 
    return math.degrees(a_rad)

def _rad(a_deg):
    return math.radians(a_deg)

def _norm_deg(a):
    """Normaliza para (-180, 180]."""
    a = (a + 180.0) % 360.0 - 180.0
    return 180.0 if abs(a + 180.0) < 1e-12 else a

def _bearing(p, q):
    """Azimute da corda PQ em graus (0°=E, CCW)."""
    return _deg(math.atan2(q[1] - p[1], q[0] - p[0]))

def _theta_from_bulge(b):
    """Ângulo central do arco (ASSINADO) em graus."""
    return _deg(4.0 * math.atan(b))

def _tangent_dir_at_start(p, q, bulge):
    """
    Direção da TANGENTE no ponto inicial (p) do segmento/ARCO p→q.
    - Linha: direção da corda.
    - Arco: φ + 90° - θ/2   (θ assinado; CCW>0, CW<0)
    """
    phi = _bearing(p, q)
    if abs(bulge) < EPS_BULGE:
        return phi  # linha
    theta = _theta_from_bulge(bulge)
    return phi + 90.0 - (theta / 2.0)

def _tangent_dir_at_end(p, q, bulge):
    """
    Direção da TANGENTE no ponto final (q) do segmento/ARCO p→q.
    - Linha: direção da corda.
    - Arco: φ + 90° + θ/2
    """
    phi = _bearing(p, q)
    if abs(bulge) < EPS_BULGE:
        return phi  # linha
    theta = _theta_from_bulge(bulge)
    return phi + 90.0 + (theta / 2.0)

def safe_add_giro_angular(msp, doc, v1_pt, ponto_az_pt, v2_pt):
    """
    Adapta a chamada do helper de giro angular para as duas variantes:
    - add_giro_angular_arc_to_dxf(doc, v1, ponto_az, v2)   # variante antiga (doc)
    - add_giro_angular_arc_to_dxf(msp, v1, ponto_az, v2)   # variante nova (msp)
    """
    if 'add_giro_angular_arc_to_dxf' not in globals():
        return

    try:
        # tenta assinatura nova (msp, ...)
        add_giro_angular_arc_to_dxf(msp, v1_pt, ponto_az_pt, v2_pt)
    except AttributeError as e:
        # se dentro do helper tentaram usar msp.modelspace(), é porque esperavam doc
        try:
            add_giro_angular_arc_to_dxf(doc, v1_pt, ponto_az_pt, v2_pt)
        except Exception:
            raise
    except TypeError:
        # se a assinatura claramente exige doc, tenta doc
        add_giro_angular_arc_to_dxf(doc, v1_pt, ponto_az_pt, v2_pt)


def _internal_angles_with_bulge(points_bulge):
    """
    Calcula ângulos internos (0–360) em todos os vértices usando tangentes reais.
    `points_bulge`: lista de dicts [{'x':..,'y':..,'bulge_next':..}, ...]
    Retorna lista em graus (float).
    """
    n = len(points_bulge)
    angs = []
    for i in range(n):
        # índices circularmente
        i_prev = (i - 1) % n
        i_next = (i + 1) % n

        p_prev = (points_bulge[i_prev]['x'], points_bulge[i_prev]['y'])
        p_curr = (points_bulge[i]['x'],     points_bulge[i]['y'])
        p_next = (points_bulge[i_next]['x'], points_bulge[i_next]['y'])

        bulge_prev = float(points_bulge[i_prev].get('bulge_next', 0.0))  # do seg (i-1)→i
        bulge_curr = float(points_bulge[i].get('bulge_next', 0.0))        # do seg i→(i+1)

        # direções tangentes no vértice i
        dir_in  = _tangent_dir_at_end(p_prev, p_curr, bulge_prev)   # chegando em Vi
        dir_out = _tangent_dir_at_start(p_curr, p_next, bulge_curr) # saindo de Vi

        # giro assinado (esquerda + / direita -)
        turn = _norm_deg(dir_out - dir_in)

        # ângulo interno da poligonal (concavo pode passar de 180)
        interno = 180.0 - turn
        # normaliza para [0, 360)
        if interno < 0.0:
            interno += 360.0
        elif interno >= 360.0:
            interno -= 360.0

        angs.append(interno)
    return angs

def create_memorial_descritivo(
    uuid_str, doc, lines, proprietario, matricula, caminho_salvar, confrontantes, ponto_az,
    dxf_file_path, area_dxf, azimute, v1, msp, dxf_filename, excel_file_path, tipo,
    giro_angular_v1_dms, distancia_az_v1, sentido_poligonal='horario', modo="ANGULO_P1_P2",
    diretorio_concluido=None,
    points_bulge=None
):
    """
    DXF já vem tratado do pipeline (sem reler aqui):
    - Usa LWPOLYLINE fechada (com bulge) fornecida por get_document_info_from_dxf.
    - Normaliza sentido, calcula ângulos internos (com concavidade), desenha arcos internos.
    - Só desenha AZ no modo ANGULO_AZ.
    - Gera Excel diretamente.
    """
    logger.info("[CMD] pontos_bulge recebidos: %s", len(points_bulge) if points_bulge else 0)

    # Garante que o estilo de texto "STANDARD" exista no DXF
    if "STANDARD" not in msp.doc.styles:
        msp.doc.styles.new("STANDARD")

    if diretorio_concluido is None:
        diretorio_concluido = caminho_salvar

    dxf_output_path = os.path.join(
        diretorio_concluido,
        f"{uuid_str}_FECHADA_{tipo}_{matricula}.dxf"
    )

    # 0) valida base
    if points_bulge is None or len(points_bulge) < 3:
        _log_error(f"[AZ] points_bulge ausente/insuficiente (type={type(points_bulge)}, len={0 if not points_bulge else len(points_bulge)}). Verifique get_document_info_from_dxf.")
        return None

    # 1) normaliza sentido
    pts = _ensure_orientation(points_bulge, sentido_poligonal)
    orient = _polygon_orientation(pts)
    _log_info(f"Sentido normalizado: {'anti-horário' if orient == +1 else 'horário'}")

    # 2) ângulos internos + concavidade
    EPS_BULGE = 1e-9

    has_any_bulge = any(abs(float(p.get('bulge_next', 0.0))) > EPS_BULGE for p in pts)

    if has_any_bulge:
        internos_deg = _internal_angles_with_bulge(pts)  # usa tangentes reais (funciona também para bulge=0)
        concavo = [a > 180.0 for a in internos_deg]      # se você quiser a flag de concavidade
        logger.info("Ângulos internos: modo BULGE-AWARE (misto retas+arcos).")
    else:
        # compatibilidade com sua rotina antiga quando não há bulge algum
        internos_deg, concavo = _internal_angles_and_concavity(pts, sentido_poligonal)
        logger.info("Ângulos internos: modo LEGADO (somente retas).")

    # 3) desenha arcos internos por dentro
    _draw_internal_angles(msp, pts, internos_deg, sentido_poligonal, raio_frac=0.10)

    # 4) desenho do AZ depende do modo
    # ANGULO_AZ  → desenha Az, linha Az–V1, arco e rótulos
    # ANGULO_P1_P2 → NÃO desenha Az/linha/arco (poligonal ABERTA já mostra amarração)
    if modo == "ANGULO_AZ" and ponto_az is not None and v1 is not None:
        dx = v1[0] - ponto_az[0]
        dy = v1[1] - ponto_az[1]
        dist = math.hypot(dx, dy)
        if dist > 1e-6:
            try:
                _desenhar_referencia_az(msp, ponto_az, v1, azimute)
            except Exception as e:
                logger.error("Erro ao desenhar referência de Az: %s", e)
        else:
            logger.warning("⚠️ Distância Az–V1 ≈ 0; desenho do Az suprimido.")

    # 5) Excel (sem reler nada)
    try:
        ordered_points_xy = [(p['x'], p['y']) for p in pts]
        total_pontos = len(ordered_points_xy)
        data = []

        if ponto_az is not None:
            ponto_az_e = f"{ponto_az[0]:,.3f}".replace(",", "").replace(".", ",")
            ponto_az_n = f"{ponto_az[1]:,.3f}".replace(",", "").replace(".", ",")
        else:
            ponto_az_e = ""
            ponto_az_n = ""

        for i in range(total_pontos):
            p2 = ordered_points_xy[i]
            p3 = ordered_points_xy[(i + 1) % total_pontos]

            dx, dy = p3[0] - p2[0], p3[1] - p2[1]
            distance = math.hypot(dx, dy)
            description = f"V{i + 1}_V{(i + 2) if i + 1 < total_pontos else 1}"
            confrontante = confrontantes[i % len(confrontantes)] if confrontantes else ""
            ang_interno_dms = _convert_to_dms_safe(internos_deg[i])

            if i == 0:
                distancia_az_v1_str = f"{float(distancia_az_v1):.2f}".replace(".", ",") if distancia_az_v1 is not None else ""
                azimute_az_v1_str   = _convert_to_dms_safe(float(azimute)) if azimute is not None else ""
                giro_v1_str         = giro_angular_v1_dms or ""
                p_az_e, p_az_n      = ponto_az_e, ponto_az_n
            else:
                distancia_az_v1_str = ""
                azimute_az_v1_str   = ""
                giro_v1_str         = ""
                p_az_e, p_az_n      = "", ""

            data.append({
                "V": f"V{i + 1}",
                "E": f"{p2[0]:,.3f}".replace(",", "").replace(".", ","),
                "N": f"{p2[1]:,.3f}".replace(",", "").replace(".", ","),
                "Z": "0,000",
                "Divisa": description,
                "Angulo Interno": ang_interno_dms,
                "Distancia(m)": f"{distance:,.2f}".replace(",", "").replace(".", ","),
                "Confrontante": confrontante,
                "ponto_AZ_E": p_az_e,
                "ponto_AZ_N": p_az_n,
                "distancia_Az_V1": distancia_az_v1_str,
                "Azimute Az_V1": azimute_az_v1_str,
                "Giro Angular Az_V1_V2": giro_v1_str
            })

            try:
                if distance > 0.01 and 'add_label_and_distance' in globals():
                    add_label_and_distance(msp, p2, p3, f"V{i + 1}", distance)
            except Exception as e:
                _log_error(f"Falha ao rotular distância do lado V{i+1}: {e}")

        
        # escreve excel
        df = pd.DataFrame(data)

        # ── Garantir as 3 colunas do ANGULO_AZ antes de salvar
        cols_novas = ["AZIMUTE_AZ_V1_GRAUS", "DISTANCIA_AZ_V1_M", "GIRO_V1_GRAUS"]
        for c in cols_novas:
            if c not in df.columns:
                df[c] = ""  # ou pd.NA

        # Garante diretório e salva
        try:
            os.makedirs(os.path.dirname(excel_file_path), exist_ok=True)
            df.to_excel(excel_file_path, index=False)
            _log_info(f"Excel escrito (primeira passagem): {os.path.abspath(excel_file_path)}")
        except Exception as e:
            _log_error(f"Falha ao salvar Excel na primeira passagem: {e}")
            raise  # deixe a main ver o stacktrace

        # Formatação openpyxl
        try:
            wb = openpyxl.load_workbook(excel_file_path)
            ws = wb.active

            for cell in ws[1]:
                cell.font = Font(bold=True)
                cell.alignment = Alignment(horizontal="center", vertical="center")

            col_widths = {
                "A": 8, "B": 15, "C": 15, "D": 0, "E": 15,
                "F": 15, "G": 15, "H": 50, "I": 15,
                "J": 15, "K": 15, "L": 20, "M": 20
            }
            for col, width in col_widths.items():
                ws.column_dimensions[col].width = width

            for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=1, max_col=ws.max_column):
                for cell in row:
                    cell.alignment = Alignment(horizontal="center", vertical="center")

            wb.save(excel_file_path)
            _log_info(f"Excel salvo e formatado: {os.path.abspath(excel_file_path)}")
        except Exception as e:
            _log_error(f"Falha ao formatar/salvar Excel com openpyxl: {e}")
            raise

        # Confirma existência
        if os.path.exists(excel_file_path):
            _log_info(f"✅ Excel confirmado em disco: {os.path.abspath(excel_file_path)}")
        else:
            _log_error(f"❌ Excel NÃO encontrado após salvar: {os.path.abspath(excel_file_path)}")


        # extras DXF (opcionais e seguros)
        try:
            v1_pt = ordered_points_xy[0]
            v2_pt = ordered_points_xy[1]
            # se existir o helper e você quiser o giro no V1 com Az:
            if 'add_giro_angular_arc_to_dxf' in globals() and ponto_az is not None:
                # padronize este helper para (msp, v1_pt, ponto_az, v2_pt)
                add_giro_angular_arc_to_dxf(msp, v1_pt, ponto_az, v2_pt)
                _log_info("Giro horário Az–V1–V2 adicionado com sucesso.")
        except Exception as e:
            _log_error(f"Erro ao adicionar giro angular: {e}")

        try:
            if "Vertices" not in msp.doc.layers:
                msp.doc.layers.add("Vertices")
        except Exception:
            pass

        # garanta a camada
        try:
            if "Vertices" not in msp.doc.layers:
                msp.doc.layers.add("Vertices")
        except Exception:
            pass

        for i, (x, y) in enumerate(ordered_points_xy):
            try:
                msp.add_circle(center=(x, y), radius=0.5, dxfattribs={"layer": "Vertices"})
                msp.add_text(
                    f"V{i + 1}",
                    dxfattribs={
                        "height": 0.3,
                        "layer": "Vertices",
                        "insert": (x + 0.30, y + 0.30)  # <<< POSIÇÃO DO RÓTULO
                    }
                )
            except Exception as e:
                logger.warning(f"Falha rotulando V{i+1}: {e}")


        # só desenhe o arco do azimute se realmente quiser no produto FECHADA
        # e se houver amarração (Az) válida:
        if modo == "ANGULO_AZ" and ponto_az is not None:
            try:
                azim = calculate_azimuth(ponto_az, v1_pt)
                _desenhar_referencia_az(msp, ponto_az, v1_pt, azim)
                _log_info("Arco do Azimute Az–V1 adicionado com sucesso.")
            except Exception as e:
                _log_error(f"Erro ao adicionar arco do azimute: {e}")

        # 6) salvar DXF final
        try:
            doc.saveas(dxf_output_path)
            logger.info("✅ DXF FECHADA salvo corretamente: %s", dxf_output_path)
        except Exception as e:
            logger.error("Erro ao salvar DXF FECHADA: %s", e)

    except Exception as e:
        _log_error(f"❌ Erro ao gerar o memorial descritivo: {e}")
        return None

    return excel_file_path











def generate_initial_text(proprietario, matricula, descricao, area, perimeter, rua, cidade, ponto_az, azimute, distancia):
    """
    Gera o texto inicial do memorial descritivo.
    """
    initial_text = (
        f"MEMORIAL DESCRITIVO\n"
        f"NOME PROPRIETÁRIO / OCUPANTE: {proprietario}\n"
        f"DESCRIÇÃO: {descricao}\n"
        f"DOCUMENTAÇÃO: MATRÍCULA {matricula}\n"
        f"ÁREA DO IMÓVEL: {area:.2f} metros quadrados\n"
        f"PERÍMETRO: {perimeter:.4f} metros\n"
        f"Área localizada na rua {rua}, município de {cidade}, com a seguinte descrição:\n"
        f"O Ponto Az está localizado nas coordenadas E {ponto_az[0]:.3f}, N {ponto_az[1]:.3f}.\n"
        f"Daí, com Azimute de {azimute} e distância de {distancia:.2f} metros, chega-se ao Vértice V1, "
        f"origem da área descrição, alinhado com a rua {rua}."
    )
    return initial_text


def generate_angular_text(az_v1, v1_v2, distancia, angulo, rua, confrontante):
    """
    Gera o texto do ângulo em V1 entre Az-V1 e V1-V2.
    """
    angular_text = (
        f"Daí, visando o Ponto “Az”, com giro angular horário de {angulo} e diste {distancia:.2f} metros, "
        f"chega-se ao Vértice V2, também alinhado com a rua {rua} e limítrofe com {confrontante}."
    )
    return angular_text


def generate_recurring_text(df, rua, confrontantes):
    """
    Gera o texto recorrente enquanto percorre os vértices da poligonal.
    """
    recurring_texts = []
    num_vertices = len(df)
    
    for i in range(1, num_vertices - 1):  # De V2 até o penúltimo vértice
        current = df.iloc[i]
        next_vertex = df.iloc[i + 1]
        confrontante = confrontantes[i % len(confrontantes)]
        
        recurring_text = (
            f"Daí, visando o Vértice {current['V']}, com giro angular horário de {current['Angulo Interno']} "
            f"e diste {current['Distancia(m)']} metros, chega-se ao Vértice {next_vertex['V']}, "
            f"também alinhado à rua {rua} e limítrofe com {confrontante}."
        )
        recurring_texts.append(recurring_text)
    
    return recurring_texts


def generate_final_text(df, rua, confrontantes):
    """
    Gera o texto final ao retornar ao V1.
    """
    last_vertex = df.iloc[-1]
    first_vertex = df.iloc[0]
    confrontante = confrontantes[-1 % len(confrontantes)]  # Confrontante do último trecho
    
    final_text = (
        f"Daí, visando o vértice {last_vertex['V']}, com giro angular de {last_vertex['Angulo Interno']} "
        f"e diste {last_vertex['Distancia(m)']} metros, chega-se ao vértice {first_vertex['V']}, "
        f"origem da presente descrição, no alinhamento da rua {rua} e próximo aos lotes de {confrontante}."
    )
    return final_text


def create_memorial_document(
    uuid_str,
    proprietario,
    matricula,
    matricula_texto,
    area_total,
    cpf,
    rgi,
    excel_file_path,
    template_path,
    output_path,
    assinatura_path,
    ponto_amarracao,
    azimute,
    distancia_amarracao_v1,
    rua,
    cidade,
    confrontantes,
    area_dxf,
    desc_ponto_amarracao,
    perimeter_dxf,
    giro_angular_v1_dms
):

    try:
        df = pd.read_excel(excel_file_path, engine='openpyxl', dtype=str)

        # Carrega confrontantes diretamente da coluna do Excel
        confrontantes = df['Confrontante'].dropna().tolist()

        # Processa colunas numéricas
        df['Distancia(m)'] = df['Distancia(m)'].str.replace(',', '.').astype(float)
        df['E'] = df['E'].str.replace(',', '').astype(float)
        df['N'] = df['N'].str.replace(',', '').astype(float)

        # Calcular perímetro e área
        perimeter = df['Distancia(m)'].sum()
        x = df['E'].values
        y = df['N'].values
        area = abs(sum(x[i] * y[(i + 1) % len(x)] - x[(i + 1) % len(x)] * y[i] for i in range(len(x))) / 2)

        doc_word = Document(template_path)

        # 🔴 Remover linhas vazias ou parágrafos indesejados no topo
        while doc_word.paragraphs and not doc_word.paragraphs[0].text.strip():
            doc_word.paragraphs[0]._element.getparent().remove(doc_word.paragraphs[0]._element)

        for para in doc_word.paragraphs:
            if "copilot" in para.text.lower():
                para._element.getparent().remove(para._element)

        set_default_font(doc_word)

        # Título
        p = doc_word.add_paragraph(style='Normal')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.add_run("MEMORIAL DESCRITIVO").bold = True

        doc_word.add_paragraph()

        # Parágrafos descritivos
        p = doc_word.add_paragraph(style='Normal')
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.add_run("Objetivo: ").bold = True
        p.add_run(f"Área destinada à servidão de passagem para execução de coletor de fundo pertencente à rede coletora de esgoto de {cidade}/RS.")

        p = doc_word.add_paragraph(style='Normal')
        p.add_run("Matrícula Número: ").bold = True
        p.add_run(f"{matricula} - {rgi}")

        area_total_formatada = f"{area_dxf:.2f}".replace(".", ",")
        p = doc_word.add_paragraph(style='Normal')
        p.add_run("Área Total do Terreno: ").bold = True
        p.add_run(area_total_formatada)

        # p = doc_word.add_paragraph(style='Normal')
        # p.add_run("Proprietário: ").bold = True
        # p.add_run(f"{proprietario} - CPF/CNPJ: {cpf}")

        p = doc_word.add_paragraph(style='Normal')
        p.add_run("Área de Servidão de Passagem: ").bold = True
        p.add_run(f"{area_dxf:.2f}".replace(".", ",") + " m")
        sup = p.add_run("2")
        sup.font.superscript = True
        sup.font.size = Pt(12)

        doc_word.add_paragraph()

        p = doc_word.add_paragraph(style='Normal')
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        #p.add_run("Descrição: ").bold = True
        p.add_run("Área com ").font.name = 'Arial'

        run1 = p.add_run(f"{area_dxf:.2f}".replace(".", ",")+" m")
        run1.font.name = 'Arial'
        run1.font.size = Pt(12)

        run2 = p.add_run("2")
        run2.font.name = 'Arial'
        run2.font.size = Pt(12)
        run2.font.superscript = True

        p.add_run(f" localizada na {rua}, município de {cidade},com a finalidade de servidão de passagem com a seguinte descrição e confrontações, onde os ângulos foram medidos no sentido horário.").font.name = 'Arial'

        doc_word.add_paragraph()
        doc_word.add_paragraph("Pontos definidos pelas Coordenadas Planas no Sistema U.T.M. – SIRGAS 2000.", style='Normal')
        doc_word.add_paragraph()

        # Coordenadas do ponto Az
        ponto_az_1 = f"{Coorde_E_ponto_Az:.2f}".replace(".", ",")
        ponto_az_2 = f"{Coorde_N_ponto_Az:.2f}".replace(".", ",")


        azimute_dms = convert_to_dms(azimuth)
        distancia_str = f"{distance:.2f}".replace(".", ",")

        # Linha: ponto de amarração
        p = doc_word.add_paragraph(style='Normal')
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        p.add_run("O ponto ")
        p.add_run("Az").bold = True
        p.add_run(f", ponto de amarração, está localizado na {desc_ponto_Az} nas coordenadas E(X) {ponto_az_1} e N(Y) {ponto_az_2}.")

        p.paragraph_format.space_after = Pt(12)  # ⬅️ Força um espaçamento abaixo do parágrafo

        
        # Linha: Azimute e distância até V1
        p = doc_word.add_paragraph(style='Normal')
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.add_run("Daí, com Azimute de ").bold = False
        p.add_run(azimute_dms).bold = False
        p.add_run(f" e distância de {distancia_str} metros, chega-se ao vértice ")
        p.add_run("V1").bold = True
        p.add_run(", origem da área descrição, alinhado com a rua " + rua + ".")
        p.paragraph_format.space_after = Pt(12)  # ⬅️ FORÇA espaçamento após esse parágrafo



        # ➤ Percorrer vértices
        for i in range(len(df)):
            current = df.iloc[i]
            next_vertex = df.iloc[(i + 1) % len(df)]
            distancia = f"{current['Distancia(m)']:.2f}".replace(".", ",")
            confrontante = current['Confrontante']
            giro_angular = current['Angulo Interno']

            p = doc_word.add_paragraph(style='Normal')
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            p.add_run("Do vértice ").bold = False
            p.add_run(current['V']).bold = True

            if i == 0:
                p.add_run(
                    f", com giro angular horário de {giro_angular_v1_dms} e distância de {distancia} metros, "
                    f"confrontando com área pertencente à {confrontante}, chega-se ao vértice "
                )
            elif next_vertex['V'] == "V1" and i == len(df) - 1:
                p.add_run(
                    f", com giro angular horário de {giro_angular} e distância de {distancia} metros, "
                    f"confrontando com área pertencente à {confrontante}, chega-se ao vértice "
                )
                p.add_run(next_vertex['V']).bold = True
                p.add_run(", origem da presente descrição.")
                doc_word.add_paragraph()
                break
            else:
                p.add_run(
                    f", com giro angular horário de {giro_angular} e distância de {distancia} metros, "
                    f"confrontando com área pertencente à {confrontante}, chega-se ao vértice "
                )

            p.add_run(next_vertex['V']).bold = True
            p.add_run(";")
            doc_word.add_paragraph()

        # Parágrafos descritivos
        p = doc_word.add_paragraph(style='Normal')
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.add_run(f"Os angulos foram medidos no sentido horário.")

        data_atual = datetime.now().strftime("%d de %B de %Y")

        # converte mês para português
        for ingles, portugues in MESES_PT_BR.items():
            if ingles in data_atual:
                data_atual = data_atual.replace(ingles, portugues)
                break
        doc_word.add_paragraph(f"\nPorto Alegre, RS, {data_atual}.", style='Normal')
        doc_word.add_paragraph("\n\n")
        doc_word.save(output_path)

        print(f"Memorial descritivo salvo em: {output_path}")

    except Exception as e:
        print(f"Erro ao criar o documento memorial: {e}")




def find_excel_file(directory, keywords):
    """
    Busca um arquivo Excel no diretório contendo todas as palavras-chave no nome.
    Se não encontrar, exibe a lista de arquivos disponíveis.
    """
    if not directory or not os.path.exists(directory):
        print(f"Erro: O diretório '{directory}' não existe ou não foi especificado corretamente.")
        return None

    excel_files = [file for file in os.listdir(directory) if file.endswith(".xlsx")]

    if not excel_files:
        print(f"Nenhum arquivo Excel encontrado no diretório: {directory}")
        return None

    for file in excel_files:
        if all(keyword.lower() in file.lower() for keyword in keywords):
            return os.path.join(directory, file)

    # Se nenhum arquivo correspondente foi encontrado, listar os arquivos disponíveis
    print(f"Nenhum arquivo Excel contendo {keywords} foi encontrado em '{directory}'.")
    print("Arquivos disponíveis no diretório:")
    for f in excel_files:
        print(f"  - {f}")

    return None



#função não pode ser usada para LINUX       
# def convert_docx_to_pdf(output_path, pdf_file_path):
#     """
#     Converte um arquivo DOCX para PDF usando a biblioteca comtypes.
#     """
#     try:
#         # Verificar se o arquivo DOCX existe antes de abrir
#         if not os.path.exists(output_path):
#             raise FileNotFoundError(f"Arquivo DOCX não encontrado: {output_path}")
        
#         print(f"Tentando converter o arquivo DOCX: {output_path} para PDF: {pdf_file_path}")

#         word = comtypes.client.CreateObject("Word.Application")
#         word.Visible = False  # Ocultar a interface do Word
#         doc = word.Documents.Open(output_path)
        
#         # Aguardar alguns segundos antes de salvar como PDF
#         import time
#         time.sleep(2)

#         #doc.SaveAs(pdf_file_path, FileFormat=17)  # 17 corresponde ao formato PDF
#         doc.Close()
#         word.Quit()
#         #print(f"Arquivo PDF salvo com sucesso em: {pdf_file_path}")
#     except FileNotFoundError as fnf_error:
#         print(f"Erro: {fnf_error}")
#     except Exception as e:
#         print(f"Erro ao converter DOCX para PDF: {e}")
#     finally:
#         try:
#             word.Quit()
#         except:
#             pass  # Garantir que o Word seja fechado



def sanitize_filename(filename):
    return re.sub(r'[<>:"/\\|?*]', '', filename)

        
def main_poligonal_fechada(uuid_str, excel_path, dxf_path, diretorio_preparado, diretorio_concluido, caminho_template, sentido_poligonal='horario'):

    caminho_salvar = diretorio_concluido 
    template_path = caminho_template 
    # Carrega dados do imóvel
    dados_imovel_excel_path = excel_path
    dados_imovel_df = pd.read_excel(dados_imovel_excel_path, sheet_name='Dados_do_Imóvel', header=None)
    dados_imovel = dict(zip(dados_imovel_df.iloc[:, 0], dados_imovel_df.iloc[:, 1]))

    # Extrai informações
    proprietario = dados_imovel.get("NOME DO PROPRIETÁRIO", "").strip()
    cpf = dados_imovel.get("CPF/CNPJ", "").strip()
    matricula = sanitize_filename(str(dados_imovel.get("DOCUMENTAÇÃO DO IMÓVEL", "")).strip())
    matricula_texto = str(dados_imovel.get("DOCUMENTAÇÃO DO IMÓVEL", "")).strip()
    descricao = dados_imovel.get("OBRA", "").strip()
    area_total = dados_imovel.get("ÁREA TOTAL DO TERRENO DOCUMENTADA", "").replace("\t", "").replace("\n", "").strip()
    cidade = dados_imovel.get("CIDADE", "").strip().capitalize()
    rgi= dados_imovel.get("RGI", "").strip().capitalize()
    rua = dados_imovel.get("LOCAL", "").strip()
    desc_ponto_Az = dados_imovel.get("AZ", "").strip()

    # Diretório para salvar resultados
    
    os.makedirs(caminho_salvar, exist_ok=True)

    # Identifica tipo (SER, REM, etc)
    dxf_filename = os.path.basename(dxf_path).upper()
    if "ETE" in dxf_filename:
        tipo = "ETE"
    elif "REM" in dxf_filename:
        tipo = "REM"
    elif "SER" in dxf_filename:
        tipo = "SER"
    elif "ACE" in dxf_filename:
        tipo = "ACE"
    else:
        logger.info("❌ Não foi possível determinar automaticamente o tipo (ETE, REM, SER ou ACE).")
        return

    padrao_fechada = os.path.join(diretorio_preparado, f"{uuid_str}_FECHADA_{tipo}*.xlsx")

    arquivos_encontrados = glob.glob(padrao_fechada)
    if not arquivos_encontrados:
        logger.info(f"❌ Arquivo de confrontantes não encontrado com o padrão: {padrao_fechada}")
        return
    confrontantes_df = pd.read_excel(arquivos_encontrados[0])
    confrontantes = confrontantes_df.iloc[:, 1].dropna().tolist()

    # DXF limpo
    # ⚠️ Substitui a limpeza anterior por apenas conversão R2010
    dxf_limpo_path = os.path.join(caminho_salvar, f"DXF_LIMPO_{matricula}.dxf")
    dxf_file_path = limpar_dxf_e_converter_r2010(dxf_path, dxf_limpo_path)


    # 📁 Procurar CONCLUIDO dentro da cidade (REPESCAGEM_*/CONCLUIDO)
    # O diretório CONCLUIDO já é passado corretamente
    diretorio_concluido_real = diretorio_concluido

   
    # # 🧭 Obter ponto de amarração anterior ao V1
    # try:
    #     ponto_amarracao, codigo_amarracao = obter_ponto_amarracao_anterior_v1(planilha_aberta_saida)
    #     logger.info(f"📌 Ponto de amarração identificado: {codigo_amarracao} com coordenadas {ponto_amarracao}")
    # except Exception as e:
    #     logger.error(f"❌ Erro ao obter ponto de amarração: {e}")
    #     return

    # 🔍 Extrair geometria do DXF
    # Extrair geometria FECHADA do DXF
    doc, lines, perimeter_dxf, area_dxf, ponto_az_dxf, msp, pts_bulge = get_document_info_from_dxf(dxf_file_path)

    # Pare aqui se não houver geometria válida
    if not (doc and lines):
        logger.info("Nenhuma linha foi encontrada ou não foi possível acessar o documento.")
        pythoncom.CoUninitialize()
        return

    logger.info(f"📐 Área da poligonal: {area_dxf:.6f} m²")

    v1 = lines[0][0]
    v2 = lines[1][0]

    # Use o ponto retornado pela função
    azimute = calculate_azimuth(ponto_az_dxf, v1)
    distancia_az_v1 = calculate_distance(ponto_az_dxf, v1)
    giro_angular_v1 = calculate_angular_turn(ponto_az_dxf, v1, v2)
    giro_angular_v1_dms = convert_to_dms(360 - giro_angular_v1)

    logger.info(f"📌 Azimute Az→V1: {azimute:.4f}°, Distância: {distancia_az_v1:.2f} m")

    # Caminho do Excel de saída
    excel_file_path = os.path.join(
        diretorio_concluido,
        f"{uuid_str}_FECHADA_{tipo}_{matricula}.xlsx"
    )
    logger.info(f"✅ Excel FECHADA salvo corretamente: {excel_file_path}")

    # 🛠 Criar memorial e Excel (passe modo e pts_bulge)
    create_memorial_descritivo(
        uuid_str, doc, lines, proprietario, matricula, caminho_salvar, confrontantes, ponto_az_dxf,
        dxf_file_path, area_dxf, azimute, v1, msp, dxf_filename, excel_file_path, tipo,
        giro_angular_v1_dms, distancia_az_v1, sentido_poligonal=sentido_poligonal,
        modo="ANGULO_P1_P2", points_bulge=pts_bulge
    )

    # 📄 Gerar DOCX (apenas uma vez)
    if excel_file_path:
        output_path_docx = os.path.join(
            diretorio_concluido,
            f"{uuid_str}_FECHADA_{tipo}_{matricula}.docx"
        )
        logger.info(f"✅ DOCX FECHADA salvo corretamente: {output_path_docx}")

        assinatura_path = r"C:\Users\Paulo\Documents\CASSINHA\MEMORIAIS DESCRITIVOS\Assinatura.jpg"
        #desc_ponto_amarracao = f"ponto {codigo_amarracao}, obtido na planilha da poligonal aberta"

        create_memorial_document(
            uuid_str=uuid_str,
            proprietario=proprietario,
            matricula=matricula,
            matricula_texto=matricula_texto,
            area_total=area_total,
            cpf=cpf,
            rgi=rgi,
            excel_file_path=excel_file_path,
            template_path=template_path,
            output_path=output_path_docx,
            assinatura_path=assinatura_path,
            ponto_amarracao=ponto_az_dxf,
            azimute=azimute,
            distancia_amarracao_v1=distancia_az_v1,
            rua=rua,
            cidade=cidade,
            confrontantes=confrontantes,
            area_dxf=area_dxf,
            desc_ponto_amarracao=desc_ponto_Az,
            perimeter_dxf=perimeter_dxf,
            giro_angular_v1_dms=giro_angular_v1_dms,
        )
    else:
        logger.info("excel_file_path não definido ou inválido.")

    logger.info("Documento do AutoCAD fechado.")


#FINAL DO CODIGO

