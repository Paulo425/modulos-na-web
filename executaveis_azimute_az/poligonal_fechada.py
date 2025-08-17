import ezdxf
import math
import csv
import os
import re
import glob
import locale
from docx import Document
from docx.shared import Inches
from datetime import datetime
from decimal import Decimal, getcontext
import pandas as pd
import locale
import openpyxl
from openpyxl.styles import Alignment, Font
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import logging


# Diretório para logs
BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
LOG_DIR = os.path.join(BASE_DIR, 'static', 'logs')
os.makedirs(LOG_DIR, exist_ok=True)

# Arquivo de log específico para poligonal_fechada
log_file = os.path.join(LOG_DIR, f'poligonal_fechada_{datetime.now().strftime("%Y%m%d_%H%M%S")}.log')

# Configuração básica do logger
logger = logging.getLogger(__name__)
logger.setLevel(logging.DEBUG)

file_handler = logging.FileHandler(log_file)
file_handler.setFormatter(logging.Formatter('%(asctime)s [%(levelname)s] %(message)s'))

# Verificar se já não existem handlers para não duplicar
if not logger.handlers:
    logger.addHandler(file_handler)

getcontext().prec = 28  # Define a precisão para 28 casas decimais

MESES_PT_BR = {
    'January': 'janeiro',
    'February': 'fevereiro',
    'March': 'março',
    'April': 'abril',
    'May': 'maio',
    'June': 'junho',
    'July': 'julho',
    'August': 'agosto',
    'September': 'setembro',
    'October': 'outubro',
    'November': 'novembro',
    'December': 'dezembro'
}
def limpar_dxf_e_converter_r2010(original_path, saida_path):
    """
    Lê um DXF original e regrava o arquivo com a versão R2010,
    preservando LWPOLYLINE (com bulge). Não cria geometria nova.
    """
    try:
        doc_antigo = ezdxf.readfile(original_path)
        msp_antigo = doc_antigo.modelspace()
        doc_novo = ezdxf.new(dxfversion='R2010')
        msp_novo = doc_novo.modelspace()

        encontrou_polilinha = False

        for entity in msp_antigo.query('LWPOLYLINE'):
            if entity.closed:
                # Pega (x, y, bulge) como tuplas
                pontos_xyb = entity.get_points('xyb')  # [(x, y, bulge), ...]
                # Regrava preservando o bulge
                msp_novo.add_lwpolyline(
                    pontos_xyb,
                    format='xyb',
                    close=True,
                    dxfattribs={'layer': entity.dxf.layer}
                )
                encontrou_polilinha = True
                break

        if not encontrou_polilinha:
            raise ValueError("Nenhuma polilinha fechada encontrada no DXF original.")

        doc_novo.saveas(saida_path)
        logger.info(f"✅ DXF convertido e salvo como R2010 em: {saida_path}")
        return saida_path

    except Exception as e:
        logger.error(f"❌ Erro ao converter DXF para R2010: {e}")
        return original_path




def get_document_info_from_dxf(dxf_file_path):
    try:
        doc = ezdxf.readfile(dxf_file_path)
        msp = doc.modelspace()

        lines = []
        perimeter_dxf = 0.0
        area_dxf = 0.0
        ponto_az = None
        ordered_points = []
        ordered_points_with_bulge = []

        # Lê a PRIMEIRA LWPOLYLINE fechada
        for entity in msp.query('LWPOLYLINE'):
            if entity.closed:
                pts_xyb = entity.get_points('xyb')  # lista de tuplas (x, y, bulge)
                # Se houver último = primeiro, remove duplicata
                if len(pts_xyb) >= 2 and (pts_xyb[0][0], pts_xyb[0][1]) == (pts_xyb[-1][0], pts_xyb[-1][1]):
                    pts_xyb = pts_xyb[:-1]

                # Monta listas
                for (x, y, b) in pts_xyb:
                    ordered_points.append((x, y))
                    ordered_points_with_bulge.append({'x': x, 'y': y, 'bulge_next': float(b or 0.0)})

                logger.info(
                    f"pts_bulge: n={len(ordered_points_with_bulge)} | exemplo={ordered_points_with_bulge[:2]}"
                )

                # Linhas + perímetro
                n = len(ordered_points)
                for i in range(n):
                    p1 = ordered_points[i]
                    p2 = ordered_points[(i + 1) % n]
                    lines.append((p1, p2))
                    perimeter_dxf += ((p2[0] - p1[0]) ** 2 + (p2[1] - p1[1]) ** 2) ** 0.5

                # Área por shoelace
                x = [p[0] for p in ordered_points]
                y = [p[1] for p in ordered_points]
                area_dxf = abs(sum(x[i] * y[(i + 1) % n] - x[(i + 1) % n] * y[i] for i in range(n)) / 2.0)
                break

        if not lines:
            logger.info("Nenhuma polilinha fechada encontrada no arquivo DXF.")
            return None, [], 0.0, 0.0, None, None, []

        # Ponto Az (TEXT → INSERT → POINT)
        def _has_az(s: str | None) -> bool:
            if not s:
                return False
            s2 = s.strip().lower()
            return (
                s2 == "az" or s2.startswith("az ") or s2.startswith("az:") or
                " az " in f" {s2} " or "azimute" in s2 or "az." in s2
            )

        # TEXT
        if ponto_az is None:
            for e in msp.query('TEXT'):
                if _has_az(e.dxf.text):
                    ins = e.dxf.insert
                    ponto_az = (ins.x, ins.y, 0.0)
                    break

        # MTEXT
        if ponto_az is None:
            for e in msp.query('MTEXT'):
                if _has_az(e.text):
                    ins = e.dxf.insert
                    ponto_az = (ins.x, ins.y, 0.0)
                    break

        # INSERT (nome do bloco) + ATTRIBs
        if ponto_az is None:
            for br in msp.query('INSERT'):
                if _has_az(br.dxf.name):
                    ins = br.dxf.insert
                    ponto_az = (ins.x, ins.y, 0.0)
                    break
                try:
                    iter_attribs = br.attribs if hasattr(br, "attribs") and not callable(br.attribs) else br.attribs()
                    for att in iter_attribs:
                        if _has_az(getattr(att.dxf, "text", None)):
                            ins = br.dxf.insert
                            ponto_az = (ins.x, ins.y, 0.0)
                            raise StopIteration
                except StopIteration:
                    break
                except Exception:
                    pass

        # POINT
        if ponto_az is None:
            for e in msp.query('POINT'):
                loc = e.dxf.location
                ponto_az = (loc.x, loc.y, 0.0)
                break

        # Fallback final
        if ponto_az is None:
            ponto_az = (ordered_points[0][0], ordered_points[0][1], 0.0)
            logger.warning("⚠️ Ponto Az não encontrado no DXF. Usando fallback (primeiro ponto).")

        logger.info(f"Linhas processadas: {len(lines)}")
        logger.info(f"Perímetro do DXF: {perimeter_dxf:.2f} m")
        logger.info(f"Área do DXF: {area_dxf:.2f} m²")

        # >>> RETORNA 7 ITENS <<<
        return doc, lines, perimeter_dxf, area_dxf, ponto_az, msp, ordered_points_with_bulge

    except Exception as e:
        logger.error(f"Erro ao obter informações do documento: {e}")
        # >>> TAMBÉM retorna 7 itens no erro <<<
        return None, [], 0.0, 0.0, None, None, []


# 🔹 Função para definir a fonte padrão
def set_default_font(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    
def calculate_point_on_line(start, end, distance):
    """
    Calcula um ponto a uma determinada distância sobre uma linha entre dois pontos.
    :param start: Coordenadas do ponto inicial (x, y).
    :param end: Coordenadas do ponto final (x, y).
    :param distance: Distância do ponto inicial ao longo da linha.
    :return: Coordenadas do ponto calculado (x, y).
    """
    dx, dy = end[0] - start[0], end[1] - start[1]
    length = math.hypot(dx, dy)  # Calcula o comprimento da linha
    if length == 0:
        raise ValueError("Ponto inicial e final são iguais, não é possível calcular um ponto na linha.")
    return (
        start[0] + (dx / length) * distance,
        start[1] + (dy / length) * distance
    )
def add_azimuth_arc_to_dxf(msp, ponto_az, v1, azimute):
    """
    Adiciona o arco do azimute ao DXF usando ezdxf.
    """
    try:
        logger.info(f"Iniciando a adição do arco de azimute. Azimute: {azimute}°")

        # Criar camada 'Azimute', se não existir
        if 'Azimute' not in msp.doc.layers:
            msp.doc.layers.new(name='Azimute', dxfattribs={'color': 1})
            logger.info("Camada 'Azimute' criada com sucesso.")

        # Traçar segmento entre Az e V1
        msp.add_line(start=ponto_az, end=v1, dxfattribs={'layer': 'Azimute'})
        logger.info(f"Segmento entre Az e V1 desenhado de {ponto_az} para {v1}")

        # Traçar segmento para o norte
        north_point = (ponto_az[0], ponto_az[1] + 2)
        msp.add_line(start=ponto_az, end=north_point, dxfattribs={'layer': 'Azimute'})
        logger.info(f"Linha para o norte desenhada com sucesso de {ponto_az} para {north_point}")

        # Calcular o ponto inicial (1 metro de Az para V1)
        # Calcular distância entre ponto Az e V1 para definir raio adaptativo
        dist = calculate_distance(ponto_az, v1)
        radius = 0.4 if dist <= 0.5 else 1.0

        # Calcular os pontos do arco com esse raio
        start_arc = calculate_point_on_line(ponto_az, v1, radius)
        end_arc = calculate_point_on_line(ponto_az, north_point, radius)

        # Traçar o arco do azimute
        msp.add_arc(
            center=ponto_az,
            radius=radius,
            start_angle=math.degrees(math.atan2(start_arc[1] - ponto_az[1], start_arc[0] - ponto_az[0])),
            end_angle=math.degrees(math.atan2(end_arc[1] - ponto_az[1], end_arc[0] - ponto_az[0])),
            dxfattribs={'layer': 'Azimute'}
        )
        logger.info(f"Arco do azimute desenhado com sucesso com valor de {azimute}° no ponto {ponto_az}")

       # Adicionar rótulo do azimute diretamente com o texto "Azimute:"
        azimuth_label = f"Azimute: {convert_to_dms(azimute)}"  # Incluir o prefixo "Azimute:"

        # Calcular a posição do rótulo
        label_position = (
            ponto_az[0] + 1.0 * math.cos(math.radians(azimute / 2)),
            ponto_az[1] + 1.0 * math.sin(math.radians(azimute / 2))
        )

        # Adicionar o texto ao desenho
        msp.add_text(
            azimuth_label,
            dxfattribs={
                'height': 0.25,
                'layer': 'Azimute',
                'insert': label_position  # Define a posição diretamente
            }
        )

        logger.info(f"Rótulo do azimute adicionado com sucesso: '{azimuth_label}' em {label_position}")


    except Exception as e:
        logger.error(f"Erro na função `add_azimuth_arc_to_dxf`: {e}")

def calculate_azimuth(p1, p2):
    """
    Calcula o azimute entre dois pontos em graus.
    Azimute é o ângulo medido no sentido horário a partir do Norte.
    """
    delta_x = p2[0] - p1[0]  # Diferença em X (Leste/Oeste)
    delta_y = p2[1] - p1[1]  # Diferença em Y (Norte/Sul)

    # Calcular o ângulo em radianos
    azimuth_rad = math.atan2(delta_x, delta_y)

    # Converter para graus
    azimuth_deg = math.degrees(azimuth_rad)

    # Garantir que o azimute esteja no intervalo [0, 360)
    if azimuth_deg < 0:
        azimuth_deg += 360

    return azimuth_deg

def calculate_distance(point1, point2):
    """
    Calcula a distância entre dois pontos em um plano 2D.
    :param point1: Tupla (x1, y1) representando o primeiro ponto.
    :param point2: Tupla (x2, y2) representando o segundo ponto.
    :return: Distância entre os pontos.
    """
    dx = point2[0] - point1[0]
    dy = point2[1] - point1[1]
    return math.sqrt(dx**2 + dy**2)

def add_north_arrow(msp, base_point, length=10):
    """
    Adiciona uma seta (linha) apontando para o Norte a partir do ponto base (Az).
    """
    # Linha vertical representando o Norte
    north_end = (base_point[0], base_point[1] + length)
    msp.add_line(start=base_point, end=north_end, dxfattribs={'layer': 'LAYOUT_AZIMUTES'})

    # Adiciona o texto "N"
    msp.add_text("N", dxfattribs={
        'height': 1.0,
        'insert': (north_end[0], north_end[1] + 1),
        'layer': 'LAYOUT_AZIMUTES'
    })



# Função para calcular azimute e distância
def calculate_azimuth_and_distance(start_point, end_point):
    dx = end_point[0] - start_point[0]
    dy = end_point[1] - start_point[1]
    distance = math.hypot(dx, dy)
    azimuth = math.degrees(math.atan2(dx, dy))
    if azimuth < 0:
        azimuth += 360
    return azimuth, distance


def add_azimuth_arc(doc, msp, ponto_az, v1, azimuth, radius=8):
    """
    Adiciona o arco geométrico representando o ângulo de Azimute entre o norte e a linha Az→V1.
    """
    try:
        # Cria a camada específica caso não exista
        if 'LAYOUT_AZIMUTES' not in doc.layers:
            doc.layers.new(name='LAYOUT_AZIMUTES', dxfattribs={'color': 5})

        # Ângulo inicial sempre aponta para o norte (90° na convenção CAD)
        start_angle = 90.0

        # O ângulo final é obtido subtraindo do azimute (pois CAD mede no sentido anti-horário)
        end_angle = 90.0 - azimuth

        # Garante que os ângulos estejam no intervalo 0-360
        if end_angle < 0:
            end_angle += 360

        # Adiciona o arco geométrico ao DXF
        msp.add_arc(
            center=ponto_az,
            radius=radius,
            start_angle=end_angle,
            end_angle=start_angle,
            dxfattribs={'layer': 'LAYOUT_AZIMUTES'}
        )

        # Adiciona o texto de rótulo próximo ao arco (já está correto)
        mid_angle_rad = math.radians((start_angle + end_angle) / 2)
        label_position = (
            ponto_az[0] + (radius + 1.5) * math.cos(mid_angle_rad),
            ponto_az[1] + (radius + 1.5) * math.sin(mid_angle_rad)
        )
        azimuth_label = f"Azimute = {convert_to_dms(azimuth)}"
        msp.add_text(
            azimuth_label,
            dxfattribs={
                'height': 1.0,
                'layer': 'LAYOUT_AZIMUTES',
                'insert': label_position
            }
        )

        logger.info(f"✅ Arco do azimute ({azimuth_label}) adicionado com sucesso.")

    except Exception as e:
        logger.error(f"❌ Erro ao adicionar arco do azimute: {e}")



# Função para converter graus decimais para DMS
def convert_to_dms(decimal_degrees):
    degrees = int(decimal_degrees)
    minutes = int(abs(decimal_degrees - degrees) * 60)
    seconds = abs((decimal_degrees - degrees - minutes / 60) * 3600)
    return f"{degrees}° {minutes}' {seconds:.2f}\""

# Função para calcular a área de uma poligonal
def calculate_polygon_area(points):
    n = len(points)
    area = 0.0
    for i in range(n):
        x1, y1 = points[i][0], points[i][1]
        x2, y2 = points[(i + 1) % n][0], points[(i + 1) % n][1]
        area += x1 * y2 - x2 * y1
    return abs(area) / 2.0

def _carregar_confrontantes(uuid_str, tipo, diretorio_preparado):
    padrao = os.path.join(diretorio_preparado, f"{uuid_str}_FECHADA_{tipo}.xlsx")
    matches = glob.glob(padrao)
    if not matches:
        logger.warning(f"Planilha FECHADA não encontrada: {padrao}")
        return []
    caminho = matches[0]
    try:
        df = pd.read_excel(caminho)
    except Exception as e:
        logger.error(f"Falha ao ler planilha FECHADA {caminho}: {e}")
        return []

    # tenta localizar a coluna confrontante de forma tolerante
    cols = {str(c).strip().lower(): c for c in df.columns}
    col_confrontante = cols.get('confrontante')
    if not col_confrontante:
        logger.warning("Coluna 'Confrontante' não encontrada; prosseguindo com lista vazia.")
        return []

    return df[col_confrontante].fillna('').astype(str).tolist()


def add_label_and_distance(doc, msp, start_point, end_point, label, distance):
    """
    Adiciona um rótulo no vértice e a distância corretamente alinhada à linha no arquivo DXF.
    
    :param doc: Objeto Drawing do ezdxf.
    :param msp: ModelSpace do ezdxf.
    :param start_point: Coordenadas do ponto inicial (x, y).
    :param end_point: Coordenadas do ponto final (x, y).
    :param label: Nome do vértice (ex: V1, V2).
    :param distance: Distância entre os pontos (em metros).
    """
    try:
        msp = doc.modelspace()

        # Criar camadas necessárias (sem alterar as que não precisam)
        for layer_name, color in [
            ("LAYOUT_VERTICES", 2),  # Vermelho para vértices
            ("LAYOUT_DISTANCIAS", 4),  # Azul para distâncias
            ("LAYOUT_AZIMUTES", 5)  # Magenta para azimutes
        ]:
            if layer_name not in doc.layers:
                doc.layers.new(name=layer_name, dxfattribs={"color": color})

        # 🔹 Adicionar círculo no ponto inicial (Vértices)
        msp.add_circle(center=start_point[:2], radius=1.0, dxfattribs={'layer': 'LAYOUT_VERTICES'})

        # 🔹 Adicionar rótulo do vértice
        text_point = (start_point[0] + 1, start_point[1])  # Posição deslocada
        msp.add_text(
            label,
            dxfattribs={'height': 0.5, 'layer': 'LAYOUT_VERTICES', 'insert': text_point}
        )

        # 🔹 Calcular o ponto médio da linha
        mid_x = (start_point[0] + end_point[0]) / 2
        mid_y = (start_point[1] + end_point[1]) / 2

        # 🔹 Vetor da linha
        dx = end_point[0] - start_point[0]
        dy = end_point[1] - start_point[1]
        length = math.hypot(dx, dy)

        # Evitar erro de divisão por zero
        if length == 0:
            return

        # 🔹 Ângulo da linha
        angle = math.degrees(math.atan2(dy, dx))

        # 🔹 Ajuste de ângulo para manter leitura correta
        if angle < -90 or angle > 90:
            angle += 180  

        # 🔹 Afastar o rótulo da linha
        offset = 0.3  # Ajuste para evitar sobreposição
        perp_x = -dy / length * offset
        perp_y = dx / length * offset
        displaced_mid_point = (mid_x + perp_x, mid_y + perp_y)

        # 🔹 Formatar a distância corretamente
        distancia_formatada = f"{distance:.2f}".replace(".", ",")

        # 🔹 Adicionar rótulo da distância corretamente alinhado
        msp.add_text(
            f"{distancia_formatada} ",
            dxfattribs={
                "height": 1.0,  # Aumenta a altura do texto
                "layer": "LAYOUT_DISTANCIAS",
                "rotation": angle,  # Alinhamento correto à linha
                "insert": displaced_mid_point
            }
        )

        print(f"✅ Distância {distancia_formatada} m adicionada corretamente em {displaced_mid_point}")

    except Exception as e:
        print(f"❌ Erro ao adicionar rótulo de distância: {e}")

def azimuth_deg(p1, p2):
    dx, dy = p2[0]-p1[0], p2[1]-p1[1]
    ang = math.degrees(math.atan2(dx, dy))  # 0° = Norte, cresce no sentido horário
    return ang + 360 if ang < 0 else ang

def calculate_angular_turn(p1, p2, p3):
    """
    Calcula o giro angular no ponto `p2` entre os segmentos `p1-p2` e `p2-p3` no sentido horário.
    Retorna o ângulo em graus.
    """
    import math
    
    dx1, dy1 = p1[0] - p2[0], p1[1] - p2[1]  # Vetor do segmento p1-p2
    dx2, dy2 = p3[0] - p2[0], p3[1] - p2[1]  # Vetor do segmento p2-p3

    angle1 = math.atan2(dy1, dx1)
    angle2 = math.atan2(dy2, dx2)

    # Calcula o ângulo horário
    angular_turn = (angle2 - angle1) % (2 * math.pi)
    angular_turn_degrees = math.degrees(angular_turn)

    return angular_turn_degrees


#     return confrontantes
def sanitize_filename(filename):
    # Substitui os caracteres inválidos por um caractere válido (ex: espaço ou underline)
    sanitized_filename = re.sub(r'[\\/*?:"<>|]', "_", filename)  # Substitui caracteres inválidos por "_"
    return sanitized_filename

def rotate_polygon_start_at_v1(lines, pts_bulge, v1_target, sentido_poligonal="", tol=1e-4):
    """
    Reordena 'lines' ([(p_i, p_{i+1})...]) e 'pts_bulge' para que o primeiro
    vértice seja o mais próximo de v1_target. Se for anti-horário, dá 1 passo extra.
    """
    if not lines:
        return lines, pts_bulge, None

    import math
    def dist2(a, b): 
        return (a[0]-b[0])**2 + (a[1]-b[1])**2

    # índice do vértice mais próximo de v1_target
    verts = [seg[0] for seg in lines]
    idx = min(range(len(verts)), key=lambda i: dist2(verts[i], v1_target))

    # roda para começar em idx
    lines = lines[idx:] + lines[:idx]
    pts_bulge = pts_bulge[idx:] + pts_bulge[:idx]
    rot_idx = idx

    # se for anti-horário, ande uma casa no “sentido horário” desejado
    if str(sentido_poligonal).lower().startswith("anti"):
        lines = lines[1:] + lines[:1]
        pts_bulge = pts_bulge[1:] + pts_bulge[:1]
        rot_idx = (idx + 1) % len(verts)

    return lines, pts_bulge, rot_idx

def _azimuth_deg(p1, p2):
    dx, dy = (p2[0] - p1[0]), (p2[1] - p1[1])
    ang = math.degrees(math.atan2(dx, dy))
    return ang + 360.0 if ang < 0.0 else ang

def _convert_to_dms_safe(deg):
    d = float(deg) % 360.0
    g = int(d)
    m = int((d - g) * 60)
    s = (d - g - m/60) * 3600
    return f"{g}° {m}' {s:.2f}\""


def add_az_marker_to_dxf(
    doc_dxf,
    ponto_az,            # (x, y) ou (x, y, z)
    v1,                  # (x, y) do V1
    azimute_deg,         # float (0..360)
    distancia_az_v1,     # float (em metros) -> novo
    *,
    v2=None,             # ⬅️ NOVO: (x, y) do V2 (vizinho correto)
    sentido=None,        # ⬅️ NOVO: 'horario' | 'anti_horario'
    draw_giro=True,      # ⬅️ NOVO: desenhar o arco do Giro Angular
    layer="Az_Marker",
    north_len=8.0,
    text_height=0.6,
    arc_radius=5.0,
    draw_minor_arc=False
):
    """Rótulo 'Az', marcador Norte, arco do azimute + rótulo
       e rótulo da distância sobre a reta Az→V1 (sem set_pos)."""

    def to_dms_string(deg):
        d = abs(deg)
        g = int(d)
        m = int((d - g) * 60)
        s = (d - g - m/60) * 3600
        sign = "-" if deg < 0 else ""
        return f"{sign}{g}°{m}'{s:.2f}\""

    msp = doc_dxf.modelspace()

    # Garante a layer
    try:
        if layer not in doc_dxf.layers:
            doc_dxf.layers.new(name=layer)
    except Exception:
        pass

    ax, ay = ponto_az[0], ponto_az[1]
    v1x, v1y = v1[0], v1[1]

    # 1) Rótulo "Az"
    az_text = msp.add_text("Az", dxfattribs={"height": text_height, "layer": layer})
    az_text.dxf.insert = (ax, ay)

    # 2) Marcador do Norte e rótulo "N"
    msp.add_line((ax, ay), (ax, ay + north_len), dxfattribs={"layer": layer})
    n_text = msp.add_text("N", dxfattribs={"height": text_height, "layer": layer})
    n_text.dxf.insert = (ax, ay + north_len + text_height * 1.2)

    # 3) Linha Az→V1 (referência)
    msp.add_line((ax, ay), (v1x, v1y), dxfattribs={"layer": layer})

    # 3.1) Rótulo da distância sobre a reta (NOVO)
    dx, dy = (v1x - ax), (v1y - ay)
    seg_len = math.hypot(dx, dy) or 1.0
    ux, uy = dx / seg_len, dy / seg_len                  # unitário ao longo da reta
    px, py = -uy, ux                                     # unitário perpendicular (esquerda)

    # ponto médio + pequeno offset perpendicular (para não "sentar" na linha)
    midx = (ax + v1x) / 2.0
    midy = (ay + v1y) / 2.0
    offset = text_height * 1.0                           # ajuste fino aqui se quiser
    mid_shift = (midx + px * offset, midy + py * offset)

    rot_deg = math.degrees(math.atan2(dy, dx)) % 360.0   # rotação do texto na direção da reta
    dist_label = f"{distancia_az_v1:.2f} "

    dist_text = msp.add_text(
        dist_label,
        dxfattribs={
            "height": text_height,
            "layer": layer,
            "rotation": rot_deg
        }
    )
    dist_text.dxf.insert = mid_shift

    # 4) Arco do azimute (de Norte até direção Az→V1) + rótulo
    ang_auto = math.degrees(math.atan2(dy, dx)) % 360    # 0°=E, 90°=N, CCW
    north_auto = 90.0

    if draw_minor_arc:
        delta_ccw = (ang_auto - north_auto) % 360
        if delta_ccw <= 180:
            start_ang, end_ang = north_auto, ang_auto
        else:
            start_ang, end_ang = ang_auto, north_auto
    else:
        # arco com extensão igual ao azimute (N→direção, sentido horário)
        start_ang, end_ang = ang_auto, north_auto

    msp.add_arc(
        center=(ax, ay),
        radius=arc_radius,
        start_angle=start_ang,
        end_angle=end_ang,
        dxfattribs={"layer": layer}
    )

    arc_len_ccw = (end_ang - start_ang) % 360
    mid_ang = (start_ang + arc_len_ccw / 2.0) % 360
    mid_rad = math.radians(mid_ang)
    label_r = arc_radius + text_height * 2.0
    label_pos = (ax + label_r * math.cos(mid_rad), ay + label_r * math.sin(mid_rad))

    az_label = f"Azimute = {to_dms_string(azimute_deg)}"
    lbl = msp.add_text(az_label, dxfattribs={"height": text_height, "layer": layer})
    lbl.dxf.insert = label_pos

    # ===== 5) GIRO ANGULAR Az–V1–V2 (pivot em V1) =====
    # Requer v2 e sentido; desenha o arco certo independente da ordem dos vértices.
    if draw_giro and (v2 is not None):
        v1x, v1y = float(v1[0]), float(v1[1])
        azx, azy  = float(ponto_az[0]), float(ponto_az[1])
        v2x, v2y  = float(v2[0]), float(v2[1])

        # ângulos (0..360) das direções a partir de V1
        a_az = math.degrees(math.atan2(azy - v1y, azx - v1x)) % 360.0
        a_v2 = math.degrees(math.atan2(v2y - v1y, v2x - v1x)) % 360.0

        # Giro HORÁRIO SEMPRE: de V1→Az para V1→V2
        giro  = (a_az - a_v2) % 360.0          # valor do giro em graus (horário)
        start = (a_az - giro) % 360.0          # add_arc é CCW ⇒ desenhe de (Az - giro) → Az
        end   = a_az

        # raio do arco do giro (use o mesmo arc_radius, ou ajuste se quiser)
        giro_radius = arc_radius

        # arco do giro em V1
        msp.add_arc(
            center=(v1x, v1y),
            radius=giro_radius,
            start_angle=start,
            end_angle=end,
            dxfattribs={"layer": layer}
        )

        # rótulo do giro — sempre horizontal
        sweep = (end - start) % 360.0
        mid   = (start + sweep/2.0) % 360.0
        mid_r = math.radians(mid)
        lbl_r = giro_radius + text_height*1.2
        lbl_pt = (v1x + lbl_r*math.cos(mid_r), v1y + lbl_r*math.sin(mid_r))

        # usa convert_to_dms se existir; senão, fallback no to_dms_string local
        try:
            giro_txt = f"Giro Angular: {convert_to_dms(giro)}"
        except NameError:
            giro_txt = f"Giro Angular: {to_dms_string(giro)}"

        t = msp.add_text(
            giro_txt,
            dxfattribs={"height": text_height, "layer": layer}
        )
        t.dxf.insert   = lbl_pt
        t.dxf.rotation = 0  # horizontal
        try:
            t.dxf.halign      = 1  # CENTER
            t.dxf.valign      = 2  # MIDDLE
            t.dxf.align_point = lbl_pt
        except Exception:
            pass      

# ==== HELPERs necessários para AZIMUTE_AZ (portados de ANGULO_AZ) ====

def _log_info(msg: str):
    try:
        logger.info(msg)
    except Exception:
        print(msg)

def _log_error(msg: str):
    try:
        logger.error(msg)
    except Exception:
        print(msg)

def _convert_to_dms_safe(deg):
    """Converte graus decimais em DMS de forma tolerante; retorna '' se inválido."""
    try:
        d = float(deg)
    except (TypeError, ValueError):
        return ""
    sign = "-" if d < 0 else ""
    d = abs(d)
    g = int(d)
    m = int((d - g) * 60)
    s = (d - g - m/60) * 3600
    return f"{sign}{g}° {m}' {s:.2f}\""

def _azimuth_deg(p, q):
    """Azimute (0..360) de p->q com 0°=Norte, sentido horário."""
    x1, y1 = float(p[0]), float(p[1])
    x2, y2 = float(q[0]), float(q[1])
    a = math.degrees(math.atan2(x2 - x1, y2 - y1))  # nota: atan2(dx, dy) p/ azimute
    if a < 0:
        a += 360.0
    return a

def _ring_area_xy(pts):
    """Área assinada do anel (lista de dicts com chaves 'x','y'). CCW > 0."""
    if not pts:
        return 0.0
    area = 0.0
    n = len(pts)
    for i in range(n):
        x1, y1 = float(pts[i]['x']), float(pts[i]['y'])
        x2, y2 = float(pts[(i + 1) % n]['x']), float(pts[(i + 1) % n]['y'])
        area += x1 * y2 - x2 * y1
    return area / 2.0

def _polygon_orientation(pts):
    """+1 = anti-horário (CCW), -1 = horário (CW), 0 = degenerado."""
    a = _ring_area_xy(pts)
    if a > 0:
        return +1
    if a < 0:
        return -1
    return 0

def _ensure_orientation(pts, sentido_poligonal):
    """
    Garante o sentido do polígono conforme 'horario' ou 'anti_horario'.
    Se inverter, inverte também o sinal de 'bulge_next' se existir (não reindexa).
    """
    if not pts:
        return []
    want_ccw = str(sentido_poligonal).lower().startswith("anti")
    ori = _polygon_orientation(pts)
    is_ccw = (ori == +1)
    if want_ccw == is_ccw:
        return pts
    rev = list(reversed(pts))
    # como não usamos bulge adiante para ângulos internos, basta flip de sinal
    for p in rev:
        if 'bulge_next' in p and p['bulge_next'] is not None:
            try:
                p['bulge_next'] = -float(p['bulge_next'])
            except Exception:
                pass
    return rev
# ==== FIM dos HELPERs ====

# Função para criar memorial descritivo
# def create_memorial_descritivo(
#     uuid_str, doc, msp, lines, proprietario, matricula, caminho_salvar,
#     excel_file_path, ponto_az, distance_az_v1, azimute_az_v1, tipo, sentido_poligonal='horario',diretorio_concluido=None, encoding='ISO-8859-1'
# ):

#     # Carregar confrontantes da planilha FECHADA
#     confrontantes_df = pd.read_excel(excel_file_path)
#     confrontantes_dict = dict(zip(confrontantes_df['Código'], confrontantes_df['Confrontante']))

#     if confrontantes_df.empty:
#         logger.error("❌ Planilha de confrontantes está vazia.")
#         return None

#     if not lines:
#         logger.error("❌ Sem linhas disponíveis no DXF.")
#         return None

#     ordered_points = [line[0] for line in lines] + [lines[-1][1]]

#     area = calculate_polygon_area(ordered_points)

#     # if area < 0:
#     #     ordered_points.reverse()
#     #     logger.info("Pontos reorganizados para sentido horário.")


#     # Agora inverter o sentido corretamente, incluindo tratamento dos arcos (bulge)
#     if sentido_poligonal == 'horario':
#         if area > 0:
#             ordered_points.reverse()
#             area = abs(area)
#             # Inverte o sentido dos arcos (bulges), se existirem
#             for ponto in ordered_points:
#                 if 'bulge' in ponto and ponto['bulge'] != 0:
#                     ponto['bulge'] *= -1
#             logger.info(f"Área da poligonal invertida para sentido horário com ajuste dos arcos: {area:.4f} m²")
#         else:
#             logger.info(f"Área da poligonal já no sentido horário: {abs(area):.4f} m²")

#     else:  # sentido_poligonal == 'anti_horario'
#         if area < 0:
#             ordered_points.reverse()
#             area = abs(area)
#             # Inverte o sentido dos arcos (bulges), se existirem
#             for ponto in ordered_points:
#                 if 'bulge' in ponto and ponto['bulge'] != 0:
#                     ponto['bulge'] *= -1
#             logger.info(f"Área da poligonal invertida para sentido anti-horário com ajuste dos arcos: {area:.4f} m²")
#         else:
#             logger.info(f"Área da poligonal já no sentido anti-horário: {abs(area):.4f} m²")


#     data = []
#     total_vertices = len(ordered_points) - 1

#     for i in range(total_vertices):
#         start_point = ordered_points[i]
#         end_point = ordered_points[i + 1]

#         azimuth, distance = calculate_azimuth_and_distance(start_point, end_point)
#         azimuth_dms = convert_to_dms(azimuth)

#         confrontante = confrontantes_df.iloc[i]['Confrontante'] if i < len(confrontantes_df) else "Desconhecido"

#         coord_e_ponto_az = f"{ponto_az[0]:.3f}".replace('.', ',') if i == 0 else ""
#         coord_n_ponto_az = f"{ponto_az[1]:.3f}".replace('.', ',') if i == 0 else ""

#         data.append({
#             "V": f"V{i + 1}",
#             "E": f"{start_point[0]:.3f}".replace('.', ','),
#             "N": f"{start_point[1]:.3f}".replace('.', ','),
#             "Z": "0,000",
#             "Divisa": f"V{i + 1}_V{1 if (i + 1) == total_vertices else i + 2}",
#             "Azimute": azimuth_dms,
#             "Distancia(m)": f"{distance:.2f}".replace('.', ','),
#             "Confrontante": confrontante,
#             "Coord_E_ponto_Az": coord_e_ponto_az,
#             "Coord_N_ponto_Az": coord_n_ponto_az,
#             "distancia_Az_V1": f"{distance_az_v1:.2f}".replace('.', ',') if i == 0 else "",
#             "Azimute Az_V1": convert_to_dms(azimute_az_v1) if i == 0 else ""
#         })

#         # Adicionar labels no DXF
#         add_label_and_distance(doc, msp, start_point, end_point, f"V{i + 1}", distance)

#     # Caminho padronizado do Excel de saída
#     excel_output_path = os.path.join(caminho_salvar, f"{uuid_str}_FECHADA_{tipo}_{matricula}.xlsx")

#     # Salvar no Excel
#     df = pd.DataFrame(data, dtype=str)
#     df.to_excel(excel_output_path, index=False)

#     # Formatar Excel
#     wb = openpyxl.load_workbook(excel_output_path)
#     ws = wb.active

#     for cell in ws[1]:
#         cell.font = Font(bold=True)
#         cell.alignment = Alignment(horizontal="center", vertical="center")

#     column_widths = {
#         "A": 8, "B": 15, "C": 15, "D": 10, "E": 20,
#         "F": 15, "G": 15, "H": 40, "I": 20, "J": 20,
#         "K": 18, "L": 18,
#     }
#     for col, width in column_widths.items():
#         ws.column_dimensions[col].width = width

#     for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=1, max_col=ws.max_column):
#         for cell in row:
#             cell.alignment = Alignment(horizontal="center", vertical="center")

#     wb.save(excel_output_path)
#     logger.info(f"✅ Excel salvo e formatado: {excel_output_path}")

#     # Adicionar arco de Azimute ao DXF
#     try:
#         msp = doc.modelspace()
#         v1 = ordered_points[0]
#         azimuth = calculate_azimuth(ponto_az, v1)
#         add_azimuth_arc(doc, msp, ponto_az, v1, azimuth)
#         logger.info("✅ Arco de azimute adicionado ao DXF.")
#     except Exception as e:
#         logger.error(f"❌ Erro ao adicionar arco de azimute: {e}")
    
#     # Adicionar linha entre ponto Az e V1 (parte faltante adicionada aqui)
#     try:
#         msp = doc.modelspace()
#         msp.add_line(start=ponto_az, end=v1, dxfattribs={'layer': 'LAYOUT_AZIMUTES'})
#         logger.info("✅ Linha Az→V1 adicionada ao DXF.")
#     except Exception as e:
#         logger.error(f"❌ Erro ao adicionar linha Az→V1: {e}")


#     # Adicionar distância entre Az e V1 no DXF
#     try:
#         msp = doc.modelspace()
#         add_label_and_distance(doc, msp, ponto_az, v1, "Az-V1", distance_az_v1)
#         logger.info(f"✅ Distância Az-V1 ({distance_az_v1:.2f} m) adicionada ao DXF.")
#     except Exception as e:
#         logger.error(f"❌ Erro ao adicionar distância Az-V1: {e}")

#     # Adicionar linha apontando para o Norte no ponto Az
#     try:
#         msp = doc.modelspace()  # É importante garantir o msp atualizado aqui também
#         add_north_arrow(msp, ponto_az)
#         logger.info("✅ Linha Norte adicionada ao DXF.")
#     except Exception as e:
#         logger.error(f"❌ Erro ao adicionar linha Norte: {e}")

#     # Salvar o DXF com as alterações
#     try:
#         dxf_output_path = os.path.join(caminho_salvar, f"{uuid_str}_FECHADA_{tipo}_{matricula}.dxf")
#         doc.saveas(dxf_output_path)
#         logger.info(f"✅ DXF atualizado salvo: {dxf_output_path}")
#     except Exception as e:
#         logger.error(f"❌ Erro ao salvar DXF atualizado: {e}")

#     return excel_output_path

def create_memorial_descritivo(
    uuid_str,
    doc,
    lines,
    proprietario,
    matricula,
    caminho_salvar,
    confrontantes,          # lista de strings (ou [])
    ponto_az,               # tupla (x, y) ou None
    dxf_file_path,
    area_dxf,
    azimute,                # float (graus decimais Az→V1)
    v1,                     # tupla (x, y)
    msp,
    dxf_filename,
    excel_file_path,        # caminho do XLSX de SAÍDA
    tipo,
    giro_angular_v1_dms,    # string DMS ou None
    distancia_az_v1,        # float ou None
    *,
    sentido_poligonal="horario",
    modo="AZIMUTE_AZ",
    points_bulge=None,      # lista de pontos com bulge (do parser)
    diretorio_concluido=None,
    # aliases aceitos (para chamadas antigas):
    azimute_az_v1=None,
    distance_az_v1=None,
):
    """
    DXF já vem tratado do pipeline (sem reler aqui):
    - Usa LWPOLYLINE fechada (com bulge) fornecida por get_document_info_from_dxf.
    - Normaliza sentido, calcula ângulos internos (com concavidade), desenha arcos internos.
    - Só desenha AZ no modo ANGULO_AZ.
    - Gera Excel diretamente.
    """
    logger.info("[CMD] pontos_bulge recebidos: %s", len(points_bulge) if points_bulge else 0)

    # Garante que o estilo de texto "STANDARD" exista no DXF
    if "STANDARD" not in msp.doc.styles:
        msp.doc.styles.new("STANDARD")

    if diretorio_concluido is None:
        diretorio_concluido = caminho_salvar

    dxf_output_path = os.path.join(
        diretorio_concluido,
        f"{uuid_str}_FECHADA_{tipo}_{matricula}.dxf"
    )

    # 0) valida base
    if points_bulge is None or len(points_bulge) < 3:
        _log_error(f"[AZ] points_bulge ausente/insuficiente (type={type(points_bulge)}, len={0 if not points_bulge else len(points_bulge)}). Verifique get_document_info_from_dxf.")
        return None

    # 1) normaliza sentido
    pts = _ensure_orientation(points_bulge, sentido_poligonal)
    orient = _polygon_orientation(pts)
    _log_info(f"Sentido normalizado: {'anti-horário' if orient == +1 else 'horário'}")

    # 2) (AZIMUTE POR LADO) helper local
    def _azimuth_deg(p1, p2):
        # 0° = Norte; cresce no sentido horário
        dx, dy = (p2[0] - p1[0]), (p2[1] - p1[1])
        ang = math.degrees(math.atan2(dx, dy))
        return ang + 360.0 if ang < 0.0 else ang

    

    # # 4) desenho do AZ depende do modo
    # # ANGULO_AZ  → desenha Az, linha Az–V1, arco e rótulos
    # # ANGULO_P1_P2 → NÃO desenha Az/linha/arco (poligonal ABERTA já mostra amarração)
    # if modo == "ANGULO_AZ" and ponto_az is not None and v1 is not None:
    #     dx = v1[0] - ponto_az[0]
    #     dy = v1[1] - ponto_az[1]
    #     dist = math.hypot(dx, dy)
    #     if dist > 1e-6:
    #         try:
    #             _desenhar_referencia_az(msp, ponto_az, v1, azimute)
    #         except Exception as e:
    #             logger.error("Erro ao desenhar referência de Az: %s", e)
    #     else:
    #         logger.warning("⚠️ Distância Az–V1 ≈ 0; desenho do Az suprimido.")

    # 5) Excel (sem reler nada)
    try:
        ordered_points_xy = [(p['x'], p['y']) for p in pts]
        total_pontos = len(ordered_points_xy)
        data = []

        if ponto_az is not None:
            ponto_az_e = f"{ponto_az[0]:,.3f}".replace(",", "").replace(".", ",")
            ponto_az_n = f"{ponto_az[1]:,.3f}".replace(",", "").replace(".", ",")
        else:
            ponto_az_e = ""
            ponto_az_n = ""

        for i in range(total_pontos):
            p2 = ordered_points_xy[i]
            p3 = ordered_points_xy[(i + 1) % total_pontos]

            dx, dy = p3[0] - p2[0], p3[1] - p2[1]
            distance = math.hypot(dx, dy)
            description = f"V{i + 1}_V{(i + 2) if i + 1 < total_pontos else 1}"
            confrontante = confrontantes[i % len(confrontantes)] if confrontantes else ""
            #ang_interno_dms = _convert_to_dms_safe(internos_deg[i])
            az_seg_deg = _azimuth_deg(p2, p3)
            az_seg_dms = _convert_to_dms_safe(az_seg_deg)

            if i == 0:
                distancia_az_v1_str = f"{float(distancia_az_v1):.2f}".replace(".", ",") if distancia_az_v1 is not None else ""
                azimute_az_v1_str   = _convert_to_dms_safe(float(azimute)) if azimute is not None else ""
                giro_v1_str         = giro_angular_v1_dms or ""
                p_az_e, p_az_n      = ponto_az_e, ponto_az_n
            else:
                distancia_az_v1_str = ""
                azimute_az_v1_str   = ""
                giro_v1_str         = ""
                p_az_e, p_az_n      = "", ""

            data.append({
                "V": f"V{i + 1}",
                "E": f"{p2[0]:,.3f}".replace(",", "").replace(".", ","),
                "N": f"{p2[1]:,.3f}".replace(",", "").replace(".", ","),
                "Z": "0,000",
                "Divisa": description,
                "Azimute": az_seg_dms,
                "Distancia(m)": f"{distance:,.2f}".replace(",", "").replace(".", ","),
                "Confrontante": confrontante,
                "ponto_AZ_E": p_az_e,
                "ponto_AZ_N": p_az_n,
                "distancia_Az_V1": distancia_az_v1_str,
                "Azimute Az_V1": azimute_az_v1_str,
                "Giro Angular Az_V1_V2": giro_v1_str
            })

            try:
                if distance > 0.01 and 'add_label_and_distance' in globals():
                    add_label_and_distance(doc, msp, p2, p3, f"V{i + 1}", distance)
            except Exception as e:
                _log_error(f"Falha ao rotular distância do lado V{i+1}: {e}")

        
        # escreve excel
        df = pd.DataFrame(data)

        # ── Garantir as 3 colunas do ANGULO_AZ antes de salvar
        #cols_novas = ["AZIMUTE_AZ_V1_GRAUS", "DISTANCIA_AZ_V1_M", "GIRO_V1_GRAUS"]
        # for c in cols_novas:
        #     if c not in df.columns:
        #         df[c] = ""  # ou pd.NA

        # Garante diretório e salva
        try:
            os.makedirs(os.path.dirname(excel_file_path), exist_ok=True)
            df.to_excel(excel_file_path, index=False)
            _log_info(f"Excel escrito (primeira passagem): {os.path.abspath(excel_file_path)}")
        except Exception as e:
            _log_error(f"Falha ao salvar Excel na primeira passagem: {e}")
            raise  # deixe a main ver o stacktrace

        # Formatação openpyxl
        try:
            wb = openpyxl.load_workbook(excel_file_path)
            ws = wb.active

            for cell in ws[1]:
                cell.font = Font(bold=True)
                cell.alignment = Alignment(horizontal="center", vertical="center")

            col_widths = {
                "A": 8, "B": 15, "C": 15, "D": 0, "E": 15,
                "F": 15, "G": 15, "H": 50, "I": 15,
                "J": 15, "K": 15, "L": 20, "M": 20
            }
            for col, width in col_widths.items():
                ws.column_dimensions[col].width = width

            for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=1, max_col=ws.max_column):
                for cell in row:
                    cell.alignment = Alignment(horizontal="center", vertical="center")

            wb.save(excel_file_path)
            _log_info(f"Excel salvo e formatado: {os.path.abspath(excel_file_path)}")
        except Exception as e:
            _log_error(f"Falha ao formatar/salvar Excel com openpyxl: {e}")
            raise

        # Confirma existência
        if os.path.exists(excel_file_path):
            _log_info(f"✅ Excel confirmado em disco: {os.path.abspath(excel_file_path)}")
        else:
            _log_error(f"❌ Excel NÃO encontrado após salvar: {os.path.abspath(excel_file_path)}")


        # # extras DXF (opcionais e seguros)
        # try:
        #     v1_pt = ordered_points_xy[0]
        #     v2_pt = ordered_points_xy[1]
            # se existir o helper e você quiser o giro no V1 com Az:
        #    if 'add_giro_angular_arc_to_dxf' in globals() and ponto_az is not None:
                # padronize este helper para (msp, v1_pt, ponto_az, v2_pt)
                #add_giro_angular_arc_to_dxf(msp, v1_pt, ponto_az, v2_pt)
                #_log_info("Giro horário Az–V1–V2 adicionado com sucesso.")
        # except Exception as e:
        #     _log_error(f"Erro ao adicionar giro angular: {e}")

        try:
            if "Vertices" not in msp.doc.layers:
                msp.doc.layers.new(name="Vertices")
        except Exception:
            pass

        # garanta a camada
        try:
            if "Vertices" not in msp.doc.layers:
                msp.doc.layers.new(name="Vertices")
        except Exception:
            pass

        for i, (x, y) in enumerate(ordered_points_xy):
            try:
                msp.add_circle(center=(x, y), radius=0.5, dxfattribs={"layer": "Vertices"})
                msp.add_text(
                    f"V{i + 1}",
                    dxfattribs={
                        "height": 0.3,
                        "layer": "Vertices",
                        "insert": (x + 0.30, y + 0.30)  # <<< POSIÇÃO DO RÓTULO
                    }
                )
            except Exception as e:
                logger.warning(f"Falha rotulando V{i+1}: {e}")


        # # só desenhe o arco do azimute se realmente quiser no produto FECHADA
        # # e se houver amarração (Az) válida:
        # if modo == "ANGULO_AZ" and ponto_az is not None:
        #     try:
        #         azim = calculate_azimuth(ponto_az, v1_pt)
        #         _desenhar_referencia_az(msp, ponto_az, v1_pt, azim)
        #         _log_info("Arco do Azimute Az–V1 adicionado com sucesso.")
        #     except Exception as e:
        #         _log_error(f"Erro ao adicionar arco do azimute: {e}")

        # 6) salvar DXF final
        try:
            doc.saveas(dxf_output_path)
            logger.info("✅ DXF FECHADA salvo corretamente: %s", dxf_output_path)
        except Exception as e:
            logger.error("Erro ao salvar DXF FECHADA: %s", e)

    except Exception as e:
        _log_error(f"❌ Erro ao gerar o memorial descritivo: {e}")
        return None

    return excel_file_path





# def create_memorial_document(
#     uuid_str, proprietario, matricula, descricao, excel_file_path, template_path, 
#     output_path, perimeter_dxf, area_dxf, desc_ponto_Az, Coorde_E_ponto_Az, Coorde_N_ponto_Az,
#     azimuth, distance, uso_solo, area_imovel, cidade, rua, comarca, rgi, caminho_salvar, tipo
# ):
#     try:
#         # Ler arquivo Excel
#         df = pd.read_excel(excel_file_path)
#         df['N'] = pd.to_numeric(df['N'].astype(str).str.replace(',', '.'), errors='coerce')
#         df['E'] = pd.to_numeric(df['E'].astype(str).str.replace(',', '.'), errors='coerce')
#         df['Distancia(m)'] = pd.to_numeric(df['Distancia(m)'].astype(str).str.replace(',', '.'), errors='coerce')

#         # Criar documento Word
#         doc_word = Document(template_path)
#         set_default_font(doc_word)

#         p1 = doc_word.add_paragraph("MEMORIAL DESCRITIVO", style='Normal')
#         p1.runs[0].bold = True
#         p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#         doc_word.add_paragraph()

#         texto_tipo = {
#             "ETE": f"Área da matrícula {matricula} destinada a {descricao} - SES de {cidade}",
#             "REM": f"Área remanescente da matrícula {matricula} destinada a {descricao}",
#             "SER": f"Área da matrícula {matricula} destinada à SERVIDÃO ADMINISTRATIVA DE ACESSO À {descricao}",
#             "ACE": f"Área da matrícula {matricula} destinada ao ACESSO DA SERVIDÃO ADMINISTRATIVA DA {descricao}",
#         }.get(tipo, "Tipo não especificado")

#         p = doc_word.add_paragraph(style='Normal')
#         p.add_run("Imóvel: ")
#         p.add_run(texto_tipo).bold = True

#         doc_word.add_paragraph(f"Matrícula: Número - {matricula} do {RI} de {comarca}", style='Normal')
#         doc_word.add_paragraph(f"Proprietário: {proprietario}", style='Normal')
#         doc_word.add_paragraph(f"Local: {rua} - {cidade}", style='Normal')
#         doc_word.add_paragraph(f"Área: {area_dxf:,.2f} m²".replace(",", "X").replace(".", ",").replace("X", "."), style='Normal')
#         doc_word.add_paragraph(f"Perímetro: {perimeter_dxf:,.2f} m".replace(",", "X").replace(".", ",").replace("X", "."), style='Normal')
#         doc_word.add_paragraph()

#         area_dxf_formatada = f"{area_dxf:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
#         texto_paragrafo = (f"Área {uso_solo} com {area_dxf_formatada} m², parte de um todo maior da Matrícula Nº {matricula} com {area_imovel} "
#                            f"do {rgi} de {comarca}, localizada na {rua}, na cidade de {cidade}, definida através do seguinte levantamento "
#                            "topográfico, onde os ângulos foram medidos no sentido horário.")
#         p = doc_word.add_paragraph(texto_paragrafo, style='Normal')
#         p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
#         doc_word.add_paragraph()

#         coord_E_ponto_Az = f"{Coorde_E_ponto_Az:.3f}".replace(".", ",")
#         coord_N_ponto_Az = f"{Coorde_N_ponto_Az:.3f}".replace(".", ",")
#         doc_word.add_paragraph(
#             f"O Ponto Az, ponto de amarração, está localizado na {desc_ponto_Az} nas coordenadas "
#             f"E(X) {coord_E_ponto_Az} e N(Y) {coord_N_ponto_Az}.", style='Normal'
#         ).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
#         doc_word.add_paragraph()

#         distance_formatada = f"{distance:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
#         p = doc_word.add_paragraph(style='Normal')
#         p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
#         p.add_run(f"Daí, com Azimute de {convert_to_dms(azimuth)} e distância de {distance_formatada} m, chega-se ao Vértice ")
#         p.add_run("V1").bold = True
#         p.add_run(", origem da descrição desta área.")
#         doc_word.add_paragraph()

#         initial = df.iloc[0]
#         coord_N_inicial = f"{initial['N']:.3f}".replace(".", ",")
#         coord_E_inicial = f"{initial['E']:.3f}".replace(".", ",")
#         doc_word.add_paragraph("Pontos definidos pelas Coordenadas Planas no Sistema U.T.M. – SIRGAS 2000.", style='Normal')
#         doc_word.add_paragraph()

#         p2 = doc_word.add_paragraph(style='Normal')
#         p2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
#         p2.add_run("Inicia-se a descrição deste perímetro no vértice ")
#         p2.add_run(f"{initial['V']}").bold = True
#         p2.add_run(f", de coordenadas N(Y) {coord_N_inicial} e E(X) {coord_E_inicial}, situado no limite com {initial['Confrontante']}.")
#         doc_word.add_paragraph()

#         for i in range(len(df)):
#             current = df.iloc[i]
#             next_point = df.iloc[(i + 1) % len(df)]

#             distancia_formatada = f"{current['Distancia(m)']:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
#             coord_N_formatada = f"{next_point['N']:.3f}".replace(".", ",")
#             coord_E_formatada = f"{next_point['E']:.3f}".replace(".", ",")

#             complemento = ", origem desta descrição," if next_point['V'] == 'V1' else ""

#             p = doc_word.add_paragraph(style='Normal')
#             p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
#             p.add_run(f"Deste, segue com azimute de {current['Azimute']} e distância de {distancia_formatada} m, "
#                       f"confrontando neste trecho com área pertencente à {current['Confrontante']}, até o vértice ")
#             p.add_run(f"{next_point['V']}").bold = True
#             p.add_run(f"{complemento} de coordenadas N(Y) {coord_N_formatada} e E(X) {coord_E_formatada};")
#             doc_word.add_paragraph()

        
#         data_atual = datetime.now().strftime("%d de %B de %Y")

#         # converte mês para português
#         for ingles, portugues in MESES_PT_BR.items():
#             if ingles in data_atual:
#                 data_atual = data_atual.replace(ingles, portugues)
#                 break
#         doc_word.add_paragraph(f"\nPorto Alegre, RS, {data_atual}.", style='Normal')
#         doc_word.add_paragraph("\n\n")

#         output_path = os.path.join(diretorio_concluido, f"{uuid_str}_FECHADA_{tipo}_{matricula}.docx")
#         doc_word.save(output_path)
#         logger.info(f"✅ Memorial descritivo salvo em: {output_path}")

#     except Exception as e:
#         logger.error(f"❌ Erro ao criar memorial descritivo: {e}")


def create_memorial_document(
    uuid_str, proprietario, matricula, descricao, excel_file_path, template_path, 
    output_path, perimeter_dxf, area_dxf, desc_ponto_Az, Coorde_E_ponto_Az, Coorde_N_ponto_Az,
    azimuth, distance, uso_solo, area_imovel, cidade, rua, comarca, rgi, caminho_salvar, tipo
):
    try:
        # ── Sanidade de caminhos ───────────────────────────────────────────────
        try:
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
        except Exception as e:
            logger.warning(f"[DOCX] Falha ao garantir diretório de saída {os.path.dirname(output_path)}: {e}")

        if not os.path.exists(template_path):
            logger.warning(f"[DOCX] Template não encontrado: {template_path}")

        if not os.path.exists(excel_file_path):
            logger.warning(f"[DOCX] XLSX base não encontrado: {excel_file_path}")

        # ── Ler Excel e garantir colunas numéricas ─────────────────────────────
        df = pd.read_excel(excel_file_path)
        for col in ("N", "E", "Distancia(m)"):
            if col not in df.columns:
                logger.warning(f"[DOCX] Coluna '{col}' ausente no XLSX ({excel_file_path}).")

        if "N" in df.columns:
            df["N"] = pd.to_numeric(df["N"].astype(str).str.replace(",", "."), errors="coerce")
        if "E" in df.columns:
            df["E"] = pd.to_numeric(df["E"].astype(str).str.replace(",", "."), errors="coerce")
        if "Distancia(m)" in df.columns:
            df["Distancia(m)"] = pd.to_numeric(df["Distancia(m)"].astype(str).str.replace(",", "."), errors="coerce")

        # ── Criar documento Word a partir do template ──────────────────────────
        doc_word = Document(template_path)
        set_default_font(doc_word)

        p1 = doc_word.add_paragraph("MEMORIAL DESCRITIVO", style="Normal")
        p1.runs[0].bold = True
        p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc_word.add_paragraph()

        texto_tipo = {
            "ETE": f"Área da matrícula {matricula} destinada a {descricao} - SES de {cidade}",
            "REM": f"Área remanescente da matrícula {matricula} destinada a {descricao}",
            "SER": f"Área da matrícula {matricula} destinada à SERVIDÃO ADMINISTRATIVA DE ACESSO À {descricao}",
            "ACE": f"Área da matrícula {matricula} destinada ao ACESSO DA SERVIDÃO ADMINISTRATIVA DA {descricao}",
        }.get(tipo, "Tipo não especificado")

        p = doc_word.add_paragraph(style="Normal")
        p.add_run("Imóvel: ")
        p.add_run(texto_tipo).bold = True

        # ⚠️ Usar 'rgi' (parâmetro), não 'RI'
        doc_word.add_paragraph(f"Matrícula: Número - {matricula} do {rgi} de {comarca}", style="Normal")
        doc_word.add_paragraph(f"Proprietário: {proprietario}", style="Normal")
        doc_word.add_paragraph(f"Local: {rua} - {cidade}", style="Normal")
        doc_word.add_paragraph(f"Área: {area_dxf:,.2f} m²".replace(",", "X").replace(".", ",").replace("X", "."), style="Normal")
        doc_word.add_paragraph(f"Perímetro: {perimeter_dxf:,.2f} m".replace(",", "X").replace(".", ",").replace("X", "."), style="Normal")
        doc_word.add_paragraph()

        area_dxf_formatada = f"{area_dxf:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        texto_paragrafo = (
            f"Área {uso_solo} com {area_dxf_formatada} m², parte de um todo maior da Matrícula Nº {matricula} com {area_imovel} "
            f"do {rgi} de {comarca}, localizada na {rua}, na cidade de {cidade}, definida através do seguinte levantamento "
            "topográfico, onde os ângulos foram medidos no sentido horário."
        )
        p = doc_word.add_paragraph(texto_paragrafo, style="Normal")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        doc_word.add_paragraph()

        coord_E_ponto_Az = f"{Coorde_E_ponto_Az:.3f}".replace(".", ",")
        coord_N_ponto_Az = f"{Coorde_N_ponto_Az:.3f}".replace(".", ",")
        doc_word.add_paragraph(
            f"O Ponto Az, ponto de amarração, está localizado na {desc_ponto_Az} nas coordenadas "
            f"E(X) {coord_E_ponto_Az} e N(Y) {coord_N_ponto_Az}.", style="Normal"
        ).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        doc_word.add_paragraph()

        distance_formatada = f"{distance:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        p = doc_word.add_paragraph(style="Normal")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.add_run(f"Daí, com Azimute de {convert_to_dms(azimuth)} e distância de {distance_formatada} m, chega-se ao Vértice ")
        p.add_run("V1").bold = True
        p.add_run(", origem da descrição desta área.")
        doc_word.add_paragraph()

        # Início da descrição pelo primeiro vértice
        if not df.empty:
            initial = df.iloc[0]
            try:
                coord_N_inicial = f"{initial['N']:.3f}".replace(".", ",")
                coord_E_inicial = f"{initial['E']:.3f}".replace(".", ",")
            except Exception:
                coord_N_inicial = coord_E_inicial = "--"

            doc_word.add_paragraph("Pontos definidos pelas Coordenadas Planas no Sistema U.T.M. – SIRGAS 2000.", style="Normal")
            doc_word.add_paragraph()

            p2 = doc_word.add_paragraph(style="Normal")
            p2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            p2.add_run("Inicia-se a descrição deste perímetro no vértice ")
            p2.add_run(f"{initial.get('V', 'V1')}").bold = True
            p2.add_run(f", de coordenadas N(Y) {coord_N_inicial} e E(X) {coord_E_inicial}, situado no limite com {initial.get('Confrontante','')}.")
            doc_word.add_paragraph()

            # Trechos
            for i in range(len(df)):
                current = df.iloc[i]
                next_point = df.iloc[(i + 1) % len(df)]

                try:
                    distancia_formatada = f"{float(current['Distancia(m)']):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
                except Exception:
                    distancia_formatada = "--"

                try:
                    coord_N_formatada = f"{float(next_point['N']):.3f}".replace(".", ",")
                    coord_E_formatada = f"{float(next_point['E']):.3f}".replace(".", ",")
                except Exception:
                    coord_N_formatada = coord_E_formatada = "--"

                complemento = ", origem desta descrição," if next_point.get("V", "") == "V1" else ""

                p = doc_word.add_paragraph(style="Normal")
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                p.add_run(
                    f"Deste, segue com azimute de {current.get('Azimute','')} e distância de {distancia_formatada} m, "
                    f"confrontando neste trecho com área pertencente à {current.get('Confrontante','')}, até o vértice "
                )
                p.add_run(f"{next_point.get('V','')}").bold = True
                p.add_run(f"{complemento} de coordenadas N(Y) {coord_N_formatada} e E(X) {coord_E_formatada};")
                doc_word.add_paragraph()

        # Data em PT-BR
        data_atual = datetime.now().strftime("%d de %B de %Y")
        for ingles, portugues in MESES_PT_BR.items():
            if ingles in data_atual:
                data_atual = data_atual.replace(ingles, portugues)
                break
        doc_word.add_paragraph(f"\nPorto Alegre, RS, {data_atual}.", style="Normal")
        doc_word.add_paragraph("\n\n")

        # ── Salvar no caminho recebido (não reconstruir) ───────────────────────
        doc_word.save(output_path)
        logger.info(f"✅ Memorial descritivo salvo em: {output_path}")

        # ── Sanidade pós-salvamento ───────────────────────────────────────────
        try:
            exists = os.path.exists(output_path)
            logger.info(f"[SANITY DOCX] Existe? {exists} -> {output_path}")
            try:
                listing = os.listdir(os.path.dirname(output_path))
                logger.info(f"[SANITY DOCX] Conteúdo do diretório (parcial): {listing[:20]}")
            except Exception as e:
                logger.warning(f"[SANITY DOCX] Falha ao listar diretório destino: {e}")
        except Exception as e:
            logger.warning(f"[SANITY DOCX] Falha na checagem de existência: {e}")

        return output_path

    except Exception as e:
        logger.error(f"❌ Erro ao criar memorial descritivo: {e}")
        return None



        
# Função principal
def main_poligonal_fechada(uuid_str, excel_path, dxf_path, diretorio_preparado, diretorio_concluido, caminho_template, sentido_poligonal='horario'):

    caminho_salvar = diretorio_concluido 
    os.makedirs(caminho_salvar, exist_ok=True)

    # 🔹 Carrega dados do imóvel
    dados_imovel_df = pd.read_excel(excel_path, sheet_name='Dados_do_Imóvel', header=None)
    dados_imovel = dict(zip(dados_imovel_df.iloc[:, 0], dados_imovel_df.iloc[:, 1]))

    # 🔹 Extrai variáveis necessárias
    proprietario = dados_imovel.get("NOME DO PROPRIETÁRIO", "").strip()
    matricula = sanitize_filename(dados_imovel.get("DOCUMENTAÇÃO DO IMÓVEL", "").strip())
    descricao = dados_imovel.get("OBRA", "").strip()
    uso_solo = dados_imovel.get("ZONA", "").strip()
    area_imovel = dados_imovel.get("ÁREA TOTAL DO TERRENO DOCUMENTADA", "").replace("\t", "").replace("\n", "").strip()
    cidade = dados_imovel.get("CIDADE", "").strip()
    rua = dados_imovel.get("LOCAL", "").strip()
    comarca = dados_imovel.get("COMARCA", "").strip()
    rgi = dados_imovel.get("RI", "").strip()
    desc_ponto_Az = dados_imovel.get("AZ", "").strip()

    # 🔹 Define tipo pela nomenclatura do DXF
    dxf_filename = os.path.basename(dxf_path).upper()

    if "ETE" in dxf_filename:
        tipo = "ETE"
    elif "REM" in dxf_filename:
        tipo = "REM"
    elif "SER" in dxf_filename:
        tipo = "SER"
    elif "ACE" in dxf_filename:
        tipo = "ACE"
    else:
        logger.error("❌ Tipo (ETE, REM, SER ou ACE) não identificado no nome do DXF.")
        return

    # 🔹 Busca planilha FECHADA correta com uuid_str
    padrao_busca = os.path.join(diretorio_preparado, f"{uuid_str}_FECHADA_{tipo}.xlsx")
    arquivos_encontrados = glob.glob(padrao_busca)

    if not arquivos_encontrados:
        logger.error(f"❌ Planilha confrontantes FECHADA não encontrada: {padrao_busca}")
        return
    confrontantes = _carregar_confrontantes(uuid_str, tipo, diretorio_preparado)

    # 4) Ler geometria do DXF ORIGINAL
    doc0, lines0, perimeter0, area0, ponto_az_dxf, msp0, pts_bulge0 = get_document_info_from_dxf(dxf_path)
    if not (doc0 and lines0):
        logger.error("❌ Nenhuma polilinha fechada encontrada no DXF original.")
        return
    if not ponto_az_dxf:
        logger.error("❌ Ponto Az não encontrado no DXF original.")
        return
    if len(lines0) < 2:
        logger.error("❌ Poligonal precisa de ao menos 2 segmentos.")
        return

    v1 = lines0[0][0]
    v_next = lines0[1][0]
    v_prev = lines0[-1][0]
    v2_for_arc = v_prev if str(sentido_poligonal).lower().startswith("anti") else v_next

    azimute_v1 = azimuth_deg(ponto_az_dxf, v1)
    distancia_az_v1 = calculate_distance(ponto_az_dxf, v1)
    giro_angular_v1 = calculate_angular_turn(ponto_az_dxf, v1, v2_for_arc)
    giro_angular_v1_dms = convert_to_dms(giro_angular_v1)

    logger.info(f"📐 Área (orig): {area0:.6f} m² | Perímetro (orig): {perimeter0:.6f} m")
    logger.info(f"📌 Az→V1: {azimute_v1:.4f}° | Dist(Az,V1): {distancia_az_v1:.2f} m | Giro@V1: {giro_angular_v1_dms}")

    Coord_E_ponto_Az = float(ponto_az_dxf[0])
    Coord_N_ponto_Az = float(ponto_az_dxf[1])

    # 5) Gerar DXF LIMPO R2010 e reler
    dxf_limpo_path = os.path.join(caminho_salvar, f"DXF_LIMPO_{matricula}.dxf")
    dxf_file_path = limpar_dxf_e_converter_r2010(dxf_path, dxf_limpo_path)

    doc, lines, perimeter_dxf, area_dxf, ponto_az, msp, pts_bulge = get_document_info_from_dxf(dxf_file_path)
    if not (doc and lines):
        logger.error("❌ Nenhuma polilinha fechada encontrada no DXF limpo.")
        return

    # 6) Reindexar para que o V1 do LIMPO seja o mesmo do Az→V1 do ORIGINAL
    lines, pts_bulge, rot_idx = rotate_polygon_start_at_v1(lines, pts_bulge, v1, sentido_poligonal)
    if rot_idx is not None:
        logger.info(f"Reindexado: V1 do DXF limpo alinhado ao vértice de Az→V1 (idx {rot_idx}).")
    else:
        logger.warning("Não foi possível reindexar a poligonal; mantendo ordem do DXF limpo.")

    # 7) Desenhar elementos de Az no DXF limpo
    try:
        add_az_marker_to_dxf(
            doc_dxf=doc,
            ponto_az=ponto_az_dxf,
            v1=v1,
            azimute_deg=azimute_v1,
            distancia_az_v1=distancia_az_v1,
            v2=v2_for_arc,
            sentido=sentido_poligonal,
            layer="Az_Marker",
            north_len=8.0,
            text_height=0.6,
            arc_radius=5.0,
            draw_minor_arc=False,
        )
    except Exception as e:
        logger.exception(f"Erro ao desenhar marcador de Az: {e}")



    excel_saida = os.path.join(diretorio_concluido, f"{uuid_str}_FECHADA_{tipo}_{matricula}.xlsx")


    # 8) Gerar a planilha FECHADA de AZIMUTE+DISTÂNCIA
    excel_resultado = create_memorial_descritivo(
        uuid_str=uuid_str,
        doc=doc,
        lines=lines,
        proprietario=proprietario,
        matricula=matricula,
        caminho_salvar=caminho_salvar,
        confrontantes=confrontantes,           # ✅ agora vai
        ponto_az=ponto_az_dxf,
        dxf_file_path=dxf_file_path,           # ✅ passe o LIMPO
        area_dxf=area_dxf,
        azimute=azimute_v1,                    # (graus decimais Az→V1)
        v1=v1,
        msp=msp,
        dxf_filename=dxf_filename,
        excel_file_path=excel_saida,           # arquivo de SAÍDA, não a planilha de entrada
        tipo=tipo,
        giro_angular_v1_dms=giro_angular_v1_dms,
        distancia_az_v1=distancia_az_v1,
        sentido_poligonal=sentido_poligonal,
        modo="AZIMUTE_AZ",
        points_bulge=pts_bulge
    )
    if not excel_resultado:
        logger.error("❌ Falha ao gerar memorial descritivo (planilha).")
        return

    # 9) Gerar DOCX
    output_docx_path = os.path.join(diretorio_concluido, f"{uuid_str}_FECHADA_{tipo}_{matricula}.docx")
    create_memorial_document(
        uuid_str=uuid_str,
        proprietario=proprietario,
        matricula=matricula,
        descricao=descricao,
        excel_file_path=excel_resultado,
        template_path=caminho_template,
        output_path=output_docx_path,
        perimeter_dxf=perimeter_dxf,
        area_dxf=area_dxf,
        desc_ponto_Az=desc_ponto_Az,
        Coorde_E_ponto_Az=Coord_E_ponto_Az,
        Coorde_N_ponto_Az=Coord_N_ponto_Az,
        azimuth=azimute_v1,
        distance=distancia_az_v1,
        uso_solo=uso_solo,
        area_imovel=area_imovel,
        cidade=cidade,
        rua=rua,
        comarca=comarca,
        rgi=rgi,
        caminho_salvar=caminho_salvar,
        tipo=tipo
    )

    logger.info("🔵 [main_poligonal_fechada] Processamento concluído com sucesso.")
    print("Processamento concluído com sucesso.")


#ATUALIZANDO
