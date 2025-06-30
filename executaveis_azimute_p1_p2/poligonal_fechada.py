import math
import csv
import pandas as pd
import glob
import re
from docx import Document
from docx.shared import Inches
from datetime import datetime
import os
import ezdxf
import time
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from decimal import Decimal, getcontext
import pandas as pd
from docx.shared import Pt
import openpyxl
from openpyxl.styles import Alignment, Font
import logging 

BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))

getcontext().prec = 28  # Define a precisão para 28 casas decimais

# Configuração manual para nomes dos meses em português (independente do locale)
MESES_PT_BR = [
    "janeiro", "fevereiro", "março", "abril", "maio", "junho",
    "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"
]

def obter_data_formatada(data):
    dia = data.day
    mes = MESES_PT_BR[data.month - 1]
    ano = data.year
    return f"{dia:02d} de {mes} de {ano}"

# Obter data atual formatada manualmente
data_atual = obter_data_formatada(datetime.now())

logger = logging.getLogger(__name__)
logger.setLevel(logging.DEBUG)  # Garanta que está em DEBUG

def is_clockwise(points):
    """
    Verifica se a poligonal está no sentido horário.
    Retorna True se for horário, False se anti-horário.
    """
    area = 0.0
    for i in range(len(points)):
        j = (i + 1) % len(points)
        area += points[i][0] * points[j][1]
        area -= points[j][0] * points[i][1]
    return area < 0

def ensure_counterclockwise(points):
    """
    Garantir que a lista de pontos esteja no sentido anti-horário.
    Se estiver no sentido horário, inverte os pontos.
    """
    if is_clockwise(points):
        points.reverse()
    return points


getcontext().prec = 28  # Define a precisão para 28 casas decimais

# Configurar locale para português do Brasil
#locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8')  # Para sistemas Linux ou Mac
# locale.setlocale(locale.LC_TIME, 'Portuguese_Brazil.1252')  # Para Windows, caso necessário

# Obter data atual formatada em português
data_atual = datetime.now().strftime("%d de %B de %Y")


# Função que processa as linhas da poligonal
# Função que processa as linhas da poligonal
def limpar_dxf_e_converter_r2010(original_path, saida_path):
    """
    Lê um DXF original e regrava o arquivo com a versão R2010,
    garantindo a preservação de entidades como ARC e CIRCLE.
    Não adiciona nenhum ponto ou geometria nova.
    """
    try:
        doc_antigo = ezdxf.readfile(original_path)
        msp_antigo = doc_antigo.modelspace()
        doc_novo = ezdxf.new(dxfversion='R2010')
        msp_novo = doc_novo.modelspace()

        encontrou_polilinha = False

        for entity in msp_antigo.query('LWPOLYLINE'):
            if entity.closed:
                pontos = [point[:2] for point in entity.get_points('xy')]
                msp_novo.add_lwpolyline(
                    pontos,
                    close=True,
                    dxfattribs={'layer': entity.dxf.layer}
                )
                encontrou_polilinha = True
                break

        if not encontrou_polilinha:
            raise ValueError("Nenhuma polilinha fechada encontrada no DXF original.")

        doc_novo.saveas(saida_path)
        logger.info(f"✅ DXF convertido e salvo como R2010 em: {saida_path}")
        return saida_path

    except Exception as e:
        logger.error(f"❌ Erro ao converter DXF para R2010: {e}")
        return original_path




def get_document_info_from_dxf(dxf_file_path):
    try:
        doc = ezdxf.readfile(dxf_file_path)  
        msp = doc.modelspace()  

        lines = []
        perimeter_dxf = 0
        area_dxf = 0
        ponto_az = None
        area_poligonal = None

        for entity in msp.query('LWPOLYLINE'):
            if entity.closed:
                points = entity.get_points('xy')  
                num_points = len(points)

                for i in range(num_points - 1):
                    start_point = (points[i][0], points[i][1])
                    end_point = (points[i + 1][0], points[i + 1][1])
                    lines.append((start_point, end_point))

                    segment_length = ((end_point[0] - start_point[0]) ** 2 + 
                                      (end_point[1] - start_point[1]) ** 2) ** 0.5
                    perimeter_dxf += segment_length

                start_point = (points[-1][0], points[-1][1])
                end_point = (points[0][0], points[0][1])
                lines.append((start_point, end_point))

                segment_length = ((end_point[0] - start_point[0]) ** 2 + 
                                  (end_point[1] - start_point[1]) ** 2) ** 0.5
                perimeter_dxf += segment_length

                x = [point[0] for point in points]
                y = [point[1] for point in points]
                area_dxf = abs(sum(x[i] * y[(i + 1) % num_points] - x[(i + 1) % num_points] * y[i] for i in range(num_points)) / 2)

                break  

        if not lines:
            logger.info("Nenhuma polilinha encontrada no arquivo DXF.")
            return None, [], 0, 0, None, None

        for entity in msp.query('TEXT'):
            if "Az" in entity.dxf.text:
                ponto_az = (entity.dxf.insert.x, entity.dxf.insert.y, 0)
                logger.info(f"Ponto Az encontrado em texto: {ponto_az}")
        
        for entity in msp.query('INSERT'):
            if "Az" in entity.dxf.name:
                ponto_az = (entity.dxf.insert.x, entity.dxf.insert.y, 0)
                logger.info(f"Ponto Az encontrado no bloco: {ponto_az}")
        
        for entity in msp.query('POINT'):
            ponto_az = (entity.dxf.location.x, entity.dxf.location.y, 0)
            logger.info(f"Ponto Az encontrado como ponto: {ponto_az}")
        

        # if not ponto_az:
        #     logger.info("Ponto Az não encontrado no arquivo DXF.")
        #     return None, lines, 0, 0, None, None

        logger.info(f"Linhas processadas: {len(lines)}")
        logger.info(f"Perímetro do DXF: {perimeter_dxf:.2f} metros")
        logger.info(f"Área do DXF: {area_dxf:.2f} metros quadrados")

        return doc, lines, perimeter_dxf, area_dxf, ponto_az, area_poligonal

    except Exception as e:
        logger.error(f"Erro ao obter informações do documento: {e}")
        return None, [], 0, 0, None, None
    
def obter_ponto_amarracao_anterior_v1(planilha_aberta_path):
    """
    Retorna o ponto imediatamente anterior ao V1 a partir da planilha de saída da poligonal aberta.
    Adaptado para colunas: 'Ponto', 'Coord_E', 'Coord_N'.
    """
    df = pd.read_excel(planilha_aberta_path, engine='openpyxl')

    if "Ponto" not in df.columns or "Coord_E" not in df.columns or "Coord_N" not in df.columns:
        raise ValueError("Planilha ABERTA não contém colunas 'Ponto', 'Coord_E' e 'Coord_N'.")

    idx_v1 = df[df["Ponto"] == "V1"].index
    if len(idx_v1) == 0:
        raise ValueError("Ponto V1 não encontrado na planilha.")
    elif idx_v1[0] == 0:
        raise ValueError("Não existe ponto anterior ao V1.")

    linha = df.iloc[idx_v1[0] - 1]
    e = float(str(linha["Coord_E"]).replace(",", "."))
    n = float(str(linha["Coord_N"]).replace(",", "."))
    codigo = linha["Ponto"]

    return (e, n), codigo


    
# 🔹 Função para definir a fonte padrão
def set_default_font(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    
def calculate_point_on_line(start, end, distance):
    """
    Calcula um ponto a uma determinada distância sobre uma linha.
    """
    dx, dy = end[0] - start[0], end[1] - start[1]
    length = math.hypot(dx, dy)
    return (
        start[0] + dx / length * distance,
        start[1] + dy / length * distance
    )

def add_azimuth_arc_to_dxf(msp, ponto_az, v1, azimute):
    """
    Adiciona o arco do azimute ao DXF usando ezdxf.
    """
    try:
        logger.info(f"Iniciando a adição do arco de azimute. Azimute: {azimute}°")

        # Criar camada 'Azimute', se não existir
        if 'Azimute' not in msp.doc.layers:
            msp.doc.layers.new(name='Azimute', dxfattribs={'color': 1})
            logger.info("Camada 'Azimute' criada com sucesso.")

        # Traçar segmento entre Az e V1
        msp.add_line(start=ponto_az, end=v1, dxfattribs={'layer': 'Azimute'})
        logger.info(f"Segmento entre Az e V1 desenhado de {ponto_az} para {v1}")

        # Traçar segmento para o norte
        north_point = (ponto_az[0], ponto_az[1] + 2)
        msp.add_line(start=ponto_az, end=north_point, dxfattribs={'layer': 'Azimute'})
        logger.info(f"Linha para o norte desenhada com sucesso de {ponto_az} para {north_point}")

        # Calcular o ponto inicial (1 metro de Az para V1)
        # Calcular distância entre ponto Az e V1 para definir raio adaptativo
        dist = calculate_distance(ponto_az, v1)
        radius = 0.4 if dist <= 0.5 else 1.0

        # Calcular os pontos do arco com esse raio
        start_arc = calculate_point_on_line(ponto_az, v1, radius)
        end_arc = calculate_point_on_line(ponto_az, north_point, radius)

        # Traçar o arco do azimute
        msp.add_arc(
            center=ponto_az,
            radius=radius,
            start_angle=math.degrees(math.atan2(start_arc[1] - ponto_az[1], start_arc[0] - ponto_az[0])),
            end_angle=math.degrees(math.atan2(end_arc[1] - ponto_az[1], end_arc[0] - ponto_az[0])),
            dxfattribs={'layer': 'Azimute'}
        )
        logger.info(f"Arco do azimute desenhado com sucesso com valor de {azimute}° no ponto {ponto_az}")

       # Adicionar rótulo do azimute diretamente com o texto "Azimute:"
        azimuth_label = f"Azimute: {convert_to_dms(azimute)}"  # Incluir o prefixo "Azimute:"

        # Calcular a posição do rótulo
        label_position = (
            ponto_az[0] + 1.0 * math.cos(math.radians(azimute / 2)),
            ponto_az[1] + 1.0 * math.sin(math.radians(azimute / 2))
        )

        # Adicionar o texto ao desenho
        msp.add_text(
            azimuth_label,
            dxfattribs={
                'height': 0.25,
                'layer': 'Azimute',
                'insert': label_position  # Define a posição diretamente
            }
        )

        logger.info(f"Rótulo do azimute adicionado com sucesso: '{azimuth_label}' em {label_position}")


    except Exception as e:
        logger.error(f"Erro na função `add_azimuth_arc_to_dxf`: {e}")

def calculate_polygon_area(points):
    """
    Calcula a área de uma poligonal fechada utilizando o método do produto cruzado com alta precisão.
    :param points: Lista de coordenadas [(x1, y1), (x2, y2), ...].
    :return: Área da poligonal.
    """
    n = len(points)
    if n < 3:
        return Decimal(0)  # Não é uma poligonal válida

    area = Decimal(0)
    for i in range(n):
        x1, y1 = Decimal(points[i][0]), Decimal(points[i][1])
        x2, y2 = Decimal(points[(i + 1) % n][0]), Decimal(points[(i + 1) % n][1])
        area += x1 * y2 - y1 * x2

    return abs(area) / 2
    
def convert_to_dms(decimal_degrees):
    """
    Converte graus decimais para o formato graus, minutos e segundos (DMS).
    """
    try:
        # Verificar se o valor é NaN
        if math.isnan(decimal_degrees):
            raise ValueError("Valor de entrada é NaN")

        degrees = int(decimal_degrees)
        minutes = int((abs(decimal_degrees) - abs(degrees)) * 60)
        seconds = round((abs(decimal_degrees) - abs(degrees) - minutes / 60) * 3600, 2)
        return f"{degrees}°{minutes}'{seconds}\""
    except Exception as e:
        logger.info(f"Erro na conversão para DMS: {e}")
        return "0°0'0.00\""  # Valor padrão em caso de erro

def calculate_distance(p1, p2):
    return math.sqrt((p2[0] - p1[0])**2 + (p2[1] - p1[1])**2)


def degrees_to_dms(angle):
    """
    Converte um ângulo em graus decimais para o formato graus, minutos e segundos (° ' ").
    """
    degrees = int(angle)
    minutes = int((angle - degrees) * 60)
    seconds = round((angle - degrees - minutes / 60) * 3600, 2)
    return f"{degrees}°{minutes}'{seconds}\""
import math

def calculate_azimuth(p1, p2):
    """
    Calcula o azimute entre dois pontos p1 e p2.
    
    :param p1: Coordenadas do ponto 1 (E, N)
    :param p2: Coordenadas do ponto 2 (E, N)
    :return: Azimute em graus decimais
    """
    delta_x = p2[0] - p1[0]
    delta_y = p2[1] - p1[1]
    
    azimuth = math.degrees(math.atan2(delta_x, delta_y)) % 360  # Garantir valor entre 0° e 360°
    
    return azimuth


def create_arrow_block(doc, block_name="ARROW"):
    """
    Cria um bloco no DXF representando uma seta sólida como um triângulo.
    """
    if block_name in doc.blocks:
        return  # O bloco já existe

    block = doc.blocks.new(name=block_name)

    # Definir o triângulo da seta
    length = 0.5  # Comprimento da seta
    base_half_length = length / 2

    tip = (0, 0)  # Ponta da seta no eixo de coordenadas
    base1 = (-base_half_length, -length)
    base2 = (base_half_length, -length)

    block.add_solid([base1, base2, tip])
import math

def add_giro_angular_arc_to_dxf(doc_dxf, v1, az, v2):

    """
    Adiciona um arco representando o giro angular horário no espaço de modelo do DXF já aberto.
    """
    try:
        msp = doc_dxf.modelspace()
        # Calcular distância entre V1–Az e V1–V2 para escolher o menor
        dist_az = calculate_distance(v1, az)
        dist_v2 = calculate_distance(v1, v2)
        min_dist = min(dist_az, dist_v2)

        radius = 0.4 if min_dist <= 0.5 else 1.0  # Raio adaptativo
        # Traçar a reta entre V1 e Az
        msp.add_line(start=v1[:2], end=az[:2])
        logger.info(f"Linha entre V1 e Az traçada com sucesso.")

        # Definir os pontos de apoio
        def calculate_displacement(point1, point2, distance):
            dx = point2[0] - point1[0]
            dy = point2[1] - point1[1]
            magnitude = math.hypot(dx, dy)
            return (
                point1[0] + (dx / magnitude) * distance,
                point1[1] + (dy / magnitude) * distance,
            )

        # Calcular os pontos de apoio
        ponto_inicial = calculate_displacement(v1, v2, radius)  # 2m na reta V1-V2
        ponto_final = calculate_displacement(v1, az, radius)   # 2m na reta V1-Az

        # Calcular os ângulos dos vetores
        angle_v2 = math.degrees(math.atan2(ponto_inicial[1] - v1[1], ponto_inicial[0] - v1[0]))
        angle_az = math.degrees(math.atan2(ponto_final[1] - v1[1], ponto_final[0] - v1[0]))

        # Calcular o giro angular no sentido horário
        giro_angular = (angle_az - angle_v2) % 360  # Garantir que o ângulo esteja no intervalo [0, 360)
        if giro_angular < 0:  # Caso negativo, ajustar para o sentido horário
            giro_angular += 360

        logger.info(f"Giro angular calculado corretamente: {giro_angular:.2f}°")

        # Traçar o arco
        msp.add_arc(center=v1[:2], radius=radius, start_angle=angle_v2, end_angle=angle_az)
        logger.info(f"Arco do giro angular traçado com sucesso.")

        # Adicionar rótulo ao arco
        label_offset = 3.0
        deslocamento_x=3
        deslocamento_y=-3
        angle_middle = math.radians((angle_v2 + angle_az) / 2)
        label_position = (
            v1[0] + (label_offset+deslocamento_x) * math.cos(angle_middle),
            v1[1] + (label_offset+deslocamento_y) * math.sin(angle_middle),
        )
        # Converter o ângulo para DMS e exibir no rótulo
        giro_angular_dms = f"{convert_to_dms(giro_angular)}"
        msp.add_text(
            giro_angular_dms,
            dxfattribs={
                'height': 0.3,
                'layer': 'Labels',
                'insert': label_position  # Define a posição do texto
            }
        )
        logger.info(f"Rótulo do giro angular ({giro_angular_dms}) adicionado com sucesso.")

    except Exception as e:
        logger.error(f"Erro ao adicionar o arco do giro angular ao DXF: {e}") 



def calculate_arc_angles(p1, p2, p3):
    try:
        # Vetores a partir de p2
        dx1, dy1 = p1[0] - p2[0], p1[1] - p2[1]  # Vetor de p2 para p1
        dx2, dy2 = p3[0] - p2[0], p3[1] - p2[1]  # Vetor de p2 para p3

        # Ângulos dos vetores em relação ao eixo X
        angle1 = math.degrees(math.atan2(dy1, dx1)) % 360
        angle2 = math.degrees(math.atan2(dy2, dx2)) % 360

        # AQUI ESTÁ A CORREÇÃO:
        internal_angle = (angle1 - angle2) % 360

        if internal_angle > 180:
            #internal_angle = 360 - internal_angle
            internal_angle = internal_angle
            start_angle = angle1
            end_angle = angle2
        else:
            start_angle = angle2
            end_angle = angle1

        return start_angle % 360, end_angle % 360

    except Exception as e:
        logger.info(f"Erro ao calcular ângulos do arco: {e}")
        return 0, 0





def insert_and_rotate_arrow(msp, center, radius, angle, rotation_adjustment, block_name="ARROW"):
    """
    Insere um bloco da seta alinhado ao arco e rotacionado conforme especificado.
    """
    arrow_tip_x = center[0] + radius * math.cos(angle)
    arrow_tip_y = center[1] + radius * math.sin(angle)
    arrow_tip = (arrow_tip_x, arrow_tip_y)

    rotation_angle = math.degrees(angle) + rotation_adjustment

    msp.add_blockref(
        block_name,
        insert=arrow_tip,
        dxfattribs={"rotation": rotation_angle}
    )
import math





def add_angle_visualization_to_dwg(msp, ordered_points, angulos_excel):
    """
    Adiciona ângulos internos no espaço de modelo do DXF usando diretamente os ângulos calculados do Excel.
    """
    try:
        total_points = len(ordered_points)
        
        for i, p2 in enumerate(ordered_points):
            if i == 0:
                logger.info("⏩ Ignorando arco e rótulo para V1")
                continue
            p1 = ordered_points[i - 1] if i > 0 else ordered_points[-1]
            p3 = ordered_points[(i + 1) % total_points]

            def calculate_displacement(base_point, direction_point, distance):
                dx = direction_point[0] - base_point[0]
                dy = direction_point[1] - base_point[1]
                magnitude = math.hypot(dx, dy)
                return (
                    base_point[0] + (dx / magnitude) * distance,
                    base_point[1] + (dy / magnitude) * distance,
                )

            # 🔁 Raio adaptativo com base na distância entre vértices adjacentes
            dist_lado = math.hypot(p3[0] - p2[0], p3[1] - p2[1])
            # 🔁 Raio adaptativo com base na distância entre vértices adjacentes
            dist_lado = math.hypot(p3[0] - p2[0], p3[1] - p2[1])

            # 🔁 Raio adaptativo com base na menor distância entre os dois lados que chegam ao vértice
            lado_antes = math.hypot(p2[0] - p1[0], p2[1] - p1[1])
            lado_depois = math.hypot(p3[0] - p2[0], p3[1] - p2[1])
            lado_menor = min(lado_antes, lado_depois)

            if lado_menor <= 0.5:
                radius = lado_menor * 0.8  # proporcional ao menor lado
            else:
                radius = 1.0


            ponto_inicial = calculate_displacement(p2, p3, radius)
            ponto_final = calculate_displacement(p2, p1, radius)

            start_angle = math.degrees(math.atan2(ponto_inicial[1] - p2[1], ponto_inicial[0] - p2[0]))
            end_angle = math.degrees(math.atan2(ponto_final[1] - p2[1], ponto_final[0] - p2[0]))

            # Garantir que o arco desenhe no sentido correto
            if end_angle < start_angle:
                end_angle += 360
            logger.info(f"--- V{i+1} ---")
            logger.info(f"Ângulo Excel: {repr(angulos_excel[i])}")
            logger.info(f"Raio: {radius:.2f}")
            logger.info(f"Start angle: {start_angle:.2f}°, End angle: {end_angle:.2f}°")

            try:
                # Adicionar o arco interno ao desenho
                msp.add_arc(
                    center=p2,
                    radius=radius,
                    start_angle=start_angle,
                    end_angle=end_angle,
                    dxfattribs={'layer': 'Internal_Arcs'}
                )

                # Agora, usa diretamente o ângulo do Excel (sem recalcular!)
                internal_angle_dms = angulos_excel[i]

                # Adicionar rótulo do ângulo interno
                label_offset = 1
                label_position = (
                    p2[0] + label_offset * math.cos(math.radians((start_angle + end_angle) / 2)),
                    p2[1] + label_offset * math.sin(math.radians((start_angle + end_angle) / 2))
                )

                msp.add_text(
                    internal_angle_dms,
                    dxfattribs={
                        'height': 0.3,
                        'layer': 'Labels',
                        'insert': label_position
                    }
                )
            except Exception as e:
                logger.error(f"Erro ao adicionar arco ou rótulo no vértice V{i+1}: {e}")

            logger.info(f"Vértice V{i+1}: Ângulo interno {internal_angle_dms}")

    except Exception as e:
        logger.error(f"Erro ao adicionar ângulos internos ao DXF: {e}")




def calculate_internal_angle(p1, p2, p3):
    try:
        dx1, dy1 = p1[0] - p2[0], p1[1] - p2[1]
        dx2, dy2 = p3[0] - p2[0], p3[1] - p2[1]

        angle1 = math.atan2(dy1, dx1)
        angle2 = math.atan2(dy2, dx2)

        internal_angle = math.degrees(angle2 - angle1) % 360
        return internal_angle

    except Exception as e:
        logger.info(f"Erro inesperado ao calcular o ângulo interno: {e}")
        return 0



def calculate_label_position(p2, start_angle, end_angle, radius=1.8):
    """
    Calcula a posição do rótulo do ângulo interno, deslocando-o para uma posição central no arco.
    
    :param p2: Ponto central do arco (vértice da poligonal)
    :param start_angle: Ângulo de início do arco
    :param end_angle: Ângulo de fim do arco
    :param radius: Raio do deslocamento para evitar sobreposição
    :return: Posição (x, y) onde o rótulo do ângulo interno será colocado
    """
    try:
        # Calcular o ângulo médio entre os dois ângulos do arco
        mid_angle = math.radians((start_angle + end_angle) / 2)

        # Calcular a posição do rótulo deslocado no ângulo médio
        label_x = p2[0] + radius * math.cos(mid_angle)
        label_y = p2[1] + radius * math.sin(mid_angle)

        return (label_x, label_y)

    except Exception as e:
        logger.info(f"Erro ao calcular posição do rótulo do ângulo: {e}")
        return p2  # Retorna o próprio ponto central caso ocorra erro


import math

def add_label_and_distance(msp, start_point, end_point, label, distance):
    """
    Adiciona rótulos e distâncias no espaço de modelo usando ezdxf.
    """
    try:
        # Calcular ponto médio
        mid_point = (
            (start_point[0] + end_point[0]) / 2,
            (start_point[1] + end_point[1]) / 2
        )

        # Vetor da linha
        dx = end_point[0] - start_point[0]
        dy = end_point[1] - start_point[1]
        length = math.hypot(dx, dy)

        # Ângulo da linha
        angle = math.degrees(math.atan2(dy, dx))

        # Corrigir para manter a leitura sempre da esquerda para a direita
        if angle < -90 or angle > 90:
            angle += 180  

        # Afastar o rótulo da linha
        offset = -0.5  # Ajuste o valor para mais afastamento
        perp_x = -dy / length * offset
        perp_y = dx / length * offset
        displaced_mid_point = (mid_point[0] + perp_x, mid_point[1] + perp_y)

        # Criar layer se não existir
        if "Distance_Labels" not in msp.doc.layers:
            msp.doc.layers.new(name="Distance_Labels", dxfattribs={"color": 2})  # Define cor para melhor visualização

        # Adicionar o rótulo da distância no LAYER correto
        msp.add_text(
            f"{distance:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".") + "",
            dxfattribs={
                "height": 0.3,
                "layer": "Distance_Labels",  # Novo LAYER
                "rotation": angle,
                "insert": displaced_mid_point
            }
        )
        logger.info(f"Distância {distance:.2f} m adicionada corretamente em {displaced_mid_point}")

    except Exception as e:
        logger.error(f"Erro ao adicionar rótulo de distância: {e}")

def calculate_angular_turn(p1, p2, p3):
    """
    Calcula o giro angular no ponto `p2` entre os segmentos `p1-p2` e `p2-p3` no sentido horário.
    Retorna o ângulo em graus.
    """
    import math
    
    dx1, dy1 = p1[0] - p2[0], p1[1] - p2[1]  # Vetor do segmento p1-p2
    dx2, dy2 = p3[0] - p2[0], p3[1] - p2[1]  # Vetor do segmento p2-p3

    angle1 = math.atan2(dy1, dx1)
    angle2 = math.atan2(dy2, dx2)

    # Calcula o ângulo horário
    angular_turn = (angle2 - angle1) % (2 * math.pi)
    angular_turn_degrees = math.degrees(angular_turn)

    return angular_turn_degrees




def create_memorial_descritivo(
        uuid_str, doc, lines, proprietario, matricula, caminho_salvar, confrontantes, ponto_amarracao,
        dxf_file_path, area_dxf, azimute, v1, msp, dxf_filename, excel_file_path, tipo,
        giro_angular_v1_dms,
        diretorio_concluido=None
    ):
    """
    Cria o memorial descritivo e atualiza o DXF com base no ponto de amarração real (anterior ao V1).
    """

    if diretorio_concluido is None:
        diretorio_concluido = caminho_salvar
    
    # Carregar a planilha de confrontantes
    confrontantes_df = pd.read_excel(excel_file_path)

    # Número de registros no arquivo
    num_registros = len(confrontantes_df)

    # Transformar o dataframe em um dicionário de código -> confrontante
    confrontantes_dict = dict(zip(confrontantes_df['Código'], confrontantes_df['Confrontante']))

    # Verificar se a planilha foi carregada corretamente
    if not confrontantes_dict:
        print("Erro ao carregar confrontantes.")
        return None

    if not lines:
        print("Nenhuma linha disponível para criar o memorial descritivo.")
        return None

#     # Coletar os pontos diretamente na ordem dos vértices no DXF
#     ordered_points = [line[0] for line in lines] + [lines[-1][1]]  # Fechando a poligonal

    # Coletar os pontos diretamente na ordem dos vértices no DXF
    ordered_points = [line[0] for line in lines] + [lines[-1][1]]  # Fechando a poligonal
    num_vertices = len(ordered_points)
    
    # Calcular a área da poligonal
    area = calculate_polygon_area(ordered_points)

    # Ajustar para o sentido horário se necessário
    if area < 0:  # Sentido horário ou antihorário
        ordered_points.reverse()  # Reorganizar para sentido horário
        print(f"Pontos reorganizados para sentido horário: {ordered_points}")
       
    # Preparar os dados para o Excel
    data = []
    for i in range(len(ordered_points) - 1):
        start_point = ordered_points[i]
        end_point = ordered_points[i + 1]

        # Calcular azimute e distância
        azimuth, distance = calculate_azimuth_and_distance(start_point, end_point)
        azimuth_dms = convert_to_dms(azimuth)

        # Buscar o confrontante
        confrontante = confrontantes_df.iloc[i]['Confrontante'] if i < len(confrontantes_df) else "Desconhecido"

        # Adicionar as coordenadas do ponto Az apenas na primeira linha
        coord_e_ponto_az = f"{ponto_az[0]:.3f}".replace('.', ',') if i == 0 else ""
        coord_n_ponto_az = f"{ponto_az[1]:.3f}".replace('.', ',') if i == 0 else ""
        
        # Corrigindo geração da coluna "Divisa"
        if i < len(ordered_points) - 2:
            divisa_label = f"V{i + 1}_V{i + 2}"
        else:  # Para o último ponto, fecha com V1
            divisa_label = f"V{i + 1}_V1"


        # Adicionar linha ao conjunto de dados
        data.append({
            "V": f"V{i + 1}",
            "E": f"{start_point[0]:.3f}".replace('.', ','),
            "N": f"{start_point[1]:.3f}".replace('.', ','),
            "Z": "0.000",
            "Divisa": divisa_label,
            "Azimute": azimuth_dms,
            "Distancia(m)": f"{distance:.2f}".replace('.', ','),
            "Confrontante": confrontante,
         })

        # Adicionar rótulos e distância ao DXF
        add_label_and_distance(doc, msp, start_point, end_point, f"V{i + 1}", distance)


    # Criar DataFrame e salvar em Excel
    df = pd.DataFrame(data, dtype=str)
    excel_output_path = os.path.join(caminho_salvar, f"FECHADA_{tipo}_Memorial_{matricula}.xlsx")
    df.to_excel(excel_output_path, index=False)

    # Formatar o Excel
    wb = openpyxl.load_workbook(excel_output_path)
    ws = wb.active

    # Formatar cabeçalho
    for cell in ws[1]:
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center")

    # Ajustar largura das colunas
    column_widths = {
        "A": 8,   # V
        "B": 15,  # E
        "C": 15,  # N
        "D": 10,  # Z
        "E": 20,  # Divisa
        "F": 15,  # Azimute
        "G": 15,  # Distancia(m)
        "H": 30,  # Confrontante
    }
    for col, width in column_widths.items():
        ws.column_dimensions[col].width = width

    # Centralizar o conteúdo das células
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=1, max_col=ws.max_column):
        for cell in row:
            cell.alignment = Alignment(horizontal="center", vertical="center")

    # Salvar o arquivo formatado
    wb.save(excel_output_path)
    print(f"Arquivo Excel salvo e formatado em: {excel_output_path}")

    

    # Salvar o arquivo DXF com as alterações
    try:
        dxf_output_path = os.path.join(
        diretorio_concluido,
        f"{uuid_str}_FECHADA_{tipo}_{matricula}.dxf"
    )

        doc.saveas(dxf_output_path)
        print(f"Arquivo DXF atualizado salvo em: {dxf_output_path}")
    except Exception as e:
        print(f"Erro ao salvar o arquivo DXF final: {e}")

    return excel_output_path







def generate_initial_text(proprietario, matricula, descricao, area, perimeter, rua, cidade, ponto_amarracao, azimute, distancia):
    """
    Gera o texto inicial do memorial descritivo.
    """
    initial_text = (
        f"MEMORIAL DESCRITIVO\n"
        f"NOME PROPRIETÁRIO / OCUPANTE: {proprietario}\n"
        f"DESCRIÇÃO: {descricao}\n"
        f"DOCUMENTAÇÃO: MATRÍCULA {matricula}\n"
        f"ÁREA DO IMÓVEL: {area:.2f} metros quadrados\n"
        f"PERÍMETRO: {perimeter:.4f} metros\n"
        f"Área localizada na rua {rua}, município de {cidade}, com a seguinte descrição:\n"
        f"O Ponto Az está localizado nas coordenadas E {ponto_amarracao[0]:.3f}, N {ponto_amarracao[1]:.3f}.\n"
        f"Daí, com Azimute de {azimute} e distância de {distancia:.2f} metros, chega-se ao Vértice V1, "
        f"origem da área descrição, alinhado com a rua {rua}."
    )
    return initial_text


def generate_angular_text(az_v1, v1_v2, distancia, angulo, rua, confrontante):
    """
    Gera o texto do ângulo em V1 entre Az-V1 e V1-V2.
    """
    angular_text = (
        f"Daí, visando o Ponto “Az”, com giro angular horário de {angulo} e diste {distancia:.2f} metros, "
        f"chega-se ao Vértice V2, também alinhado com a rua {rua} e limítrofe com {confrontante}."
    )
    return angular_text


def generate_recurring_text(df, rua, confrontantes):
    """
    Gera o texto recorrente enquanto percorre os vértices da poligonal.
    """
    recurring_texts = []
    num_vertices = len(df)
    
    for i in range(1, num_vertices - 1):  # De V2 até o penúltimo vértice
        current = df.iloc[i]
        next_vertex = df.iloc[i + 1]
        confrontante = confrontantes[i % len(confrontantes)]
        
        recurring_text = (
            f"Daí, visando o Vértice {current['V']}, com giro angular horário de {current['Angulo Interno']} "
            f"e diste {current['Distancia(m)']} metros, chega-se ao Vértice {next_vertex['V']}, "
            f"também alinhado à rua {rua} e limítrofe com {confrontante}."
        )
        recurring_texts.append(recurring_text)
    
    return recurring_texts


def generate_final_text(df, rua, confrontantes):
    """
    Gera o texto final ao retornar ao V1.
    """
    last_vertex = df.iloc[-1]
    first_vertex = df.iloc[0]
    confrontante = confrontantes[-1 % len(confrontantes)]  # Confrontante do último trecho
    
    final_text = (
        f"Daí, visando o vértice {last_vertex['V']}, com giro angular de {last_vertex['Angulo Interno']} "
        f"e diste {last_vertex['Distancia(m)']} metros, chega-se ao vértice {first_vertex['V']}, "
        f"origem da presente descrição, no alinhamento da rua {rua} e próximo aos lotes de {confrontante}."
    )
    return final_text


def create_memorial_document(
    uuid_str,
    proprietario,
    matricula,
    matricula_texto,
    area_total,
    cpf,
    rgi,
    excel_file_path,
    template_path,
    output_path,
    assinatura_path,
    ponto_amarracao,
    azimute,
    distancia_amarracao_v1,
    rua,
    cidade,
    confrontantes,
    area_dxf,
    desc_ponto_amarracao,
    perimeter_dxf,
    giro_angular_v1_dms
):

    try:
        # Ler o arquivo Excel
        df = pd.read_excel(excel_file_path)
        df['N'] = pd.to_numeric(df['N'].astype(str).str.replace(',', '.'), errors='coerce')
        df['E'] = pd.to_numeric(df['E'].astype(str).str.replace(',', '.'), errors='coerce')
        df['Distancia(m)'] = pd.to_numeric(df['Distancia(m)'].astype(str).str.replace(',', '.'), errors='coerce')
        # Criar o documento Word
        doc_word = Document(template_path)
        set_default_font(doc_word)  # 🔹 Aplica a fonte Arial 12 ao documento
        # Adiciona o preâmbulo centralizado com texto em negrito
        p1 = doc_word.add_paragraph(style='Normal')
        run = p1.add_run("MEMORIAL DESCRITIVO")
        run.bold = True
        p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # Método recomendado para garantir espaço visível:
        p_espaco = doc_word.add_paragraph()
        p_espaco.add_run('\n')  # Garantindo espaçamento vertical.


        #doc_word.add_paragraph(f"Imóvel: Área da matrícula {matricula} destinada a {descricao} ", style='Normal')
        p = doc_word.add_paragraph(style='Normal')
        p.add_run("Imóvel: ")

        if tipo == "ETE":
            texto = f"Área da matrícula {matricula} destinada a {descricao} - SES de {cidade}"
        elif tipo == "REM":
            texto = f"Área remanescente da matrícula {matricula} destinada a {descricao}"
        elif tipo == "SER":
            texto = f"Área da matrícula {matricula} destinada a SERVIDÃO ADMINISTRATIVA ACESSO À {descricao}"
        elif tipo == "ACE":
            texto = f"Área da matrícula {matricula} destinada ao ACESSO DA SERVIDÃO ADMINISTRATIVA DA {descricao}"
        else:
            texto = "Tipo não especificado"

        run_bold = p.add_run(texto)
        run_bold.bold = True

        doc_word.add_paragraph(f"Matrícula: Número - {matricula} do {RI} de {comarca} ", style='Normal')
        doc_word.add_paragraph(f"Proprietário: {proprietario}", style='Normal')
        doc_word.add_paragraph(f"Local: {rua} - {cidade}", style='Normal')
        
        doc_word.add_paragraph(f"Área: {area_dxf:,.2f} m²".replace(",", "X").replace(".", ",").replace("X", "."), style='Normal')
        doc_word.add_paragraph(f"Perímetro: {perimeter_dxf:,.2f} m".replace(",", "X").replace(".", ",").replace("X", "."), style='Normal')
        # Pula uma linha antes deste parágrafo
        doc_word.add_paragraph()
        
        
        # Primeiro, formate corretamente a área uma única vez antes do parágrafo:
        area_dxf_formatada = f"{area_dxf:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

        # Cria o parágrafo com variáveis formatadas (em uma única linha no Python para clareza)
        texto_paragrafo = (f"Área {uso_solo} com {area_dxf_formatada} m², parte de um todo maior da Matrícula N° {matricula} com {area_imovel} do {RI} de {comarca}, localizada na {rua}, na cidade de {cidade}, definida através do seguinte levantamento topográfico, onde os ângulos foram medidos no sentido horário.")
        
        # Cria o parágrafo e remove qualquer indentação especial (recuo pendente ou primeira linha)
        p = doc_word.add_paragraph(texto_paragrafo, style='Normal')
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # alinhamento justificado, se desejar
        doc_word.add_paragraph()  # Linha em branco após o parágrafo
        
        # Remove indentação/recuos
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Pt(0)
        p.paragraph_format.right_indent = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.keep_together = True
        
        # Removendo explicitamente recuo pendente (esse é o ajuste essencial!)
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.hanging_indent = None

        # Formata coordenadas individualmente (sem milhar)
        coord_E_ponto_Az = f"{Coorde_E_ponto_Az:.3f}".replace(".", ",")
        coord_N_ponto_Az = f"{Coorde_N_ponto_Az:.3f}".replace(".", ",")
        
        # Primeiro parágrafo ajustado corretamente
        # p = doc_word.add_paragraph(
        #     f"O Ponto Az, ponto de amarração, está localizado na {desc_ponto_Az} nas coordenadas "
        #     f"E(X) {coord_E_ponto_Az} e N(Y) {coord_N_ponto_Az}.",
        #     style='Normal'
        # )
        # p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        #doc_word.add_paragraph()  # Linha em branco após o primeiro parágrafo
        
        # Formatação correta da distância (com ponto do milhar)
        distance_formatada = f"{distance:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        
        # Segundo parágrafo ajustado corretamente
        # Cria o parágrafo vazio inicialmente
        p = doc_word.add_paragraph(style='Normal')
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        
        # Primeira parte sem negrito até "Vértice"
        # p.add_run(f"Daí, com Azimute de {convert_to_dms(azimuth)} e distância de {distance_formatada} m, chega-se ao Vértice ")
        
        # # Insere o vértice V1 em negrito
        # run_v1 = p.add_run("V1")
        # run_v1.bold = True  # V1 em negrito
        
        # # Restante do texto normal
        # p.add_run(", origem da descrição desta área.")

        #doc_word.add_paragraph()  # Linha em branco após o primeiro parágrafo
        
        # Início da descrição do perímetro
        initial = df.iloc[0]
        coord_N_inicial = f"{initial['N']:.3f}".replace(".", ",")
        coord_E_inicial = f"{initial['E']:.3f}".replace(".", ",")
        
        # Primeiro parágrafo
        p1 = doc_word.add_paragraph("Pontos definidos pelas Coordenadas Planas no Sistema U.T.M. – SIRGAS 2000.", style='Normal')
        p1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        
        # Pula uma linha
        doc_word.add_paragraph()
        
        # Segundo parágrafo
        # Cria o parágrafo vazio inicialmente
        p2 = doc_word.add_paragraph(style='Normal')
        p2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        
        # Texto inicial sem negrito até "vértice"
        p2.add_run("Inicia-se a descrição deste perímetro no vértice ")
        
        # Insere o vértice inicial em negrito
        run_v_inicial = p2.add_run(f"{initial['V']}")
        run_v_inicial.bold = True  # Define negrito
        
        # Restante do texto sem negrito
        p2.add_run(
            f", de coordenadas N(Y) {coord_N_inicial} e E(X) {coord_E_inicial}, "
            f"situado no limite com {initial['Confrontante']}."
        )


        doc_word.add_paragraph()  # Linha em branco após o parágrafo

        # Descrição dos segmentos
        num_points = len(df)
        for i in range(num_points):
            current = df.iloc[i]
            next_index = (i + 1) % num_points
            next_point = df.iloc[next_index]
        
            distancia_formatada = f"{current['Distancia(m)']:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
            coord_N_formatada = f"{next_point['N']:.3f}".replace(".", ",")
            coord_E_formatada = f"{next_point['E']:.3f}".replace(".", ",")
        
            # Checa se o próximo vértice é V1, para inserir texto especial
            if next_point['V'] == 'V1':
                complemento = ", origem desta descrição,"
            else:
                complemento = ""
        
            # Criação do parágrafo inicialmente vazio
            p = doc_word.add_paragraph(style='Normal')
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        
            # Primeira parte do texto sem negrito
            # 🔹 Limpeza explícita da variável confrontante antes de usá-la
            confrontante_limpo = str(current['Confrontante']).replace('\n', ' ').replace('\r', ' ').replace('\t', ' ').strip()
            
            p.add_run(
                f"Deste, segue com azimute de {current['Azimute']} e distância de {distancia_formatada} m, "
                f"confrontando neste trecho com área pertencente à {confrontante_limpo}, até o vértice "
            )

        
            # Adiciona o rótulo do vértice em negrito
            run_vertice = p.add_run(f"{next_point['V']}")
            run_vertice.bold = True  # Aqui é definido o negrito para o vértice
        
            # Completa o restante do texto sem negrito
            p.add_run(
                f"{complemento} de coordenadas N(Y) {coord_N_formatada} e E(X) {coord_E_formatada};"
            )
        
            # Adiciona uma linha em branco após cada parágrafo
            doc_word.add_paragraph()

        # Adicionar data e assinatura
        data_atual = datetime.now().strftime("%d de %B de %Y")
        doc_word.add_paragraph(f"\n Porto Alegre, RS, {data_atual}.", style='Normal')
        doc_word.add_paragraph("\n\n")
        output_path = os.path.normpath(os.path.join(caminho_salvar, f"FECHADA_{tipo}_Memorial_MAT_{matricula}.docx"))
        doc_word.save(output_path)
        print(f"Memorial descritivo salvo em: {output_path}")
    except Exception as e:
        print(f"Erro ao criar o documento memorial: {e}")




def find_excel_file(directory, keywords):
    """
    Busca um arquivo Excel no diretório contendo todas as palavras-chave no nome.
    Se não encontrar, exibe a lista de arquivos disponíveis.
    """
    if not directory or not os.path.exists(directory):
        logger.info(f"Erro: O diretório '{directory}' não existe ou não foi especificado corretamente.")
        return None

    excel_files = [file for file in os.listdir(directory) if file.endswith(".xlsx")]

    if not excel_files:
        logger.info(f"Nenhum arquivo Excel encontrado no diretório: {directory}")
        return None

    for file in excel_files:
        if all(keyword.lower() in file.lower() for keyword in keywords):
            return os.path.join(directory, file)

    # Se nenhum arquivo correspondente foi encontrado, listar os arquivos disponíveis
    logger.info(f"Nenhum arquivo Excel contendo {keywords} foi encontrado em '{directory}'.")
    logger.info("Arquivos disponíveis no diretório:")
    for f in excel_files:
        logger.info(f"  - {f}")

    return None



        
# def convert_docx_to_pdf(output_path, pdf_file_path):
#     """
#     Converte um arquivo DOCX para PDF usando a biblioteca comtypes.
#     """
#     try:
#         # Verificar se o arquivo DOCX existe antes de abrir
#         if not os.path.exists(output_path):
#             raise FileNotFoundError(f"Arquivo DOCX não encontrado: {output_path}")
        
#         logger.info(f"Tentando converter o arquivo DOCX: {output_path} para PDF: {pdf_file_path}")

#         word = comtypes.client.CreateObject("Word.Application")
#         word.Visible = False  # Ocultar a interface do Word
#         doc = word.Documents.Open(output_path)
        
#         # Aguardar alguns segundos antes de salvar como PDF
#         import time
#         time.sleep(2)

#         doc.SaveAs(pdf_file_path, FileFormat=17)  # 17 corresponde ao formato PDF
#         doc.Close()
#         word.Quit()
#         logger.info(f"Arquivo PDF salvo com sucesso em: {pdf_file_path}")
#     except FileNotFoundError as fnf_error:
#         logger.error(f"Erro: {fnf_error}")
#     except Exception as e:
#         logger.info(f"Erro ao converter DOCX para PDF: {e}")
#     finally:
#         try:
#             word.Quit()
#         except:
#             pass  # Garantir que o Word seja fechado



def sanitize_filename(filename):
    return re.sub(r'[<>:"/\\|?*]', '', filename)

        
def main_poligonal_fechada(uuid_str, excel_path, dxf_path, diretorio_preparado, diretorio_concluido, caminho_template):

    caminho_salvar = diretorio_concluido 
    template_path = caminho_template 
    # Carrega dados do imóvel
    dados_imovel_excel_path = excel_path
    dados_imovel_df = pd.read_excel(dados_imovel_excel_path, sheet_name='Dados_do_Imóvel', header=None)
    dados_imovel = dict(zip(dados_imovel_df.iloc[:, 0], dados_imovel_df.iloc[:, 1]))

    # Extrai informações
    proprietario = dados_imovel.get("NOME DO PROPRIETÁRIO", "").strip()
    cpf = dados_imovel.get("CPF/CNPJ", "").strip()
    matricula = sanitize_filename(str(dados_imovel.get("DOCUMENTAÇÃO DO IMÓVEL", "")).strip())
    matricula_texto = str(dados_imovel.get("DOCUMENTAÇÃO DO IMÓVEL", "")).strip()
    descricao = dados_imovel.get("OBRA", "").strip()
    area_total = dados_imovel.get("ÁREA TOTAL DO TERRENO DOCUMENTADA", "").replace("\t", "").replace("\n", "").strip()
    cidade = dados_imovel.get("CIDADE", "").strip().capitalize()
    rgi= dados_imovel.get("RGI", "").strip().capitalize()
    rua = dados_imovel.get("LOCAL", "").strip()
    desc_ponto_Az = dados_imovel.get("AZ", "").strip()

    # Diretório para salvar resultados
    
    os.makedirs(caminho_salvar, exist_ok=True)

    # Identifica tipo (SER, REM, etc)
    dxf_filename = os.path.basename(dxf_path).upper()
    if "ETE" in dxf_filename:
        tipo = "ETE"
    elif "REM" in dxf_filename:
        tipo = "REM"
    elif "SER" in dxf_filename:
        tipo = "SER"
    elif "ACE" in dxf_filename:
        tipo = "ACE"
    else:
        logger.info("❌ Não foi possível determinar automaticamente o tipo (ETE, REM, SER ou ACE).")
        return

    padrao_fechada = os.path.join(diretorio_preparado, f"{uuid_str}_FECHADA_{tipo}*.xlsx")

    arquivos_encontrados = glob.glob(padrao_fechada)
    if not arquivos_encontrados:
        logger.info(f"❌ Arquivo de confrontantes não encontrado com o padrão: {padrao_fechada}")
        return
    confrontantes_df = pd.read_excel(arquivos_encontrados[0])
    confrontantes = confrontantes_df.iloc[:, 1].dropna().tolist()

    # DXF limpo
    # ⚠️ Substitui a limpeza anterior por apenas conversão R2010
    dxf_limpo_path = os.path.join(caminho_salvar, f"DXF_LIMPO_{matricula}.dxf")
    dxf_file_path = limpar_dxf_e_converter_r2010(dxf_path, dxf_limpo_path)


    # 🔍 Buscar planilha que COMEÇA com ABERTA_{TIPO} no diretório CONCLUIDO
    padrao_aberta = os.path.join(diretorio_concluido, f"{uuid_str}_ABERTA_{tipo}*.xlsx")
    planilhas_aberta = glob.glob(padrao_aberta)

    if not planilhas_aberta:
        logger.info(f"❌ Nenhuma planilha encontrada começando com 'ABERTA_{tipo}' no diretório: {diretorio_concluido}")
        return

    planilha_aberta_saida = planilhas_aberta[0]
    logger.info(f"📄 Planilha ABERTA localizada: {planilha_aberta_saida}")

   
    if not planilhas_aberta:
        logger.info(f"❌ Nenhuma planilha encontrada contendo 'ABERTA' e '{tipo}' no nome dentro de: {diretorio_concluido}")
        return

    planilha_aberta_saida = planilhas_aberta[0]
    logger.info(f"📄 Planilha ABERTA localizada: {planilha_aberta_saida}")


    # 📁 Procurar CONCLUIDO dentro da cidade (REPESCAGEM_*/CONCLUIDO)
    # O diretório CONCLUIDO já é passado corretamente
    diretorio_concluido_real = diretorio_concluido

   
    # 🧭 Obter ponto de amarração anterior ao V1
    try:
        ponto_amarracao, codigo_amarracao = obter_ponto_amarracao_anterior_v1(planilha_aberta_saida)
        logger.info(f"📌 Ponto de amarração identificado: {codigo_amarracao} com coordenadas {ponto_amarracao}")
    except Exception as e:
        logger.error(f"❌ Erro ao obter ponto de amarração: {e}")
        return

    # 🔍 Extrair geometria do DXF
    doc, lines, perimeter_dxf, area_dxf, _, area_poligonal = get_document_info_from_dxf(dxf_file_path)
    if not doc or not ponto_amarracao:
        logger.info("Erro ao processar o arquivo DXF.")
        return

    try:
        doc_dxf = ezdxf.readfile(dxf_file_path)
        msp = doc_dxf.modelspace()
    except Exception as e:
        logger.error(f"Erro ao abrir o arquivo DXF para edição: {e}")
        return

    if doc and lines:
        logger.info(f"Área da poligonal obtida (do DXF): {area_dxf:.6f} m²")
        logger.info(f"Perímetro da poligonal (do DXF): {perimeter_dxf:.6f} metros")

        # Cálculo com base no ponto real
        v1 = lines[0][0]
        azimute = calculate_azimuth(ponto_amarracao, v1)
        # Cálculo do giro angular no vértice V1 (do ponto externo para dentro da poligonal)
        v2 = lines[1][0]  # V2 é o segundo ponto da poligonal
        giro_angular_v1 = calculate_angular_turn(ponto_amarracao, v1, v2)
        giro_angular_v1_dms = convert_to_dms(360 - giro_angular_v1)

        distancia_az_v1 = calculate_distance(ponto_amarracao, v1)

        # Caminho do Excel de saída
        excel_file_path = os.path.join(
            diretorio_concluido,
            f"{uuid_str}_FECHADA_{tipo}_{matricula}.xlsx"
        )


        logger.info(f"✅ Excel FECHADA salvo corretamente: {excel_file_path}")

        # 🛠 Criar memorial e Excel
        create_memorial_descritivo(
            uuid_str, doc, lines, proprietario, matricula, caminho_salvar, confrontantes, ponto_amarracao,
            dxf_file_path, area_dxf, azimute, v1, msp, dxf_filename, excel_file_path, tipo,giro_angular_v1_dms
        )

        # 📄 Gerar DOCX
        if excel_file_path:
            # template_path = os.path.join(BASE_DIR, "templates_doc", "Memorial_modelo_padrao.docx")
            output_path_docx = os.path.join(
                diretorio_concluido,
                f"{uuid_str}_FECHADA_{tipo}_{matricula}.docx"
            )

            logger.info(f"✅ DOCX FECHADA salvo corretamente: {output_path_docx}")

            assinatura_path = r"C:\Users\Paulo\Documents\CASSINHA\MEMORIAIS DESCRITIVOS\Assinatura.jpg"

            desc_ponto_amarracao = f"ponto {codigo_amarracao}, obtido na planilha da poligonal aberta"

            create_memorial_document(
                uuid_str=uuid_str,
                proprietario=proprietario,
                matricula=matricula,  # usado para salvar arquivos
                matricula_texto=matricula_texto,  # usado no Word
                area_total=area_total,
                cpf=cpf,
                rgi=rgi,
                excel_file_path=excel_file_path,
                template_path=template_path,
                output_path=output_path_docx,
                assinatura_path=assinatura_path,
                ponto_amarracao=ponto_amarracao,
                azimute=azimute,
                distancia_amarracao_v1=distancia_az_v1,
                rua=rua,
                cidade=cidade,
                confrontantes=confrontantes,
                area_dxf=area_dxf,
                desc_ponto_amarracao=desc_ponto_amarracao,
                perimeter_dxf=perimeter_dxf,
                giro_angular_v1_dms=giro_angular_v1_dms,
               
            )




            # 🧾 Converter para PDF
            # time.sleep(2)
            # if os.path.exists(output_path_docx):
            #     pdf_file_path = os.path.join(caminho_salvar, f"FECHADA_{tipo}_Memorial_{matricula}.pdf")
            #     convert_docx_to_pdf(output_path_docx, pdf_file_path)
            #     logger.info(f"Arquivo PDF salvo em: {pdf_file_path}")
            # else:
            #     logger.info(f"Erro: O arquivo DOCX '{output_path_docx}' não foi encontrado.")
        else:
            logger.info("excel_file_path não definido ou inválido.")
        logger.info("Documento do AutoCAD fechado.")
    else:
        logger.info("Nenhuma linha foi encontrada ou não foi possível acessar o documento.")
        pythoncom.CoUninitialize()
    #atualizado agora





