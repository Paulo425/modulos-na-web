import math
import csv
import pandas as pd
import glob
import re
from docx import Document
from docx.shared import Inches
from datetime import datetime
import os
import ezdxf
import time
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from decimal import Decimal, getcontext
import pandas as pd
from docx.shared import Pt
import openpyxl
from openpyxl.styles import Alignment, Font
import logging 
EPS_BULGE = 1e-9


BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))

getcontext().prec = 28  # Define a precisão para 28 casas decimais

# Configuração manual para nomes dos meses em português (independente do locale)
MESES_PT_BR = {
    'January': 'janeiro',
    'February': 'fevereiro',
    'March': 'março',
    'April': 'abril',
    'May': 'maio',
    'June': 'junho',
    'July': 'julho',
    'August': 'agosto',
    'September': 'setembro',
    'October': 'outubro',
    'November': 'novembro',
    'December': 'dezembro'
}


logger = logging.getLogger(__name__)
logger.setLevel(logging.DEBUG)  # Garanta que está em DEBUG


import re

def convert_dms_to_decimal(dms_str):
    matches = re.findall(r"(\d+)[°\s]+(\d+)'[\s]*(\d+(?:,\d+|\.\d+)?)", dms_str)
    if not matches:
        raise ValueError(f"Formato de ângulo inválido: {dms_str}")
    graus, minutos, segundos = matches[0]
    graus = float(graus)
    minutos = float(minutos)
    segundos = float(segundos.replace(',', '.'))
    return graus + minutos / 60 + segundos / 3600


def is_clockwise(points):
    """
    Verifica se a poligonal está no sentido horário.
    Retorna True se for horário, False se anti-horário.
    """
    area = 0.0
    for i in range(len(points)):
        j = (i + 1) % len(points)
        area += points[i][0] * points[j][1]
        area -= points[j][0] * points[i][1]
    return area < 0

def ensure_counterclockwise(points):
    """
    Garantir que a lista de pontos esteja no sentido anti-horário.
    Se estiver no sentido horário, inverte os pontos.
    """
    if is_clockwise(points):
        points.reverse()
    return points


getcontext().prec = 28  # Define a precisão para 28 casas decimais

# Configurar locale para português do Brasil
#locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8')  # Para sistemas Linux ou Mac
# locale.setlocale(locale.LC_TIME, 'Portuguese_Brazil.1252')  # Para Windows, caso necessário

# Obter data atual formatada em português
data_atual = datetime.now().strftime("%d de %B de %Y")


# Função que processa as linhas da poligonal
# Função que processa as linhas da poligonal
def limpar_dxf_e_converter_r2010(original_path, saida_path):
    """
    Lê um DXF original e regrava o arquivo com a versão R2010,
    preservando LWPOLYLINE (com bulge). Não cria geometria nova.
    """
    try:
        doc_antigo = ezdxf.readfile(original_path)
        msp_antigo = doc_antigo.modelspace()
        doc_novo = ezdxf.new(dxfversion='R2010')
        msp_novo = doc_novo.modelspace()

        encontrou_polilinha = False

        for entity in msp_antigo.query('LWPOLYLINE'):
            if entity.closed:
                # Pega (x, y, bulge) como tuplas
                pontos_xyb = entity.get_points('xyb')  # [(x, y, bulge), ...]
                # Regrava preservando o bulge
                msp_novo.add_lwpolyline(
                    pontos_xyb,
                    format='xyb',
                    close=True,
                    dxfattribs={'layer': entity.dxf.layer}
                )
                encontrou_polilinha = True
                break

        if not encontrou_polilinha:
            raise ValueError("Nenhuma polilinha fechada encontrada no DXF original.")

        doc_novo.saveas(saida_path)
        logger.info(f"✅ DXF convertido e salvo como R2010 em: {saida_path}")
        return saida_path

    except Exception as e:
        logger.error(f"❌ Erro ao converter DXF para R2010: {e}")
        return original_path


def obter_pontos_ordenados_do_dxf(dxf_file):
    doc = ezdxf.readfile(dxf_file)
    msp = doc.modelspace()
    points = [(point.dxf.location.x, point.dxf.location.y) for point in msp.query('POINT')]
    return points
    
def calcular_area_poligonal(pontos):
    """Calcula a área da poligonal fechada usando a fórmula shoelace."""
    n = len(pontos)
    area = 0.0
    for i in range(n):
        x1, y1 = pontos[i]
        x2, y2 = pontos[(i + 1) % n]
        area += (x1 * y2) - (x2 * y1)
    return abs(area) / 2



def get_document_info_from_dxf(dxf_file_path):
    try:
        doc = ezdxf.readfile(dxf_file_path)
        msp = doc.modelspace()

        lines = []
        perimeter_dxf = 0.0
        area_dxf = 0.0
        ponto_az = None
        ordered_points = []
        ordered_points_with_bulge = []

        # Lê a PRIMEIRA LWPOLYLINE fechada
        for entity in msp.query('LWPOLYLINE'):
            if entity.closed:
                pts_xyb = entity.get_points('xyb')  # lista de tuplas (x, y, bulge)
                # Se houver último = primeiro, remove duplicata
                if len(pts_xyb) >= 2 and (pts_xyb[0][0], pts_xyb[0][1]) == (pts_xyb[-1][0], pts_xyb[-1][1]):
                    pts_xyb = pts_xyb[:-1]

                # Monta listas
                for (x, y, b) in pts_xyb:
                    ordered_points.append((x, y))
                    ordered_points_with_bulge.append({'x': x, 'y': y, 'bulge_next': float(b or 0.0)})

                logger.info(
                    f"pts_bulge: n={len(ordered_points_with_bulge)} | exemplo={ordered_points_with_bulge[:2]}"
                )

                # Linhas + perímetro
                n = len(ordered_points)
                for i in range(n):
                    p1 = ordered_points[i]
                    p2 = ordered_points[(i + 1) % n]
                    lines.append((p1, p2))
                    perimeter_dxf += ((p2[0] - p1[0]) ** 2 + (p2[1] - p1[1]) ** 2) ** 0.5

                # Área por shoelace
                x = [p[0] for p in ordered_points]
                y = [p[1] for p in ordered_points]
                area_dxf = abs(sum(x[i] * y[(i + 1) % n] - x[(i + 1) % n] * y[i] for i in range(n)) / 2.0)
                break

        if not lines:
            logger.info("Nenhuma polilinha fechada encontrada no arquivo DXF.")
            return None, [], 0.0, 0.0, None, None, []

        # Ponto Az (TEXT → INSERT → POINT)
        for entity in msp.query('TEXT'):
            if "Az" in (entity.dxf.text or ""):
                ponto_az = (entity.dxf.insert.x, entity.dxf.insert.y, 0.0)
                break

        if ponto_az is None:
            for entity in msp.query('INSERT'):
                if "Az" in (entity.dxf.name or ""):
                    ponto_az = (entity.dxf.insert.x, entity.dxf.insert.y, 0.0)
                    break

        if ponto_az is None:
            for entity in msp.query('POINT'):
                ponto_az = (entity.dxf.location.x, entity.dxf.location.y, 0.0)
                break

        if ponto_az is None:
            # Fallback: primeiro vértice
            ponto_az = (ordered_points[0][0], ordered_points[0][1], 0.0)
            logger.warning("⚠️ Ponto Az não encontrado no DXF. Usando fallback (primeiro ponto).")

        logger.info(f"Linhas processadas: {len(lines)}")
        logger.info(f"Perímetro do DXF: {perimeter_dxf:.2f} m")
        logger.info(f"Área do DXF: {area_dxf:.2f} m²")

        # >>> RETORNA 7 ITENS <<<
        return doc, lines, perimeter_dxf, area_dxf, ponto_az, msp, ordered_points_with_bulge

    except Exception as e:
        logger.error(f"Erro ao obter informações do documento: {e}")
        # >>> TAMBÉM retorna 7 itens no erro <<<
        return None, [], 0.0, 0.0, None, None, []




    
def obter_ponto_amarracao_anterior_v1(planilha_aberta_path):
    """
    Retorna o ponto imediatamente anterior ao V1 a partir da planilha de saída da poligonal aberta.
    Adaptado para colunas: 'Ponto', 'Coord_E', 'Coord_N'.
    """
    df = pd.read_excel(planilha_aberta_path, engine='openpyxl')

    if "Ponto" not in df.columns or "Coord_E" not in df.columns or "Coord_N" not in df.columns:
        raise ValueError("Planilha ABERTA não contém colunas 'Ponto', 'Coord_E' e 'Coord_N'.")

    idx_v1 = df[df["Ponto"] == "V1"].index
    if len(idx_v1) == 0:
        raise ValueError("Ponto V1 não encontrado na planilha.")
    elif idx_v1[0] == 0:
        raise ValueError("Não existe ponto anterior ao V1.")

    linha = df.iloc[idx_v1[0] - 1]
    e = float(str(linha["Coord_E"]).replace(",", "."))
    n = float(str(linha["Coord_N"]).replace(",", "."))
    codigo = linha["Ponto"]

    return (e, n), codigo


    
# 🔹 Função para definir a fonte padrão
def set_default_font(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    
def calculate_point_on_line(a, b, dist):
    """Retorna o ponto a uma distância 'dist' ao longo do segmento A->B."""
    ax, ay = float(a[0]), float(a[1])
    bx, by = float(b[0]), float(b[1])
    dx, dy = bx - ax, by - ay
    L = math.hypot(dx, dy)
    if L <= 1e-9:
        return (ax, ay)
    t = dist / L
    return (ax + dx * t, ay + dy * t)

def add_azimuth_arc_to_dxf(msp, ponto_az, v1, azimute):
    """Desenha segmento Az–V1, linha para o norte, arco do azimute e rótulo."""
    try:
        if (ponto_az is None) or (v1 is None):
            logger.error("Az ou V1 ausentes; abortando desenho do azimute.")
            return
        if math.hypot(v1[0]-ponto_az[0], v1[1]-ponto_az[1]) <= 1e-9:
            logger.error("Distância Az–V1 zero; abortando desenho.")
            return

        logger.info("Iniciando a adição do arco de azimute. Azimute: %s°", azimute)

        if 'Azimute' not in msp.doc.layers:
            msp.doc.layers.new(name='Azimute', dxfattribs={'color': 1})
            logger.info("Camada 'Azimute' criada com sucesso.")

        # segmento Az–V1
        msp.add_line(start=(ponto_az[0], ponto_az[1]), end=(v1[0], v1[1]), dxfattribs={'layer': 'Azimute'})
        logger.info("Segmento entre Az e V1 desenhado de %s para %s", ponto_az, v1)

        # segmento norte
        north_point = (ponto_az[0], ponto_az[1] + 2.0)
        msp.add_line(start=(ponto_az[0], ponto_az[1]), end=north_point, dxfattribs={'layer': 'Azimute'})
        logger.info("Linha para o norte desenhada com sucesso de %s para %s", ponto_az, north_point)

        # arco
        dist = calculate_distance(ponto_az, v1)
        radius = 0.4 if dist <= 0.5 else 1.0
        start_arc = calculate_point_on_line(ponto_az, v1, radius)
        end_arc   = calculate_point_on_line(ponto_az, north_point, radius)
        msp.add_arc(
            center=(ponto_az[0], ponto_az[1]),
            radius=radius,
            start_angle=math.degrees(math.atan2(start_arc[1]-ponto_az[1], start_arc[0]-ponto_az[0])),
            end_angle  =math.degrees(math.atan2(end_arc[1]  -ponto_az[1], end_arc[0]  -ponto_az[0])),
            dxfattribs={'layer': 'Azimute'}
        )
        logger.info("Arco do azimute desenhado com sucesso.")

        # rótulo do azimute
        azimuth_label = f"Azimute: {convert_to_dms(azimute)}"
        label_position = (
            ponto_az[0] + 1.0 * math.cos(math.radians(azimute / 2.0)),
            ponto_az[1] + 1.0 * math.sin(math.radians(azimute / 2.0))
        )
        msp.add_text(azimuth_label, dxfattribs={'height': 0.25, 'layer': 'Azimute'}).set_pos(label_position)
        logger.info("Rótulo do azimute adicionado em %s", label_position)

    except Exception:
        logger.exception("Erro na função `add_azimuth_arc_to_dxf`")


# def add_distance_label(msp, p0, p1, dist):
#     """Rótulo simples da distância no ponto médio do segmento p0–p1."""
#     try:
#         mid = ((p0[0]+p1[0])/2.0, (p0[1]+p1[1])/2.0)
#         txt = f"{dist:,.2f} m".replace(",", "X").replace(".", ",").replace("X",".")
#         if 'Azimute' not in msp.doc.layers:
#             msp.doc.layers.new(name='Azimute', dxfattribs={'color': 1})
#         msp.add_text(txt, dxfattribs={'height': 0.25, 'layer': 'Azimute'}).set_pos(mid)
#         logger.info("Distância %.2f m adicionada corretamente em %s", dist, mid)
#     except Exception:
#         logger.exception("Erro ao adicionar rótulo de distância")


def _desenhar_referencia_az(msp, ponto_az, v1, azimute):
    """Wrapper: desenha arco + rótulo da distância Az–V1 com guards."""
    dx, dy = v1[0]-ponto_az[0], v1[1]-ponto_az[1]
    dist = math.hypot(dx, dy)
    if dist <= 1e-9:
        logger.warning("⚠️ Distância Az–V1 ≈ 0; desenho do Az suprimido.")
        return
    add_azimuth_arc_to_dxf(msp, ponto_az, v1, azimute)
    add_distance_label(msp, (ponto_az[0], ponto_az[1]), (v1[0], v1[1]), dist)





def calculate_polygon_area(points):
    """
    Calcula a área de uma poligonal fechada utilizando o método do produto cruzado com alta precisão.
    :param points: Lista de coordenadas [(x1, y1), (x2, y2), ...].
    :return: Área da poligonal.
    """
    n = len(points)
    if n < 3:
        return Decimal(0)  # Não é uma poligonal válida

    area = Decimal(0)
    for i in range(n):
        x1, y1 = Decimal(points[i][0]), Decimal(points[i][1])
        x2, y2 = Decimal(points[(i + 1) % n][0]), Decimal(points[(i + 1) % n][1])
        area += x1 * y2 - y1 * x2

    return abs(area) / 2
    
def convert_to_dms(decimal_degrees):
    """
    Converte graus decimais para o formato graus, minutos e segundos (DMS).
    """
    try:
        # Verificar se o valor é NaN
        if math.isnan(decimal_degrees):
            raise ValueError("Valor de entrada é NaN")

        degrees = int(decimal_degrees)
        minutes = int((abs(decimal_degrees) - abs(degrees)) * 60)
        seconds = round((abs(decimal_degrees) - abs(degrees) - minutes / 60) * 3600, 2)
        return f"{degrees}°{minutes}'{seconds}\""
    except Exception as e:
        logger.info(f"Erro na conversão para DMS: {e}")
        return "0°0'0.00\""  # Valor padrão em caso de erro

def calculate_distance(p1, p2):
    return math.sqrt((p2[0] - p1[0])**2 + (p2[1] - p1[1])**2)


def degrees_to_dms(angle):
    """
    Converte um ângulo em graus decimais para o formato graus, minutos e segundos (° ' ").
    """
    degrees = int(angle)
    minutes = int((angle - degrees) * 60)
    seconds = round((angle - degrees - minutes / 60) * 3600, 2)
    return f"{degrees}°{minutes}'{seconds}\""
import math

def calculate_azimuth(p1, p2):
    """
    Calcula o azimute entre dois pontos p1 e p2.
    
    :param p1: Coordenadas do ponto 1 (E, N)
    :param p2: Coordenadas do ponto 2 (E, N)
    :return: Azimute em graus decimais
    """
    delta_x = p2[0] - p1[0]
    delta_y = p2[1] - p1[1]
    
    azimuth = math.degrees(math.atan2(delta_x, delta_y)) % 360  # Garantir valor entre 0° e 360°
    
    return azimuth


def create_arrow_block(doc, block_name="ARROW"):
    """
    Cria um bloco no DXF representando uma seta sólida como um triângulo.
    """
    if block_name in doc.blocks:
        return  # O bloco já existe

    block = doc.blocks.new(name=block_name)

    # Definir o triângulo da seta
    length = 0.5  # Comprimento da seta
    base_half_length = length / 2

    tip = (0, 0)  # Ponta da seta no eixo de coordenadas
    base1 = (-base_half_length, -length)
    base2 = (base_half_length, -length)

    block.add_solid([base1, base2, tip])
import math

def add_giro_angular_arc_to_dxf(doc_dxf, v1, az, v2):

    """
    Adiciona um arco representando o giro angular horário no espaço de modelo do DXF já aberto.
    """
    try:
        msp = doc_dxf.modelspace()
        # Calcular distância entre V1–Az e V1–V2 para escolher o menor
        dist_az = calculate_distance(v1, az)
        dist_v2 = calculate_distance(v1, v2)
        min_dist = min(dist_az, dist_v2)

        radius = 0.4 if min_dist <= 0.5 else 1.0  # Raio adaptativo
        # Traçar a reta entre V1 e Az
        msp.add_line(start=v1[:2], end=az[:2])
        logger.info(f"Linha entre V1 e Az traçada com sucesso.")

        # Definir os pontos de apoio
        def calculate_displacement(point1, point2, distance):
            dx = point2[0] - point1[0]
            dy = point2[1] - point1[1]
            magnitude = math.hypot(dx, dy)
            return (
                point1[0] + (dx / magnitude) * distance,
                point1[1] + (dy / magnitude) * distance,
            )

        # Calcular os pontos de apoio
        ponto_inicial = calculate_displacement(v1, v2, radius)  # 2m na reta V1-V2
        ponto_final = calculate_displacement(v1, az, radius)   # 2m na reta V1-Az

        # Calcular os ângulos dos vetores
        angle_v2 = math.degrees(math.atan2(ponto_inicial[1] - v1[1], ponto_inicial[0] - v1[0]))
        angle_az = math.degrees(math.atan2(ponto_final[1] - v1[1], ponto_final[0] - v1[0]))

        # Calcular o giro angular no sentido horário
        giro_angular = (angle_az - angle_v2) % 360  # Garantir que o ângulo esteja no intervalo [0, 360)
        if giro_angular < 0:  # Caso negativo, ajustar para o sentido horário
            giro_angular += 360

        logger.info(f"Giro angular calculado corretamente: {giro_angular:.2f}°")

        # Traçar o arco
        msp.add_arc(center=v1[:2], radius=radius, start_angle=angle_v2, end_angle=angle_az)
        logger.info(f"Arco do giro angular traçado com sucesso.")

        # Adicionar rótulo ao arco
        label_offset = 3.0
        deslocamento_x=3
        deslocamento_y=-3
        angle_middle = math.radians((angle_v2 + angle_az) / 2)
        label_position = (
            v1[0] + (label_offset+deslocamento_x) * math.cos(angle_middle),
            v1[1] + (label_offset+deslocamento_y) * math.sin(angle_middle),
        )
        # Converter o ângulo para DMS e exibir no rótulo
        giro_angular_dms = f"{convert_to_dms(giro_angular)}"
        msp.add_text(
            giro_angular_dms,
            dxfattribs={
                'height': 0.3,
                'layer': 'Labels',
                'insert': label_position  # Define a posição do texto
            }
        )
        logger.info(f"Rótulo do giro angular ({giro_angular_dms}) adicionado com sucesso.")

    except Exception as e:
        logger.error(f"Erro ao adicionar o arco do giro angular ao DXF: {e}") 



def calculate_arc_angles(p1, p2, p3):
    try:
        # Vetores a partir de p2
        dx1, dy1 = p1[0] - p2[0], p1[1] - p2[1]  # Vetor de p2 para p1
        dx2, dy2 = p3[0] - p2[0], p3[1] - p2[1]  # Vetor de p2 para p3

        # Ângulos dos vetores em relação ao eixo X
        angle1 = math.degrees(math.atan2(dy1, dx1)) % 360
        angle2 = math.degrees(math.atan2(dy2, dx2)) % 360

        # AQUI ESTÁ A CORREÇÃO:
        internal_angle = (angle1 - angle2) % 360

        if internal_angle > 180:
            #internal_angle = 360 - internal_angle
            internal_angle = internal_angle
            start_angle = angle1
            end_angle = angle2
        else:
            start_angle = angle2
            end_angle = angle1

        return start_angle % 360, end_angle % 360

    except Exception as e:
        logger.info(f"Erro ao calcular ângulos do arco: {e}")
        return 0, 0





def insert_and_rotate_arrow(msp, center, radius, angle, rotation_adjustment, block_name="ARROW"):
    """
    Insere um bloco da seta alinhado ao arco e rotacionado conforme especificado.
    """
    arrow_tip_x = center[0] + radius * math.cos(angle)
    arrow_tip_y = center[1] + radius * math.sin(angle)
    arrow_tip = (arrow_tip_x, arrow_tip_y)

    rotation_angle = math.degrees(angle) + rotation_adjustment

    msp.add_blockref(
        block_name,
        insert=arrow_tip,
        dxfattribs={"rotation": rotation_angle}
    )
import math




# def add_angle_visualization_to_dwg(msp, ordered_points, angulos_decimais):
#     try:
#         total_points = len(ordered_points)

#         for i, p2 in enumerate(ordered_points):
#             p1 = ordered_points[i - 1]
#             p3 = ordered_points[(i + 1) % total_points]

#             # Vetores adjacentes ao vértice atual
#             vec_p2_p1 = (p1[0] - p2[0], p1[1] - p2[1])
#             vec_p2_p3 = (p3[0] - p2[0], p3[1] - p2[1])

#             # Ângulo interno exato (excel)
#             internal_angle = angulos_decimais[i]

#             # Tamanho adequado do arco (proporcional ao menor lado adjacente)
#             lado1 = math.hypot(*vec_p2_p1)
#             lado2 = math.hypot(*vec_p2_p3)
#             raio = min(lado1, lado2) * 0.1

#             # Função auxiliar para calcular ângulo polar em graus
#             def polar_angle(v):
#                 return math.degrees(math.atan2(v[1], v[0])) % 360

#             angle1 = polar_angle(vec_p2_p1)
#             angle2 = polar_angle(vec_p2_p3)

#             # Produto vetorial para determinar lado interno
#             cross = vec_p2_p1[0] * vec_p2_p3[1] - vec_p2_p1[1] * vec_p2_p3[0]

#             if cross < 0:  # Arco interno no sentido horário
#                 start_angle = angle1
#                 end_angle = (start_angle - internal_angle) % 360
#             else:          # Arco interno no sentido anti-horário
#                 start_angle = angle1
#                 end_angle = (start_angle + internal_angle) % 360

#             # desenhar arco corretamente no lado interno
#             msp.add_arc(
#                 center=p2,
#                 radius=raio,
#                 start_angle=start_angle,
#                 end_angle=end_angle,
#                 dxfattribs={'layer': 'Internal_Arcs'}
#             )

#             # posição ideal do rótulo
#             angulo_meio = (start_angle + end_angle) / 2 % 360
#             distancia_label = raio + 1.0
#             posicao_label = (
#                 p2[0] + distancia_label * math.cos(math.radians(angulo_meio)),
#                 p2[1] + distancia_label * math.sin(math.radians(angulo_meio))
#             )

#             internal_angle_dms = convert_to_dms(internal_angle)

#             # texto com rotação correta
#             msp.add_text(
#                 internal_angle_dms,
#                 dxfattribs={
#                     'height': 0.7,
#                     'layer': 'Labels',
#                     'insert': posicao_label,
#                     'rotation': angulo_meio if angulo_meio <= 180 else angulo_meio - 180
#                 }
#             )

#             print(f"Vértice V{i+1}: Ângulo interno {internal_angle_dms}")

#     except Exception as e:
#         print(f"Erro ao adicionar ângulos internos ao DXF: {e}")

import math

def _xy(pt):
    # aceita tuple/list (x,y) OU dict com chaves 'x','y' ou 'X','Y'
    if isinstance(pt, dict):
        if 'x' in pt and 'y' in pt:
            return float(pt['x']), float(pt['y'])
        if 'X' in pt and 'Y' in pt:
            return float(pt['X']), float(pt['Y'])
        # se houver outras chaves, tente mapear por índice
        return float(pt.get(0)), float(pt.get(1))
    return float(pt[0]), float(pt[1])

def _ang(p_from, p_to):
    dx = p_to[0] - p_from[0]
    dy = p_to[1] - p_from[1]
    return (math.degrees(math.atan2(dy, dx)) + 360.0) % 360.0

def add_angle_visualization_to_dwg(msp, ordered_points, angulos_decimais, sentido_poligonal):
    """
    Desenha os arcos de ângulos internos SEMPRE para dentro.
    Premissa: 'ordered_points' já está no sentido desejado (horário ou anti_horário).
    'sentido_poligonal' deve ser 'horario' ou 'anti_horario'.
    """
    try:
        n = len(ordered_points)
        if n < 3:
            return  # nada a fazer

        for i in range(n):
            # vizinhos (p1 <- p2 -> p3)
            p2 = _xy(ordered_points[i])
            p1 = _xy(ordered_points[i - 1])
            p3 = _xy(ordered_points[(i + 1) % n])

            # ângulos das DIREÇÕES adjacentes no vértice
            ang_in  = _ang(p2, p1)  # direção de entrada (segmento p1->p2, visto a partir de p2)
            ang_out = _ang(p2, p3)  # direção de saída   (segmento p2->p3, visto a partir de p2)

            interno = float(angulos_decimais[i])

            # raio: 10% do menor lado adjacente (com piso de 1 cm)
            lado1 = math.hypot(p1[0] - p2[0], p1[1] - p2[1])
            lado2 = math.hypot(p3[0] - p2[0], p3[1] - p2[1])
            raio = max(0.01, 0.10 * min(lado1, lado2))

            # ===== REGRA FUNDAMENTAL (ezdxf varre SEMPRE CCW de start->end) =====
            if sentido_poligonal == 'anti_horario':
                # para ficar "por dentro", varra CCW de SAÍDA para ENTRADA
                start_angle = ang_out
                end_angle   = (start_angle + interno) % 360.0
                mid_angle   = (start_angle + interno / 2.0) % 360.0
            else:
                # 'horario' → varra CCW de ENTRADA para SAÍDA (minor arc é o interno)
                start_angle = ang_in
                end_angle   = (start_angle + interno) % 360.0
                mid_angle   = (start_angle + interno / 2.0) % 360.0

            # desenha arco
            msp.add_arc(
                center=(p2[0], p2[1]),
                radius=raio,
                start_angle=start_angle,
                end_angle=end_angle,
                dxfattribs={'layer': 'Internal_Arcs'}
            )

            # rótulo no meio do arco
            distancia_label = raio + 1.0
            pos_label = (
                p2[0] + distancia_label * math.cos(math.radians(mid_angle)),
                p2[1] + distancia_label * math.sin(math.radians(mid_angle)),
            )

            # se já tem sua função utilitária, mantenha
            ang_texto = convert_to_dms(interno) if 'convert_to_dms' in globals() else f"{interno:.4f}°"

            msp.add_text(
                ang_texto,
                dxfattribs={
                    'height': 0.7,
                    'layer': 'Labels',
                    'insert': pos_label,
                    # rotação amigável
                    'rotation': mid_angle if mid_angle <= 180 else mid_angle - 180
                }
            )

    except Exception as e:
        print(f"Erro ao adicionar ângulos internos ao DXF: {e}")








def calculate_internal_angle(p1, p2, p3):
    try:
        # Vetores a partir do ponto central (p2)
        dx1, dy1 = p1[0] - p2[0], p1[1] - p2[1]
        dx2, dy2 = p3[0] - p2[0], p3[1] - p2[1]

        # Ângulos dos vetores em relação ao eixo X
        angle1 = math.atan2(dy1, dx1)
        angle2 = math.atan2(dy2, dx2)

        # Ângulo formado entre os vetores
        internal_angle = (angle2 - angle1) % (2 * math.pi)

        # Se o ângulo for maior que 180°, use o suplementar
        if internal_angle > math.pi:
            internal_angle = (2 * math.pi) - internal_angle

        # Retornar o ângulo em graus
        return math.degrees(internal_angle)

    except Exception as e:
        print(f"Erro inesperado ao calcular o ângulo interno: {e}")
        return 0



def calculate_label_position(p2, start_angle, end_angle, radius=1.8):
    """
    Calcula a posição do rótulo do ângulo interno, deslocando-o para uma posição central no arco.
    
    :param p2: Ponto central do arco (vértice da poligonal)
    :param start_angle: Ângulo de início do arco
    :param end_angle: Ângulo de fim do arco
    :param radius: Raio do deslocamento para evitar sobreposição
    :return: Posição (x, y) onde o rótulo do ângulo interno será colocado
    """
    try:
        # Calcular o ângulo médio entre os dois ângulos do arco
        mid_angle = math.radians((start_angle + end_angle) / 2)

        # Calcular a posição do rótulo deslocado no ângulo médio
        label_x = p2[0] + radius * math.cos(mid_angle)
        label_y = p2[1] + radius * math.sin(mid_angle)

        return (label_x, label_y)

    except Exception as e:
        logger.info(f"Erro ao calcular posição do rótulo do ângulo: {e}")
        return p2  # Retorna o próprio ponto central caso ocorra erro


import math

def add_label_and_distance(msp, start_point, end_point, label, distance):
    """
    Adiciona rótulos e distâncias no espaço de modelo usando ezdxf.
    """
    try:
        # Calcular ponto médio
        mid_point = (
            (start_point[0] + end_point[0]) / 2,
            (start_point[1] + end_point[1]) / 2
        )

        # Vetor da linha
        dx = end_point[0] - start_point[0]
        dy = end_point[1] - start_point[1]
        length = math.hypot(dx, dy)

        # Ângulo da linha
        angle = math.degrees(math.atan2(dy, dx))

        # Corrigir para manter a leitura sempre da esquerda para a direita
        if angle < -90 or angle > 90:
            angle += 180  

        # Afastar o rótulo da linha
        offset = -0.5  # Ajuste o valor para mais afastamento
        perp_x = -dy / length * offset
        perp_y = dx / length * offset
        displaced_mid_point = (mid_point[0] + perp_x, mid_point[1] + perp_y)

        # Criar layer se não existir
        if "Distance_Labels" not in msp.doc.layers:
            msp.doc.layers.new(name="Distance_Labels", dxfattribs={"color": 2})  # Define cor para melhor visualização

        # Adicionar o rótulo da distância no LAYER correto
        msp.add_text(
            f"{distance:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".") + "",
            dxfattribs={
                "height": 0.3,
                "layer": "Distance_Labels",  # Novo LAYER
                "rotation": angle,
                "insert": displaced_mid_point
            }
        )
        logger.info(f"Distância {distance:.2f} m adicionada corretamente em {displaced_mid_point}")

    except Exception as e:
        logger.error(f"Erro ao adicionar rótulo de distância: {e}")

def calculate_angular_turn(p1, p2, p3):
    """
    Calcula o giro angular no ponto `p2` entre os segmentos `p1-p2` e `p2-p3` no sentido horário.
    Retorna o ângulo em graus.
    """
    import math
    
    dx1, dy1 = p1[0] - p2[0], p1[1] - p2[1]  # Vetor do segmento p1-p2
    dx2, dy2 = p3[0] - p2[0], p3[1] - p2[1]  # Vetor do segmento p2-p3

    angle1 = math.atan2(dy1, dx1)
    angle2 = math.atan2(dy2, dx2)

    # Calcula o ângulo horário
    angular_turn = (angle2 - angle1) % (2 * math.pi)
    angular_turn_degrees = math.degrees(angular_turn)

    return angular_turn_degrees

# ==== HELPERS_ANGULOS_DXF_BEGIN ====
def _log_info(msg):
    try:
        logger.info(msg)
    except Exception:
        print(msg)

def _log_error(msg):
    try:
        logger.error(msg)
    except Exception:
        print(msg)

def _convert_to_dms_safe(graus):
    if "convert_to_dms" in globals():
        return convert_to_dms(graus)
    g = float(graus)
    d = int(g)
    m_f = abs((g - d) * 60.0)
    m = int(m_f)
    s = (m_f - m) * 60.0
    return f"{abs(d):02d}°{m:02d}'{s:06.3f}\"{'S' if g<0 else ''}"

def _angle_deg(dx, dy):
    return (math.degrees(math.atan2(dy, dx)) + 360.0) % 360.0

def _chord_angle(pA, pB):
    return _angle_deg(pB['x'] - pA['x'], pB['y'] - pA['y'])

def _angle_diff_abs(a, b):
    d = (b - a + 540.0) % 360.0 - 180.0
    return abs(d)

def _bulge_tangents_deg(pA, pB, bulge):
    alpha = _chord_angle(pA, pB)
    if abs(bulge) < 1e-12:
        return alpha, alpha
    theta = 4.0 * math.degrees(math.atan(bulge))
    s = 1.0 if bulge > 0 else -1.0
    offset = 90.0 - (abs(theta) / 2.0)
    tan_start = (alpha + s * offset) % 360.0
    tan_end   = (alpha - s * offset) % 360.0
    return tan_start, tan_end

def _polygon_orientation(pts_xyb) -> int:
    # +1 CCW, -1 CW
    area2 = 0.0
    n = len(pts_xyb)
    for i in range(n):
        x1, y1 = pts_xyb[i]['x'], pts_xyb[i]['y']
        x2, y2 = pts_xyb[(i+1) % n]['x'], pts_xyb[(i+1) % n]['y']
        area2 += x1*y2 - x2*y1
    return +1 if area2 > 0 else -1

def _ensure_orientation(points, sentido_desejado):
    pts = list(points)
    orient = _polygon_orientation(pts)
    target = +1 if sentido_desejado == 'anti_horario' else -1
    if orient != target:
        pts.reverse()
        for p in pts:
            p['bulge_next'] = -p.get('bulge_next', 0.0)
    return pts

def _extract_poly_points_with_bulge(doc_dxf):
    msp = doc_dxf.modelspace()
    for e in msp.query("LWPOLYLINE"):
        if not e.closed:
            continue
        pts = []
        for v in list(e):
            x, y = v.dxf.x, v.dxf.y
            bulge_next = float(v.dxf.bulge or 0.0)
            pts.append({'x': x, 'y': y, 'bulge_next': bulge_next})
        return pts
    raise ValueError("Nenhuma LWPOLYLINE fechada encontrada.")

def _chain_lines_closed(msp, tol=1e-6):
    lines = list(msp.query("LINE"))
    if not lines:
        raise ValueError("Não há LWPOLYLINE e nem LINEs no DXF.")
    segs = [((l.dxf.start.x, l.dxf.start.y), (l.dxf.end.x, l.dxf.end.y)) for l in lines]

    def _close(a, b):
        return (abs(a[0]-b[0]) <= tol) and (abs(a[1]-b[1]) <= tol)

    used = [False]*len(segs)
    path = [segs[0][0], segs[0][1]]
    used[0] = True

    changed = True
    while changed:
        changed = False
        for i, (s, e) in enumerate(segs):
            if used[i]:
                continue
            if _close(path[-1], s):
                path.append(e); used[i] = True; changed = True
            elif _close(path[-1], e):
                path.append(s); used[i] = True; changed = True
            elif _close(path[0], e):
                path.insert(0, s); used[i] = True; changed = True
            elif _close(path[0], s):
                path.insert(0, e); used[i] = True; changed = True

    if not _close(path[0], path[-1]):
        raise ValueError("LINEs não formam anel fechado (ou tolerância insuficiente).")

    path = path[:-1]  # remove duplicado final
    pts = [{'x': x, 'y': y, 'bulge_next': 0.0} for (x, y) in path]
    return pts

def _ensure_poly_from_dxf(doc_dxf):
    msp = doc_dxf.modelspace()
    try:
        pts = _extract_poly_points_with_bulge(doc_dxf)
        return pts, None
    except Exception:
        _log_info("Nenhuma LWPOLYLINE fechada; tentando unir LINEs...")
        pts = _chain_lines_closed(msp)
        xy = [(p['x'], p['y']) for p in pts]
        e = msp.add_lwpolyline(xy, format="xy", dxfattribs={"closed": True})
        _log_info("LWPOLYLINE criada a partir de LINEs.")
        return pts, e

def _internal_angles_and_concavity(pts_xyb, sentido_poligonal):
    import math
    n = len(pts_xyb)
    if n < 3:
        return [], []

    s = _polygon_orientation(pts_xyb)  # +1 CCW, -1 CW
    internos_deg = []
    concavo = []

    EPS = 1e-12
    for i in range(n):
        p_prev = pts_xyb[(i-1) % n]
        p      = pts_xyb[i]
        p_next = pts_xyb[(i+1) % n]

        ux = p['x'] - p_prev['x']
        uy = p['y'] - p_prev['y']
        vx = p_next['x'] - p['x']
        vy = p_next['y'] - p['y']

        # evita degenerações
        nu = math.hypot(ux, uy)
        nv = math.hypot(vx, vy)
        if nu < EPS or nv < EPS:
            internos_deg.append(0.0)
            concavo.append(False)
            continue
        ux /= nu; uy /= nu
        vx /= nv; vy /= nv

        cross = ux*vy - uy*vx
        dot   = ux*vx + uy*vy
        theta = math.atan2(cross, dot)  # (-pi, pi]

        internal = math.pi - s*theta
        # normaliza para [0, 2pi)
        while internal < 0:
            internal += 2*math.pi
        while internal >= 2*math.pi:
            internal -= 2*math.pi

        internos_deg.append(math.degrees(internal))

        # côncavo se o giro tiver sinal oposto à orientação
        concavo.append((s*theta) < 0)

    return internos_deg, concavo

def _draw_internal_angles(msp, points, internos_deg, sentido_poligonal, raio_frac=0.10):
    n = len(points)
    for i in range(n):
        p1 = points[i - 1]
        p2 = points[i]
        p3 = points[(i + 1) % n]

        ang_in  = _chord_angle(p2, p1)
        ang_out = _chord_angle(p2, p3)
        interno = internos_deg[i]

        lado1 = math.hypot(p1['x'] - p2['x'], p1['y'] - p2['y'])
        lado2 = math.hypot(p3['x'] - p2['x'], p3['y'] - p2['y'])
        raio = max(0.01, raio_frac * min(lado1, lado2))

        if sentido_poligonal == 'anti_horario':
            start = ang_out
        else:
            start = ang_in
        end = (start + interno) % 360.0
        mid = (start + interno / 2.0) % 360.0

        try:
            if "Internal_Arcs" not in msp.doc.layers:
                msp.doc.layers.add("Internal_Arcs")
        except Exception:
            pass

        msp.add_arc(
            center=(p2['x'], p2['y']),
            radius=raio,
            start_angle=start,
            end_angle=end,
            dxfattribs={'layer': 'Internal_Arcs'}
        )

        pos = (
            p2['x'] + (raio + 1.0) * math.cos(math.radians(mid)),
            p2['y'] + (raio + 1.0) * math.sin(math.radians(mid)),
        )
        texto = _convert_to_dms_safe(interno)
        try:
            if "Labels" not in msp.doc.layers:
                msp.doc.layers.add("Labels")
        except Exception:
            pass

        msp.add_text(
            texto,
            dxfattribs={
                'height': 0.7,
                'layer': 'Labels',
                'insert': pos,
                'rotation': 0,              # ← força horizontal
            }
        )

def _dist2(a, b):
    return (a[0]-b[0])**2 + (a[1]-b[1])**2

def escolher_ponto_az_externo(v1_xy, ponto_az_dxf, pontos_aberta, ponto_amarracao):
    """
    Decide o 'ponto AZ externo' a usar nos cálculos/desenho:
    - Prioriza ponto Az do DXF (se existir e ≠ V1),
    - senão pega o ponto externo mais próximo de V1 vindo da poligonal ABERTA,
    - senão usa o ponto de amarração,
    - senão retorna None.
    Todos como tuplas (x,y,0.0).
    """
    EPS2 = 1e-12
    vx, vy = float(v1_xy[0]), float(v1_xy[1])

    # 1) Az do DXF
    if ponto_az_dxf is not None:
        ax, ay = float(ponto_az_dxf[0]), float(ponto_az_dxf[1])
        if _dist2((ax, ay), (vx, vy)) > EPS2:
            return (ax, ay, 0.0)

    # 2) ponto externo mais próximo (ABERTA)
    if pontos_aberta:
        # cada item pode ser (x,y) ou (x,y,...) → normalize
        candidatos = []
        for p in pontos_aberta:
            px, py = float(p[0]), float(p[1])
            if _dist2((px, py), (vx, vy)) > EPS2:
                candidatos.append((px, py))
        if candidatos:
            px, py = min(candidatos, key=lambda q: _dist2(q, (vx, vy)))
            return (px, py, 0.0)

    # 3) ponto de amarração
    if ponto_amarracao is not None:
        px, py = float(ponto_amarracao[0]), float(ponto_amarracao[1])
        if _dist2((px, py), (vx, vy)) > EPS2:
            return (px, py, 0.0)

    return None

def add_distance_label(msp, p1, p2, distancia):
    try:
        mid = ((p1[0]+p2[0])/2.0, (p1[1]+p2[1])/2.0)
        msp.add_text(
            f"{distancia:,.2f}".replace(",", "").replace(".", ","),
            dxfattribs={"height": 0.25, "layer": "Azimute", "insert": (mid[0], mid[1])}
        )
        logger.info("Rótulo de distância adicionado com sucesso.")
    except Exception as e:
        logger.error(f"Erro ao adicionar rótulo de distância: {e}")


# ==== HELPERS_ANGULOS_DXF_END ====
#DAQUI PARA BAIXO HELPERS RELATIVOS A EXISTENCIA DE BULGE NA POLIGONAL




def _deg(a_rad): 
    return math.degrees(a_rad)

def _rad(a_deg):
    return math.radians(a_deg)

def _norm_deg(a):
    """Normaliza para (-180, 180]."""
    a = (a + 180.0) % 360.0 - 180.0
    return 180.0 if abs(a + 180.0) < 1e-12 else a

def _bearing(p, q):
    """Azimute da corda PQ em graus (0°=E, CCW)."""
    return _deg(math.atan2(q[1] - p[1], q[0] - p[0]))

def _theta_from_bulge(b):
    """Ângulo central do arco (ASSINADO) em graus."""
    return _deg(4.0 * math.atan(b))

def _tangent_dir_at_start(p, q, bulge):
    """
    Direção da TANGENTE no ponto inicial (p) do segmento/ARCO p→q.
    - Linha: direção da corda.
    - Arco: φ + 90° - θ/2   (θ assinado; CCW>0, CW<0)
    """
    phi = _bearing(p, q)
    if abs(bulge) < EPS_BULGE:
        return phi  # linha
    theta = _theta_from_bulge(bulge)
    return phi + 90.0 - (theta / 2.0)

def _tangent_dir_at_end(p, q, bulge):
    """
    Direção da TANGENTE no ponto final (q) do segmento/ARCO p→q.
    - Linha: direção da corda.
    - Arco: φ + 90° + θ/2
    """
    phi = _bearing(p, q)
    if abs(bulge) < EPS_BULGE:
        return phi  # linha
    theta = _theta_from_bulge(bulge)
    return phi + 90.0 + (theta / 2.0)

def _internal_angles_with_bulge(points_bulge):
    """
    Calcula ângulos internos (0–360) em todos os vértices usando tangentes reais.
    `points_bulge`: lista de dicts [{'x':..,'y':..,'bulge_next':..}, ...]
    Retorna lista em graus (float).
    """
    n = len(points_bulge)
    angs = []
    for i in range(n):
        # índices circularmente
        i_prev = (i - 1) % n
        i_next = (i + 1) % n

        p_prev = (points_bulge[i_prev]['x'], points_bulge[i_prev]['y'])
        p_curr = (points_bulge[i]['x'],     points_bulge[i]['y'])
        p_next = (points_bulge[i_next]['x'], points_bulge[i_next]['y'])

        bulge_prev = float(points_bulge[i_prev].get('bulge_next', 0.0))  # do seg (i-1)→i
        bulge_curr = float(points_bulge[i].get('bulge_next', 0.0))        # do seg i→(i+1)

        # direções tangentes no vértice i
        dir_in  = _tangent_dir_at_end(p_prev, p_curr, bulge_prev)   # chegando em Vi
        dir_out = _tangent_dir_at_start(p_curr, p_next, bulge_curr) # saindo de Vi

        # giro assinado (esquerda + / direita -)
        turn = _norm_deg(dir_out - dir_in)

        # ângulo interno da poligonal (concavo pode passar de 180)
        interno = 180.0 - turn
        # normaliza para [0, 360)
        if interno < 0.0:
            interno += 360.0
        elif interno >= 360.0:
            interno -= 360.0

        angs.append(interno)
    return angs



# def create_memorial_descritivo(
#         uuid_str, doc, lines, proprietario, matricula, caminho_salvar, confrontantes, ponto_az,
#         dxf_file_path, area_dxf, azimute, v1, msp, dxf_filename, excel_file_path, tipo,
#         giro_angular_v1_dms, distancia_az_v1, sentido_poligonal='horario',
#         diretorio_concluido=None
#     ):
#     """
#     Cria o memorial descritivo e atualiza o DXF com base no ponto de amarração real (anterior ao V1).
#     """
#     azimute_az_v1=azimute
#     ponto_amarracao=ponto_az
#     distance_az_v1=distancia_az_v1


#     if diretorio_concluido is None:
#         diretorio_concluido = caminho_salvar

#     if not lines:
#         logger.info("Nenhuma linha disponível para criar o memorial descritivo.")
#         return None

#     dxf_file_path = dxf_file_path.strip('"')
#     dxf_output_path = os.path.join(
#         diretorio_concluido,
#         f"{uuid_str}_FECHADA_{tipo}_{matricula}.dxf"
#     )

#     logger.info(f"✅ DXF FECHADA salvo corretamente: {dxf_output_path}")

#     try:
#         doc_dxf = ezdxf.readfile(dxf_file_path)
#         msp = doc_dxf.modelspace()
#     except Exception as e:
#         logger.error(f"Erro ao abrir o arquivo DXF para edição: {e}")
#         return None

#     ordered_points = [line[0] for line in lines]
#     if ordered_points[-1] != lines[-1][1]:
#         ordered_points.append(lines[-1][1])

#     ordered_points = ensure_counterclockwise(ordered_points)
#     area = calcular_area_poligonal(ordered_points)
#     # Agora inverter o sentido corretamente, incluindo tratamento dos arcos (bulge)
#     if sentido_poligonal == 'horario':
#         if area > 0:
#             ordered_points.reverse()
#             area = abs(area)
#             # Inverte o sentido dos arcos (bulges), se existirem
#             for ponto in ordered_points:
#                 if 'bulge' in ponto and ponto['bulge'] != 0:
#                     ponto['bulge'] *= -1
#             logger.info(f"Área da poligonal invertida para sentido horário com ajuste dos arcos: {area:.4f} m²")
#         else:
#             logger.info(f"Área da poligonal já no sentido horário: {abs(area):.4f} m²")

#     else:  # sentido_poligonal == 'anti_horario'
#         if area < 0:
#             ordered_points.reverse()
#             area = abs(area)
#             # Inverte o sentido dos arcos (bulges), se existirem
#             for ponto in ordered_points:
#                 if 'bulge' in ponto and ponto['bulge'] != 0:
#                     ponto['bulge'] *= -1
#             logger.info(f"Área da poligonal invertida para sentido anti-horário com ajuste dos arcos: {area:.4f} m²")
#         else:
#             logger.info(f"Área da poligonal já no sentido anti-horário: {abs(area):.4f} m²")

# #teste
#     # Cálculo de distância V1–Ponto de Amarração
#     #distance_amarracao_v1 = calculate_distance(ponto_amarracao, ordered_points[0])

#     # Corrigir fechamento duplicado
#     tolerancia = 0.001
#     if math.isclose(ordered_points[0][0], ordered_points[-1][0], abs_tol=tolerancia) and \
#        math.isclose(ordered_points[0][1], ordered_points[-1][1], abs_tol=tolerancia):
#         ordered_points.pop()

#     try:
#         data = []
#         total_pontos = len(ordered_points)

#         for i in range(total_pontos):
#             p1 = ordered_points[i - 1] if i > 0 else ordered_points[-1]
#             p2 = ordered_points[i]
#             p3 = ordered_points[(i + 1) % total_pontos]

#             internal_angle = calculate_internal_angle(p1, p2, p3)
#             internal_angle_dms = convert_to_dms(internal_angle)

#             description = f"V{i + 1}_V{(i + 2) if i + 1 < total_pontos else 1}"
#             dx = p3[0] - p2[0]
#             dy = p3[1] - p2[1]
#             distance = math.hypot(dx, dy)
#             confrontante = confrontantes[i % len(confrontantes)]

#             ponto_az_e = f"{ponto_az[0]:,.3f}".replace(",", "").replace(".", ",") if i == 0 else ""
#             ponto_az_n = f"{ponto_az[1]:,.3f}".replace(",", "").replace(".", ",") if i == 0 else ""
#             distancia_az_v1_str = f"{distance_az_v1:.2f}".replace(".", ",") if i == 0 else ""
#             azimute_az_v1_str = convert_to_dms(azimute_az_v1) if i == 0 else ""
#             giro_v1_str = giro_angular_v1_dms if i == 0 else ""

#             data.append({
#                 "V": f"V{i + 1}",
#                 "E": f"{p2[0]:,.3f}".replace(",", "").replace(".", ","),
#                 "N": f"{p2[1]:,.3f}".replace(",", "").replace(".", ","),
#                 "Z": "0,000",
#                 "Divisa": description,
#                 "Angulo Interno": internal_angle_dms,
#                 "Distancia(m)": f"{distance:,.2f}".replace(",", "").replace(".", ","),
#                 "Confrontante": confrontante,
#                 "ponto_AZ_E": ponto_az_e,
#                 "ponto_AZ_N": ponto_az_n,
#                 "distancia_Az_V1": distancia_az_v1_str,
#                 "Azimute Az_V1": azimute_az_v1_str,
#                 "Giro Angular Az_V1_V2": giro_v1_str
#             })

#             if distance > 0.01:
#                 add_label_and_distance(msp, p2, p3, f"V{i + 1}", distance)

#         df = pd.DataFrame(data)
#         df.to_excel(excel_file_path, index=False)
#         # Formatar Excel
#         wb = openpyxl.load_workbook(excel_file_path)
#         ws = wb.active

#         for cell in ws[1]:
#             cell.font = Font(bold=True)
#             cell.alignment = Alignment(horizontal="center", vertical="center")

#         col_widths = {
#             "A": 8, "B": 15, "C": 15, "D": 0, "E": 15,
#             "F": 15, "G": 15, "H": 50, "I": 15,
#             "J": 15, "K": 15, "L": 20, "M": 20
#         }

#         for col, width in col_widths.items():
#             ws.column_dimensions[col].width = width

#         # Corpo
#         for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=1, max_col=ws.max_column):
#             for cell in row:
#                 cell.alignment = Alignment(horizontal="center", vertical="center")

#         wb.save(excel_file_path)

#         # Ler de volta o Excel recém-gerado para pegar ângulos internos exatos
#         df_angulos = pd.read_excel(excel_file_path)
#         angulos_excel_dms = df_angulos["Angulo Interno"].tolist()

#         # Converter DMS para decimal (necessário para o desenho no DXF)
#         angulos_excel_decimal = [convert_dms_to_decimal(a) for a in angulos_excel_dms]

#         # Agora desenha os arcos no DXF com esses ângulos
#         add_angle_visualization_to_dwg(msp, ordered_points, angulos_excel_decimal, sentido_poligonal=sentido_poligonal)


#         logger.info(f"Arquivo Excel salvo em: {excel_file_path}")

         

#         # ➕ Giro Angular
#         try:
#             v1 = ordered_points[0]
#             v2 = ordered_points[1]
#             add_giro_angular_arc_to_dxf(doc_dxf, v1, ponto_az, v2)
#             print("Giro horário Az-V1-V2 adicionado com sucesso.")
#         except Exception as e:
#             print(f"Erro ao adicionar giro angular: {e}")

#         # ➕ Camada e rótulo de vértices
#         if "Vertices" not in msp.doc.layers:
#             msp.doc.layers.add("Vertices", dxfattribs={"color": 1})

#         for i, vertex in enumerate(ordered_points):
#             msp.add_circle(center=vertex, radius=0.5, dxfattribs={"layer": "Vertices"})
#             label_pos = (vertex[0] + 0.3, vertex[1] + 0.3)
#             msp.add_text(f"V{i + 1}", dxfattribs={
#                 "height": 0.3,
#                 "layer": "Vertices",
#                 "insert": label_pos
#             })

#         # ➕ Adicionar arco e rótulo do Azimute
#         try:
#             azimute = calculate_azimuth(ponto_az, v1)
#             add_azimuth_arc_to_dxf(msp, ponto_az, v1, azimute)
#             print("Arco do Azimute Az-V1 adicionado com sucesso.")
#         except Exception as e:
#             print(f"Erro ao adicionar arco do azimute: {e}")

#         # ➕ Adicionar distância Az–V1
#         try:
#             distancia_az_v1 = calculate_distance(ponto_az, v1)
#             add_label_and_distance(msp, ponto_az, v1, "", distancia_az_v1)
#             print(f"Distância Az-V1 ({distancia_az_v1:.2f} m) adicionada com sucesso.")
#         except Exception as e:
#             print(f"Erro ao adicionar distância entre Az e V1: {e}")

#         # ➕ Salvar DXF
#         doc_dxf.saveas(dxf_output_path)
#         print(f"📁 Arquivo DXF final salvo em: {dxf_output_path}")

#     except Exception as e:
#         print(f"❌ Erro ao gerar o memorial descritivo: {e}")
#         return None

#     return excel_file_path

def create_memorial_descritivo(
    uuid_str, doc, lines, proprietario, matricula, caminho_salvar, confrontantes, ponto_az,
    dxf_file_path, area_dxf, azimute, v1, msp, dxf_filename, excel_file_path, tipo,
    giro_angular_v1_dms, distancia_az_v1, sentido_poligonal='horario', modo="ANGULO_P1_P2",
    diretorio_concluido=None,
    points_bulge=None
):
    """
    DXF já vem tratado do pipeline (sem reler aqui):
    - Usa LWPOLYLINE fechada (com bulge) fornecida por get_document_info_from_dxf.
    - Normaliza sentido, calcula ângulos internos (com concavidade), desenha arcos internos.
    - Só desenha AZ no modo ANGULO_AZ.
    - Gera Excel diretamente.
    """
    logger.info("[CMD] pontos_bulge recebidos: %s", len(points_bulge) if points_bulge else 0)

    # Garante que o estilo de texto "STANDARD" exista no DXF
    if "STANDARD" not in msp.doc.styles:
        msp.doc.styles.new("STANDARD")

    if diretorio_concluido is None:
        diretorio_concluido = caminho_salvar

    dxf_output_path = os.path.join(
        diretorio_concluido,
        f"{uuid_str}_FECHADA_{tipo}_{matricula}.dxf"
    )

    # 0) valida base
    if points_bulge is None or len(points_bulge) < 3:
        _log_error(f"[AZ] points_bulge ausente/insuficiente (type={type(points_bulge)}, len={0 if not points_bulge else len(points_bulge)}). Verifique get_document_info_from_dxf.")
        return None

    # 1) normaliza sentido
    pts = _ensure_orientation(points_bulge, sentido_poligonal)
    orient = _polygon_orientation(pts)
    _log_info(f"Sentido normalizado: {'anti-horário' if orient == +1 else 'horário'}")

    # 2) ângulos internos + concavidade
    EPS_BULGE = 1e-9

    has_any_bulge = any(abs(float(p.get('bulge_next', 0.0))) > EPS_BULGE for p in pts)

    if has_any_bulge:
        internos_deg = _internal_angles_with_bulge(pts)  # usa tangentes reais (funciona também para bulge=0)
        concavo = [a > 180.0 for a in internos_deg]      # se você quiser a flag de concavidade
        logger.info("Ângulos internos: modo BULGE-AWARE (misto retas+arcos).")
    else:
        # compatibilidade com sua rotina antiga quando não há bulge algum
        internos_deg, concavo = _internal_angles_and_concavity(pts, sentido_poligonal)
        logger.info("Ângulos internos: modo LEGADO (somente retas).")

    # 3) desenha arcos internos por dentro
    _draw_internal_angles(msp, pts, internos_deg, sentido_poligonal, raio_frac=0.10)

    # 4) desenho do AZ depende do modo
    # ANGULO_AZ  → desenha Az, linha Az–V1, arco e rótulos
    # ANGULO_P1_P2 → NÃO desenha Az/linha/arco (poligonal ABERTA já mostra amarração)
    if modo == "ANGULO_AZ" and ponto_az is not None and v1 is not None:
        dx = v1[0] - ponto_az[0]
        dy = v1[1] - ponto_az[1]
        dist = math.hypot(dx, dy)
        if dist > 1e-6:
            try:
                _desenhar_referencia_az(msp, ponto_az, v1, azimute)
            except Exception as e:
                logger.error("Erro ao desenhar referência de Az: %s", e)
        else:
            logger.warning("⚠️ Distância Az–V1 ≈ 0; desenho do Az suprimido.")

    # 5) Excel (sem reler nada)
    try:
        ordered_points_xy = [(p['x'], p['y']) for p in pts]
        total_pontos = len(ordered_points_xy)
        data = []

        if ponto_az is not None:
            ponto_az_e = f"{ponto_az[0]:,.3f}".replace(",", "").replace(".", ",")
            ponto_az_n = f"{ponto_az[1]:,.3f}".replace(",", "").replace(".", ",")
        else:
            ponto_az_e = ""
            ponto_az_n = ""

        for i in range(total_pontos):
            p2 = ordered_points_xy[i]
            p3 = ordered_points_xy[(i + 1) % total_pontos]

            dx, dy = p3[0] - p2[0], p3[1] - p2[1]
            distance = math.hypot(dx, dy)
            description = f"V{i + 1}_V{(i + 2) if i + 1 < total_pontos else 1}"
            confrontante = confrontantes[i % len(confrontantes)] if confrontantes else ""
            ang_interno_dms = _convert_to_dms_safe(internos_deg[i])

            if i == 0:
                distancia_az_v1_str = f"{float(distancia_az_v1):.2f}".replace(".", ",") if distancia_az_v1 is not None else ""
                azimute_az_v1_str   = _convert_to_dms_safe(float(azimute)) if azimute is not None else ""
                giro_v1_str         = giro_angular_v1_dms or ""
                p_az_e, p_az_n      = ponto_az_e, ponto_az_n
            else:
                distancia_az_v1_str = ""
                azimute_az_v1_str   = ""
                giro_v1_str         = ""
                p_az_e, p_az_n      = "", ""

            data.append({
                "V": f"V{i + 1}",
                "E": f"{p2[0]:,.3f}".replace(",", "").replace(".", ","),
                "N": f"{p2[1]:,.3f}".replace(",", "").replace(".", ","),
                "Z": "0,000",
                "Divisa": description,
                "Angulo Interno": ang_interno_dms,
                "Distancia(m)": f"{distance:,.2f}".replace(",", "").replace(".", ","),
                "Confrontante": confrontante,
                "ponto_AZ_E": p_az_e,
                "ponto_AZ_N": p_az_n,
                "distancia_Az_V1": distancia_az_v1_str,
                "Azimute Az_V1": azimute_az_v1_str,
                "Giro Angular Az_V1_V2": giro_v1_str
            })

            try:
                if distance > 0.01 and 'add_label_and_distance' in globals():
                    add_label_and_distance(msp, p2, p3, f"V{i + 1}", distance)
            except Exception as e:
                _log_error(f"Falha ao rotular distância do lado V{i+1}: {e}")

        
        # escreve excel
        df = pd.DataFrame(data)

        # ── Garantir as 3 colunas do ANGULO_AZ antes de salvar
        cols_novas = ["AZIMUTE_AZ_V1_GRAUS", "DISTANCIA_AZ_V1_M", "GIRO_V1_GRAUS"]
        for c in cols_novas:
            if c not in df.columns:
                df[c] = ""  # ou pd.NA

        # Garante diretório e salva
        try:
            os.makedirs(os.path.dirname(excel_file_path), exist_ok=True)
            df.to_excel(excel_file_path, index=False)
            _log_info(f"Excel escrito (primeira passagem): {os.path.abspath(excel_file_path)}")
        except Exception as e:
            _log_error(f"Falha ao salvar Excel na primeira passagem: {e}")
            raise  # deixe a main ver o stacktrace

        # Formatação openpyxl
        try:
            wb = openpyxl.load_workbook(excel_file_path)
            ws = wb.active

            for cell in ws[1]:
                cell.font = Font(bold=True)
                cell.alignment = Alignment(horizontal="center", vertical="center")

            col_widths = {
                "A": 8, "B": 15, "C": 15, "D": 0, "E": 15,
                "F": 15, "G": 15, "H": 50, "I": 15,
                "J": 15, "K": 15, "L": 20, "M": 20
            }
            for col, width in col_widths.items():
                ws.column_dimensions[col].width = width

            for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=1, max_col=ws.max_column):
                for cell in row:
                    cell.alignment = Alignment(horizontal="center", vertical="center")

            wb.save(excel_file_path)
            _log_info(f"Excel salvo e formatado: {os.path.abspath(excel_file_path)}")
        except Exception as e:
            _log_error(f"Falha ao formatar/salvar Excel com openpyxl: {e}")
            raise

        # Confirma existência
        if os.path.exists(excel_file_path):
            _log_info(f"✅ Excel confirmado em disco: {os.path.abspath(excel_file_path)}")
        else:
            _log_error(f"❌ Excel NÃO encontrado após salvar: {os.path.abspath(excel_file_path)}")


        # extras DXF (opcionais e seguros)
        try:
            v1_pt = ordered_points_xy[0]
            v2_pt = ordered_points_xy[1]
            # se existir o helper e você quiser o giro no V1 com Az:
            if 'add_giro_angular_arc_to_dxf' in globals() and ponto_az is not None:
                # padronize este helper para (msp, v1_pt, ponto_az, v2_pt)
                add_giro_angular_arc_to_dxf(msp, v1_pt, ponto_az, v2_pt)
                _log_info("Giro horário Az–V1–V2 adicionado com sucesso.")
        except Exception as e:
            _log_error(f"Erro ao adicionar giro angular: {e}")

        try:
            if "Vertices" not in msp.doc.layers:
                msp.doc.layers.add("Vertices")
        except Exception:
            pass

        # garanta a camada
        try:
            if "Vertices" not in msp.doc.layers:
                msp.doc.layers.add("Vertices")
        except Exception:
            pass

        for i, (x, y) in enumerate(ordered_points_xy):
            try:
                msp.add_circle(center=(x, y), radius=0.5, dxfattribs={"layer": "Vertices"})
                msp.add_text(
                    f"V{i + 1}",
                    dxfattribs={
                        "height": 0.3,
                        "layer": "Vertices",
                        "insert": (x + 0.30, y + 0.30)  # <<< POSIÇÃO DO RÓTULO
                    }
                )
            except Exception as e:
                logger.warning(f"Falha rotulando V{i+1}: {e}")


        # só desenhe o arco do azimute se realmente quiser no produto FECHADA
        # e se houver amarração (Az) válida:
        if modo == "ANGULO_AZ" and ponto_az is not None:
            try:
                azim = calculate_azimuth(ponto_az, v1_pt)
                _desenhar_referencia_az(msp, ponto_az, v1_pt, azim)
                _log_info("Arco do Azimute Az–V1 adicionado com sucesso.")
            except Exception as e:
                _log_error(f"Erro ao adicionar arco do azimute: {e}")

        # 6) salvar DXF final
        try:
            doc.saveas(dxf_output_path)
            logger.info("✅ DXF FECHADA salvo corretamente: %s", dxf_output_path)
        except Exception as e:
            logger.error("Erro ao salvar DXF FECHADA: %s", e)

    except Exception as e:
        _log_error(f"❌ Erro ao gerar o memorial descritivo: {e}")
        return None

    return excel_file_path

def rotate_polygon_start_at_v1(lines, pts_bulge, v1_target, sentido_poligonal="", tol=1e-4):
    """
    Reordena 'lines' ([(p_i, p_{i+1})...]) e 'pts_bulge' para que o primeiro
    vértice seja o mais próximo de v1_target. Se for anti-horário, dá 1 passo extra.
    """
    if not lines:
        return lines, pts_bulge, None

    import math
    def dist2(a, b): 
        return (a[0]-b[0])**2 + (a[1]-b[1])**2

    # índice do vértice mais próximo de v1_target
    verts = [seg[0] for seg in lines]
    idx = min(range(len(verts)), key=lambda i: dist2(verts[i], v1_target))

    # roda para começar em idx
    lines = lines[idx:] + lines[:idx]
    pts_bulge = pts_bulge[idx:] + pts_bulge[:idx]
    rot_idx = idx

    # se for anti-horário, ande uma casa no “sentido horário” desejado
    if str(sentido_poligonal).lower().startswith("anti"):
        lines = lines[1:] + lines[:1]
        pts_bulge = pts_bulge[1:] + pts_bulge[:1]
        rot_idx = (idx + 1) % len(verts)

    return lines, pts_bulge, rot_idx






def generate_initial_text(proprietario, matricula, descricao, area, perimeter, rua, cidade, ponto_amarracao, azimute, distancia):
    """
    Gera o texto inicial do memorial descritivo.
    """
    initial_text = (
        f"MEMORIAL DESCRITIVO\n"
        f"NOME PROPRIETÁRIO / OCUPANTE: {proprietario}\n"
        f"DESCRIÇÃO: {descricao}\n"
        f"DOCUMENTAÇÃO: MATRÍCULA {matricula}\n"
        f"ÁREA DO IMÓVEL: {area:.2f} metros quadrados\n"
        f"PERÍMETRO: {perimeter:.4f} metros\n"
        f"Área localizada na rua {rua}, município de {cidade}, com a seguinte descrição:\n"
        f"O Ponto Az está localizado nas coordenadas E {ponto_amarracao[0]:.3f}, N {ponto_amarracao[1]:.3f}.\n"
        f"Daí, com Azimute de {azimute} e distância de {distancia:.2f} metros, chega-se ao Vértice V1, "
        f"origem da área descrição, alinhado com a rua {rua}."
    )
    return initial_text


def generate_angular_text(az_v1, v1_v2, distancia, angulo, rua, confrontante):
    """
    Gera o texto do ângulo em V1 entre Az-V1 e V1-V2.
    """
    angular_text = (
        f"Daí, visando o Ponto “Az”, com giro angular horário de {angulo} e diste {distancia:.2f} metros, "
        f"chega-se ao Vértice V2, também alinhado com a rua {rua} e limítrofe com {confrontante}."
    )
    return angular_text


def generate_recurring_text(df, rua, confrontantes):
    """
    Gera o texto recorrente enquanto percorre os vértices da poligonal.
    """
    recurring_texts = []
    num_vertices = len(df)
    
    for i in range(1, num_vertices - 1):  # De V2 até o penúltimo vértice
        current = df.iloc[i]
        next_vertex = df.iloc[i + 1]
        confrontante = confrontantes[i % len(confrontantes)]
        
        recurring_text = (
            f"Daí, visando o Vértice {current['V']}, com giro angular horário de {current['Angulo Interno']} "
            f"e diste {current['Distancia(m)']} metros, chega-se ao Vértice {next_vertex['V']}, "
            f"também alinhado à rua {rua} e limítrofe com {confrontante}."
        )
        recurring_texts.append(recurring_text)
    
    return recurring_texts


def generate_final_text(df, rua, confrontantes):
    """
    Gera o texto final ao retornar ao V1.
    """
    last_vertex = df.iloc[-1]
    first_vertex = df.iloc[0]
    confrontante = confrontantes[-1 % len(confrontantes)]  # Confrontante do último trecho
    
    final_text = (
        f"Daí, visando o vértice {last_vertex['V']}, com giro angular de {last_vertex['Angulo Interno']} "
        f"e diste {last_vertex['Distancia(m)']} metros, chega-se ao vértice {first_vertex['V']}, "
        f"origem da presente descrição, no alinhamento da rua {rua} e próximo aos lotes de {confrontante}."
    )
    return final_text


def create_memorial_document(
    uuid_str,
    proprietario,
    matricula,
    matricula_texto,
    area_total,
    cpf,
    rgi,
    excel_file_path,
    template_path,
    output_path,
    assinatura_path,
    ponto_amarracao,
    azimute,
    distancia_amarracao_v1,
    rua,
    cidade,
    confrontantes,
    area_dxf,
    desc_ponto_amarracao,
    perimeter_dxf,
    giro_angular_v1_dms
):

    try:
        # Ler a planilha gerada
        df = pd.read_excel(excel_file_path, engine='openpyxl', dtype=str)
        logger.info("AREA_TOTAL_DENTRO_DOCUMENT: %s", area_total)


        # Corrigir vírgulas em valores numéricos
        df['Distancia(m)'] = df['Distancia(m)'].str.replace(',', '.').astype(float)
        df['E'] = df['E'].str.replace(',', '').astype(float)
        df['N'] = df['N'].str.replace(',', '').astype(float)

        # Calcular perímetro e área se necessário
        perimeter = df['Distancia(m)'].sum()
        x = df['E'].values
        y = df['N'].values
        area = abs(sum(x[i] * y[(i + 1) % len(x)] - x[(i + 1) % len(x)] * y[i] for i in range(len(x))) / 2)

        # Criar documento Word
        doc_word = Document(template_path)
        # 🔴 Remove parágrafos vazios iniciais (inclusive espaços e quebras invisíveis)
        while doc_word.paragraphs and not doc_word.paragraphs[0].text.strip():
            p_element = doc_word.paragraphs[0]._element
            p_element.getparent().remove(p_element)
        
        set_default_font(doc_word)

        v1 = (df.iloc[0]['E'], df.iloc[0]['N'])

        # Converter para formatos amigáveis
        area_dxf_formatada = f"{area_dxf:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        ponto_amarracao_1 = f"{ponto_amarracao[0]:.3f}".replace(".", ",")
        ponto_amarracao_2 = f"{ponto_amarracao[1]:.3f}".replace(".", ",")
        distancia_str = f"{distancia_amarracao_v1:.2f}".replace(".", ",")
        azimute_dms = convert_to_dms(azimute)


        
        
        # 🔴 Remove parágrafos indesejados como o Copilot
        for para in doc_word.paragraphs:
            if "copilot" in para.text.lower():
                p_element = para._element
                p_element.getparent().remove(p_element)

        # ⬇️ Agora insere o título sem espaçamento extra
        p = doc_word.add_paragraph(style='Normal')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.add_run("MEMORIAL DESCRITIVO").bold = True
                
        
        doc_word.add_paragraph()  # Parágrafo vazio para pular uma linha
        p = doc_word.add_paragraph(style='Normal')
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.add_run("Objetivo: ").bold = True
        p.add_run(f"Área destinada à servidão de passagem para execução de coletor de fundo pertencente à rede coletora de esgoto de {cidade}/RS.")


        p = doc_word.add_paragraph(style='Normal')
        p.add_run("Matrícula Número: ").bold = True
        p.add_run(f"{matricula_texto} - {rgi}")

        area_total_formatada = str(area_total).replace(".", ",")
        p = doc_word.add_paragraph(style='Normal')
        p.add_run("Área Total do Terreno: ").bold = True
        p.add_run(area_total_formatada)



        p = doc_word.add_paragraph(style='Normal')
        p.add_run("Proprietário: ").bold = True
        p.add_run(f"{proprietario} - CPF/CNPJ: {cpf}")


        p = doc_word.add_paragraph(style='Normal')
        p.add_run("Área de Servidão de Passagem: ").bold = True
        run1 = p.add_run(f"{area_dxf_formatada} m"); run1.font.name = 'Arial'; run1.font.size = Pt(12)
        run2 = p.add_run("2"); run2.font.name = 'Arial'; run2.font.size = Pt(12); run2.font.superscript = True


        

        doc_word.add_paragraph()  # Parágrafo vazio para pular uma linha
        p = doc_word.add_paragraph(style='Normal')
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        #p.add_run("Descrição: ").bold = True
        p.add_run("Área com ").font.name = 'Arial'

        run1 = p.add_run(f"{area_dxf_formatada} m")
        run1.font.name = 'Arial'
        run1.font.size = Pt(12)

        run2 = p.add_run("2")
        run2.font.name = 'Arial'
        run2.font.size = Pt(12)
        run2.font.superscript = True

        p.add_run(f" localizada na {rua}, município de {cidade}, com a finalidade de servidão de passagem com a seguinte descrição e confrontações, onde os ângulos foram medidos no sentido horário.").font.name = 'Arial'








        # 🔵 Texto opcional com o ponto de amarração real
        doc_word.add_paragraph()  # Parágrafo vazio para pular uma linha
        doc_word.add_paragraph(
            f"Pontos definidos pelas Coordenadas Planas no Sistema U.T.M. – SIRGAS 2000.",
            style='Normal'
        )
        #doc_word.add_paragraph(
        #    f"Daí, com azimute de {azimute_dms} e distância de {distancia_str} metros, chega-se ao Vértice V1, origem da área descrita.",
        #    style='Normal'
        #)

        # Descrição sequencial dos vértices
        for i in range(len(df)):
            current = df.iloc[i]
            next_vertex = df.iloc[(i + 1) % len(df)]
            distancia = f"{current['Distancia(m)']:.2f}".replace(".", ",")
            confrontante = current['Confrontante']
            giro_angular = current['Angulo Interno']

            if i == 0:  # Primeira linha é sempre V1
                p = doc_word.add_paragraph(style='Normal')
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                p.add_run("Do vértice ").bold = False
                p.add_run(f"{current['V']}").bold = True
                p.add_run(
                    f", com giro angular horário de {giro_angular_v1_dms} e distância de {distancia} metros, "
                    f"confrontando com área pertencente à {confrontante}, chega-se ao vértice "
                )
                p.add_run(f"{next_vertex['V']}").bold = True
                p.add_run(";")
                doc_word.add_paragraph()

            elif next_vertex['V'] == "V1" and i == len(df) - 1:
                p = doc_word.add_paragraph(style='Normal')
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                p.add_run("Do vértice ").bold = False
                p.add_run(f"{current['V']}").bold = True
                p.add_run(
                    f", com giro angular horário de {giro_angular} e distância de {distancia} metros, "
                    f"confrontando com área pertencente à {confrontante}, chega-se ao vértice "
                )
                p.add_run(f"{next_vertex['V']}").bold = True
                p.add_run(", origem da presente descrição.")
                doc_word.add_paragraph()

            else:
                p = doc_word.add_paragraph(style='Normal')
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                p.add_run("Do vértice ").bold = False
                p.add_run(f"{current['V']}").bold = True
                p.add_run(
                    f", com giro angular horário de {giro_angular} e distância de {distancia} metros, "
                    f"confrontando com área pertencente à {confrontante}, chega-se ao vértice "
                )
                p.add_run(f"{next_vertex['V']}").bold = True
                p.add_run(";")
                doc_word.add_paragraph()


        doc_word.add_paragraph()  # Parágrafo vazio para pular uma linha
        p = doc_word.add_paragraph(style='Normal')
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.add_run(f"Os angulos foram medidos no sentido horário.")


        # Rodapé
        data_atual = datetime.now().strftime("%d de %B de %Y")

        # converte mês para português
        for ingles, portugues in MESES_PT_BR.items():
            if ingles in data_atual:
                data_atual = data_atual.replace(ingles, portugues)
                break
        doc_word.add_paragraph(f"\nPorto Alegre, RS, {data_atual}.", style='Normal')
        doc_word.add_paragraph("\n\n")

        doc_word.save(output_path)
        logger.info(f"Memorial descritivo salvo em: {output_path}")

    except Exception as e:
        logger.error(f"Erro ao criar o documento memorial: {e}")




def find_excel_file(directory, keywords):
    """
    Busca um arquivo Excel no diretório contendo todas as palavras-chave no nome.
    Se não encontrar, exibe a lista de arquivos disponíveis.
    """
    if not directory or not os.path.exists(directory):
        logger.info(f"Erro: O diretório '{directory}' não existe ou não foi especificado corretamente.")
        return None

    excel_files = [file for file in os.listdir(directory) if file.endswith(".xlsx")]

    if not excel_files:
        logger.info(f"Nenhum arquivo Excel encontrado no diretório: {directory}")
        return None

    for file in excel_files:
        if all(keyword.lower() in file.lower() for keyword in keywords):
            return os.path.join(directory, file)

    # Se nenhum arquivo correspondente foi encontrado, listar os arquivos disponíveis
    logger.info(f"Nenhum arquivo Excel contendo {keywords} foi encontrado em '{directory}'.")
    logger.info("Arquivos disponíveis no diretório:")
    for f in excel_files:
        logger.info(f"  - {f}")

    return None



        
# def convert_docx_to_pdf(output_path, pdf_file_path):
#     """
#     Converte um arquivo DOCX para PDF usando a biblioteca comtypes.
#     """
#     try:
#         # Verificar se o arquivo DOCX existe antes de abrir
#         if not os.path.exists(output_path):
#             raise FileNotFoundError(f"Arquivo DOCX não encontrado: {output_path}")
        
#         logger.info(f"Tentando converter o arquivo DOCX: {output_path} para PDF: {pdf_file_path}")

#         word = comtypes.client.CreateObject("Word.Application")
#         word.Visible = False  # Ocultar a interface do Word
#         doc = word.Documents.Open(output_path)
        
#         # Aguardar alguns segundos antes de salvar como PDF
#         import time
#         time.sleep(2)

#         doc.SaveAs(pdf_file_path, FileFormat=17)  # 17 corresponde ao formato PDF
#         doc.Close()
#         word.Quit()
#         logger.info(f"Arquivo PDF salvo com sucesso em: {pdf_file_path}")
#     except FileNotFoundError as fnf_error:
#         logger.error(f"Erro: {fnf_error}")
#     except Exception as e:
#         logger.info(f"Erro ao converter DOCX para PDF: {e}")
#     finally:
#         try:
#             word.Quit()
#         except:
#             pass  # Garantir que o Word seja fechado



def sanitize_filename(filename):
    return re.sub(r'[<>:"/\\|?*]', '', filename)

        
def main_poligonal_fechada(uuid_str, excel_path, dxf_path, diretorio_preparado, diretorio_concluido, caminho_template, sentido_poligonal='horario'):

    caminho_salvar = diretorio_concluido 
    template_path = caminho_template 
    # Carrega dados do imóvel
    dados_imovel_excel_path = excel_path
    dados_imovel_df = pd.read_excel(dados_imovel_excel_path, sheet_name='Dados_do_Imóvel', header=None)
    dados_imovel = dict(zip(dados_imovel_df.iloc[:, 0], dados_imovel_df.iloc[:, 1]))

    # Extrai informações
    proprietario = dados_imovel.get("NOME DO PROPRIETÁRIO", "").strip()
    cpf = dados_imovel.get("CPF/CNPJ", "").strip()
    matricula = sanitize_filename(str(dados_imovel.get("DOCUMENTAÇÃO DO IMÓVEL", "")).strip())
    matricula_texto = str(dados_imovel.get("DOCUMENTAÇÃO DO IMÓVEL", "")).strip()
    descricao = dados_imovel.get("OBRA", "").strip()
    area_total = dados_imovel.get("ÁREA TOTAL DO TERRENO DOCUMENTADA", "").replace("\t", "").replace("\n", "").strip()
    cidade = dados_imovel.get("CIDADE", "").strip().capitalize()
    rgi= dados_imovel.get("RGI", "").strip().capitalize()
    rua = dados_imovel.get("LOCAL", "").strip()
    desc_ponto_Az = dados_imovel.get("AZ", "").strip()

    # Diretório para salvar resultados
    
    os.makedirs(caminho_salvar, exist_ok=True)

    # Identifica tipo (SER, REM, etc)
    dxf_filename = os.path.basename(dxf_path).upper()
    if "ETE" in dxf_filename:
        tipo = "ETE"
    elif "REM" in dxf_filename:
        tipo = "REM"
    elif "SER" in dxf_filename:
        tipo = "SER"
    elif "ACE" in dxf_filename:
        tipo = "ACE"
    else:
        logger.info("❌ Não foi possível determinar automaticamente o tipo (ETE, REM, SER ou ACE).")
        return

    padrao_fechada = os.path.join(diretorio_preparado, f"{uuid_str}_FECHADA_{tipo}*.xlsx")

    arquivos_encontrados = glob.glob(padrao_fechada)
    if not arquivos_encontrados:
        logger.info(f"❌ Arquivo de confrontantes não encontrado com o padrão: {padrao_fechada}")
        return
    confrontantes_df = pd.read_excel(arquivos_encontrados[0])
    confrontantes = confrontantes_df.iloc[:, 1].dropna().tolist()

    # DXF limpo
    # ⚠️ Substitui a limpeza anterior por apenas conversão R2010
    dxf_limpo_path = os.path.join(caminho_salvar, f"DXF_LIMPO_{matricula}.dxf")
    dxf_file_path = limpar_dxf_e_converter_r2010(dxf_path, dxf_limpo_path)


    # 🔍 Buscar planilha que COMEÇA com ABERTA_{TIPO} no diretório CONCLUIDO
    padrao_aberta = os.path.join(diretorio_concluido, f"{uuid_str}_ABERTA_{tipo}*.xlsx")
    planilhas_aberta = glob.glob(padrao_aberta)

    if not planilhas_aberta:
        logger.info(f"❌ Nenhuma planilha encontrada começando com 'ABERTA_{tipo}' no diretório: {diretorio_concluido}")
        return

    planilha_aberta_saida = planilhas_aberta[0]
    logger.info(f"📄 Planilha ABERTA localizada: {planilha_aberta_saida}")

   
    if not planilhas_aberta:
        logger.info(f"❌ Nenhuma planilha encontrada contendo 'ABERTA' e '{tipo}' no nome dentro de: {diretorio_concluido}")
        return

    planilha_aberta_saida = planilhas_aberta[0]
    logger.info(f"📄 Planilha ABERTA localizada: {planilha_aberta_saida}")


    # 📁 Procurar CONCLUIDO dentro da cidade (REPESCAGEM_*/CONCLUIDO)
    # O diretório CONCLUIDO já é passado corretamente
    diretorio_concluido_real = diretorio_concluido

   
    # 🧭 Obter ponto de amarração anterior ao V1
    try:
        ponto_amarracao, codigo_amarracao = obter_ponto_amarracao_anterior_v1(planilha_aberta_saida)
        logger.info(f"📌 Ponto de amarração identificado: {codigo_amarracao} com coordenadas {ponto_amarracao}")
    except Exception as e:
        logger.error(f"❌ Erro ao obter ponto de amarração: {e}")
        return

    # 🔍 Extrair geometria do DXF
    # Extrair geometria FECHADA do DXF
    doc, lines, perimeter_dxf, area_dxf, ponto_az_dxf, msp, pts_bulge = get_document_info_from_dxf(dxf_file_path)

    # Pare aqui se não houver geometria válida
    if not (doc and lines):
        logger.info("Nenhuma linha foi encontrada ou não foi possível acessar o documento.")
        pythoncom.CoUninitialize()
        return

    logger.info(f"📐 Área da poligonal: {area_dxf:.6f} m²")

    # >>> Fixar V1 no vértice tocado pela poligonal aberta (usa ponto_amarracao)
    def _dist2(a, b):
        dx = a[0] - b[0]; dy = a[1] - b[1]
        return dx*dx + dy*dy

    cand1 = lines[0][0]      # V1 atual
    cand2 = lines[-1][0]     # “V28” (vizinho anterior)
    target_v1 = cand1 if _dist2(cand1, ponto_amarracao) <= _dist2(cand2, ponto_amarracao) else cand2

    lines, pts_bulge, _ = rotate_polygon_start_at_v1(lines, pts_bulge, target_v1, sentido_poligonal)

    # v1 = lines[0][0]
    # v2 = lines[1][0]

    v1 = lines[0][0]
    v_next = lines[1][0]   # V2 (seguindo a ordem da polyline)
    v_prev = lines[-1][0]  # V28 (vizinho “anterior”)

    # Use o vizinho conforme o sentido escolhido na UI
    if str(sentido_poligonal).lower().startswith("anti"):
        v2_for_arc = v_prev     # anti-horário → vizinho “anterior” (V28)
    else:
        v2_for_arc = v_next     # horário → vizinho “seguinte” (V2)

    if ponto_az_dxf and v1 and abs(ponto_az_dxf[0]-v1[0])<1e-6 and abs(ponto_az_dxf[1]-v1[1])<1e-6:
        logger.warning("⚠️ Ponto Az parece ser fallback (igual ao V1). Verifique o DXF original.")

    # Use o ponto retornado pela função
    azimute = calculate_azimuth(ponto_az_dxf, v1)
    distancia_az_v1 = calculate_distance(ponto_az_dxf, v1)
    giro_angular_v1 = calculate_angular_turn(ponto_az_dxf, v1, v2_for_arc)
    giro_angular_v1_dms = convert_to_dms(360 - giro_angular_v1)

    logger.info(f"📌 Azimute Az→V1: {azimute:.4f}°, Distância: {distancia_az_v1:.2f} m")

    # Caminho do Excel de saída
    excel_file_path = os.path.join(
        diretorio_concluido,
        f"{uuid_str}_FECHADA_{tipo}_{matricula}.xlsx"
    )
    logger.info(f"✅ Excel FECHADA salvo corretamente: {excel_file_path}")

    # 🛠 Criar memorial e Excel (passe modo e pts_bulge)
    create_memorial_descritivo(
        uuid_str, doc, lines, proprietario, matricula, caminho_salvar, confrontantes, ponto_az_dxf,
        dxf_file_path, area_dxf, azimute, v1, msp, dxf_filename, excel_file_path, tipo,
        giro_angular_v1_dms, distancia_az_v1, sentido_poligonal=sentido_poligonal,
        modo="ANGULO_P1_P2", points_bulge=pts_bulge
    )

    # 📄 Gerar DOCX (apenas uma vez)
    if excel_file_path:
        output_path_docx = os.path.join(
            diretorio_concluido,
            f"{uuid_str}_FECHADA_{tipo}_{matricula}.docx"
        )
        logger.info(f"✅ DOCX FECHADA salvo corretamente: {output_path_docx}")

        assinatura_path = r"C:\Users\Paulo\Documents\CASSINHA\MEMORIAIS DESCRITIVOS\Assinatura.jpg"
        desc_ponto_amarracao = f"ponto {codigo_amarracao}, obtido na planilha da poligonal aberta"

        create_memorial_document(
            uuid_str=uuid_str,
            proprietario=proprietario,
            matricula=matricula,
            matricula_texto=matricula_texto,
            area_total=area_total,
            cpf=cpf,
            rgi=rgi,
            excel_file_path=excel_file_path,
            template_path=template_path,
            output_path=output_path_docx,
            assinatura_path=assinatura_path,
            ponto_amarracao=ponto_amarracao,
            azimute=azimute,
            distancia_amarracao_v1=distancia_az_v1,
            rua=rua,
            cidade=cidade,
            confrontantes=confrontantes,
            area_dxf=area_dxf,
            desc_ponto_amarracao=desc_ponto_amarracao,
            perimeter_dxf=perimeter_dxf,
            giro_angular_v1_dms=giro_angular_v1_dms,
        )
    else:
        logger.info("excel_file_path não definido ou inválido.")

    logger.info("Documento do AutoCAD fechado.")





