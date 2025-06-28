import os
import gc
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import ezdxf
import glob
import logging
from ezdxf.addons import Importer

BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
TEMPLATE_PADRAO = os.path.join(BASE_DIR, 'templates', 'template_padrao.docx')

logger = logging.getLogger(__name__)

# Função otimizada para extrair conteúdo DOCX
def extrair_conteudo_docx(docx_path):
    logger.info(f"📂 Extraindo conteúdo de: {docx_path}")
    if not os.path.exists(docx_path):
        logger.error("❌ Arquivo DOCX não encontrado.")
        return []

    doc = Document(docx_path)
    conteudo = [
        (p.text, p.alignment, [(run.text, run.bold, run.italic, run.font.size, run.font.name) for run in p.runs])
        for p in doc.paragraphs
    ]
    logger.info(f"📌 {len(conteudo)} parágrafos extraídos do DOCX.")
    return conteudo

# Função otimizada para inserir conteúdo entre parágrafos com template
def inserir_conteudo_entre_paragrafos_com_template(template_path, doc_fechado_path, conteudo_aberto,
                                                   output_docx_path, paragrafo_inicial, paragrafo_final):
    doc_final = Document(template_path)
    doc_fechado = Document(doc_fechado_path)

    ja_inserido = False

    for par in doc_fechado.paragraphs:
        novo_par = doc_final.add_paragraph()
        novo_par.alignment = par.alignment

        for run in par.runs:
            novo_run = novo_par.add_run(run.text)
            novo_run.bold = run.bold
            novo_run.font.size = Pt(12)
            novo_run.font.name = 'Arial'

        if paragrafo_final in par.text and not ja_inserido:
            doc_final.add_paragraph()

            for texto, alinhamento, runs in conteudo_aberto:
                novo_par_aberto = doc_final.add_paragraph()
                novo_par_aberto.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                for texto_run, bold, italic, size, font_name in runs:
                    run_novo = novo_par_aberto.add_run(texto_run)
                    run_novo.bold = bold
                    run_novo.italic = italic
                    run_novo.font.size = Pt(12)
                    run_novo.font.name = 'Arial'

                doc_final.add_paragraph()

            ja_inserido = True

    doc_final.save(output_docx_path)
    logger.info(f"✅ Documento final salvo em: {output_docx_path}")

# Função otimizada para unificar DXFs
def unificar_dxf(dxf_aberto_path, dxf_fechado_path, output_dxf_unificado):
    dxf_aberto = ezdxf.readfile(dxf_aberto_path)
    dxf_fechado = ezdxf.readfile(dxf_fechado_path)

    dxf_unificado = ezdxf.new(dxfversion='R2010')
    importer_aberto = Importer(dxf_aberto, dxf_unificado)
    importer_aberto.import_modelspace()
    importer_aberto.finalize()

    importer_fechado = Importer(dxf_fechado, dxf_unificado)
    importer_fechado.import_modelspace()
    importer_fechado.finalize()

    auditor = dxf_unificado.audit()
    if auditor.errors:
        logger.warning(f"⚠️ Problemas corrigidos pelo auditor DXF: {auditor.errors}")

    dxf_unificado.saveas(output_dxf_unificado)
    logger.info(f"✅ DXF unificado salvo em: {output_dxf_unificado}")

# Função principal final otimizada com UUID (🚨 Código corrigido aqui!)
def main_unir_poligonais(diretorio_concluido):
    tipos = ['ETE', 'REM', 'SER', 'ACE']

    for tipo in tipos:
        logger.info(f"🚩 Tipo atual: {tipo}")

        doc_aberto = glob.glob(os.path.join(diretorio_concluido, f"*ABERTA*{tipo}*.docx"))
        doc_fechado = glob.glob(os.path.join(diretorio_concluido, f"*FECHADA*{tipo}*.docx"))
        dxf_aberto = glob.glob(os.path.join(diretorio_concluido, f"*ABERTA*{tipo}*.dxf"))
        dxf_fechado = glob.glob(os.path.join(diretorio_concluido, f"*FECHADA*{tipo}*.dxf"))

        logger.info(f"📂 Encontrados - DOC aberto: {doc_aberto}, DOC fechado: {doc_fechado}, DXF aberto: {dxf_aberto}, DXF fechado: {dxf_fechado}")

        if not (doc_aberto and doc_fechado and dxf_aberto and dxf_fechado):
            logger.warning(f"⚠️ Arquivos incompletos para tipo {tipo}. Pulando...")
            continue

        doc_fechado_path = doc_fechado[0]
        nome_base = os.path.splitext(os.path.basename(doc_fechado_path))[0].replace('FECHADA_', '').replace(tipo, '').strip('_')

        output_dxf_path = os.path.join(diretorio_concluido, f"{tipo}_{nome_base}_FINAL.dxf")
        output_docx_path = os.path.join(diretorio_concluido, f"{tipo}_{nome_base}_FINAL.docx")

        logger.info(f"✅ Preparando criação de arquivos finais: DXF={output_dxf_path}, DOCX={output_docx_path}")

        conteudo_aberto = extrair_conteudo_docx(doc_aberto[0])

        inserir_conteudo_entre_paragrafos_com_template(
            TEMPLATE_PADRAO,
            doc_fechado[0],
            conteudo_aberto,
            output_docx_path,
            "Pontos definidos pelas Coordenadas Planas no Sistema U.T.M. – SIRGAS 2000.",
            "Pontos definidos pelas Coordenadas Planas no Sistema U.T.M. – SIRGAS 2000."
        )

        logger.info(f"✅ Documento DOCX criado em: {output_docx_path}")

        unificar_dxf(dxf_aberto[0], dxf_fechado[0], output_dxf_path)

        logger.info(f"✅ Arquivo DXF criado em: {output_dxf_path}")

    logger.info("✅ Unificação concluída para todos os tipos disponíveis.")


