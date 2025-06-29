import ezdxf
import math
import pandas as pd
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Alignment
from docx.shared import Pt
import os 
from ezdxf.enums import TextEntityAlignment
import gc
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import glob
import locale
from datetime import datetime
from decimal import getcontext
import logging

BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))

getcontext().prec = 28  # Define a precisão para 28 casas decimais

try:
    locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8')  # Para Render (Linux)
except locale.Error:
    try:
        locale.setlocale(locale.LC_TIME, 'Portuguese_Brazil.1252')  # Para Windows
    except locale.Error:
        locale.setlocale(locale.LC_TIME, 'C')  # Fallback universal

# Obter data atual formatada
data_atual = datetime.now().strftime("%d de %B de %Y")

logger = logging.getLogger(__name__)
logger.setLevel(logging.DEBUG)  # Garanta que está em DEBUG

# 🔹 Função para definir a fonte padrão
def set_default_font(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

# 🔹 Função para calcular azimute
def calculate_azimuth(p1, p2):
    delta_x = p2[0] - p1[0]
    delta_y = p2[1] - p1[1]
    azimuth_rad = math.atan2(delta_x, delta_y)
    azimuth_deg = math.degrees(azimuth_rad)
    return azimuth_deg % 360

# 🔹 Função para calcular distância
def calculate_distance(p1, p2):
    return math.hypot(p2[0] - p1[0], p2[1] - p1[1])

# 🔹 Formata coordenadas para três casas decimais
def format_coordinate(value, decimal_places=3):
    return f"{value:.{decimal_places}f}".replace(".", ",")

# 🔹 Função para calcular ponto em linha
def calculate_point_on_line(start, end, distance):
    dx, dy = end[0] - start[0], end[1] - start[1]
    length = math.hypot(dx, dy)
    if length == 0:
        return start
    return (start[0] + (dx / length) * distance, start[1] + (dy / length) * distance)

# 🔹 Função para converter graus decimais para DMS
def convert_to_dms(decimal_degrees):
    degrees = int(decimal_degrees)
    minutes = int(abs(decimal_degrees - degrees) * 60)
    seconds = abs((decimal_degrees - degrees - minutes / 60) * 3600)
    return f"{degrees}° {minutes}' {seconds:.2f}\""


#🔹 Adicionar rótulo de azimute no LAYER 'Azimutes'
def add_azimuth_label(msp, ponto_inicio, ponto_destino):
    try:
        azimuth = calculate_azimuth(ponto_inicio, ponto_destino)
        mid_x = (ponto_inicio[0] + ponto_destino[0]) / 2
        mid_y = (ponto_inicio[1] + ponto_destino[1]) / 2

        msp.add_text(
            f"{azimuth:.2f}°",
            dxfattribs={'height': 1.0, 'layer': 'Azimutes', 'insert': (mid_x, mid_y)}
        )
        logger.info(f"✅ Rótulo de azimute '{convert_to_dms(azimuth)}' adicionado entre {ponto_inicio} e {ponto_destino}")

    except Exception as e:
        logger.error(f"❌ Erro ao adicionar rótulo de azimute: {e}")
        


def add_mtext_distance(msp, ponto_inicio, ponto_destino):
    try:
        distance = calculate_distance(ponto_inicio, ponto_destino)
        distancia_formatada = f"{distance:.2f}".replace(".", ",")

        # Posição intermediária
        mid_x = (ponto_inicio[0] + ponto_destino[0]) / 2
        mid_y = (ponto_inicio[1] + ponto_destino[1]) / 2
        dx = ponto_destino[0] - ponto_inicio[0]
        dy = ponto_destino[1] - ponto_inicio[1]
        length = math.hypot(dx, dy)
        angle = math.degrees(math.atan2(dy, dx))

        if angle < -90 or angle > 90:
            angle += 180

        offset = 1.5
        perp_x = -dy / length * offset
        perp_y = dx / length * offset
        displaced_mid_point = (mid_x + perp_x, mid_y + perp_y)

        msp.add_text(
            distancia_formatada,
            dxfattribs={
                "height": 1.2,
                "layer": "Distancias",
                "insert": displaced_mid_point,
                "rotation": angle
            }
        )

        logger.info(f"✅ Texto simples da distância '{distancia_formatada}' adicionado com sucesso.")
    except Exception as e:
        logger.error(f"❌ Erro ao adicionar distância como texto: {e}")





def add_dimension(msp, ponto_inicio, ponto_destino):
    """
    Adiciona uma cota (DIMENSION) entre dois pontos no DXF, garantindo alinhamento correto.
    
    :param msp: ModelSpace do DXF
    :param ponto_inicio: Coordenadas (x, y) do ponto inicial
    :param ponto_destino: Coordenadas (x, y) do ponto final
    """
    try:
        # Criar uma cota alinhada (Aligned Dimension)
        dim = msp.add_aligned_dim(
            p1=ponto_inicio, 
            p2=ponto_destino,
            distance=1.5,  # Distância do texto em relação à linha
            dxfattribs={'layer': 'Distancias'}
        )
        
        # Atualizar para que apareça corretamente no DXF
        dim.render()

        logger.info(f"✅ Cota de dimensão adicionada entre {ponto_inicio} e {ponto_destino}")

    except Exception as e:
        logger.error(f"❌ Erro ao adicionar cota de dimensão: {e}")

def add_north_line(msp, current_vertex, length=2):
    """
    Desenha uma linha vertical apontando para o norte a partir do vértice atual.
    
    :param msp: ModelSpace do DXF
    :param current_vertex: Coordenadas (x, y) do ponto de origem
    :param length: Comprimento da linha para o norte (padrão = 2 unidades)
    """
    try:
        # 🔹 Calcular o ponto para cima (Norte)
        north_point = (current_vertex[0], current_vertex[1] + length)

        # 🔹 Adicionar a linha no ModelSpace
        msp.add_line(start=current_vertex, end=north_point, dxfattribs={'layer': 'Norte'})
        logger.info(f"✅ Linha para o norte desenhada de {current_vertex} para {north_point}")

    except Exception as e:
        logger.error(f"❌ Erro ao adicionar linha para o norte: {e}")


def add_azimuth_and_distance(msp, ponto_inicio, ponto_destino):
    """
    Desenha uma linha entre dois pontos, mas NÃO adiciona rótulo para evitar duplicação.
    """
    try:
        azimuth = calculate_azimuth(ponto_inicio, ponto_destino)

        # 🔹 Apenas desenhar a linha, SEM adicionar texto!
        msp.add_line(start=ponto_inicio, end=ponto_destino, dxfattribs={'layer': 'Amarrações'})

        logger.info(f"✅ Linha adicionada entre {ponto_inicio} e {ponto_destino}")

    except Exception as e:
        logger.error(f"❌ Erro na função add_azimuth_and_distance: {e}")





# 🔹 Função para adicionar arco de azimute
def add_azimuth_arc_aberta(msp, current_vertex, next_vertex):
    """
    Desenha um arco representando o azimute e adiciona uma linha para o norte no vértice.
    """
    try:
        # 🔹 Calcular azimute entre os pontos
        azimuth = calculate_azimuth(current_vertex, next_vertex)

        # 🔹 Desenhar linha para o norte antes do arco
        add_north_line(msp, current_vertex, length=3)  # Linha de referência para o Norte

        # 🔹 Calcular pontos de referência para o arco
        north_point = (current_vertex[0], current_vertex[1] + 2)
        start_arc = calculate_point_on_line(current_vertex, north_point, 1)
        end_arc = calculate_point_on_line(current_vertex, next_vertex, 1)

        # 🔹 Calcular ângulos do arco
        start_angle = math.degrees(math.atan2(end_arc[1] - current_vertex[1], end_arc[0] - current_vertex[0]))
        end_angle = math.degrees(math.atan2(start_arc[1] - current_vertex[1], start_arc[0] - current_vertex[0]))

        if end_angle < start_angle:
            end_angle += 360

        # 🔹 Desenhar o arco do azimute
        msp.add_arc(
            center=current_vertex,
            radius=1,
            start_angle=start_angle,
            end_angle=end_angle,
            dxfattribs={'layer': 'Azimute'}
        )

        # 🔹 Adicionar rótulo do azimute
        azimuth_text = convert_to_dms(azimuth)
        label_position = (
            current_vertex[0] + 1.5 * math.cos(math.radians(azimuth / 2)),
            current_vertex[1] + 1.5 * math.sin(math.radians(azimuth / 2))
        )
        msp.add_text(azimuth_text, dxfattribs={'height': 0.5, 'layer': 'Azimute', 'insert': label_position})

        logger.info(f"✅ Arco do azimute e linha para o norte adicionados no ponto {current_vertex}")

    except Exception as e:
        logger.error(f"❌ Erro na função add_azimuth_arc_aberta: {e}")

# 🔹 Função para criar um arquivo Excel com os pontos da poligonal
def create_excel_from_points(pontos_abertos, ponto_V1, output_excel_path):
    """
    Cria um arquivo Excel com os pontos da poligonal aberta, garantindo que não haja pontos extras.
    Agora adiciona as colunas 'Segmento' e 'Azimute' com base no desenho geométrico.
    """
    pontos_abertos.append(ponto_V1)  # Adiciona V1 como último ponto
    data = []

    for i in range(len(pontos_abertos)):
        coord_E, coord_N = pontos_abertos[i]
        ponto_nome = f"P{i + 1}" if i < len(pontos_abertos) - 1 else "V1"

        segmento = ""
        azimute = ""

        if i < len(pontos_abertos) - 1:
            next_coord_E, next_coord_N = pontos_abertos[i + 1]
            next_ponto_nome = f"P{i + 2}" if i < len(pontos_abertos) - 2 else "V1"

            # Define segmento
            segmento = f"{ponto_nome}-{next_ponto_nome}"

            # Calcula azimute e formata para DMS
            azimuth_deg = calculate_azimuth((coord_E, coord_N), (next_coord_E, next_coord_N))
            azimute = convert_to_dms(azimuth_deg)
            distancia_metros = calculate_distance((coord_E, coord_N), (next_coord_E, next_coord_N))
            distancia_str = f"{distancia_metros:.2f}".replace(".", ",")


        # Se for o último ponto (V1), não há distância
        if i == len(pontos_abertos) - 1:
            data.append([ponto_nome, format_coordinate(coord_E, 3), format_coordinate(coord_N, 3), "", "", ""])
        else:
            data.append([ponto_nome, format_coordinate(coord_E, 3), format_coordinate(coord_N, 3), segmento, azimute, distancia_str])



    # Criar DataFrame com as novas colunas
    df = pd.DataFrame(data, columns=["Ponto", "Coord_E", "Coord_N", "Segmento", "Azimute", "Distancia"])


    # Criar planilha Excel
    wb = Workbook()
    ws = wb.active
    ws.title = "Pontos"

    # Adicionar cabeçalhos
    ws.append(["Ponto", "Coord_E", "Coord_N", "Segmento", "Azimute", "Distancia"])


    # Adicionar dados ao Excel
    for row in df.itertuples(index=False):
        ws.append(row)

    # Ajustar a largura das colunas para melhor legibilidade
    for col in ws.columns:
        col_letter = col[0].column_letter
        ws.column_dimensions[col_letter].width = 30
        for cell in col:
            cell.alignment = Alignment(horizontal="center", vertical="center")

    # Salvar o Excel corrigido
    wb.save(output_excel_path)
    wb.close()  # <- Fecha corretamente o arquivo para liberar o uso

    logger.info(f"✅ Arquivo Excel salvo corretamente em: {output_excel_path}")


# 🔹 Gera o documento DOCX com memorial descritivo
def create_memorial_document_aberta(excel_file_path, confrontantes_file_path, desc_ponto_P1, output_doc_path):
    """
    Gera o documento DOCX baseado nos pontos da poligonal aberta (incluindo V1 no Excel).
    """
    try:
        #logger.info(f"📂 Lendo confrontantes do arquivo: {confrontantes_file_path}")
        #xls_confrontantes = pd.ExcelFile(confrontantes_file_path)
        #logger.info(f"📂 Abas disponíveis no arquivo de confrontantes: {xls_confrontantes.sheet_names}")

        # 🔹 Lendo a planilha corrigida
        df = pd.read_excel(excel_file_path, usecols=["Ponto", "Coord_E", "Coord_N", "Segmento", "Azimute"])

        # 📌 Detectar a primeira aba automaticamente
        #xls = pd.ExcelFile(excel_file_path)
        #primeira_aba = xls.sheet_names[0]
        #logger.info(f"📂 Usando a aba: {primeira_aba}")
        # 📌 Detectar automaticamente o nome da primeira aba
        primeira_aba = pd.ExcelFile(confrontantes_file_path).sheet_names[0]
        # 📌 Carregar os confrontantes corretamente
        # Carregar os confrontantes usando a aba detectada
        df_confrontantes = pd.read_excel(confrontantes_file_path, sheet_name=primeira_aba, usecols=["Código", "Confrontante"])

        # 🔹 Limpeza dos dados
        df_confrontantes.columns = df_confrontantes.columns.str.strip()
        df_confrontantes["Código"] = df_confrontantes["Código"].astype(str).str.strip()

        df["Azimute"] = df["Azimute"].fillna("")  

        # Criando o documento
        doc_word = Document()
        set_default_font(doc_word)
        #doc_word.add_paragraph("MEMORIAL DESCRITIVO DA POLIGONAL ABERTA\n", style='Heading 1')

        # 🔹 Primeiro ponto (P1)
        nome_p1 = df.iloc[0]["Ponto"]
        coord_e_p1 = df.iloc[0]["Coord_E"]
        coord_n_p1 = df.iloc[0]["Coord_N"]

        p1_paragraph = doc_word.add_paragraph(style='Normal')
        p1_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p1_paragraph.add_run("O Ponto ")
        run_p1 = p1_paragraph.add_run(nome_p1)
        run_p1.bold = True
        p1_paragraph.add_run(
            f", ponto de amarração inicial, está localizado na {desc_ponto_P1}, "
            f"nas coordenadas N(Y) {coord_n_p1} e E(X) {coord_e_p1} ."
        )

        # 🔹 Processar pontos
        for i in range(len(df) - 1):
            row_atual = df.iloc[i]
            row_proximo = df.iloc[i + 1]

            nome_atual = row_atual["Ponto"]
            coord_e_atual = row_atual["Coord_E"]
            coord_n_atual = row_atual["Coord_N"]
            azimuth = row_atual["Azimute"]

            nome_proximo = row_proximo["Ponto"]
            coord_e_proximo = row_proximo["Coord_E"]
            coord_n_proximo = row_proximo["Coord_N"]

            confrontante = df_confrontantes.loc[
                df_confrontantes["Código"].str.strip() == nome_atual.strip(), 
                "Confrontante"
            ].values
            confrontante = confrontante[0] if len(confrontante) > 0 else "Desconhecido"

            distance = calculate_distance(
                (float(str(coord_e_atual).replace(",", ".")), float(str(coord_n_atual).replace(",", "."))),
                (float(str(coord_e_proximo).replace(",", ".")), float(str(coord_n_proximo).replace(",", ".")))
            )

            p = doc_word.add_paragraph(style='Normal')
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            p.add_run(
                f"Daí, com Azimute de {azimuth} e distância de {format_coordinate(distance, 2)} metros, "
                f"confrontando com área de {confrontante}, chega-se ao ponto "
            )
            run_ponto = p.add_run(nome_proximo)
            run_ponto.bold = True
            p.add_run(
                f" nas coordenadas N(Y) {coord_n_proximo} e E(X) {coord_e_proximo} ."
            )

       

        doc_word.save(output_doc_path)
        logger.info(f"✅ Memorial descritivo salvo corretamente em: {output_doc_path}")

    except Exception as e:
        logger.error(f"❌ Erro ao criar o documento memorial: {e}")

    
# 🔹 Processa a poligonal e gera os 3 arquivos
def process_poligonal_aberta_e_fechada(dxf_file_path, output_file_path, output_excel_path, confrontantes_file_path, output_doc_path, desc_ponto_P1):


    """
    Processa os pontos da poligonal aberta, cria o Excel e o DOCX sem pontos intermediários.
    Garante que a poligonal aberta se conecte corretamente ao primeiro vértice da poligonal fechada (V1).
    """
    try:
        # Ler o arquivo DXF
        doc = ezdxf.readfile(dxf_file_path)
        msp = doc.modelspace()

        # 🔹 Coletar apenas os pontos da poligonal aberta
        pontos_abertos = []
        for entity in msp.query('POINT'):
            pontos_abertos.append((entity.dxf.location.x, entity.dxf.location.y))

        # 🔹 Coletar também textos como pontos
        for entity in msp.query('TEXT'):
            pontos_abertos.append((entity.dxf.insert.x, entity.dxf.insert.y))

        # 🔹 Remover pontos duplicados e manter apenas os primeiros 4 pontos
        pontos_abertos = list(dict.fromkeys(pontos_abertos))  # Mantém todos os pontos únicos


        if len(pontos_abertos) < 2:
            logger.info("❌ Erro: Menos de 2 pontos encontrados na poligonal aberta.")
            return

       
        ponto_V1 = None
        for entity in msp.query('LWPOLYLINE'):
            if entity.closed:
                pontos_fechados = list(entity.get_points('xy'))
        
                # 🔸 Remover vértice duplicado do final se for igual ao primeiro
                if pontos_fechados[0] == pontos_fechados[-1]:
                    pontos_fechados.pop()
                    logger.info("⚠️ Vértice duplicado no final da poligonal fechada removido com sucesso.")
        
                # Agora, garantimos que ponto_V1 seja válido e único
                ponto_V1 = pontos_fechados[0][:2]
                break
        if not ponto_V1:
            logger.info("❌ Erro: Nenhum ponto V1 encontrado para conexão.")
            return

        # 🔹 Criar o Excel corretamente
        create_excel_from_points(pontos_abertos, ponto_V1, output_excel_path)
        # 🔹 Adicionar círculos e rótulos para os vértices no LAYER 'Vertices'
        if "Vertices" not in doc.layers:
            doc.layers.new(name="Vertices", dxfattribs={"color": 3})  # Criar camada 'Vertices' se não existir

        #ordered_points = pontos_abertos + [ponto_V1]  # Garantir que V1 seja incluído
        ordered_points = pontos_abertos   # Não inclui o V1 seja incluído
        for i, vertex in enumerate(ordered_points):
            msp.add_circle(center=vertex, radius=0.5, dxfattribs={"layer": "Vertices"})  # Ajuste o raio conforme necessário



        # 🔹 Adicionar rótulos de distância e azimute entre os pontos da poligonal aberta
        for i in range(len(pontos_abertos) - 1):
            add_azimuth_and_distance(msp, pontos_abertos[i], pontos_abertos[i + 1])
            add_azimuth_arc_aberta(msp, pontos_abertos[i], pontos_abertos[i + 1])

            distancia = calculate_distance(pontos_abertos[i], pontos_abertos[i + 1])

            if distancia > 0:
                add_mtext_distance(msp, pontos_abertos[i], pontos_abertos[i + 1])  # ✅ Garante que só é chamado uma vez

        # 🔹 Garantir que o último trecho (P último → V1) também receba o rótulo alinhado
        if ponto_V1 not in pontos_abertos:  
            distancia_v1 = calculate_distance(pontos_abertos[-1], ponto_V1)  # Último segmento real

            if distancia_v1 > 0:
                add_azimuth_and_distance(msp, pontos_abertos[-1], ponto_V1)
                add_azimuth_arc_aberta(msp, pontos_abertos[-1], ponto_V1)
                #add_mtext_distance(msp, pontos_abertos[-1], ponto_V1)  # ✅ Apenas uma vez
        # 🔹 Salvar o arquivo DXF corrigido
        doc.saveas(output_file_path)
        logger.info(f"✅ Arquivo DXF atualizado salvo corretamente em: {output_file_path}")
        

        # Criar o documento DOCX corretamente
        logger.info(f"📂 Verificando confrontantes_file_path antes de chamar a função: {confrontantes_file_path}")
        create_memorial_document_aberta(output_excel_path, confrontantes_file_path, desc_ponto_P1, output_doc_path)



    except Exception as e:
        logger.error(f"❌ Erro ao processar o DXF: {e}")
        
# 🔹 Função principal
def main_poligonal_aberta(uuid_str, excel_path, dxf_path, diretorio_preparado, diretorio_concluido):
    logger.info("🔹 Executando poligonal aberta com as variáveis definidas:")
    logger.info(f"Excel: {excel_path}")
    logger.info(f"DXF: {dxf_path}")
    logger.info(f"Preparado: {diretorio_preparado}")
    logger.info(f"Concluído: {diretorio_concluido}")

    dxf_file_path = dxf_path
    output_folder = diretorio_concluido

    diretorio_confrontantes = diretorio_preparado
    # Definindo automaticamente o tipo com base no nome do arquivo DXF
    tipo = ""
    dxf_filename = os.path.basename(dxf_file_path).upper()
    if "ETE" in dxf_filename:
        tipo = "ETE"
    elif "REM" in dxf_filename:
        tipo = "REM"
    elif "SER" in dxf_filename:
        tipo = "SER"
    elif "ACE" in dxf_filename:
        tipo = "ACE"
    else:
        logger.info("❌ Não foi possível determinar automaticamente o tipo (ETE, REM, SER ou ACE).")
        return
    
    # Busca dinâmica do arquivo confrontante correto na pasta PREPARADO
    padrao_aberta = os.path.join(diretorio_preparado, f"{uuid_str}_ABERTA_{tipo}*.xlsx")

    arquivos_encontrados = glob.glob(padrao_aberta)

    if not arquivos_encontrados:
        logger.info(f"❌ Arquivo de confrontantes não encontrado com o padrão: {padrao_aberta}")
        return

    confrontantes_file_path = arquivos_encontrados[0]  # Definido automaticamente!

   
    #confrontantes_file_path = input("Digite o caminho do arquivo de confrontantes (Excel): ").strip('"')
    #desc_ponto_P1 = input("Descreva o ponto P1: ")
    # 🔹 Lê automaticamente a descrição do ponto P1 da planilha confrontantes
    try:
        df_confrontantes = pd.read_excel(confrontantes_file_path, sheet_name=0, usecols=["Confrontante"])
        desc_ponto_P1 = df_confrontantes.iloc[0]["Confrontante"].strip()
    except Exception as e:
        logger.error(f"❌ Erro ao ler descrição automática do ponto P1: {e}")
        return

    # 🔹 Extração do ponto P1 a partir da aba 'Dados_do_Imóvel' (campo 'AZ')
    try:
        df_dados_imovel = pd.read_excel(excel_path, sheet_name='Dados_do_Imóvel', header=None)
        dados_dict = dict(zip(df_dados_imovel.iloc[:, 0], df_dados_imovel.iloc[:, 1]))
        desc_ponto_P1 = dados_dict.get("AZ", "").strip()
        logger.info(f"📌 Descrição do ponto P1 (AZ) extraída da aba Dados_do_Imóvel: {desc_ponto_P1}")
    except Exception as e:
        logger.error(f"❌ Erro ao extrair descrição do ponto P1 (AZ): {e}")
        return
        
    dxf_filename = os.path.splitext(os.path.basename(dxf_file_path))[0]  # Obtém o nome base do DXF

    # 🔹 Alteração aqui, adicionando "ABERTA_" aos nomes dos arquivos:
    output_excel_path = os.path.join(output_folder, f"{uuid_str}_ABERTA_{tipo}_{dxf_filename}.xlsx")
    output_doc_path = os.path.join(output_folder, f"{uuid_str}_ABERTA_{tipo}_{dxf_filename}.docx")
    output_dxf_path = os.path.join(output_folder, f"{uuid_str}_ABERTA_{tipo}_{dxf_filename}.dxf")


    process_poligonal_aberta_e_fechada(
        dxf_file_path, output_dxf_path, output_excel_path, confrontantes_file_path, output_doc_path, desc_ponto_P1
    )
