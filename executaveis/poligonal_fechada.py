import ezdxf
import math
import csv
import os
import re
import glob
import logging
from docx import Document
from docx.shared import Inches
from datetime import datetime
from decimal import Decimal, getcontext
import pandas as pd
import locale
import openpyxl
from openpyxl.styles import Alignment, Font
from docx.shared import Pt
from ezdxf.math import Vec3
from ezdxf.math import bulge_to_arc, ConstructionArc, Vec2
from shapely.geometry import Polygon
import traceback
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import sys
import time
import uuid


BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))

try:
    from ezdxf.math import Vec3 as Vector
except ImportError:
    from ezdxf.math import Vector

getcontext().prec = 28  # Define a precisão para 28 casas decimais

# Correção definitiva do locale para Windows:
try:
    locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8')  # para Render (Linux)
except locale.Error:
    try:
        locale.setlocale(locale.LC_TIME, 'Portuguese_Brazil.1252')  # para Windows
    except locale.Error:
        locale.setlocale(locale.LC_TIME, 'C')  # fallback neutro

# Exemplo da data:
data_atual = datetime.now().strftime("%d de %B de %Y")

def obter_data_em_portugues():
    meses_pt = {
        "January": "janeiro", "February": "fevereiro", "March": "março",
        "April": "abril", "May": "maio", "June": "junho",
        "July": "julho", "August": "agosto", "September": "setembro",
        "October": "outubro", "November": "novembro", "December": "dezembro"
    }
    data = datetime.now()
    mes_en = data.strftime("%B")
    mes_pt = meses_pt.get(mes_en, mes_en)
    return f"{data.day:02d} de {mes_pt} de {data.year}"

# Configuração de logger adicional
logger = logging.getLogger(__name__)
logger.setLevel(logging.DEBUG)

fh = logging.FileHandler(os.path.join(BASE_DIR, 'static', 'logs', f'poligonal_{datetime.now().strftime("%Y%m%d_%H%M%S")}.log'), encoding="utf-8")
fh.setLevel(logging.DEBUG)
formatter = logging.Formatter('%(asctime)s [%(levelname)s] %(message)s')
fh.setFormatter(formatter)
logger.addHandler(fh)


def limpar_dxf_e_inserir_ponto_az(original_path, saida_path):
    try:
        doc_antigo = ezdxf.readfile(original_path)
        msp_antigo = doc_antigo.modelspace()
        doc_novo = ezdxf.new(dxfversion='R2010')
        msp_novo = doc_novo.modelspace()

        pontos_polilinha = None
        bulges_polilinha = None
        ponto_inicial_real = None

        for entity in msp_antigo.query('LWPOLYLINE'):
            if entity.closed:
                # 🔧 Leitura com verificação de duplicatas
                pontos_polilinha_raw = entity.get_points('xyseb')
                ponto_inicial_real = (float(pontos_polilinha_raw[0][0]), float(pontos_polilinha_raw[0][1]))
                pontos_polilinha = []
                bulges_polilinha = []

                tolerancia = 1e-6  # Tolerância para considerar pontos idênticos

                for pt in pontos_polilinha_raw:
                    x, y, *_, bulge = pt
                    x, y = float(x), float(y)
                    if not pontos_polilinha:
                        # Primeiro ponto
                        pontos_polilinha.append((x, y))
                        bulges_polilinha.append(bulge)
                    else:
                        x_ant, y_ant = pontos_polilinha[-1]
                        if math.hypot(x - x_ant, y - y_ant) > tolerancia:
                            pontos_polilinha.append((x, y))
                            bulges_polilinha.append(bulge)
                        else:
                            print(f"⚠️ Ponto duplicado consecutivo removido: {(x, y)}")

                # 🔍 Verificação extra para ponto final duplicado
                if len(pontos_polilinha) > 2 and math.hypot(
                    pontos_polilinha[0][0] - pontos_polilinha[-1][0],
                    pontos_polilinha[0][1] - pontos_polilinha[-1][1]
                ) < tolerancia:
                    print("⚠️ Último ponto é igual ao primeiro — removendo ponto final duplicado.")
                    pontos_polilinha.pop()
                    bulges_polilinha.pop()

                # 🔍 Verificação extra para P1 == P2
                if len(pontos_polilinha) > 1 and math.hypot(
                    pontos_polilinha[0][0] - pontos_polilinha[1][0],
                    pontos_polilinha[0][1] - pontos_polilinha[1][1]
                ) < tolerancia:
                    print("⚠️ Primeiro ponto é igual ao segundo — removendo o segundo ponto duplicado.")
                    pontos_polilinha.pop(1)
                    bulges_polilinha.pop(1)

                break

        if pontos_polilinha is None:
            raise ValueError("Nenhuma polilinha fechada encontrada no DXF original.")

        # if calculate_signed_area(pontos_polilinha) < 0:
        #     pontos_polilinha.reverse()
        #     bulges_polilinha.reverse()
        #     bulges_polilinha = [-b for b in bulges_polilinha]

        pontos_com_bulge = [
            (pontos_polilinha[i][0], pontos_polilinha[i][1], bulges_polilinha[i])
            for i in range(len(pontos_polilinha))
        ]

        msp_novo.add_lwpolyline(
            pontos_com_bulge,
            format='xyb',
            close=True,
            dxfattribs={'layer': 'DIVISA_PROJETADA'}
        )

        # Não desenha mais o ponto Az, mas retorna as coordenadas de V1 como ponto_az válido
        ponto_az = pontos_polilinha[0]

        doc_novo.saveas(saida_path)
        print(f"✅ DXF limpo salvo em: {saida_path}")
        
        return saida_path, ponto_az, ponto_inicial_real

    except Exception as e:
        print(f"❌ Erro ao limpar DXF: {e}")
        return original_path, None, None


def calculate_signed_area(points):
    area = 0
    for i in range(len(points)):
        x1, y1 = points[i]
        x2, y2 = points[(i + 1) % len(points)]
        area += (x1 * y2) - (x2 * y1)
    return area / 2


def calculate_area_with_arcs(points, arcs):
    """
    Calcula a área de uma poligonal que inclui segmentos de linha reta e arcos.
    """
    num_points = len(points)
    area_linear = 0.0
    for i in range(num_points):
        x1, y1 = points[i]
        x2, y2 = points[(i + 1) % num_points]
        area_linear += x1 * y2 - x2 * y1
    area_linear = abs(area_linear) / 2.0

    area_arcs = 0.0
    for arc in arcs:
        start_point = (float(arc['start_point'][0]), float(arc['start_point'][1]))
        end_point = (float(arc['end_point'][0]), float(arc['end_point'][1]))
        center = (float(arc['center'][0]), float(arc['center'][1]))
        radius = arc['radius']
        start_angle = math.radians(arc['start_angle'])
        end_angle = math.radians(arc['end_angle'])

        delta_angle = end_angle - start_angle
        if delta_angle <= 0:
            delta_angle += 2 * math.pi

        sector_area = 0.5 * radius**2 * delta_angle
        triangle_area = 0.5 * abs((start_point[0] - center[0]) * (end_point[1] - center[1]) -
                                  (end_point[0] - center[0]) * (start_point[1] - center[1]))
        segment_area = sector_area - triangle_area

        if delta_angle > math.pi:
            area_arcs -= segment_area
        else:
            area_arcs += segment_area

    area_total = area_linear + area_arcs
    return area_total


def bulge_to_arc_length(start_point, end_point, bulge):
    dx = end_point[0] - start_point[0]
    dy = end_point[1] - start_point[1]
    chord_length = math.hypot(dx, dy)

    sagitta = (bulge * chord_length) / 2
    radius = ((chord_length / 2) ** 2 + sagitta ** 2) / (2 * abs(sagitta))
    angle = 4 * math.atan(abs(bulge))
    arc_length = radius * angle
    return arc_length, radius, angle





def get_document_info_from_dxf(dxf_file_path):
    """
    Lê a primeira LWPOLYLINE FECHADA do DXF e retorna:
      - doc (ezdxf document)
      - lines: list[ (start_xy, end_xy) ]
      - arcs:  list[ {start_point, end_point, center, radius, start_angle, end_angle, length, bulge} ]
      - perimeter_dxf (float)
      - area_dxf (float)
    """
    try:
        import ezdxf
        import math
        import traceback
        from shapely.geometry import Polygon

        doc = ezdxf.readfile(dxf_file_path)
        msp = doc.modelspace()

        lines = []
        arcs = []
        perimeter_dxf = 0.0
        area_dxf = 0.0

        found = False

        for entity in msp.query('LWPOLYLINE'):
            # Aceita apenas a primeira polilinha FECHADA (comportamento original)
            if not getattr(entity, "closed", False):
                continue

            # Tenta obter pontos detalhados; se não der, cai para 'xyb'
            try:
                polyline_points = entity.get_points('xyseb')  # x, y, start_width, end_width, bulge
            except TypeError:
                polyline_points = entity.get_points('xyb')    # x, y, bulge

            num_points = len(polyline_points)
            if num_points < 2:
                continue

            boundary_points = []

            for i in range(num_points):
                # Desempacota ponta A
                ptA = polyline_points[i]
                if len(ptA) == 5:
                    x_start, y_start, _sA, _eA, bA = ptA
                elif len(ptA) == 3:
                    x_start, y_start, bA = ptA
                else:
                    # formato inesperado
                    continue

                # Desempacota ponta B (i+1, módulo num_points)
                ptB = polyline_points[(i + 1) % num_points]
                if len(ptB) == 5:
                    x_end, y_end, _sB, _eB, _bB_unused = ptB
                elif len(ptB) == 3:
                    x_end, y_end, _bB_unused = ptB
                else:
                    continue

                # Normaliza bulge
                try:
                    bulge = float(bA or 0.0)
                except Exception:
                    bulge = 0.0

                start_point = (float(x_start), float(y_start))
                end_point   = (float(x_end),   float(y_end))

                dx = end_point[0] - start_point[0]
                dy = end_point[1] - start_point[1]
                chord_length = math.hypot(dx, dy)

                # Ignora segmento degenerado (corda ~0), mas preserva contorno
                if chord_length < 1e-9:
                    boundary_points.append((start_point[0], start_point[1]))
                    continue

                if abs(bulge) > 1e-12:
                    # Relação bulge ↔ ângulo central: bulge = tan(delta/4)
                    delta = 4.0 * math.atan(bulge)
                    sin_half = math.sin(delta / 2.0)

                    # Se ângulo muito pequeno, trata como linha
                    if abs(sin_half) < 1e-12:
                        lines.append((start_point, end_point))
                        boundary_points.append((start_point[0], start_point[1]))
                        perimeter_dxf += chord_length
                        continue

                    # Raio do arco
                    radius = chord_length / (2.0 * abs(sin_half))
                    arc_length = abs(radius * delta)

                    # Distância do centro ao meio da corda: d = sqrt(R^2 - (c/2)^2)
                    half_c = chord_length / 2.0
                    sq = radius * radius - half_c * half_c
                    if sq < 0.0:
                        sq = 0.0  # estabiliza numérico
                    offset_dist = math.sqrt(sq)

                    # Vetor perpendicular normalizado à corda
                    nx, ny = -dy / chord_length, dx / chord_length
                    # Direção do centro conforme sinal do bulge
                    if bulge < 0:
                        nx, ny = -nx, -ny

                    # Centro do círculo
                    midx = (start_point[0] + end_point[0]) / 2.0
                    midy = (start_point[1] + end_point[1]) / 2.0
                    cx = midx + nx * offset_dist
                    cy = midy + ny * offset_dist

                    # Ângulos
                    start_angle = math.atan2(start_point[1] - cy, start_point[0] - cx)
                    end_angle   = start_angle + delta

                    # Armazena arco
                    arcs.append({
                        'start_point': (start_point[0], start_point[1]),
                        'end_point':   (end_point[0],   end_point[1]),
                        'center':      (cx, cy),
                        'radius':      radius,
                        'start_angle': math.degrees(start_angle),
                        'end_angle':   math.degrees(end_angle),
                        'length':      arc_length,
                        'bulge':       bulge,
                    })

                    # Densifica contorno para perímetro/área
                    num_arc_points = 64
                    for t in range(num_arc_points):
                        # t ∈ [0,1) para evitar repetir endpoint; o próximo segmento entra com seu ponto inicial
                        frac = t / float(num_arc_points)
                        ang = start_angle + delta * frac
                        px = cx + radius * math.cos(ang)
                        py = cy + radius * math.sin(ang)
                        boundary_points.append((px, py))

                    perimeter_dxf += arc_length

                else:
                    # Linha reta
                    lines.append((start_point, end_point))
                    boundary_points.append((start_point[0], start_point[1]))
                    perimeter_dxf += chord_length

            # (Opcional) fecha o contorno explicitamente
            if boundary_points and boundary_points[0] != boundary_points[-1]:
                boundary_points.append(boundary_points[0])

            # Área exata do desenho
            if len(boundary_points) >= 3:
                polygon = Polygon(boundary_points)
                area_dxf = polygon.area
            else:
                area_dxf = 0.0

            found = True
            break  # mantém comportamento original: processa a primeira polilinha fechada

        if not found and not lines and not arcs:
            print("Nenhuma polilinha fechada encontrada no arquivo DXF.")
            return None, [], [], 0, 0

        print(f"Linhas processadas: {len(lines)}")
        print(f"Arcos processados:  {len(arcs)}")
        print(f"Perímetro do DXF:   {perimeter_dxf:.2f} m")
        print(f"Área do DXF:        {area_dxf:.2f} m²")

        return doc, lines, arcs, perimeter_dxf, area_dxf

    except Exception as e:
        print(f"Erro ao obter informações do documento: {e}")
        import traceback
        traceback.print_exc()
        return None, [], [], 0, 0





# 🔹 Função para definir a fonte padrão
def set_default_font(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)


def add_arc_labels(doc, msp, start_point, end_point, radius, length, label):
    try:
        mid_point = ((float(start_point[0]) + float(end_point[0]))/2, (float(start_point[1]) + float(end_point[1]))/2)

        label_radius = f"R={radius:.2f}".replace('.', ',')
        label_length = f"C={length:.2f}".replace('.', ',')

        msp.add_text(
            label_radius,
            dxfattribs={
                "height": 1.0,
                "layer": "LAYOUT_DISTANCIAS",
                "insert": (mid_point[0], mid_point[1])
            }
        )

        msp.add_text(
            label_length,
            dxfattribs={
                "height": 1.0,
                "layer": "LAYOUT_DISTANCIAS",
                "insert": (mid_point[0], mid_point[1] - 2)
            }
        )

        print(f"✅ Rótulos {label_radius} e {label_length} adicionados corretamente no DXF.")

    except Exception as e:
        print(f"❌ Erro ao adicionar rótulos dos arcos: {e}")


def calculate_point_on_line(start, end, distance):
    dx, dy = end[0] - start[0], end[1] - start[1]
    length = math.hypot(dx, dy)
    if length == 0:
        raise ValueError("Ponto inicial e final são iguais, não é possível calcular um ponto na linha.")
    return (
        start[0] + (dx / length) * distance,
        start[1] + (dy / length) * distance
    )


def calculate_azimuth(p1, p2):
    delta_x = p2[0] - p1[0]
    delta_y = p2[1] - p1[1]
    azimuth_rad = math.atan2(delta_x, delta_y)
    azimuth_deg = math.degrees(azimuth_rad)
    if azimuth_deg < 0:
        azimuth_deg += 360
    return azimuth_deg


def calculate_distance(point1, point2):
    dx = point2[0] - point1[0]
    dy = point2[1] - point1[1]
    return math.sqrt(dx**2 + dy**2)


def calculate_azimuth_and_distance(start_point, end_point):
    dx = end_point[0] - start_point[0]
    dy = end_point[1] - start_point[1]
    distance = math.hypot(dx, dy)
    azimuth = math.degrees(math.atan2(dx, dy))
    if azimuth < 0:
        azimuth += 360
    return azimuth, distance


def add_azimuth_arc(doc, msp, ponto_az, v1, azimuth):
    try:
        if 'LAYOUT_AZIMUTES' not in doc.layers:
            doc.layers.new(name='LAYOUT_AZIMUTES', dxfattribs={'color': 5})

        msp.add_line(start=ponto_az, end=v1, dxfattribs={'layer': 'LAYOUT_AZIMUTES'})

        azimuth_label = f"Azimute = {convert_to_dms(azimuth)}"
        label_position = (
            ponto_az[0] + 1.5 * math.cos(math.radians(azimuth / 2)),
            ponto_az[1] + 1.5 * math.sin(math.radians(azimuth / 2))
        )
        msp.add_text(
            azimuth_label,
            dxfattribs={'height': 0.5, 'layer': 'LAYOUT_AZIMUTES', 'insert': label_position}
        )

        print(f"Rótulo do azimute ({azimuth_label}) adicionado com sucesso em {label_position}")

    except Exception as e:
        print(f"Erro ao adicionar arco do azimute: {e}")


def convert_to_dms(decimal_degrees):
    degrees = int(decimal_degrees)
    minutes = int(abs(decimal_degrees - degrees) * 60)
    seconds = abs((decimal_degrees - degrees - minutes / 60) * 3600)
    return f"{degrees}° {minutes}' {seconds:.2f}\""


def calculate_polygon_area(points):
    n = len(points)
    area = 0.0
    for i in range(n):
        x1, y1 = points[i][0], points[i][1]
        x2, y2 = points[(i + 1) % n][0], points[(i + 1) % n][1]
        area += x1 * y2 - x2 * y1
    return area / 2.0


def add_label_and_distance(doc, msp, start_point, end_point, label, distance):
    import math

    # Lê defaults globais se existirem (senão usa estes)
    H_TXT_VERT = globals().get('H_TXT_VERT', 0.50)
    H_TXT_DIST = globals().get('H_TXT_DIST', 0.30)
    R_CIRCLE   = globals().get('R_CIRCLE', 0.30)

    try:
        start_point = (float(start_point[0]), float(start_point[1]))
        end_point   = (float(end_point[0]),   float(end_point[1]))

        # Camadas
        layers = [
            ("LAYOUT_VERTICES", 2),   # Amarelo
            ("LAYOUT_DISTANCIAS", 4)  # Azul
        ]
        for layer_name, color in layers:
            if layer_name not in doc.layers:
                doc.layers.new(name=layer_name, dxfattribs={"color": color})

        # Círculo + rótulo do vértice (sempre no start)
        msp.add_circle(center=start_point, radius=R_CIRCLE, dxfattribs={'layer': 'LAYOUT_VERTICES'})
        msp.add_text(
            label,
            dxfattribs={
                'height': H_TXT_VERT,
                'layer': 'LAYOUT_VERTICES',
                'insert': (start_point[0] + H_TXT_VERT, start_point[1] + H_TXT_VERT),
            }
        )

        # Distância posicionada pela corda (bom para LINHAS)
        mid_x = (start_point[0] + end_point[0]) / 2
        mid_y = (start_point[1] + end_point[1]) / 2
        dx, dy = end_point[0] - start_point[0], end_point[1] - start_point[1]
        length = (dx ** 2 + dy ** 2) ** 0.5
        if length == 0:
            return

        angle = math.degrees(math.atan2(dy, dx))
        if angle < -90 or angle > 90:
            angle += 180

        # leve deslocamento ortogonal
        offset = H_TXT_DIST
        perp_x, perp_y = (-dy / length) * offset, (dx / length) * offset
        mid_point_displaced = (mid_x + perp_x, mid_y + perp_y)

        distancia_formatada = f"{distance:.2f} ".replace('.', ',')
        msp.add_text(
            distancia_formatada,
            dxfattribs={
                'height': H_TXT_DIST,
                'layer': 'LAYOUT_DISTANCIAS',
                'rotation': angle,
                'insert': mid_point_displaced,
            }
        )

        print(f"✅ DEBUG: '{label}' e distância '{distancia_formatada}' inseridos em {start_point} e {mid_point_displaced} com ângulo {angle:.2f}°")

    except Exception as e:
        print(f"❌ ERRO GRAVE ao adicionar rótulo '{label}' e distância: {e}")





def sanitize_filename(filename):
    sanitized_filename = re.sub(r'[\\/*?:"<>|]', "_", filename)
    return sanitized_filename



    # --- Fallback simples para anotar um segmento quando o método nativo falhar ---
def _fallback_anotar_segmento(
    msp, start_point, end_point, label, distance,
    H_TXT_VERT, H_TXT_DIST, R_CIRCLE, logger,
    is_arc=False, arc_radius=None, arc_bulge=None, **__
):
    import math
    try:
        start = (float(start_point[0]), float(start_point[1]))
        end   = (float(end_point[0]),   float(end_point[1]))

        doc = msp.doc
        for layer_name, color in (("LAYOUT_VERTICES", 2), ("LAYOUT_DISTANCIAS", 4)):
            if layer_name not in doc.layers:
                doc.layers.new(name=layer_name, dxfattribs={"color": color})

        # vértice
        msp.add_circle(center=start, radius=R_CIRCLE, dxfattribs={'layer': 'LAYOUT_VERTICES'})
        msp.add_text(label, dxfattribs={
            'height': H_TXT_VERT, 'layer': 'LAYOUT_VERTICES',
            'insert': (start[0] + H_TXT_VERT, start[1] + H_TXT_VERT)
        })

        dist_txt = f"{distance:.2f} ".replace('.', ',')
        dx, dy = end[0]-start[0], end[1]-start[1]
        c = math.hypot(dx, dy)

        # LINHA (ou arco sem bulge): usa meio da corda
        if (not is_arc) or (arc_bulge is None) or (abs(float(arc_bulge)) < 1e-12) or (c < 1e-12):
            if c < 1e-12: return
            ang = math.degrees(math.atan2(dy, dx))
            if ang < -90 or ang > 90: ang += 180
            mid = ((start[0]+end[0])/2.0, (start[1]+end[1])/2.0)
            pos = (mid[0] + (-dy/c)*H_TXT_DIST, mid[1] + (dx/c)*H_TXT_DIST)
            msp.add_text(dist_txt, dxfattribs={
                'height': H_TXT_DIST, 'layer': 'LAYOUT_DISTANCIAS',
                'rotation': ang, 'insert': pos
            })
            return

        # ARCO: centro → ponto médio do arco → tangente
        b = float(arc_bulge)
        theta = 4.0 * math.atan(abs(b))
        R = float(arc_radius) if (arc_radius and arc_radius > 0) else (c/2.0)/max(1e-12, math.sin(theta/2.0))
        nx, ny = -dy/c, dx/c
        if b < 0: nx, ny = -nx, -ny
        offset = math.sqrt(max(0.0, R*R - (c/2.0)*(c/2.0)))
        midx, midy = (start[0]+end[0])/2.0, (start[1]+end[1])/2.0
        cx, cy = midx + nx*offset, midy + ny*offset

        ang_start = math.atan2(start[1]-cy, start[0]-cx)
        sentido = 1.0 if b > 0 else -1.0
        ang_mid = ang_start + sentido*(theta/2.0)
        px, py = cx + R*math.cos(ang_mid), cy + R*math.sin(ang_mid)

        ang_tan = ang_mid + sentido*(math.pi/2.0)
        ang_deg = math.degrees(ang_tan)
        if ang_deg < -90 or ang_deg > 90: ang_deg += 180

        msp.add_text(dist_txt, dxfattribs={
            'height': H_TXT_DIST, 'layer': 'LAYOUT_DISTANCIAS',
            'rotation': ang_deg, 'insert': (px, py)
        })

    except Exception as e:
        (logger.warning if logger else print)(f"[fallback] Falha ao anotar segmento: {e}")





def create_memorial_descritivo(
    doc, msp, lines, proprietario, matricula, caminho_salvar, arcs=None,
    excel_file_path=None, ponto_az=None, distance_az_v1=None,
    azimute_az_v1=None, ponto_inicial_real=None, tipo=None, uuid_prefix=None,
    diretorio_concluido=None, sentido_poligonal='horario'
):
    """
    Cria o memorial descritivo no DXF e salva planilha Excel.
    Segmentos esperados:
      - lines: list[ (p1, p2) ]
      - arcs : list[ {'start_point','end_point','radius','length','bulge'} ]
    """
    import os, math, logging
    import pandas as pd
    import openpyxl
    from openpyxl.styles import Font, Alignment

    logger = logging.getLogger(__name__)

    # Defaults
    H_TXT_VERT = globals().get('H_TXT_VERT', 1.20)
    H_TXT_DIST = globals().get('H_TXT_DIST', 1.20)
    R_CIRCLE   = globals().get('R_CIRCLE', 0.60)

    # --- Confrontantes -------------------------------------------------------
    if excel_file_path:
        try:
            confrontantes_df = pd.read_excel(excel_file_path)
            if "Código" in confrontantes_df.columns and "Confrontante" in confrontantes_df.columns:
                confrontantes_dict = dict(zip(confrontantes_df['Código'], confrontantes_df['Confrontante']))
            else:
                confrontantes_dict = {}
        except Exception as e:
            logger.warning(f"Erro ao carregar confrontantes: {e}")
            confrontantes_dict = {}
    else:
        confrontantes_dict = {}

    if (not lines) and (not arcs):
        logger.warning("Nenhuma geometria (linhas/arcos) disponível.")
        return None

    # --- Monta elementos unificados -----------------------------------------
    elementos = []
    for p1, p2 in (lines or []):
        elementos.append(('line', (p1, p2)))

    for arc in (arcs or []):
        bulge = float(arc.get('bulge', 0.0) or 0.0)
        elementos.append((
            'arc',
            (arc['start_point'], arc['end_point'], float(arc['radius']), float(arc['length']), bulge)
        ))

    if not elementos:
        logger.warning("Não foi possível montar a lista de elementos.")
        return None

    # --- Sequenciamento contínuo --------------------------------------------
    def same_pt(a, b, tol=1e-6):
        return abs(a[0]-b[0]) <= tol and abs(a[1]-b[1]) <= tol

    # Se pedir ponto inicial, tenta começar por ele (start OR end)
    if ponto_inicial_real:
        for i, elemento in enumerate(elementos):
            a, b = elemento[1][0], elemento[1][1]
            if same_pt(a, ponto_inicial_real) or same_pt(b, ponto_inicial_real):
                elementos = [elementos[i]] + elementos[:i] + elementos[i+1:]
                break

    sequencia_completa = []
    ponto_atual = elementos[0][1][0]

    while elementos:
        for i, elemento in enumerate(elementos):
            tipo_segmento, dados = elemento
            start_point, end_point = dados[0], dados[1]

            if same_pt(ponto_atual, start_point):
                sequencia_completa.append(elemento)
                ponto_atual = end_point
                elementos.pop(i)
                break
            elif same_pt(ponto_atual, end_point):
                # Inverte sentido do segmento
                if tipo_segmento == 'line':
                    elementos[i] = ('line', (end_point, start_point))
                else:
                    # (start, end, radius, length, bulge) -> invertendo bulge
                    radius, length, bulge = float(dados[2]), float(dados[3]), float(dados[4] or 0.0)
                    elementos[i] = ('arc', (end_point, start_point, radius, length, -bulge))
                sequencia_completa.append(elementos[i])
                ponto_atual = start_point
                elementos.pop(i)
                break
        else:
            # Fallback: consome a próxima peça para evitar loop infinito
            seg = elementos.pop(0)
            sequencia_completa.append(seg)
            ponto_atual = seg[1][1]

    # --- Orientação (CW/CCW) ------------------------------------------------
    # área assinada simples (suficiente p/ sinal)
    pontos_para_area = [seg[1][0] for seg in sequencia_completa]
    pontos_para_area.append(sequencia_completa[-1][1][1])

    simple_ordered_points = [(float(pt[0]), float(pt[1])) for pt in pontos_para_area]
    area_tmp = calculate_signed_area(simple_ordered_points)

    _sentido = (sentido_poligonal or "").strip().lower().replace("-", "_")
    def _reverter_seq(seq):
        seq.reverse()
        for i, (t, d) in enumerate(seq):
            if t == 'line':
                a, b = d
                seq[i] = ('line', (b, a))
            elif t == 'arc':
                a, b, r, L, bu = d
                seq[i] = ('arc', (b, a, float(r), float(L), -float(bu or 0.0)))

    if _sentido == 'horario':  # CW
        if area_tmp > 0:
            _reverter_seq(sequencia_completa)
            area_tmp = abs(area_tmp)
            logger.info("Sentido ajustado para horário (CW).")
    else:  # anti_horario (CCW)
        if area_tmp < 0:
            _reverter_seq(sequencia_completa)
            area_tmp = abs(area_tmp)
            logger.info("Sentido ajustado para anti-horário (CCW).")

    # --- Geometria oficial: cria UMA LWPOLYLINE (xyb) ------------------------
    xyb_points = []
    for tipo_segmento, dados in sequencia_completa:
        x, y = float(dados[0][0]), float(dados[0][1])
        b = 0.0 if tipo_segmento == 'line' else float(dados[4] or 0.0)
        xyb_points.append((x, y, b))

    try:
        pl = msp.add_lwpolyline(xyb_points, format='xyb', close=True)
        pl.dxf.layer = "POLIGONAL_MEMORIAL"
        pl.dxf.color = 4
    except Exception as e:
        logger.warning(f"Falha ao criar LWPOLYLINE com bulge: {e}")

    # --- (Opcional) área densificada respeitando arcos reais -----------------
    def _area_seq_densificada(seq, n=64):
        pts = []
        for tipo, d in seq:
            a, b = d[0], d[1]
            if tipo == 'line':
                pts.append(a)
            else:
                r, bu = float(d[2]), float(d[4] or 0.0)
                dx, dy = b[0]-a[0], b[1]-a[1]
                c = math.hypot(dx, dy)
                if c < 1e-12 or abs(bu) < 1e-12:
                    pts.append(a); continue
                theta = 4.0*math.atan(bu)
                nx, ny = -dy/c, dx/c
                if bu < 0: nx, ny = -nx, -ny
                # distância do centro ao meio da corda
                offset = max(0.0, r*r - (c/2.0)*(c/2.0)) ** 0.5
                midx, midy = (a[0]+b[0])/2.0, (a[1]+b[1])/2.0
                cx, cy = midx + nx*offset, midy + ny*offset
                ang0 = math.atan2(a[1]-cy, a[0]-cx)
                for k in range(n):
                    t = k/(n-1)
                    ang = ang0 + theta*t
                    px, py = cx + r*math.cos(ang), cy + r*math.sin(ang)
                    pts.append((px, py))
        if pts and pts[0] != pts[-1]:
            pts.append(pts[0])
        area2 = 0.0
        for i in range(len(pts)-1):
            x1,y1 = pts[i]; x2,y2 = pts[i+1]
            area2 += x1*y2 - x2*y1
        return abs(area2)/2.0

    try:
        area_curva = _area_seq_densificada(sequencia_completa, n=64)
        logger.info(f"Área densificada (arcos): {area_curva:.4f} m²")
    except Exception as e:
        logger.warning(f"Falha área densificada: {e}")

    # --- Rótulos + Planilha ---------------------------------------------------
    data = []
    num_vertices = len(sequencia_completa)

    for idx, (tipo_segmento, dados) in enumerate(sequencia_completa):
        label = f"V{idx + 1}"
        start_point = dados[0]
        end_point   = dados[1]

        if tipo_segmento == "line":
            azimuth, distance = calculate_azimuth_and_distance(start_point, end_point)
            azimute_excel   = convert_to_dms(azimuth)
            distancia_excel = f"{distance:.2f}".replace(".", ",")
            try:
                add_label_and_distance(doc, msp, start_point, end_point, label, distance)
            except Exception as e:
                logger.warning(f"[native-line] Falha ao anotar {label}: {e}")
                _fallback_anotar_segmento(
                    msp, start_point, end_point, label, distance,
                    H_TXT_VERT, H_TXT_DIST, R_CIRCLE, logger, is_arc=False
                )
        else:
            radius = float(dados[2]); length = float(dados[3]); bulge = float(dados[4] or 0.0)
            distance = length
            try:
                _fallback_anotar_segmento(
                    msp, start_point, end_point, label, distance,
                    H_TXT_VERT, H_TXT_DIST, R_CIRCLE, logger,
                    is_arc=True, arc_radius=radius, arc_bulge=bulge
                )
            except Exception as e:
                logger.warning(f"[fallback-arc] Falha ao anotar {label}: {e}")
            azimute_excel   = f"R={radius:.2f}".replace(".", ",")
            distancia_excel = f"C={distance:.2f}".replace(".", ",")

        next_label = f"V{(idx + 2) if (idx + 1) < num_vertices else 1}"
        confrontante = confrontantes_dict.get(label, "Desconhecido")
        divisa = f"{label}_{next_label}"

        data.append({
            "V": label,
            "E": f"{start_point[0]:.3f}".replace('.', ','),
            "N": f"{start_point[1]:.3f}".replace('.', ','),
            "Z": "0.000",
            "Divisa": divisa,
            "Azimute": azimute_excel,
            "Distancia(m)": distancia_excel,
            "Confrontante": confrontante,
        })

    logger.info(f"Anotações inseridas no DXF: {len(sequencia_completa)} segmentos.")

    # --- Nomes/arquivos -------------------------------------------------------
    try:
        _uuid_from_path = os.path.basename(os.path.dirname(caminho_salvar))
        if not uuid_prefix or len(_uuid_from_path) == 8:
            uuid_prefix = _uuid_from_path
    except Exception:
        pass

    if not tipo and excel_file_path:
        base_x = os.path.basename(excel_file_path).upper()
        for _t in ("ETE", "REM", "SER", "ACE"):
            if _t in base_x:
                tipo = _t
                break
    if not tipo:
        tipo = "MEM"

    # Excel
    df = pd.DataFrame(data, dtype=str)
    matricula_sanit = sanitize_filename(matricula) if isinstance(matricula, str) else str(matricula)
    excel_output_path = os.path.join(caminho_salvar, f"{uuid_prefix}_{tipo}_{matricula_sanit}.xlsx")
    df.to_excel(excel_output_path, index=False)

    wb = openpyxl.load_workbook(excel_output_path)
    ws = wb.active
    for cell in ws[1]:
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center")
    for col, width in {"A":8,"B":15,"C":15,"D":10,"E":20,"F":15,"G":15,"H":30}.items():
        ws.column_dimensions[col].width = width
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=1, max_col=ws.max_column):
        for cell in row:
            cell.alignment = Alignment(horizontal="center", vertical="center")
    wb.save(excel_output_path)
    print(f"Arquivo Excel salvo e formatado em: {excel_output_path}")

    # DXF
    try:
        dxf_output_path = os.path.join(caminho_salvar, f"{uuid_prefix}_{tipo}_{matricula_sanit}.dxf")
        doc.saveas(dxf_output_path)
        print(f"Arquivo DXF salvo em: {dxf_output_path}")
    except Exception as e:
        print(f"Erro ao salvar DXF: {e}")

    return excel_output_path



def create_memorial_document(
    proprietario, matricula, descricao, area_terreno, excel_file_path=None, template_path=None, output_path=None,
    perimeter_dxf=None, area_dxf=None, desc_ponto_Az=None, Coorde_E_ponto_Az=None, Coorde_N_ponto_Az=None,
    azimuth=None, distance=None, comarca=None, RI=None, rua=None, uuid_prefix=None
):
    try:
        def pular_linhas(doc, n_linhas):
            for _ in range(n_linhas):
                doc.add_paragraph("")

        if excel_file_path:
            df = pd.read_excel(excel_file_path)
        else:
            df = None

        doc_word = Document(template_path)
        set_default_font(doc_word)

        p1 = doc_word.add_paragraph(style='Normal')
        run = p1.add_run("MEMORIAL DESCRITIVO INDIVIDUAL")
        run.bold = True
        p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc_word.add_paragraph("")

        p = doc_word.add_paragraph(style='Normal')
        run = p.add_run("NOME PROPRIETÁRIO / OCUPANTE: ")
        run.bold = True
        p.add_run(f"{proprietario}")

        p = doc_word.add_paragraph(style='Normal')
        run = p.add_run("DESCRIÇÃO: ")
        run.bold = True
        p.add_run(f"{descricao}")

        p = doc_word.add_paragraph(style='Normal')
        run = p.add_run("DOCUMENTAÇÃO: ")
        run.bold = True
        p.add_run(f"{matricula} do {RI} da Comarca de  {comarca}")

        p = doc_word.add_paragraph(style='Normal')
        run = p.add_run("ENDEREÇO: ")
        run.bold = True
        p.add_run(f"{rua}")

        p = doc_word.add_paragraph(style='Normal')
        run = p.add_run("ÁREA LEVANTAMENTO: ")
        run.bold = True
        area_dxf_formatada = f"{round(area_dxf, 2):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        p.add_run(f"{area_dxf_formatada} metros quadrados")

        p = doc_word.add_paragraph(style='Normal')
        run = p.add_run("PERÍMETRO LEVANTAMENTO: ")
        run.bold = True
        perimeter_dxf_formatado = f"{round(perimeter_dxf, 2):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        p.add_run(f"{perimeter_dxf_formatado} metros")

        doc_word.add_paragraph("")

        if df is not None and len(df) > 0:
            initial = df.iloc[0]
            p = doc_word.add_paragraph(
                "Pontos definidos pelas Coordenadas Planas no Sistema U.T.M. – ZONA 22S – Meridiano 51ºW, georreferenciadas ao Sistema Geodésico Brasileiro – SIRGAS 2000.",
                style='Normal'
            )
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            p2 = doc_word.add_paragraph(style='Normal')
            p2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            p2.add_run("Inicia-se a descrição deste perímetro no vértice ")
            run_v_inicial = p2.add_run(f"{initial['V']}")
            run_v_inicial.bold = True
            p2.add_run(
                f", de coordenadas N(Y) {initial['N']} e E(X) {initial['E']}, "
                f"situado no limite com {initial['Confrontante']}."
            )

            num_points = len(df)
            for i in range(num_points):
                current = df.iloc[i]
                next_index = (i + 1) % num_points
                next_point = df.iloc[next_index]

                azimute = current['Azimute']
                distancia = current['Distancia(m)']
                confrontante = current['Confrontante']
                destino = next_point['V']
                coord_n = next_point['N']
                coord_e = next_point['E']

                if isinstance(azimute, str) and azimute.startswith("R=") and isinstance(distancia, str) and distancia.startswith("C="):
                    texto_paragrafo = (
                        f"Deste, segue com raio de {azimute[2:]}m e distância de {distancia[2:]}m, "
                        f"confrontando neste trecho com {confrontante}, até o vértice "
                    )
                    restante_texto = f", de coordenadas N(Y) {coord_n} e E(X) {coord_e};"
                else:
                    texto_paragrafo = (
                        f"Deste, segue com azimute de {azimute} e distância de {distancia} m, "
                        f"confrontando neste trecho com {confrontante}, até o vértice "
                    )
                    if next_index == 0:
                        restante_texto = f", origem desta descrição de coordenadas N(Y) {coord_n} e E(X) {coord_e};"
                    else:
                        restante_texto = f", de coordenadas N(Y) {coord_n} e E(X) {coord_e};"

                p = doc_word.add_paragraph(style='Normal')
                p.add_run(texto_paragrafo)
                run_v = p.add_run(destino)
                run_v.bold = True
                p.add_run(restante_texto)
        else:
            doc_word.add_paragraph("\nDescrição do perímetro não incluída neste memorial.", style='Normal')
            pular_linhas(doc_word, 8)

        paragrafo_fechamento = doc_word.add_paragraph(
            f"Fechando-se assim o perímetro com {str(round(perimeter_dxf, 2)).replace('.', ',')} m "
            f"e a área com {str(round(area_dxf, 2)).replace('.', ',')} m².",
            style='Normal'
        )
        paragrafo_fechamento.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        doc_word.add_paragraph("")
        doc_word.add_paragraph("")

        data_atual = obter_data_em_portugues()
        # Centralizar data
        paragrafo_data = doc_word.add_paragraph(f"Paraná, {data_atual}.", style='Normal')
        paragrafo_data.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc_word.add_paragraph("")
        doc_word.add_paragraph("")

        infos_engenheiro = [
            "____________________",
            "Rodrigo Luis Schmitz",
            "Técnico em Agrimensura",
            "CFT: 045.300.139-44"
        ]

        for info in infos_engenheiro:
            paragrafo = doc_word.add_paragraph(info, style='Normal')
            paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc_word.save(output_path)
        print(f"Memorial descritivo salvo em: {output_path}")

    except Exception as e:
        print(f"Erro ao criar o documento memorial: {e}")


# Função principal
def main_poligonal_fechada(caminho_excel, caminho_dxf, pasta_preparado, pasta_concluido, caminho_template, sentido_poligonal='horario'):
    print("\n🔹 Carregando dados do imóvel")
    logger.info("Iniciando processamento da poligonal fechada")
    uuid_prefix = os.path.basename(os.path.dirname(os.path.normpath(pasta_concluido)))

    uu_from_preparado = os.path.basename(os.path.dirname(os.path.normpath(pasta_preparado)))
    uu_from_concluido = os.path.basename(os.path.dirname(os.path.normpath(pasta_concluido)))
    if uu_from_preparado != uu_from_concluido:
        logger.warning(f"pasta_preparado e pasta_concluido com UUIDs distintos: {uu_from_preparado} != {uu_from_concluido}")

    if uu_from_preparado != uuid_prefix:
        logger.warning(
            f"pasta_preparado não pertence ao mesmo UUID ({uu_from_preparado} != {uuid_prefix}). "
            "Verifique preparar_arquivos() para usar /tmp/<UUID>/PREPARADO."
        )

    print(f"[DEBUG poligonal_fechada] UUID recebido: {uuid_prefix}")

    try:
        dados_df = pd.read_excel(caminho_excel, sheet_name='Dados_do_Imóvel', header=None)
        dados_dict = dict(zip(dados_df.iloc[:, 0], dados_df.iloc[:, 1]))

        proprietario = str(dados_dict.get("NOME DO PROPRIETÁRIO", "")).strip()
        matricula = str(dados_dict.get("DOCUMENTAÇÃO DO IMÓVEL", "")).strip()
        descricao = str(dados_dict.get("DESCRIÇÃO", "")).strip()
        area_terreno = str(dados_dict.get("ÁREA TOTAL DO TERRENO DOCUMENTADA", "")).strip()
        comarca = str(dados_dict.get("COMARCA", "")).strip()
        RI = str(dados_dict.get("CRI", "")).strip()
        rua = str(dados_dict.get("ENDEREÇO/LOCAL", "")).strip()

        nome_dxf = os.path.basename(caminho_dxf).upper()
        tipo = next((t for t in ["ETE", "REM", "SER", "ACE"] if t in nome_dxf), None)

        if not tipo:
            msg = "❌ Tipo do projeto não identificado no nome do DXF."
            print(msg)
            logger.warning(msg)
            return

        print(f"📁 Tipo identificado: {tipo}")
        logger.info(f"Tipo identificado: {tipo}")

        # Novo padrão correto
        padrao = os.path.join(pasta_preparado, f"FECHADA_*_{tipo}.xlsx")
        lista_encontrada = glob.glob(padrao)

        if not lista_encontrada:
            print(f"❌ Arquivo de confrontantes esperado não encontrado para tipo {tipo}")
            logger.warning(f"Nenhuma planilha FECHADA_*_{tipo}.xlsx encontrada.")
            return

        excel_confrontantes = lista_encontrada[0]
        print(f"✅ Confrontante carregado: {excel_confrontantes}")
        logger.info(f"Planilha de confrontantes usada: {excel_confrontantes}")

        # nome_limpo_dxf = f"DXF_LIMPO_{sanitize_filename(matricula)}.dxf"
        # caminho_dxf_limpo = os.path.join(pasta_concluido, nome_limpo_dxf)

        # dxf_resultado, ponto_az, ponto_inicial = limpar_dxf_e_inserir_ponto_az(caminho_dxf, caminho_dxf_limpo)
        # logger.info(f"DXF limpo salvo em: {caminho_dxf_limpo}")

        # >>> PATCH: nome do DXF limpo com UUID + TIPO + matrícula

        uuid_exec = os.path.basename(os.path.dirname(pasta_concluido))
        safe_uuid = sanitize_filename(uuid_exec)[:8]  # use a mesma variável que você loga em "[DEBUG] UUID recebido"
        safe_tipo = sanitize_filename(tipo)           # "ETE", "REM", etc.
        safe_mat  = sanitize_filename(matricula)

        nome_limpo_dxf   = f"{safe_uuid}_{safe_tipo}_{safe_mat}_LIMPO.dxf"
        caminho_dxf_limpo = os.path.join(pasta_concluido, nome_limpo_dxf)

        dxf_resultado, ponto_az, ponto_inicial = limpar_dxf_e_inserir_ponto_az(caminho_dxf, caminho_dxf_limpo)
        logger.info(f"DXF limpo salvo em: {caminho_dxf_limpo}")

        # (opcional) remover arquivo legado sem UUID
        legado = os.path.join(pasta_concluido, f"DXF_LIMPO_{safe_mat}.dxf")
        if os.path.exists(legado):
            try:
                os.remove(legado)
                logger.info(f"DXF limpo legado removido: {legado}")
            except Exception as e:
                logger.warning(f"Não foi possível remover legado {legado}: {e}")
        # <<< PATCH


       


        if not ponto_az or not ponto_inicial:
            msg = "❌ Não foi possível identificar o ponto Az ou inicial."
            print(msg)
            logger.error(msg)
            return

        doc, linhas, arcos, perimeter_dxf, area_dxf = get_document_info_from_dxf(dxf_resultado)
        if not doc or not linhas:
            msg = "❌ Documento DXF inválido ou vazio."
            print(msg)
            logger.error(msg)
            return

        msp = doc.modelspace()
        v1 = linhas[0][0]
        distance_az_v1 = calculate_distance(ponto_az, v1)
        azimute_az_v1 = calculate_azimuth(ponto_az, v1)
        azimuth = calculate_azimuth(ponto_az, v1)
        distance = math.hypot(v1[0] - ponto_az[0], v1[1] - ponto_az[1])

        print(f"📏 Azimute: {azimuth:.2f}°, Distância Az-V1: {distance:.2f}m")
        logger.info(f"Azimute: {azimuth:.2f}°, Distância Az-V1: {distance:.2f}m")

        excel_output = create_memorial_descritivo(
            doc=doc,
            msp=msp,
            lines=linhas,
            arcs=arcos,
            proprietario=proprietario,
            matricula=matricula,
            caminho_salvar=pasta_concluido,
            excel_file_path=excel_confrontantes,
            ponto_az=ponto_az,
            distance_az_v1=distance_az_v1,
            azimute_az_v1=azimute_az_v1,
            ponto_inicial_real=ponto_inicial,
            tipo=tipo,
            uuid_prefix=uuid_prefix,
            sentido_poligonal=sentido_poligonal
        )
        try:
            if os.path.exists(caminho_dxf_limpo):
                os.remove(caminho_dxf_limpo)
                logger.info(f"DXF LIMPO removido após gerar DXF final: {caminho_dxf_limpo}")
        except Exception as e:
            logger.warning(f"Não foi possível remover DXF LIMPO: {e}")

        if excel_output:
            output_docx = os.path.join(pasta_concluido, f"{uuid_prefix}_{tipo}_{sanitize_filename(matricula)}.docx")
            create_memorial_document(
                proprietario=proprietario,
                matricula=matricula,
                descricao=descricao,
                area_terreno=area_terreno,
                excel_file_path=excel_output,
                template_path=caminho_template,
                output_path=output_docx,
                perimeter_dxf=perimeter_dxf,
                area_dxf=area_dxf,
                desc_ponto_Az="",
                Coorde_E_ponto_Az=ponto_az[0],
                Coorde_N_ponto_Az=ponto_az[1],
                azimuth=azimuth,
                distance=distance,
                comarca=comarca,
                RI=RI,
                rua=rua,
                uuid_prefix=uuid_prefix
            )

            print(f"📄 Memorial gerado com sucesso: {output_docx}")
            logger.info(f"Memorial gerado com sucesso: {output_docx}")
        else:
            msg = "❌ Falha ao gerar o memorial descritivo."
            print(msg)
            logger.error(msg)

    except Exception as e:
        logger.exception("Erro inesperado durante a execução da poligonal fechada")
