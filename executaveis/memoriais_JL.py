# ⚠️ Atenção:
# A partir de agosto/2025, o módulo `memorial_azimute_jl.py` foi suprimido.
# Todas as funcionalidades de geração de memorial foram consolidadas em `memoriais_JL.py`.
import os
import math
import traceback
from datetime import datetime
from decimal import Decimal, getcontext
import locale
import re
import pandas as pd
import openpyxl
from openpyxl.styles import Alignment, Font
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
import ezdxf
from shapely.geometry import Polygon
try:
    from ezdxf.math import Vec3 as Vector
except ImportError:
    from ezdxf.math import Vector

getcontext().prec = 28  # Define a precisão para 28 casas decimais

# Força o locale para português em sistemas Windows ou Linux

try:
    locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8')  # Linux/Mac padrão
except locale.Error:
    try:
        locale.setlocale(locale.LC_TIME, 'Portuguese_Brazil.1252')  # Windows
    except locale.Error as e:
        print(f"⚠️ Locale não pôde ser definido: {e}")
        locale.setlocale(locale.LC_TIME, '')  # fallback para padrão

# Agora sim a data será em português
data_atual = datetime.now().strftime("%d de %B de %Y")



def obter_data_em_portugues():
    meses_pt = {
        "January": "janeiro", "February": "fevereiro", "March": "março",
        "April": "abril", "May": "maio", "June": "junho",
        "July": "julho", "August": "agosto", "September": "setembro",
        "October": "outubro", "November": "novembro", "December": "dezembro"
    }
    data = datetime.now()
    mes_en = data.strftime("%B")
    mes_pt = meses_pt.get(mes_en, mes_en)
    return f"{data.day:02d} de {mes_pt} de {data.year}"

# 🔧 Garante que a variável log é segura para uso (fallback para DummyLog)
def _ensure_log(log):
    if log is None:
        class DummyLog:
            def write(self, msg): pass
        return DummyLog()
    return log


def sanitize_filename(filename):
    sanitized_filename = re.sub(r'[\\/*?:"<>|]', "_", filename)
    return sanitized_filename

def calculate_distance(point1, point2):
    dx = point2[0] - point1[0]
    dy = point2[1] - point1[1]
    return math.sqrt(dx**2 + dy**2)

def calculate_azimuth_and_distance(start_point, end_point):
    dx = end_point[0] - start_point[0]
    dy = end_point[1] - start_point[1]
    distance = math.hypot(dx, dy)
    azimuth = math.degrees(math.atan2(dx, dy))
    if azimuth < 0:
        azimuth += 360
    return azimuth, distance

def calculate_azimuth(p1, p2):
    delta_x = p2[0] - p1[0]
    delta_y = p2[1] - p1[1]
    azimuth_rad = math.atan2(delta_x, delta_y)
    azimuth_deg = math.degrees(azimuth_rad)
    if azimuth_deg < 0:
        azimuth_deg += 360
    return azimuth_deg

def limpar_dxf_e_inserir_ponto_az(original_path, saida_path):
    try:
        doc_antigo = ezdxf.readfile(original_path)
        msp_antigo = doc_antigo.modelspace()
        doc_novo = ezdxf.new(dxfversion='R2010')
        msp_novo = doc_novo.modelspace()

        pontos_polilinha = None
        bulges_polilinha = None
        ponto_inicial_real = None

        for entity in msp_antigo.query('LWPOLYLINE'):
            if entity.closed:
                # 🔧 Leitura com verificação de duplicatas
                pontos_polilinha_raw = entity.get_points('xyseb')
                ponto_inicial_real = (float(pontos_polilinha_raw[0][0]), float(pontos_polilinha_raw[0][1]))
                pontos_polilinha = []
                bulges_polilinha = []

                tolerancia = 1e-6  # Tolerância para considerar pontos idênticos

                for pt in pontos_polilinha_raw:
                    x, y, *_, bulge = pt
                    x, y = float(x), float(y)
                    if not pontos_polilinha:
                        # Primeiro ponto
                        pontos_polilinha.append((x, y))
                        bulges_polilinha.append(bulge)
                    else:
                        x_ant, y_ant = pontos_polilinha[-1]
                        if math.hypot(x - x_ant, y - y_ant) > tolerancia:
                            pontos_polilinha.append((x, y))
                            bulges_polilinha.append(bulge)
                        else:
                            print(f"⚠️ Ponto duplicado consecutivo removido: {(x, y)}")

                # 🔍 Verificação extra para ponto final duplicado
                if len(pontos_polilinha) > 2 and math.hypot(
                    pontos_polilinha[0][0] - pontos_polilinha[-1][0],
                    pontos_polilinha[0][1] - pontos_polilinha[-1][1]
                ) < tolerancia:
                    print("⚠️ Último ponto é igual ao primeiro — removendo ponto final duplicado.")
                    pontos_polilinha.pop()
                    bulges_polilinha.pop()

                # 🔍 Verificação extra para P1 == P2
                if len(pontos_polilinha) > 1 and math.hypot(
                    pontos_polilinha[0][0] - pontos_polilinha[1][0],
                    pontos_polilinha[0][1] - pontos_polilinha[1][1]
                ) < tolerancia:
                    print("⚠️ Primeiro ponto é igual ao segundo — removendo o segundo ponto duplicado.")
                    pontos_polilinha.pop(1)
                    bulges_polilinha.pop(1)

                break

        if pontos_polilinha is None:
            raise ValueError("Nenhuma polilinha fechada encontrada no DXF original.")

        if calculate_signed_area(pontos_polilinha) < 0:
            pontos_polilinha.reverse()
            bulges_polilinha.reverse()
            bulges_polilinha = [-b for b in bulges_polilinha]

        pontos_com_bulge = [
            (pontos_polilinha[i][0], pontos_polilinha[i][1], bulges_polilinha[i])
            for i in range(len(pontos_polilinha))
        ]

        msp_novo.add_lwpolyline(
            pontos_com_bulge,
            format='xyb',
            close=True,
            dxfattribs={'layer': 'DIVISA_PROJETADA'}
        )

        # Não desenha mais o ponto Az, mas retorna as coordenadas de V1 como ponto_az válido
        ponto_az = pontos_polilinha[0]

        doc_novo.saveas(saida_path)
        print(f"✅ DXF limpo salvo em: {saida_path}")
        
        return saida_path, ponto_az, ponto_inicial_real

    except Exception as e:
        print(f"❌ Erro ao limpar DXF: {e}")
        return original_path, None, None

def calculate_signed_area(points):
    area = 0
    for i in range(len(points)):
        x1, y1 = points[i]
        x2, y2 = points[(i + 1) % len(points)]
        area += (x1 * y2) - (x2 * y1)
    return area / 2

def get_document_info_from_dxf(dxf_file_path):
    try:
        doc = ezdxf.readfile(dxf_file_path)
        msp = doc.modelspace()

        lines = []
        arcs = []
        perimeter_dxf = 0
        area_dxf = 0.0

        for entity in msp.query('LWPOLYLINE'):
            if entity.closed:  # ✅ correção (era is_closed)
                polyline_points = entity.get_points('xyseb')
                num_points = len(polyline_points)

                boundary_points = []

                for i in range(num_points):
                    x_start, y_start, _, _, bulge = polyline_points[i]
                    x_end, y_end, _, _, _ = polyline_points[(i + 1) % num_points]

                    start_point = (float(x_start), float(y_start))
                    end_point = (float(x_end), float(y_end))

                    if bulge != 0:
                        dx = end_point[0] - start_point[0]
                        dy = end_point[1] - start_point[1]
                        chord_length = math.hypot(dx, dy)
                        sagitta = (bulge * chord_length) / 2
                        radius = ((chord_length / 2)**2 + sagitta**2) / (2 * abs(sagitta))
                        angle_span_rad = 4 * math.atan(abs(bulge))
                        arc_length = radius * angle_span_rad

                        mid_x = (start_point[0] + end_point[0]) / 2
                        mid_y = (start_point[1] + end_point[1]) / 2
                        chord_midpoint = (mid_x, mid_y)

                        offset_dist = math.sqrt(radius**2 - (chord_length / 2)**2)
                        dx = float(end_point[0]) - float(start_point[0])
                        dy = float(end_point[1]) - float(start_point[1])

                        length = math.hypot(dx, dy)
                        perp_vector = (-dy / length, dx / length)

                        if bulge < 0:
                            perp_vector = (-perp_vector[0], -perp_vector[1])

                        center_x = chord_midpoint[0] + perp_vector[0] * offset_dist
                        center_y = chord_midpoint[1] + perp_vector[1] * offset_dist
                        center = (center_x, center_y)

                        start_angle = math.atan2(start_point[1] - center[1], start_point[0] - center[0])
                        end_angle = start_angle + (angle_span_rad if bulge > 0 else -angle_span_rad)

                        arcs.append({
                            'start_point': (start_point[0], start_point[1]),
                            'end_point': (end_point[0], end_point[1]),
                            'center': (center[0], center[1]),
                            'radius': radius,
                            'start_angle': math.degrees(start_angle),
                            'end_angle': math.degrees(end_angle),
                            'length': arc_length,
                            'bulge': float(bulge),  # <-- ADICIONE ESTA LINHA
                            'sweep_degrees': math.degrees(end_angle - start_angle)  # <-- (OPCIONAL, ajuda no debug)
                        })

                        num_arc_points = 100
                        for t in range(num_arc_points):
                            angle = start_angle + (end_angle - start_angle) * t / num_arc_points
                            arc_x = center[0] + radius * math.cos(angle)
                            arc_y = center[1] + radius * math.sin(angle)
                            boundary_points.append((arc_x, arc_y))

                        segment_length = arc_length
                        perimeter_dxf += segment_length
                    else:
                        lines.append((start_point, end_point))
                        boundary_points.append((start_point[0], start_point[1]))
                        dx = end_point[0] - start_point[0]
                        dy = end_point[1] - start_point[1]
                        segment_length = math.hypot(dx, dy)
                        perimeter_dxf += segment_length

                polygon = Polygon(boundary_points)
                area_dxf = polygon.area  # área exata do desenho
                break

        if not lines and not arcs:
            print("Nenhuma polilinha fechada encontrada no arquivo DXF.")
            return None, [], [], 0, 0  # ✅ retorna sempre 5 itens

        print(f"Linhas processadas: {len(lines)}")
        print(f"Arcos processados: {len(arcs)}")
        print(f"Perímetro do DXF: {perimeter_dxf:.2f} metros")
        print(f"Área do DXF: {area_dxf:.2f} metros quadrados")

        return doc, lines, arcs, perimeter_dxf, area_dxf

    except Exception as e:
        print(f"Erro ao obter informações do documento: {e}")
        traceback.print_exc()
        return None, [], [], 0, 0  # ✅ retorna sempre 5 itens


def calculate_area_with_arcs(points, arcs):
    """
    Calcula a área de uma poligonal que inclui segmentos de linha reta e arcos.

    :param points: Lista de tuplas (x, y) representando os vértices dos segmentos de linha reta.
    :param arcs: Lista de dicionários, cada um contendo informações sobre um arco:
                 {'start_point': (x1, y1), 'end_point': (x2, y2), 'center': (xc, yc),
                  'radius': r, 'start_angle': a1, 'end_angle': a2, 'length': l}
    :return: Área total da poligonal.
    """
    # Calcular a área usando a fórmula do polígono (Shoelace) para os segmentos de linha reta
    num_points = len(points)
    area_linear = 0.0
    for i in range(num_points):
        x1, y1 = points[i]
        x2, y2 = points[(i + 1) % num_points]
        area_linear += x1 * y2 - x2 * y1
    area_linear = abs(area_linear) / 2.0

    # Calcular a área dos arcos
    area_arcs = 0.0
    for arc in arcs:
#         start_point = Vec2(arc['start_point'])
#         end_point = Vec2(arc['end_point'])
#         center = Vec2(arc['center'])
       
        start_point = (float(arc['start_point'][0]), float(arc['start_point'][1]))
        end_point = (float(arc['end_point'][0]), float(arc['end_point'][1]))
        center = (float(arc['center'][0]), float(arc['center'][1]))

        radius = arc['radius']
        start_angle = math.radians(arc['start_angle'])
        end_angle = math.radians(arc['end_angle'])

        # Determinar a variação do ângulo
        delta_angle = end_angle - start_angle
        if delta_angle <= 0:
            delta_angle += 2 * math.pi

        # Área do setor circular
        sector_area = 0.5 * radius**2 * delta_angle

        # Área do triângulo formado pelo centro e os pontos inicial e final do arco
#         triangle_area = 0.5 * abs((start_point.x - center.x) * (end_point.y - center.y) -
#                                   (end_point.x - center.x) * (start_point.y - center.y))
        triangle_area = 0.5 * abs((start_point[0] - center[0]) * (end_point[1] - center[1]) -
                                  (end_point[0] - center[0]) * (start_point[1] - center[1]))

        # Área do segmento circular
        segment_area = sector_area - triangle_area

        # Determinar a orientação do arco para adicionar ou subtrair a área
        # Aqui assumimos que a orientação positiva indica um arco no sentido anti-horário
        # e negativa para o sentido horário. Isso pode precisar de ajustes conforme a definição dos dados.
        if delta_angle > math.pi:
            area_arcs -= segment_area  # Arco no sentido horário
        else:
            area_arcs += segment_area  # Arco no sentido anti-horário

    # Área total
    area_total = area_linear + area_arcs
    return area_total

def bulge_to_arc_length(start_point, end_point, bulge):
    #chord_length = (end_point - start_point).magnitude
    dx = end_point[0] - start_point[0]
    dy = end_point[1] - start_point[1]
    chord_length = math.hypot(dx, dy)

    sagitta = (bulge * chord_length) / 2
    radius = ((chord_length / 2) ** 2 + sagitta ** 2) / (2 * abs(sagitta))
    angle = 4 * math.atan(abs(bulge))
    arc_length = radius * angle
    return arc_length, radius, angle

import math
from shapely.geometry import Polygon



# 🔹 Função para definir a fonte padrão
def set_default_font(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

def add_arc_labels(doc, msp, start_point, end_point, radius, length, label, log=None):
    log = _ensure_log(log)

    
    try:
        #mid_point = Vec2((start_point[0] + end_point[0])/2, (start_point[1] + end_point[1])/2)
        mid_point = ((float(start_point[0]) + float(end_point[0]))/2, (float(start_point[1]) + float(end_point[1]))/2)

        label_radius = f"R={radius:.2f}".replace('.', ',')
        label_length = f"C={length:.2f}".replace('.', ',')

        msp.add_text(
            label_radius,
            dxfattribs={
                "height": 1.0,
                "layer": "LAYOUT_DISTANCIAS",
                "insert": (mid_point[0], mid_point[1])
            }
        )

        msp.add_text(
            label_length,
            dxfattribs={
                "height": 1.0,
                "layer": "LAYOUT_DISTANCIAS",
                "insert": (mid_point[0], mid_point[1] - 2)
            }
        )

        print(f"✅ Rótulos {label_radius} e {label_length} adicionados corretamente no DXF.")
        
        log.write(f"✅ Rótulos {label_radius} e {label_length} adicionados corretamente no DXF.\n")

    except Exception as e:
        print(f"❌ Erro ao adicionar rótulos dos arcos: {e}")
        
        log.write(f"❌ Erro ao adicionar rótulos dos arcos: {e}\n")


def calculate_point_on_line(start, end, distance):
    """
    Calcula um ponto a uma determinada distância sobre uma linha entre dois pontos.
    :param start: Coordenadas do ponto inicial (x, y).
    :param end: Coordenadas do ponto final (x, y).
    :param distance: Distância do ponto inicial ao longo da linha.
    :return: Coordenadas do ponto calculado (x, y).
    """
    dx, dy = end[0] - start[0], end[1] - start[1]
    length = math.hypot(dx, dy)  # Calcula o comprimento da linha
    if length == 0:
        raise ValueError("Ponto inicial e final são iguais, não é possível calcular um ponto na linha.")
    return (
        start[0] + (dx / length) * distance,
        start[1] + (dy / length) * distance
    )


def calculate_azimuth(p1, p2):
    """
    Calcula o azimute entre dois pontos em graus.
    Azimute é o ângulo medido no sentido horário a partir do Norte.
    """
    delta_x = p2[0] - p1[0]  # Diferença em X (Leste/Oeste)
    delta_y = p2[1] - p1[1]  # Diferença em Y (Norte/Sul)

    # Calcular o ângulo em radianos
    azimuth_rad = math.atan2(delta_x, delta_y)

    # Converter para graus
    azimuth_deg = math.degrees(azimuth_rad)

    # Garantir que o azimute esteja no intervalo [0, 360)
    if azimuth_deg < 0:
        azimuth_deg += 360

    return azimuth_deg

def calculate_distance(point1, point2):
    """
    Calcula a distância entre dois pontos em um plano 2D.
    :param point1: Tupla (x1, y1) representando o primeiro ponto.
    :param point2: Tupla (x2, y2) representando o segundo ponto.
    :return: Distância entre os pontos.
    """
    dx = point2[0] - point1[0]
    dy = point2[1] - point1[1]
    return math.sqrt(dx**2 + dy**2)




# Função para calcular azimute e distância
def calculate_azimuth_and_distance(start_point, end_point):
    dx = end_point[0] - start_point[0]
    dy = end_point[1] - start_point[1]
    distance = math.hypot(dx, dy)
    azimuth = math.degrees(math.atan2(dx, dy))
    if azimuth < 0:
        azimuth += 360
    return azimuth, distance


def add_azimuth_arc(doc, msp, ponto_az, v1, azimuth, log=None):
    log = _ensure_log(log)
    """
    Adiciona o arco do azimute no ModelSpace.
    """
    

    try:
        if 'LAYOUT_AZIMUTES' not in doc.layers:
            doc.layers.new(name='LAYOUT_AZIMUTES', dxfattribs={'color': 5})

        # Traçar segmento entre Az e V1
        msp.add_line(start=ponto_az, end=v1, dxfattribs={'layer': 'LAYOUT_AZIMUTES'})

        # Adicionar rótulo do azimute
        azimuth_label = f"Azimute = {convert_to_dms(azimuth)}"
        label_position = (
            ponto_az[0] + 1.5 * math.cos(math.radians(azimuth / 2)),
            ponto_az[1] + 1.5 * math.sin(math.radians(azimuth / 2))
        )
        msp.add_text(
            azimuth_label,
            dxfattribs={'height': 0.5, 'layer': 'LAYOUT_AZIMUTES', 'insert': label_position}
        )

        print(f"Rótulo do azimute ({azimuth_label}) adicionado com sucesso em {label_position}")
        
        log.write(f"Rótulo do azimute ({azimuth_label}) adicionado com sucesso em {label_position}\n")

    except Exception as e:
        print(f"Erro ao adicionar arco do azimute: {e}")
        
        log.write(f"Erro ao adicionar arco do azimute: {e}\n")


# Função para converter graus decimais para DMS
def convert_to_dms(decimal_degrees):
    degrees = int(decimal_degrees)
    minutes = int(abs(decimal_degrees - degrees) * 60)
    seconds = abs((decimal_degrees - degrees - minutes / 60) * 3600)
    return f"{degrees}° {minutes}' {seconds:.2f}\""

# Função para calcular a área de uma poligonal
def calculate_polygon_area(points):
    n = len(points)
    area = 0.0
    for i in range(n):
        x1, y1 = points[i][0], points[i][1]
        x2, y2 = points[(i + 1) % n][0], points[(i + 1) % n][1]
        area += x1 * y2 - x2 * y1
    return area / 2.0


def add_label_and_distance(doc, msp, start_point, end_point, label, distance,log=None):
    log = _ensure_log(log)
    

    try:
        # Garantir pontos como tuplas de float
        start_point = (float(start_point[0]), float(start_point[1]))
        end_point = (float(end_point[0]), float(end_point[1]))

        # Criar camadas se ainda não existirem
        layers = [
            ("LAYOUT_VERTICES", 2),   # Amarelo
            ("LAYOUT_DISTANCIAS", 4)  # Azul
        ]
        for layer_name, color in layers:
            if layer_name not in doc.layers:
                doc.layers.new(name=layer_name, dxfattribs={"color": color})

        # Adicionar círculo no ponto inicial
        msp.add_circle(center=start_point, radius=0.5, dxfattribs={'layer': 'LAYOUT_VERTICES'})

        # Adicionar rótulo do vértice (ex: V1, V2...)
        ofset_x = 0.3
        ofset_y = 0.3
        msp.add_text(
            label,
            dxfattribs={
                'height': 0.5,
                'layer': 'LAYOUT_VERTICES',
                'insert': (start_point[0] + ofset_x, start_point[1] + ofset_y)
            }
        )

        # Calcular ponto médio e vetor do segmento
        mid_x = (start_point[0] + end_point[0]) / 2
        mid_y = (start_point[1] + end_point[1]) / 2
        dx = end_point[0] - start_point[0]
        dy = end_point[1] - start_point[1]
        length = (dx ** 2 + dy ** 2) ** 0.5

        if length == 0:
            return  # evita divisão por zero

        # Cálculo do ângulo em graus
        angle = math.degrees(math.atan2(dy, dx))

        # Corrige para manter a leitura em pé
        if angle < -90 or angle > 90:
            angle += 180

        # Deslocar o rótulo perpendicularmente ao segmento
        offset = 0.5  # pode ajustar para mais ou menos afastamento

        angle_rad = math.atan2(dy, dx)
        angle_deg = math.degrees(angle_rad)

        # Define o lado correto baseado no ângulo do segmento
        if -90 <= angle_deg <= 90:
            side_factor = 1  # acima da linha
        else:
            side_factor = -1  # abaixo da linha

        perp_x = -dy / length * offset
        perp_y = dx / length * offset
        mid_point_displaced = (mid_x + perp_x, mid_y + perp_y)

        # Formatar distância
        distancia_formatada = f"{distance:.2f} m".replace('.', ',')

        # Inserir texto da distância rotacionado
        msp.add_text(
            distancia_formatada,
            dxfattribs={
                'height': 0.5,
                'layer': 'LAYOUT_DISTANCIAS',
                'rotation': angle,
                'insert': mid_point_displaced
            }
        )

        print(f"✅ DEBUG: '{label}' e distância '{distancia_formatada}' inseridos em {start_point} e {mid_point_displaced} com ângulo {angle:.2f}°")
        
        log.write(f"✅ DEBUG: '{label}' e distância '{distancia_formatada}' inseridos em {start_point} e {mid_point_displaced} com ângulo {angle:.2f}°\n")

    except Exception as e:
        print(f"❌ ERRO GRAVE ao adicionar rótulo '{label}' e distância: {e}")
        print(f"❌ ERRO GRAVE ao adicionar rótulo '{label}' e distância: {e}")



def sanitize_filename(filename):
    # Substitui os caracteres inválidos por um caractere válido (ex: espaço ou underline)
    sanitized_filename = re.sub(r'[\\/*?:"<>|]', "_", filename)  # Substitui caracteres inválidos por "_"
    return sanitized_filename
        
def _fallback_anotar_segmento(msp, start_point, end_point, label, distancia_m,
                              H_TXT_VERT, H_TXT_DIST, R_CIRCLE, log, is_arc=False):
    try:
        # marcador de vértice (círculo sempre visível)
        msp.add_circle(center=start_point, radius=R_CIRCLE,
                       dxfattribs={"layer": "ANOTACOES_DECOPA"})
        # rótulo do vértice
        msp.add_text(str(label), dxfattribs={"height": H_TXT_VERT, "layer": "ANOTACOES_DECOPA"}).set_pos(
            (start_point[0] + R_CIRCLE*1.2, start_point[1] + R_CIRCLE*1.2), align="LEFT"
        )
        # distância no meio da corda
        mid = ((start_point[0] + end_point[0]) / 2.0, (start_point[1] + end_point[1]) / 2.0)
        texto_dist = f"{distancia_m:.2f} m"
        off_dx, off_dy = (R_CIRCLE*1.8, R_CIRCLE*1.8) if not is_arc else (R_CIRCLE*2.2, R_CIRCLE*2.2)
        msp.add_text(texto_dist, dxfattribs={"height": H_TXT_DIST, "layer": "ANOTACOES_DECOPA"}).set_pos(
            (mid[0] + off_dx, mid[1] + off_dy), align="LEFT"
        )
        return True
    except Exception as e:
        log.write(f"[fallback] Falha ao anotar {label}: {e}")
        return False      


# Função para criar memorial descritivo
def create_memorial_descritivo(
    doc,
    msp,
    lines,
    proprietario,
    matricula,
    caminho_salvar,
    arcs=None,
    excel_file_path=None,
    ponto_az=None,
    distance_az_v1=None,
    azimute_az_v1=None,
    ponto_inicial_real=None,   # ✅ já existia
    tipo=None,                 # ✅ novo: alinha com a chamada
    uuid_prefix=None,          # ✅ novo: alinha com a chamada
    encoding='ISO-8859-1',
    boundary_points=None,
    log=None,
    sentido_poligonal="horario"
):

    """
    Cria o memorial descritivo diretamente no arquivo DXF e salva os dados em uma planilha Excel.
    """
    log = _ensure_log(log)
    
    assert hasattr(log, 'write'), "log não possui método write"

    if excel_file_path:
        try:
            confrontantes_df = pd.read_excel(excel_file_path)
            confrontantes_dict = dict(zip(confrontantes_df['Código'], confrontantes_df['Confrontante']))
        except Exception as e:
            print(f"Erro ao carregar arquivo de confrontantes: {e}")
            
            log.write(f"Erro ao carregar arquivo de confrontantes: {e}\n")
            confrontantes_dict = {}
    else:
        confrontantes_dict = {}

    if (not lines) and (not arcs):
        print("Nenhuma geometria (linhas ou arcos) disponível para criar o memorial descritivo.")
        return None

    # Criar uma única lista sequencial de pontos ordenados (linhas e arcos)
    elementos = []
    for p1, p2 in (lines or []):
        elementos.append(('line', (p1, p2)))

    for arc in (arcs or []):
        elementos.append(('arc', (arc['start_point'], arc['end_point'], arc['bulge'], arc['radius'])))

    # Sequenciar os segmentos corretamente
    sequencia_completa = []

    # 🔁 Reordena elementos para começar pelo ponto original do desenho
    if ponto_inicial_real:
        for i, elemento in enumerate(elementos):
            if math.hypot(elemento[1][0][0] - ponto_inicial_real[0], elemento[1][0][1] - ponto_inicial_real[1]) < 1e-6:
                elementos = [elementos[i]] + elementos[:i] + elementos[i+1:]
                break

    ponto_atual = elementos[0][1][0]  # Primeiro ponto do primeiro segmento
    while elementos:
        for i, elemento in enumerate(elementos):
            tipo_segmento, dados = elemento
            start_point, end_point = dados[0], dados[1]

            def _eq_pt(a, b, tol=1e-6):
                return abs(a[0]-b[0]) < tol and abs(a[1]-b[1]) < tol

            if _eq_pt(ponto_atual, start_point):
                sequencia_completa.append(elemento)
                ponto_atual = end_point
                elementos.pop(i)
                break
            elif _eq_pt(ponto_atual, end_point):
                # Inverte a direção do segmento para manter continuidade
                if tipo_segmento == 'line':
                    elementos[i] = ('line', (end_point, start_point))
                else:
                    elementos[i] = ('arc', (end_point, start_point, -dados[2], dados[3]))
                sequencia_completa.append(elementos[i])
                ponto_atual = start_point
                elementos.pop(i)
                break
        else:
            if elementos:
                ponto_atual = elementos[0][1][0]

    # Lista de pontos sequenciais simples para área (garante polígono fechado)
    pontos_para_area = [seg[1][0] for seg in sequencia_completa]
    pontos_para_area.append(sequencia_completa[-1][1][1])  # Fecha o polígono

    simple_ordered_points = [(float(pt[0]), float(pt[1])) for pt in pontos_para_area]
    area_tmp = calculate_signed_area(simple_ordered_points)


    # Normaliza o valor vindo da rota/formulário
    _sentido = (sentido_poligonal or "").strip().lower().replace("-", "_")

    # Regra da área assinada: CCW (anti-horário) => área > 0 ; CW (horário) => área < 0
    # Se o usuário pediu "horario" e a área veio > 0 (CCW), invertemos.
    # Se pediu "anti_horario" e a área veio < 0 (CW), invertemos.

    def _reverter_sequencia_completa(seq):
        """
        Inverte a ordem dos segmentos e troca start/end.
        Para arco: inverte também o sinal do bulge; raio permanece o mesmo.
        """
        seq.reverse()
        for i, (tipo_segmento, dados) in enumerate(seq):
            if tipo_segmento == 'line':
                start, end = dados  # ((x1,y1), (x2,y2))
                seq[i] = ('line', (end, start))
            elif tipo_segmento == 'arc':
                start, end, bulge, radius = dados
                # Reverte endpoints e inverte o sinal do bulge para manter a mesma geometria (sentido oposto)
                seq[i] = ('arc', (end, start, -bulge, radius))
          
            else:
                # fallback genérico: tenta apenas trocar start/end se houver
                try:
                    start, end = dados[0], dados[1]
                    novos = list(dados)
                    novos[0], novos[1] = end, start
                    seq[i] = (tipo_segmento, tuple(novos))
                except Exception:
                    pass


    if _sentido == 'horario':
        if area_tmp > 0:
            _reverter_sequencia_completa(sequencia_completa)
            area_tmp = abs(area_tmp)
            log.write(f"Área invertida para sentido horário (CW); linhas/arcos ajustados. |Área|={area_tmp:.4f} m²")
        else:
            log.write(f"Área já coerente com sentido horário (CW). |Área|={abs(area_tmp):.4f} m²")
    else:  # trata como 'anti_horario'
        if area_tmp < 0:
            _reverter_sequencia_completa(sequencia_completa)
            area_tmp = abs(area_tmp)
            log.write(f"Área invertida para sentido anti-horário (CCW); linhas/arcos ajustados. |Área|={area_tmp:.4f} m²")
        else:
            log.write(f"Área já coerente com sentido anti-horário (CCW). |Área|={abs(area_tmp):.4f} m²")




    # Continuação após inverter corretamente
    data = []
    num_vertices = len(sequencia_completa)  # captura a quantidade correta antes do loop
    anot_count = 0
    for idx, (tipo_segmento, dados) in enumerate(sequencia_completa):
        start_point = dados[0]
        end_point   = dados[1]

        if tipo_segmento == "line":
            azimuth, distance = calculate_azimuth_and_distance(start_point, end_point)
            azimute_excel    = convert_to_dms(azimuth)
            distancia_excel  = f"{distance:.2f}".replace(".", ",")
        elif tipo_segmento == "arc":
            # dados = (start, end, bulge, radius)
            bulge  = dados[2]
            radius = dados[3]
            theta  = 4.0 * math.atan(abs(bulge))  # ângulo central (rad)
            distance = radius * theta             # comprimento do arco
            azimute_excel   = f"R={radius:.2f}".replace(".", ",")
            distancia_excel = f"C={distance:.2f}".replace(".", ",")

        # label = f"V{idx + 1}"
        # # Usa a MESMA rotina de anotação para linhas e arcos (passando o comprimento correto)
        # add_label_and_distance(doc, msp, start_point, end_point, label, distance)

        label = f"V{idx + 1}"
        _is_arc = (tipo_segmento == "arc")

        # 1) TENTA o caminho original (que já funcionava no sentido horário)
        ok_native = False
        try:
            add_label_and_distance(doc, msp, start_point, end_point, label, distance)
            ok_native = True
        except Exception as e:
            log.write(f"[native] Falha ao anotar {label} com add_label_and_distance: {e}")

        # 2) Se falhar, usa fallback robusto para ESTE segmento
        if not ok_native:
            _ = _fallback_anotar_segmento(
                msp, start_point, end_point, label, distance,
                H_TXT_VERT, H_TXT_DIST, R_CIRCLE, log, is_arc=_is_arc
            )

        anot_count += 1


        confrontante = confrontantes_dict.get(f"V{idx + 1}", "Desconhecido")
        divisa = f"V{idx + 1}_V{idx + 2}" if idx + 1 < num_vertices else f"V{idx + 1}_V1"

        data.append({
            "V": label,
            "E": f"{start_point[0]:.3f}".replace('.', ','),
            "N": f"{start_point[1]:.3f}".replace('.', ','),
            "Z": "0.000",
            "Divisa": divisa,
            "Azimute": azimute_excel,
            "Distancia(m)": distancia_excel,
            "Confrontante": confrontante,
        })
    log.write(f"Anotações inseridas no DXF: {anot_count} segmentos (linhas+arcos)")


    # Deriva o UUID do caminho de saída se não vier preenchido
    try:
        _uuid_from_path = os.path.basename(os.path.dirname(caminho_salvar))
        if not uuid_prefix or len(_uuid_from_path) == 8:
            uuid_prefix = _uuid_from_path
    except Exception:
        pass

    # Tenta deduzir o tipo (ETE/REM/SER/ACE) se não vier
    if not tipo and excel_file_path:
        base_x = os.path.basename(excel_file_path).upper()
        for _t in ("ETE", "REM", "SER", "ACE"):
            if _t in base_x:
                tipo = _t
                break

    df = pd.DataFrame(data, dtype=str)

    matricula_sanit = sanitize_filename(matricula) if isinstance(matricula, str) else str(matricula)
    excel_output_path = os.path.join(caminho_salvar, f"{uuid_prefix}_{tipo}_{matricula_sanit}.xlsx")
    df.to_excel(excel_output_path, index=False)

    wb = openpyxl.load_workbook(excel_output_path)
    ws = wb.active

    for cell in ws[1]:
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center")

    column_widths = {
        "A": 8, "B": 15, "C": 15, "D": 10, "E": 20, "F": 15,
        "G": 15, "H": 30, "I": 20, "J": 20, "K": 15, "L": 15
    }
    for col, width in column_widths.items():
        ws.column_dimensions[col].width = width

    for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=1, max_col=ws.max_column):
        for cell in row:
            cell.alignment = Alignment(horizontal="center", vertical="center")

    wb.save(excel_output_path)
    print(f"Arquivo Excel salvo e formatado em: {excel_output_path}")

    try:
        dxf_output_path = os.path.join(caminho_salvar, f"{uuid_prefix}_{tipo}_{matricula_sanit}.dxf")
        doc.saveas(dxf_output_path)
        print(f"Arquivo DXF salvo em: {dxf_output_path}")
    except Exception as e:
        print(f"Erro ao salvar DXF: {e}")

    return excel_output_path







def create_memorial_document(
    proprietario, matricula, descricao, excel_file_path=None, template_path=None, output_path=None,
    perimeter_dxf=None, area_dxf=None, desc_ponto_Az=None, Coorde_E_ponto_Az=None, Coorde_N_ponto_Az=None,
    azimuth=None, distance=None, log=None
):
    log = _ensure_log(log)
    
    try:
        # 🔍 Verificação do template
        print(f"🔎 Caminho do template: {template_path}")
        
        log.write(f"🔎 Template path: {template_path}\n")
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"❌ Template não encontrado: {template_path}")

        # 🔍 Verificação do diretório de saída
        output_dir = os.path.dirname(output_path)
        print(f"📁 Caminho de saída do DOCX: {output_path}")
        
        log.write(f"📁 Output DOCX path: {output_path}\n")
        if not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)

        # Função para pular linhas
        def pular_linhas(doc, n_linhas):
            for _ in range(n_linhas):
                doc.add_paragraph("")

        

        # Ler o arquivo Excel se for informado
        if excel_file_path:
            df = pd.read_excel(excel_file_path)
        else:
            df = None

        # Criar o documento Word carregando o template
        doc_word = Document(template_path)
        set_default_font(doc_word)  # Fonte Arial 12
        # Criar o documento Word carregando o template
        doc_word = Document(template_path)
        set_default_font(doc_word)  # Fonte Arial 12

        # Adiciona o preâmbulo centralizado com a variável "descricao"
        p1 = doc_word.add_paragraph(style='Normal')
        p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run1 = p1.add_run("MEMORIAL DESCRITIVO INDIVIDUAL")
        run1.bold = True

        p2 = doc_word.add_paragraph(f"({descricao})", style='Normal')
        p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc_word.add_paragraph("")  # Espaço em branco

        # Dados da obra (conteúdo fixo, alinhado normalmente à esquerda)
        dados_obra = [
            "DADOS DA OBRA",
            "Rodovia: BR-277/PR",
            "Trecho: Ponte s/ Rio Emboguaçu – Ponte Internacional Brasil / Paraguai (2ª Ponte)",
            "Sub-trecho: Entr. BR-277(km 722,6) (Acesso a 2ª Ponte Rio Paraná) – Inicio Ponte Internacional Brasil / Paraguai (2ª Ponte)",
            "Segmento: km 00,00 a km 15,00",
            "Lote: Único",
            ""
        ]

        for linha in dados_obra:
            if ":" in linha:
                negrito, restante = linha.split(":", 1)
                par = doc_word.add_paragraph(style='Normal')
                run_bold = par.add_run(negrito + ":")
                run_bold.bold = True
                par.add_run(restante)  # texto normal
            else:
                par = doc_word.add_paragraph(style='Normal')
                run = par.add_run(linha)
                if linha.strip() == "DADOS DA OBRA":
                    run.bold = True

        # Adicionar dados do proprietário
        par1 = doc_word.add_paragraph(style='Normal')
        run1 = par1.add_run("NOME PROPRIETÁRIO / OCUPANTE:")
        run1.bold = True
        par1.add_run(f" {proprietario}")

        par2 = doc_word.add_paragraph(style='Normal')
        run2 = par2.add_run("DOCUMENTAÇÃO:")
        run2.bold = True
        par2.add_run(f" MATRÍCULA {matricula}")

        # ÁREA DO IMÓVEL (m²)
        par3 = doc_word.add_paragraph(style='Normal')
        run3a = par3.add_run("ÁREA DO IMÓVEL (m")
        run3a.bold = True
        run3b = par3.add_run("2")
        run3b.bold = True
        run3b.font.superscript = True
        run3c = par3.add_run("):")
        run3c.bold = True
        area_formatada = f"{area_dxf:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        par3.add_run(f" {area_formatada}")

        doc_word.add_paragraph("")  # Uma linha em branco para separar


        # Descrição do perímetro, somente se o arquivo Excel foi fornecido
        if df is not None:
            initial = df.iloc[0]
            par_intro = doc_word.add_paragraph(style='Normal')
            par_intro.add_run("Pontos definidos pelas Coordenadas Planas no Sistema U.T.M. – SIRGAS 2000.\n\n")
            par_intro.add_run("Inicia-se a descrição deste perímetro no vértice ")
            
            run_v = par_intro.add_run(str(initial['V']))
            run_v.bold = True
            
            par_intro.add_run(f", de coordenadas N(Y) {initial['N']} e E(X) {initial['E']}, situado no limite com {initial['Confrontante']}.")



            num_points = len(df)
            for i in range(num_points):
                current = df.iloc[i]
                next_index = (i + 1) % num_points
                next_point = df.iloc[next_index]

                azimute = current['Azimute']
                distancia = current['Distancia(m)']
                confrontante = current['Confrontante']
                destino = next_point['V']
                coord_n = next_point['N']
                coord_e = next_point['E']

                if azimute.startswith("R=") and distancia.startswith("C="):
                    par = doc_word.add_paragraph(style='Normal')
                    par.add_run(f"Deste, segue com raio de {azimute[2:]}m e distância de {distancia[2:]}m, ")
                    par.add_run(f"confrontando neste trecho com {confrontante}, até o vértice ")
                    run_destino = par.add_run(str(destino))
                    run_destino.bold = True
                    par.add_run(f", de coordenadas N(Y) {coord_n} e E(X) {coord_e};")

                else:
                    par = doc_word.add_paragraph(style='Normal')
                    par.add_run(f"Deste, segue com azimute de {azimute} e distância de {distancia} m, ")
                    par.add_run(f"confrontando neste trecho com {confrontante}, até o vértice ")
                    run_destino = par.add_run(str(destino))
                    run_destino.bold = True
                    par.add_run(f", de coordenadas N(Y) {coord_n} e E(X) {coord_e};")

        else:
            # Caso não haja Excel, pode deixar espaço para preenchimento manual
            doc_word.add_paragraph("Descrição do perímetro não incluída neste memorial.", style='Normal')
            pular_linhas(doc_word, 8)

        doc_word.add_paragraph("")  # Uma linha em branco para separar
   
        # Adicionar o fechamento do perímetro e área
        # Formata perímetro e área com separador de milhar (ponto) e decimal (vírgula)
        perimetro_formatado = f"{perimeter_dxf:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        area_formatada = f"{area_dxf:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

        paragrafo_fechamento = doc_word.add_paragraph(
            f"Fechando-se assim o perímetro com {perimetro_formatado} m "
            f"e a área com {area_formatada} m².",
            style='Normal'
        )
        paragrafo_fechamento.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        doc_word.add_paragraph("")
        doc_word.add_paragraph("")
        # Adicionar data
        
        data_atual = obter_data_em_portugues()
        
       # Centralizar data
        paragrafo_data = doc_word.add_paragraph(f"Paraná, {data_atual}.", style='Normal')
        paragrafo_data.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc_word.add_paragraph("")  # Uma linha em branco para separar
        doc_word.add_paragraph("")  # Uma linha em branco para separar
        doc_word.add_paragraph("")  # Uma linha em branco para separar

        # Adicionar a imagem da assinatura centralizada
        assinatura = doc_word.add_paragraph()
        assinatura.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        #run = assinatura.add_run()
        #run.add_picture(r"C:\Users\Paulo\OneDrive\Documentos\JL_ADICIONAIS\TEMPLATE_MEMORIAL\assinatura_engenheiro.jpg", width=Inches(2.0))

        # Adicionar informações do engenheiro centralizadas
        infos_engenheiro = [
            "____________________",
            "Rodrigo Luis Schmitz",
            "Técnico em Agrimensura",
            "CFT: 045.300.139-44"
        ]

        for info in infos_engenheiro:
            paragrafo = doc_word.add_paragraph(info, style='Normal')
            paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


        # Salvar o documento
        doc_word.save(output_path)
        print(f"Memorial descritivo salvo em: {output_path}")
        
        log.write(f"Memorial descritivo salvo em: {output_path}\n")

    except Exception as e:
        print(f"Erro ao criar o documento memorial: {e}")
        
        log.write(f"Erro ao criar o documento memorial: {e}\n")


        
def convert_docx_to_pdf(output_path, pdf_file_path):
    """
    Converte um arquivo DOCX para PDF usando a biblioteca comtypes.
    """
    try:
        # Verificar se o arquivo DOCX existe antes de abrir
        if not os.path.exists(output_path):
            raise FileNotFoundError(f"Arquivo DOCX não encontrado: {output_path}")
           
        
        word = comtypes.client.CreateObject("Word.Application")
        word.Visible = False  # Ocultar a interface do Word
        doc = word.Documents.Open(output_path)
        doc.SaveAs(pdf_file_path, FileFormat=17)  # 17 corresponde ao formato PDF
        doc.Close()
        word.Quit()
        print(f"Arquivo PDF salvo em: {pdf_file_path}")
    except FileNotFoundError as fnf_error:
        print(f"Erro: {fnf_error}")
    except Exception as e:
        print(f"Erro ao converter DOCX para PDF: {e}")
    finally:
        try:
            word.Quit()
        except:
            pass  # Garantir que o Word seja fechado


