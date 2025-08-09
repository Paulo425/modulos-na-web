import pyproj
import folium
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
import contextily as ctx  # Adicionar no início do arquivo
import pandas as pd  # Adicionar no início do arquivo
import os
import math
import pandas as pd
import numpy
import matplotlib
matplotlib.use('Agg')  # ← fundamental essa linha antes do pyplot
import matplotlib.pyplot as plt
import matplotlib.ticker
import unicodedata
import scipy.stats
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from tqdm import tqdm
from datetime import datetime
from docx.oxml.shared import OxmlElement
from lxml import etree
from docx.oxml.ns import qn, nsdecls
import folium
from folium import plugins
import io
import base64
from PIL import Image
import contextily as ctx
import pandas as pd
# Módulos do python-docx para manipular parágrafos e XML
from docx.text.paragraph import Paragraph
from docx.oxml.shared import OxmlElement
from lxml import etree

from docx.oxml.ns import nsdecls, qn

from typing import Union

# Para seleção de múltiplas fotos e da planilha (file dialog).



from docx.oxml.ns import nsdecls
from lxml import etree

from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH

from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import numpy


from docx.oxml.shared import OxmlElement
from lxml import etree

from docx.oxml.ns import nsdecls, qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, Inches

from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import nsdecls, qn
from docx.oxml.shared import OxmlElement
from lxml import etree
import logging



from uuid import uuid4
import fitz  # PyMuPDF
from pathlib import Path
from itertools import chain
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn



import numpy as np
import sys


logger = logging.getLogger("meu_app_logger")
# Para garantir que o logger esteja configurado se o main.py executar separadamente:
if not logger.handlers:
    file_handler = logging.FileHandler('flask_app.log', encoding='utf-8')
    file_handler.setLevel(logging.DEBUG)
    file_formatter = logging.Formatter('%(asctime)s %(levelname)s : %(message)s')
    file_handler.setFormatter(file_formatter)

    console_handler = logging.StreamHandler(sys.stdout)
    console_handler.setLevel(logging.INFO)
    console_formatter = logging.Formatter('%(levelname)s : %(message)s')
    console_handler.setFormatter(console_formatter)

    logger.addHandler(file_handler)
    logger.addHandler(console_handler)

logger.info("✅ Logger MAIN.py inicializado corretamente!")

def normaliza_sim_nao(valor):
    try:
        import pandas as pd
        if pd.isna(valor):
            return None
    except Exception:
        pass
    if valor is None:
        return None
    s = str(valor).strip().upper()
    s = (s.replace("Ã", "A")
         .replace("Â", "A")
         .replace("Á", "A")
         .replace("Ó", "O")
         .replace("Õ", "O"))
    if s in {"S", "SIM", "TRUE", "1", "YES"}:
        return "SIM"
    if s in {"N", "NAO", "NÃO", "FALSE", "0", "NO"}:
        return "NAO"
    return None

def get_multi(dic_or_series, *keys):
    for k in keys:
        if k in dic_or_series:
            v = dic_or_series.get(k)
            try:
                import pandas as pd
                if pd.isna(v):
                    continue
            except Exception:
                pass
            if v is not None and str(v).strip() != "":
                return v
    return None

def fator_pavimentacao(valor):
    norm = normaliza_sim_nao(valor)
    return 1.0 if norm == "SIM" else 0.9

def fator_acessibilidade(valor):
    norm = normaliza_sim_nao(valor)
    return 1.0 if norm == "SIM" else 0.9


def adicionar_paragrafo_apos(paragrafo):
    """
    Cria e retorna um novo parágrafo imediatamente após o parágrafo informado.
    Compatível com python-docx.
    """
    novo_par_xml = OxmlElement('w:p')
    paragrafo._p.addnext(novo_par_xml)
    return Paragraph(novo_par_xml, paragrafo._parent)
###############################################################################
# FUNÇÕES DE SUPORTE GERAIS
###############################################################################
 
def inserir_paragrafo_apos(paragrafo, texto=''):
    novo_p = OxmlElement('w:p')
    paragrafo._p.addnext(novo_p)
    novo_paragrafo = Paragraph(novo_p, paragrafo._parent)
    if texto:
        novo_paragrafo.add_run(texto)
    return novo_paragrafo
###############################################################################
# FUNÇÕES DE SUPORTE GERAIS
###############################################################################
def remover_acentos(texto):
    """
    Remove acentos de uma string.
    """
    return ''.join(
        caractere 
        for caractere in unicodedata.normalize('NFD', texto)
        if unicodedata.category(caractere) != 'Mn'
    )

def formatar_moeda_brasil(valor):
    """
    Formata o valor em Real brasileiro, trocando ponto por vírgula
    e inserindo ponto a cada milhar.
    Exemplo: 12345.6 => 'R$ 12.345,60'
    """
    formato_texto = f"{valor:,.2f}"
    formato_texto = formato_texto.replace(",", "X").replace(".", ",").replace("X", ".")
    return f"R$ {formato_texto}"

def formatar_numero_brasileiro(valor):
    """
    Formata um número em padrão brasileiro, com separador de milhar e decimal invertidos.
    Exemplo: 12345.6 => '12.345,60'
    """
    formato_texto = f"{valor:,.2f}"
    formato_texto = formato_texto.replace(",", "X").replace(".", ",").replace("X", ".")
    return formato_texto

# ---------- NOVA FUNÇÃO ----------   (Ctrl + F  →  formatar_area_brasil)
def formatar_area_brasil(valor):
    """
    Formata áreas em m² no padrão brasileiro.
    • 200       → "200"
    • 2000      → "2.000"  
    • 80234,89  → "80.234,89"
    """
    try:
        num = float(str(valor))
    except Exception:
        return str(valor)

    # Se for inteiro, não precisa de casas decimais
    if num.is_integer():
        return f"{int(num):,}".replace(",", ".")
    else:
        # Para números com decimais
        txt = f"{num:,.2f}"
        txt = txt.replace(",", "X").replace(".", ",").replace("X", ".")
        return txt

def converter_valor_brasileiro_para_float(valor_str):
    """
    Converte um valor no formato brasileiro para float.
    Ex: "200,00" -> 200.0
    Ex: "1.200,00" -> 1200.0
    """
    # Remove primeiro os pontos (separadores de milhar)
    valor_str = valor_str.replace(".", "")
    # Depois converte a vírgula para ponto decimal
    valor_str = valor_str.replace(",", ".")
    return float(valor_str)


# ─────────────────────────────────────────────────────────────────────────────
# Distância entre dois pontos geográficos (Haversine) → resultado em quilômetros
# ─────────────────────────────────────────────────────────────────────────────
from math import radians, sin, cos, sqrt, atan2

def haversine_km(lat1, lon1, lat2, lon2):
    """
    Calcula a distância entre (lat1, lon1) e (lat2, lon2) em km.
    Entradas e saídas em float. Retorna 0.0 se qualquer coordenada faltar.
    """
    try:
        φ1, λ1, φ2, λ2 = map(radians, [float(lat1), float(lon1),
                                       float(lat2), float(lon2)])
    except Exception:
        return 0.0                     # coordenada ausente ou inválida

    dφ   = φ2 - φ1
    dλ   = λ2 - λ1
    a    = sin(dφ/2)**2 + cos(φ1)*cos(φ2)*sin(dλ/2)**2
    c    = 2 * atan2(sqrt(a), sqrt(1-a))
    R_km = 6371.0088                  # raio médio da Terra
    return R_km * c

###############################################################################
# FATORES ESPECÍFICOS (para homogeneização)
###############################################################################
def fator_aproveitamento(aproveitamento_texto):
    """
    Se aproveitamento for 'urbano' ou 'urbana', retorna 1.00
    Se for 'rural', retorna 0.80
    Caso contrário, retorna 1.00
    """
    if isinstance(aproveitamento_texto, str):
        valor = aproveitamento_texto.strip().lower()
        if valor in ["urbano", "urbana"]:
            return 1.00
        elif valor == "rural":
            return 0.80
    return 1.00

def fator_topografia(texto_topografia):
    """
    Se 'sim' em 'BOA TOPOGRAFIA?', retorna 1.10; senão, 1.00
    """
    if isinstance(texto_topografia, str):
        if "sim" in remover_acentos(texto_topografia.lower()):
            return 1.10
    return 1.00

def fator_pedologia(texto_pedologia):
    """
    Se 'sim' em 'PEDOLOGIA ALAGÁVEL?', retorna 0.70; senão, 1.00
    """
    if isinstance(texto_pedologia, str):
        if "sim" in remover_acentos(texto_pedologia.lower()):
            return 0.70
    return 1.00

def fator_pavimentacao(texto_pavimentacao):
    """
    Se 'sim' em 'PAVIMENTACAO?', retorna 1.00; senão, 0.90
    """
    if isinstance(texto_pavimentacao, str):
        if "sim" in remover_acentos(texto_pavimentacao.lower()):
            return 1.00
        else:
            return 0.90
    return 1.00

def fator_esquina(texto_esquina):
    """
    Se 'sim' em 'ESQUINA?', retorna 1.15; senão, 1.00
    """
    if isinstance(texto_esquina, str):
        if "sim" in remover_acentos(texto_esquina.lower()):
            return 1.15
    return 1.00

def fator_acessibilidade(texto_acessibilidade):
    """
    Se 'sim' em 'ACESSIBILIDADE?', retorna 1.00; senão, 0.90
    """
    if isinstance(texto_acessibilidade, str):
        if "sim" in remover_acentos(texto_acessibilidade.lower()):
            return 1.00
        else:
            return 0.90
    return 1.00


###############################################################################
# INTERVALO DE CONFIANÇA (IC) VIA BOOTSTRAP DA MEDIANA - 80%
###############################################################################
def intervalo_confianca_bootstrap_mediana(valores_numericos, numero_amostras=1000, nivel_confianca=0.80):
    """
    Calcula o intervalo de confiança (IC) para a mediana via bootstrap.
    Retorna (limite_inferior, limite_superior).
    """
    array_valores = numpy.array(valores_numericos)
    quantidade = len(array_valores)
    lista_medianas = []
    for _ in range(numero_amostras):
        amostra_sorteada = numpy.random.choice(array_valores, size=quantidade, replace=True)
        lista_medianas.append(numpy.median(amostra_sorteada))
    array_medianas = numpy.array(lista_medianas)
    limite_inferior = numpy.percentile(array_medianas, (1 - nivel_confianca) / 2 * 100)
    limite_superior = numpy.percentile(array_medianas, (1 + nivel_confianca) / 2 * 100)
    return limite_inferior, limite_superior


###############################################################################
# CLASSIFICAÇÃO DO GRAU DE PRECISÃO (ABNT NBR 14653)
###############################################################################
def classificar_grau_de_precisao(amplitude_ic80):
    """
    Classifica o resultado quanto à precisão, de acordo com a amplitude do IC 80%
    e a tabela da ABNT NBR 14653.
    """
    if amplitude_ic80 <= 30:
        return "GRAU III"
    elif amplitude_ic80 <= 40:
        return "GRAU II"
    elif amplitude_ic80 <= 50:
        return "GRAU I"
    else:
        return "NÃO CLASSIFICADO"


# ###############################################################################
# # GRÁFICO DE DENSIDADE (KDE)
# ###############################################################################
# def gerar_grafico_densidade_kernel(valores_homogeneizados, nome_arquivo):
#     """
#     Gera o gráfico de densidade (KDE) e salva em 'nome_arquivo'.
#     Se o conjunto de dados tiver menos de 2 elementos, salva um gráfico com uma mensagem de aviso.
#     """
#     from scipy.stats import gaussian_kde
#     import numpy as np
#     import matplotlib.pyplot as plt

#     array_valores = np.array(valores_homogeneizados, dtype=float)
#     if array_valores.size < 2:
#         # Dados insuficientes para calcular o KDE.
#         plt.figure(figsize=(8, 6))
#         plt.text(0.5, 0.5, "Dados insuficientes para calcular KDE", 
#                  horizontalalignment='center', verticalalignment='center', 
#                  transform=plt.gca().transAxes, fontsize=12)
#         plt.title("Histograma de Densidade de Kernel (KDE)")
#         plt.xlabel("Valores Homogeneizados")
#         plt.ylabel("Densidade")
#         plt.tight_layout()
#         plt.savefig(nome_arquivo, bbox_inches='tight')
#         plt.close()
#         return

#     media_valores = np.mean(array_valores)
#     mediana_valores = np.median(array_valores)

#     eixo_x = np.linspace(array_valores.min(), array_valores.max(), 300)
#     funcao_densidade = gaussian_kde(array_valores)
#     valores_densidade = funcao_densidade(eixo_x)

#     plt.figure(figsize=(8, 6))
#     plt.fill_between(eixo_x, valores_densidade, alpha=0.6)
#     plt.title("Histograma de Densidade de Kernel (KDE)")
#     plt.xlabel("Valores Homogeneizados")
#     plt.ylabel("Densidade")
#     plt.axvline(x=media_valores, color='red', linestyle='--', label=f"Média: {media_valores:,.2f}")
#     plt.axvline(x=mediana_valores, color='green', linestyle='-', label=f"Mediana: {mediana_valores:,.2f}")
#     plt.legend()
#     plt.tight_layout()
#     plt.savefig(nome_arquivo, bbox_inches='tight')
#     plt.close()


###############################################################################
# CÁLCULO DE FATORES BÁSICOS
###############################################################################
def calcular_fator_area(area_do_avaliando, area_da_amostra, usar_fator_area):
    """
    Calcula o fator área se usar_fator_area for True.
    Mantém a fórmula padrão: (Área da Amostra / Área do Avaliado)^(1/4),
    limitado a [0.5, 1.4].
    """
    if not usar_fator_area or area_do_avaliando <= 0 or area_da_amostra <= 0:
        return 1.0
    fator_calculado = (area_da_amostra / area_do_avaliando) ** 0.25
    return max(0.5, min(1.4, fator_calculado))

def limitar_fator(x):
    """
    Limita o valor do fator x ao intervalo [0.50, 2.0].
    """
    return max(0.50, min(1.4, x))

def calcular_fator_oferta(oferta_aplicada, usar_fator_oferta):
    """
    Retorna 0.9 se usar_fator_oferta e oferta_aplicada forem True; senão, 1.0.
    (Fator fixo)
    """
    return 0.9 if (usar_fator_oferta and oferta_aplicada) else 1.0



###############################################################################
# GRÁFICOS DE ADERÊNCIA E DISPERSÃO
###############################################################################
def gerar_grafico_aderencia_totais(dataframe, valores_homogeneizados_unitarios, nome_arquivo_imagem):
    """
    Gera um gráfico comparando os VALORES TOTAIS ORIGINAIS de cada amostra 
    com os VALORES TOTAIS ESTIMADOS, calculados a partir do valor unitário homogeneizado (R$/m²)
    multiplicado pela área de cada amostra.
    """
    import numpy as np
    import matplotlib.pyplot as plt
    from scipy.stats import linregress
    import matplotlib.ticker as ticker

    # 1) Obter os valores totais originais
    valores_originais_totais = dataframe["VALOR TOTAL"].tolist()

    # 2) Calcular os valores estimados
    valores_estimados_totais = []
    for i, valor_unit in enumerate(valores_homogeneizados_unitarios):
        area = dataframe.iloc[i]["AREA TOTAL"]
        if area > 0:
            valor_total_estimado = valor_unit * area
        else:
            valor_total_estimado = 0.0
        valores_estimados_totais.append(valor_total_estimado)

    x = np.array(valores_originais_totais, dtype=float)
    y = np.array(valores_estimados_totais, dtype=float)

    fig, ax = plt.subplots(figsize=(8, 6))
    ax.scatter(x, y, color='blue', label='Amostras')

    if x.size > 0 and y.size > 0:
        limite_min = min(np.min(x), np.min(y))
        limite_max = max(np.max(x), np.max(y))
    else:
        limite_min, limite_max = 0, 1

    if len(x) >= 2 and len(y) >= 2:
        slope, intercept, r_value, p_value, std_err = linregress(x, y)
        x_fit = np.linspace(limite_min, limite_max, 100)
        y_fit = slope * x_fit + intercept
        ax.plot(x_fit, y_fit, 'r-', label=f'Reta Ajustada (R² = {r_value**2:.2f})')
    else:
        ax.text(0.5, 0.5, "Dados insuficientes para regressão", 
                horizontalalignment='center', verticalalignment='center', 
                transform=ax.transAxes, fontsize=12, color='red')

    ax.set_title("Gráfico de Aderência - Valores Totais")
    ax.set_xlabel("Valor Total Original (R$)")
    ax.set_ylabel("Valor Total Estimado (R$)")
    ax.legend()
    ax.grid(True)
    ax.tick_params(axis='x', rotation=45)

    def formatar_valor_em_reais(valor, pos):
        return formatar_moeda_brasil(valor)

    formatador = ticker.FuncFormatter(formatar_valor_em_reais)
    ax.xaxis.set_major_formatter(formatador)
    ax.yaxis.set_major_formatter(formatador)

    fig.tight_layout()
    fig.savefig(nome_arquivo_imagem, bbox_inches='tight')
    plt.close(fig)


def gerar_grafico_dispersao_mediana(valores_homogeneizados, nome_arquivo):
    """
    Gera um gráfico de dispersão simples (index vs. valores homogeneizados)
    e destaca a mediana com uma linha horizontal.
    """
    import numpy as np
    import matplotlib.pyplot as plt

    arr = np.array(valores_homogeneizados, dtype=float)
    if arr.size < 1:
        plt.figure()
        plt.text(0.5, 0.5, "Sem valores para exibir", 
                 ha='center', va='center', 
                 transform=plt.gca().transAxes, fontsize=12)
        plt.title("Dispersão dos Valores Homogeneizados")
        plt.savefig(nome_arquivo, bbox_inches='tight')
        plt.close()
        return

    indices = np.arange(1, len(arr) + 1)

    plt.figure(figsize=(8, 6))
    plt.scatter(indices, arr, marker='o', label="Valores Homogeneizados")
    mediana = np.median(arr)
    plt.axhline(y=mediana, color='r', linestyle='--', label=f"Mediana: {mediana:,.2f}")

    plt.xlabel("Índice da Amostra")
    plt.ylabel("Valor Unitário Homogeneizado (R$/m²)")
    plt.title("Gráfico de Dispersão dos Valores Homogeneizados")
    plt.legend()
    plt.tight_layout()
    plt.savefig(nome_arquivo, bbox_inches='tight')
    plt.close()

####################################################################################################################
# MAPA DE AMOSTRAS DE MERCADO – versão 2025-05-29-B  (anti-sobreposição rigoroso)
#   • ponto “CENTRO MUNICÍPIO”      (opcional, como no código-fonte original)
#   • zoom adaptativo em 2 estágios (até ×16 tiles) → nitidez elevada
#   • DPI alto (default 700) e figure-size constante  → impressão / Word nítidos
#   • Amostras: laranja-claro, bolinha grande
#   • Legenda com posição automática: procura o canto sem pontos por baixo
#   • Margem automática 8 % (mantém rótulos 100 % dentro da moldura)
#   • Z-order ajustado → a estrela SEMPRE fica por cima de qualquer amostra
####################################################################################################################
def gerar_mapa_amostras(
    dataframe_amostras,
    dados_avaliando,
    nome_png : str = "mapa_amostras.png",
    width_in : float = 6.3,
    height_in: float = 9.0,
    dpi      : int   = 700,
    sharp    : int   = 2,            # 0=rápido · 1=4× · 2=16× tiles
):

    def _placeholder(path_png: str, msg="Mapa não disponível") -> str:
            from PIL import Image, ImageDraw, ImageFont
            W, H = int(width_in * 300), int(height_in * 300)
            img  = Image.new("RGB", (W, H), "#f0f0f0")
            draw = ImageDraw.Draw(img)
            try:
                font = ImageFont.truetype("arial.ttf", 46)
            except Exception:
                font = ImageFont.load_default()
            draw.multiline_text(
                (W // 2, H // 2),
                textwrap.fill(msg, 40),
                fill="#333", font=font,
                align="center", anchor="mm"
            )
            img.save(path_png, dpi=(300, 300))
            return str(Path(path_png).resolve())
    # ------------------------------------------------------------------ #
    # IMPORTS
    # ------------------------------------------------------------------ #
    import math, warnings, textwrap
    from pathlib import Path

    import pandas as pd
    import matplotlib.pyplot as plt
    from shapely.geometry import Point

    try:
        import geopandas as gpd
        import contextily as ctx
    except ImportError as e:     # fallback → imagem cinza
        return _placeholder(
            nome_png,
            f"Instale 'geopandas' + 'contextily' para gerar o mapa.\n{e}"
        )

    # ------------------------------------------------------------------ #
    # HELPERS
    # ------------------------------------------------------------------ #
    

    def _p(v):
        """Converte número/str → float ou None."""
        if pd.isna(v):                   return None
        if isinstance(v, (int, float)):  return float(v)
        try: return float(str(v).replace("°", "").replace(",", ".").strip())
        except Exception: return None

    # converte extensão (m) → escala de zoom base (0-18)
    def _zoom_base(lado_m):
        lado_por_tile = lado_m / 256
        return math.log2(156543.03 / lado_por_tile)

    # ------------------------------------------------------------------ #
    # COLETA DE PONTOS
    # ------------------------------------------------------------------ #
    
    coords = []
    for _, r in dataframe_amostras.iterrows():
        lat, lon = _p(r.get("LATITUDE")), _p(r.get("LONGITUDE"))
        if lat is not None and lon is not None:  # não use "if lat and lon"
            # usa idx se existir; cai para AM (da planilha) como fallback
            idx = int(r.get("idx", r.get("AM", 0)) or 0)
            rotulo = f"AM {idx:02d}" if idx > 0 else "AM"
            coords.append(dict(lat=lat, lon=lon, label=rotulo, tipo="amostra"))

    lat_av, lon_av = _p(dados_avaliando.get("LATITUDE")), _p(dados_avaliando.get("LONGITUDE"))
    if lat_av is not None and lon_av is not None:
        coords.append(dict(lat=lat_av, lon=lon_av, label="AVALIANDO", tipo="avaliando"))

    # Centro da cidade (usa geopy se disponível + nome da cidade)
    try:
        from geopy.geocoders import Nominatim
        if dados_avaliando.get("CIDADE"):
            geoloc = Nominatim(user_agent="app_centro").geocode(
                f"{dados_avaliando['CIDADE']}, Brasil", timeout=6
            )
            if geoloc:
                coords.append(dict(lat=geoloc.latitude, lon=geoloc.longitude,
                                   label="CENTRO", tipo="centro"))
    except Exception as e:
        warnings.warn(f"Centro município não incluído – geopy: {e}")

    if not coords:
        return _placeholder(nome_png, "Não há coordenadas válidas.")

    # ------------------------------------------------------------------ #
    # GDF → 3857
    # ------------------------------------------------------------------ #
    gdf = (
        gpd.GeoDataFrame(
            coords,
            geometry=[Point(c["lon"], c["lat"]) for c in coords],
            crs="EPSG:4326"
        )
        .to_crs(epsg=3857)
    )
    g_am = gdf[gdf.tipo == "amostra"]
    g_av = gdf[gdf.tipo == "avaliando"]
    g_ct = gdf[gdf.tipo == "centro"]

    # ------------------------------------------------------------------ #
    # FIGURE
    # ------------------------------------------------------------------ #
    fig, ax = plt.subplots(figsize=(width_in, height_in), dpi=dpi)
    plt.subplots_adjust(top=0.995, bottom=0.005, left=0.01, right=0.99)

    # --- estilos (tamanhos e z-orders) ---------------------------------------
    size_center    = 46          # 10 % maior do que antes
    size_amostras  = size_center # amostras = centro
    size_avaliando = 192         # estrela grande e destacada

    z_basemap  = 1
    z_amostra  = 3
    z_centro   = 4
    z_aval     = 5
    z_labels   = 6

    # --- plotagem ------------------------------------------------------------
    if not g_am.empty:
        g_am.plot(ax=ax, marker="o", color="#FFB347",   # laranja-claro
                  edgecolor="k", linewidth=0.3,
                  markersize=size_amostras,
                  label="Amostras",
                  zorder=z_amostra)

    if not g_av.empty:
        g_av.plot(ax=ax, marker="*", color="yellow",
                  edgecolor="k", linewidth=0.4,
                  markersize=size_avaliando,
                  label="Imóvel Avaliando",
                  zorder=z_aval)

    if not g_ct.empty:
        g_ct.plot(ax=ax, marker="o", color="red",
                  edgecolor="k", linewidth=0.3,
                  markersize=size_center,
                  label="Centro Município",
                  zorder=z_centro)

    # rótulos (para todos os pontos)
    for _, r in gdf.iterrows():
        ax.annotate(
            r.label,
            (r.geometry.x, r.geometry.y),
            xytext=(3, 3), textcoords="offset points",
            fontsize=5, weight="bold", zorder=z_labels,
            bbox=dict(boxstyle="round,pad=0.2",
                      facecolor=("white" if r.tipo == "centro" else "#FFEFD5"),
                      edgecolor="#333", alpha=.85)
        )

    # ------------------------------------------------------------------ #
    # ENQUADRAMENTO COM MARGEM (8 %)
    # ------------------------------------------------------------------ #
    xmin, ymin, xmax, ymax = gdf.total_bounds
    dx, dy = xmax - xmin, ymax - ymin
    if dx == 0 or dy == 0:        # todos pontos sobrepostos?
        dx = dy = max(dx, dy, 1)

    xmin, xmax = xmin - dx*0.08, xmax + dx*0.08
    ymin, ymax = ymin - dy*0.08, ymax + dy*0.08

    # mantém aspecto da figura
    fig_ratio  = height_in / width_in
    data_ratio = (ymax - ymin) / (xmax - xmin)
    if data_ratio > fig_ratio:    # acrescenta largura
        extra = (ymax - ymin)/fig_ratio - (xmax - xmin)
        xmin -= extra/2; xmax += extra/2
    else:                         # acrescenta altura
        extra = (xmax - xmin)*fig_ratio - (ymax - ymin)
        ymin -= extra/2; ymax += extra/2

    ax.set_xlim(xmin, xmax)
    ax.set_ylim(ymin, ymax)
    ax.set_aspect("equal")

    # ------------------------------------------------------------------ #
    # ZOOM DINÂMICO
    # ------------------------------------------------------------------ #
    lado_m   = max(xmax - xmin, ymax - ymin)
    zoom_lvl = int(round(_zoom_base(lado_m)))
    if lado_m > 13_000: sharp += 1           # área grande? +nitidez
    zoom_lvl = max(6, min(18, zoom_lvl + sharp))

    # ------------------------------------------------------------------ #
    # BASEMAP
    # ------------------------------------------------------------------ #
    try:
        ctx.add_basemap(
            ax,
            crs=gdf.crs,
            source=ctx.providers.Esri.WorldImagery,
            attribution="",
            zoom=zoom_lvl,
            zorder=z_basemap
        )
    except Exception as e:
        warnings.warn(f"Basemap falhou ({e}) – grade simples exibida.")
        ax.grid(True, alpha=0.3, zorder=0)

    ax.axis("off")

    # ------------------------------------------------------------------ #
    # LEGENDA  –  escolhe automaticamente o canto “livre”                #
    # ------------------------------------------------------------------ #
    # 1. converte pontos → coordenadas Axes (0-1)
    trans_axes = ax.transAxes.inverted()
    pts_axes   = [
        trans_axes.transform(ax.transData.transform((r.geometry.x, r.geometry.y)))
        for _, r in gdf.iterrows()
    ]

    # 2. configurações de “caixa” (larg≈0 .28 × alt≈0 .18  em Axes)
    box_w, box_h = 0.28, 0.18
    candidates = [
        ("upper left",  (0.02, 0.98-box_h)),
        ("upper right", (1-box_w-0.02, 0.98-box_h)),
        ("lower left",  (0.02, 0.02)),
        ("lower right", (1-box_w-0.02, 0.02)),
    ]

    def _box_free(x0, y0):
        x1, y1 = x0+box_w, y0+box_h
        for (x, y) in pts_axes:
            if x0 <= x <= x1 and y0 <= y <= y1:
                return False
        return True

    for loc, (bx, by) in candidates:
        if _box_free(bx, by):
            legend_loc, legend_anchor = loc, (bx, by+box_h)
            break
    else:                   # se todos ocupados, usa upper left mesmo
        legend_loc, legend_anchor = "upper left", (0.02, 0.98)

    leg = ax.legend(
        fontsize=6,
        loc=legend_loc,
        bbox_to_anchor=legend_anchor,
        frameon=True,
        framealpha=0.9,
        facecolor="white",
        edgecolor="#444",
        borderpad=0.6,
        labelspacing=0.7,
        handletextpad=0.9,
        borderaxespad=0.4
    )
    for handle in leg.legend_handles:   # marcadores maiores na legenda
        handle.set_sizes([60])

    # ------------------------------------------------------------------ #
    # SALVAR
    # ------------------------------------------------------------------ #
    fig.savefig(nome_png, dpi=dpi, bbox_inches="tight", pad_inches=0)
    plt.close(fig)
    return str(Path(nome_png).resolve())



import os
import fitz  # PyMuPDF
from uuid import uuid4
from pathlib import Path

def salvar_pdf_como_png(caminho_pdf, pasta_saida="static/temp", dpi=200):
    """
    Converte todas as páginas de um PDF em imagens PNG usando PyMuPDF (fitz).
    Retorna uma lista com os caminhos dos arquivos gerados.
    """
    caminhos_imagens = []
    try:
        pdf = fitz.open(caminho_pdf)
        nome_base = Path(caminho_pdf).stem

        for i in range(pdf.page_count):
            pagina = pdf.load_page(i)
            pix = pagina.get_pixmap(dpi=dpi)
            nome_arquivo = f"{nome_base}_{uuid4().hex[:6]}_{i}.png"
            caminho_completo = os.path.join(pasta_saida, nome_arquivo)
            pix.save(caminho_completo)
            caminhos_imagens.append(caminho_completo)
            logger.info(f"✅ Página {i+1}/{pdf.page_count} salva: {caminho_completo}")

        pdf.close()
        return caminhos_imagens

    except Exception as e:
        logger.error(f"❌ Erro ao converter PDF com fitz: {e}", exc_info=True)
        return []



###############################################################################
# TABELA DE AMOSTRAS HOMOGENEIZADAS
###############################################################################
def calcular_detalhes_amostras(dataframe_amostras_validas, dados_avaliando, fatores_do_usuario, finalidade_do_laudo):
    """
    Monta uma lista de dicionários com os detalhes e o 'Valor Total Homogeneizado'
    para gerar a tabela final.
    
    As colunas "VU" e "VUH" são calculadas, respectivamente, com os valores unitários originais
    e os valores unitários homogenizados.
    """
    import math

    lista_detalhes = []
    area_do_avaliando = float(dados_avaliando.get("AREA TOTAL", 0))

    # Fatores do Avaliado
    f_avaliado_aprov = fator_aproveitamento(dados_avaliando.get("APROVEITAMENTO", "URBANO"))
    f_avaliado_topog = fator_topografia(dados_avaliando.get("BOA TOPOGRAFIA?", "NÃO"))
    f_avaliado_pedol = fator_pedologia(dados_avaliando.get("PEDOLOGIA ALAGÁVEL? ", "NÃO"))
    f_avaliado_pavim = fator_pavimentacao(dados_avaliando.get("PAVIMENTACAO?", "NÃO"))
    f_avaliado_esq   = fator_esquina(dados_avaliando.get(" ESQUINA?", "NÃO"))
    f_avaliado_acess = fator_acessibilidade(dados_avaliando.get("ACESSIBILIDADE?", "NÃO"))

    for indice, linha in dataframe_amostras_validas.iterrows():
        identificador_amostra = str(linha.get("AM", indice + 1))
        valor_total_amostra = linha["VALOR TOTAL"]
        area_amostra = float(linha.get("AREA TOTAL", 0))

        # Fatores básicos
        fator_area = calcular_fator_area(area_do_avaliando, area_amostra, fatores_do_usuario["area"])
        fator_oferta = calcular_fator_oferta(True, fatores_do_usuario["oferta"])

        # Cálculo dos fatores individuais (f_avaliado / f_amostra)
        f_sample_aprov = fator_aproveitamento(linha.get("APROVEITAMENTO", "URBANO"))
        if fatores_do_usuario["aproveitamento"] and f_sample_aprov != 0:
            fator_aproveitamento_calculado = f_avaliado_aprov / f_sample_aprov
        else:
            fator_aproveitamento_calculado = 1.0

        f_sample_topog = fator_topografia(linha.get("BOA TOPOGRAFIA?", "NÃO"))
        if fatores_do_usuario["topografia"] and f_sample_topog != 0:
            fator_topografia_calculado = f_avaliado_topog / f_sample_topog
        else:
            fator_topografia_calculado = 1.0

        f_sample_pedol = fator_pedologia(linha.get("PEDOLOGIA ALAGÁVEL? ", "NÃO"))
        if fatores_do_usuario["pedologia"] and f_sample_pedol != 0:
            fator_pedologia_calculado = f_avaliado_pedol / f_sample_pedol
        else:
            fator_pedologia_calculado = 1.0

        f_sample_pavim = fator_pavimentacao(linha.get("PAVIMENTACAO?", "NÃO"))
        if fatores_do_usuario["pavimentacao"] and f_sample_pavim != 0:
            fator_pavimentacao_calculado = f_avaliado_pavim / f_sample_pavim
        else:
            fator_pavimentacao_calculado = 1.0

        f_sample_esq = fator_esquina(linha.get(" ESQUINA?", "NÃO"))
        if fatores_do_usuario["esquina"] and f_sample_esq != 0:
            fator_esquina_calculado = f_avaliado_esq / f_sample_esq
        else:
            fator_esquina_calculado = 1.0

        f_sample_acess = fator_acessibilidade(linha.get("ACESSIBILIDADE?", "NÃO"))
        if fatores_do_usuario["acessibilidade"] and f_sample_acess != 0:
            fator_acessibilidade_calculado = f_avaliado_acess / f_sample_acess
        else:
            fator_acessibilidade_calculado = 1.0

        # Fator localização
        if fatores_do_usuario.get("localizacao_mesma_regiao", False):
            fator_localizacao_calculado = 1.0
        else:
            try:
                distancia_amostra = float(linha.get("DISTANCIA CENTRO", 0))
                distancia_avaliando = float(dados_avaliando.get("DISTANCIA CENTRO", 0))
                if distancia_amostra > 0 and distancia_avaliando > 0:
                    fator_item_comparativo = 1 / math.pow(distancia_amostra, 1/10)
                    fator_bem_avaliando = 1 / math.pow(distancia_avaliando, 1/10)
                    fator_localizacao_calculado = fator_bem_avaliando / fator_item_comparativo
                    fator_localizacao_calculado = limitar_localizacao(fator_bem_avaliando / fator_item_comparativo)
                else:
                    fator_localizacao_calculado = 1.0
            except:
                fator_localizacao_calculado = 1.0

        # Cálculo do valor total homogenizado aplicando todos os fatores
        valor_total_homogeneizado = (
            valor_total_amostra *
            fator_area *
            fator_oferta *
            fator_localizacao_calculado *
            fator_aproveitamento_calculado *
            fator_topografia_calculado *
            fator_pedologia_calculado *
            fator_pavimentacao_calculado *
            fator_esquina_calculado *
            fator_acessibilidade_calculado
        )

        # Cálculo dos valores unitários:
        # VU  => Valor unitário original (da planilha)
        # VUH => Valor unitário homogenizado
        if area_amostra > 0:
            vu = valor_total_amostra / area_amostra
            vuh = valor_total_homogeneizado / area_amostra
        else:
            vu = 0.0
            vuh = 0.0

        linha_detalhes = {
            "AM": identificador_amostra,
            "AREA": formatar_numero_brasileiro(area_amostra),
            "VU": formatar_moeda_brasil(vu),
            "FA": f"{fator_area:.2f}",
            "FO": f"{fator_oferta:.2f}",
            "FAP": f"{fator_aproveitamento_calculado:.2f}",
            "FT": f"{fator_topografia_calculado:.2f}",
            "FP": f"{fator_pedologia_calculado:.2f}",
            "FPA": f"{fator_pavimentacao_calculado:.2f}",
            "FE": f"{fator_esquina_calculado:.2f}",
            "FAC": f"{fator_acessibilidade_calculado:.2f}",
            "FL": f"{fator_localizacao_calculado:.2f}",
            "VUH": formatar_moeda_brasil(vuh)
        }
        lista_detalhes.append(linha_detalhes)

    return lista_detalhes
    
##############################################################################################################
#MONTAGEM DA TABELA DE AMOSTRAS HOMOGENEIZADAS
##############################################################################################################
# --------------------------------------------------------------
# >>>  inserir_tabela_amostras_calculadas
# --------------------------------------------------------------
def inserir_tabela_amostras_calculadas(documento, lista_detalhes, col_widths=None):
    """
    Insere, após o marcador [tabelaSimilares], a tabela de amostras
    homogeneizadas com:
    • Cabeçalhos: fundo azul‑claro
    • Coluna VUH inteira: fundo verde‑claro
    • Fatores limitados ao intervalo [0.50, 2.00] com 2 casas decimais
    """
    from docx.shared      import Pt, Inches
    from docx.oxml        import parse_xml
    from docx.oxml.ns     import nsdecls
    from docx.enum.text   import WD_ALIGN_PARAGRAPH
    from docx.enum.table  import WD_TABLE_ALIGNMENT

    if not lista_detalhes:
        return

    # ---- Larguras padrão (pol) -------------------------------------------
    if col_widths is None:
        col_widths = [0.6, 1.2, 1.5] + [0.6]*9 + [1.5]

    nomes    = [
        "AM","AREA","VU",
        "FA","FO","FAP","FT","FP","FPA",
        "FE","FAC","FL","VUH"
    ]
    fatores  = {"FA","FO","FAP","FT","FP","FPA","FE","FAC","FL"}

    # ---- Sombras ----------------------------------------------------------
    def _shading(fill_hex):   # cria um <w:shd ... w:fill="XXXXXX"/>
        return etree.fromstring(
            r'<w:shd {} w:val="clear" w:fill="{}"/>'.format(nsdecls('w'), fill_hex)
        )
    azul  = "BDD7EE"   # cabeçalhos
    verde = "C6E0B4"   # VUH

    # ---- Procura o marcador ----------------------------------------------
    for par in documento.paragraphs:
        if "[tabelaSimilares]" not in par.text:
            continue

        par.text = par.text.replace("[tabelaSimilares]", "")

        rows = len(lista_detalhes) + 1
        tbl  = documento.add_table(rows=rows, cols=len(nomes))
        tbl.style, tbl.alignment, tbl.allow_autofit = "Table Grid", WD_TABLE_ALIGNMENT.CENTER, False

        # Larguras
        for ci, w in enumerate(col_widths):
            for r in tbl.rows:
                r.cells[ci].width = Inches(w)

        # Cabeçalho
        hdr = tbl.rows[0]
        for ci, rotulo in enumerate(nomes):
            c = hdr.cells[ci]
            c.text = rotulo
            c._tc.get_or_add_tcPr().append(_shading(azul))
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name, run.font.size, run.font.bold = "Arial", Pt(9), True

        # Dados
        for li, am in enumerate(lista_detalhes, start=1):
            for ci, campo in enumerate(nomes):
                val = am.get(campo, "")
                if campo in fatores:
                    try:
                        val = f"{limitar_fator(float(val)):.2f}"
                    except Exception:
                        val = str(val)
                cell = tbl.rows[li].cells[ci]
                cell.text = str(val)

                # pinta a coluna VUH de verde
                if campo == "VUH":
                    cell._tc.get_or_add_tcPr().append(_shading(verde))

                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name  = "Arial"
                    run.font.size  = Pt(8 if campo in {"VU","VUH"} else 9)

        # Reposiciona a tabela logo após o parágrafo do marcador
        par._p.addnext(tbl._element)

        # Legenda
        leg1 = inserir_paragrafo_apos(par, "")
        leg2 = inserir_paragrafo_apos(
            leg1,
            "Legendas:\n"
            "- AM = Amostra\n"
            "- AREA = Área do Imóvel (m²)\n"
            "- VU = Valor Unitário Ofertado\n"
            "- FA = Fator Área\n"
            "- FO = Fator Oferta\n"
            "- FAP = Fator Aproveitamento\n"
            "- FT = Fator Topografia\n"
            "- FP = Fator Pedologia\n"
            "- FPA = Fator Pavimentação\n"
            "- FE = Fator Esquina\n"
            "- FAC = Fator Acessibilidade\n"
            "- FL = Fator Localização\n"
            "- VUH = Valor Unitário Homogeneizado\n"
        )
        for run in leg2.runs:
            run.font.name, run.font.size = "Arial", Pt(9)
        break
# --------------------------------------------------------------
# <<<  inserir_tabela_amostras_calculadas
# --------------------------------------------------------------



# #######################################################################
# # FUNÇÕES DE FORMATAÇÃO
# #######################################################################
# def inserir_tabela_amostras_originais(documento, dataframe):
#     """
#     Substitui o placeholder [amostras original] pela tabela de amostras originais,
#     com as colunas: AM, VALOR TOTAL, ÁREA TOTAL (m²), VALOR UNITÁRIO (R$/m²), CIDADE, FONTE.
#     Agora, deixamos um espaço um pouco maior entre as linhas.
#     """
#     from docx.shared import Pt, Inches
#     from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL
#     from docx.enum.text import WD_ALIGN_PARAGRAPH
#     from docx.oxml.shared import OxmlElement
#     from lxml import etree

#     from docx.oxml.ns import nsdecls, qn

#     logger.info(f"🔎 DataFrame recebido em inserir_tabela_amostras_originais:\n{dataframe.head()}")
#     logger.info(f"🔎 Colunas recebidas: {list(dataframe.columns)}")


#     # Ajuste conforme as larguras desejadas (em polegadas) para cada coluna
#     col_widths = [0.2, 1.3, 1.1, 0.8, 2.0, 2.9]

#     # Títulos visíveis no cabeçalho
#     colunas_visiveis = [
#         "AM",
#         "VALOR TOTAL",
#         "ÁREA TOTAL (m²)",
#         "VALOR UNITÁRIO (R$/m²)",
#         "CIDADE",
#         "FONTE"
#     ]

#     # Colunas correspondentes do DataFrame (caso precise filtrar ou renomear)
#     colunas_df = [
#         "idx",
#         "VALOR TOTAL",
#         "AREA TOTAL",
#         "valor_unitario",
#         "cidade",
#         "fonte"
#     ]


#     # Localiza o parágrafo onde o placeholder [amostras original] está
#     paragrafo_alvo = None
#     for paragrafo in documento.paragraphs:
#         if "[amostras original]" in paragrafo.text:
#             paragrafo_alvo = paragrafo
#             break

#     # Se não encontrou o placeholder, não faz nada
#     if not paragrafo_alvo:
#         return

#     # Remove o texto do placeholder
#     paragrafo_alvo.text = paragrafo_alvo.text.replace("[amostras original]", "")

#     # Número de linhas = registros do dataframe + 1 (para o cabeçalho)
#     num_linhas = len(dataframe) + 1
#     # Número de colunas = quantidade de títulos visíveis
#     num_colunas = len(colunas_visiveis)

#     # Cria a tabela
#     tabela = documento.add_table(rows=num_linhas, cols=num_colunas, style="Table Grid")
#     tabela.allow_autofit = False
#     tabela.alignment = WD_TABLE_ALIGNMENT.CENTER

#     # Função para centralizar verticalmente a célula
#     def set_vertical_alignment(cell):
#         tcPr = cell._tc.get_or_add_tcPr()
#         vAlign = OxmlElement('w:vAlign')
#         vAlign.set(qn('w:val'), "center")
#         tcPr.append(vAlign)

#     # --- Cabeçalho ---
#     for c, titulo_exib in enumerate(colunas_visiveis):
#         cell_header = tabela.rows[0].cells[c]
#         cell_header.text = titulo_exib

#         # Fundo azul claro no cabeçalho
#         shading_xml = etree.fromstring(
#             f'<w:shd {nsdecls("w")} w:fill="BDD7EE" w:val="clear"/>'
#         )
#         cell_header._tc.get_or_add_tcPr().append(shading_xml)

#         # Formatação da fonte do cabeçalho
#         for run in cell_header.paragraphs[0].runs:
#             run.font.name = "Arial"
#             run.font.size = Pt(10)
#             run.font.bold = True

#         # Alinhamento horizontal e vertical do cabeçalho
#         cell_header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
#         set_vertical_alignment(cell_header)

#     # --- Linhas de dados ---
#     for i, (_, row) in enumerate(dataframe.iterrows(), start=1):
#         # Monta a lista de valores (na mesma ordem das colunas do cabeçalho)
#         valores_linha = []

#         # AM
#         am_str = str(row.get("AM", ""))
#         valores_linha.append(am_str)

#         # VALOR TOTAL (exemplo de formatação de moeda)
#         try:
#             vt_str = formatar_moeda_brasil(float(row["VALOR TOTAL"]))
#         except:
#             vt_str = str(row.get("VALOR TOTAL", ""))
#         valores_linha.append(vt_str)

#         # ÁREA TOTAL
#         try:
#             area_str = formatar_numero_brasileiro(float(row["AREA TOTAL"]))
#         except:
#             area_str = str(row.get("AREA TOTAL", ""))
#         valores_linha.append(area_str)

#         # VALOR UNITÁRIO
#         try:
#             vu_str = formatar_moeda_brasil(float(row["VALOR UNITARIO"]))
#         except:
#             vu_str = str(row.get("VALOR UNITARIO", ""))
#         valores_linha.append(vu_str)

#         # CIDADE
#         cidade_str = str(row.get("CIDADE", ""))
#         valores_linha.append(cidade_str)

#         # FONTE
#         fonte_str = str(row.get("FONTE", ""))
#         valores_linha.append(fonte_str)

#         # Preenche as células
#         for col_index, valor_cel in enumerate(valores_linha):
#             cell_data = tabela.rows[i].cells[col_index]
#             cell_data.text = valor_cel

#             # Formatação da fonte das células de dados
#             for run in cell_data.paragraphs[0].runs:
#                 run.font.name = "Arial"
#                 run.font.size = Pt(8)
#                 run.font.bold = False

#             # Alinhamento horizontal
#             cell_data.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

#             # Espaçamento vertical dentro da célula
#             paragraph_format = cell_data.paragraphs[0].paragraph_format
#             paragraph_format.space_before = Pt(2)
#             paragraph_format.space_after = Pt(2)

#             # Alinhamento vertical
#             set_vertical_alignment(cell_data)

#     # --- Ajuste de altura das linhas e largura das colunas ---
#     for row_index in range(num_linhas):
#         if row_index == 0:
#             # Aumenta a altura da linha do cabeçalho
#             tabela.rows[row_index].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
#             tabela.rows[row_index].height = Pt(40)
#         else:
#             # Aumenta a altura das linhas de dados
#             tabela.rows[row_index].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
#             tabela.rows[row_index].height = Pt(26)

#         # Ajusta a largura de cada coluna
#         for col_index, w_inch in enumerate(col_widths):
#             tabela.rows[row_index].cells[col_index].width = Inches(w_inch)

#     # Insere a tabela logo depois do parágrafo alvo
#     paragrafo_alvo._p.addnext(tabela._element)



###############################################################################
# >>>>>>> TEXTO DETALHADO DE SANEAMENTO (CHAUVENET) <<<<<<
###############################################################################
def gerar_texto_saneamento_chauvenet_extremamente_detalhado(
    dataframe_inicial,
    dataframe_filtrado,
    indices_excluidos,
    amostras_excluidas,
    media,
    desvio_padrao,
    menor_valor,
    maior_valor,
    mediana_valor
):
    """
    Gera um texto completo sobre o critério de Chauvenet e as estatísticas.
    """
    n_inicial = len(dataframe_inicial)
    n_filtrado = len(dataframe_filtrado)
    n_eliminadas = len(indices_excluidos)

    valores_filtrados_para_IC = dataframe_filtrado["VALOR TOTAL"].values
    if len(valores_filtrados_para_IC) > 1:
        limite_inf_ic, limite_sup_ic = intervalo_confianca_bootstrap_mediana(valores_filtrados_para_IC, 1000, 0.80)
        amplitude_ic = ((limite_sup_ic - limite_inf_ic) / numpy.median(valores_filtrados_para_IC)) * 100.0
    else:
        amplitude_ic = 0.0

    if desvio_padrao > 0:
        ds_menor = abs(media - menor_valor) / desvio_padrao
        ds_maior = abs(maior_valor - mediana_valor) / desvio_padrao
    else:
        ds_menor = 0.0
        ds_maior = 0.0

    if n_eliminadas > 0:
        identificacoes_excluidas = ", ".join(amostras_excluidas)
    else:
        identificacoes_excluidas = "Nenhuma"

    texto_exemplo = (
        f"-SANEAMENTO DOS DADOS AMOSTRAIS (CRITÉRIO DE CHAUVENET)\n"
        f"Quantidade de Amostras Válidas: {n_filtrado} unid.\n\n"
        f"-TESTANDO A AMOSTRA DE VALOR MAIS REDUZIDO-\n"
        f"D/S calc. = (Média - Menor Valor) / Desvio Padrão\n"
        f"D/S calc. para o MENOR valor = {ds_menor:.4f}\n\n"
        f"-TESTANDO A AMOSTRA DE VALOR MAIS ELEVADO-:\n"
        f"D/S calc. = (Maior Valor - Mediana) / Desvio Padrão\n"
        f"D/S calc. para o MAIOR valor = {ds_maior:.4f}\n\n"
        f"-CONCLUSÃO-:\n"
        f"* ALGUMAS AMOSTRAS PODEM NÃO SER VÁLIDAS (caso não satisfaçam o critério de Chauvenet) *\n\n"
        f"-RESUMO ESTATÍSTICO DA AVALIAÇÃO-\n"
        f"Quantidade de Amostras Válidas (Utilizadas no cálculo): {n_filtrado}\n"
        f"Quantidade de Amostras Eliminadas pelo critério de Chauvenet: {n_eliminadas}\n"
        f"Identificação das Amostras Eliminadas: {identificacoes_excluidas}\n"
    )

    return texto_exemplo

def inserir_texto_saneamento_no_placeholder(documento, marcador_placeholder, texto_saneamento):
    """
    Substitui o placeholder [texto_relatorio_resumo_saneamento] por um texto explicativo.
    """
    for paragrafo in documento.paragraphs:
        if marcador_placeholder in paragrafo.text:
            paragrafo.text = paragrafo.text.replace(marcador_placeholder, "")
            paragrafo_titulo = inserir_paragrafo_apos(paragrafo, "• SANEAMENTO DOS DADOS AMOSTRAIS\n")
            for execucao in paragrafo_titulo.runs:
                execucao.font.name = "Arial"
                execucao.font.size = Pt(12)
                execucao.font.bold = True
            paragrafo_titulo.paragraph_format.line_spacing = 1.15

            paragrafo_saneamento = inserir_paragrafo_apos(paragrafo_titulo, "")
            linhas = texto_saneamento.split("\n")
            for linha_texto in linhas:
                execucao_linha = paragrafo_saneamento.add_run(linha_texto + "\n")
                execucao_linha.font.name = "Arial"
                execucao_linha.font.size = Pt(12)
                execucao_linha.font.bold = False

            paragrafo_saneamento.paragraph_format.line_spacing = 1.15
            break


###############################################################################
# MEMÓRIA DE CÁLCULO DETALHADA
###############################################################################
def gerar_lista_memoria_calculo(dataframe_amostras, dados_avaliando, fatores_do_usuario, finalidade_do_laudo):
    import math
    
    lista_memoria_completa = []
    area_do_avaliando = float(dados_avaliando.get("AREA TOTAL", 0))

    # Fatores do Avaliado (utilizando as funções auxiliares já definidas)
    f_avaliado_aprov = fator_aproveitamento(dados_avaliando.get("APROVEITAMENTO", "URBANO"))
    f_avaliado_topog = fator_topografia(dados_avaliando.get("BOA TOPOGRAFIA?", "NÃO"))
    f_avaliado_pedol = fator_pedologia(dados_avaliando.get("PEDOLOGIA ALAGÁVEL? ", "NÃO"))
    f_avaliado_pavim = fator_pavimentacao(dados_avaliando.get("PAVIMENTACAO?", "NÃO"))
    f_avaliado_esq   = fator_esquina(dados_avaliando.get(" ESQUINA?", "NÃO"))
    f_avaliado_acess = fator_acessibilidade(dados_avaliando.get("ACESSIBILIDADE?", "NÃO"))

    for indice, linha in dataframe_amostras.iterrows():
        identificador_amostra = str(linha.get("AM", indice+1))
        valor_total = linha["VALOR TOTAL"]
        area_da_amostra = float(linha.get("AREA TOTAL", 0))

        # 1) Cálculo dos fatores básicos
        fator_area = calcular_fator_area(area_do_avaliando, area_da_amostra, fatores_do_usuario["area"])
        fator_oferta = calcular_fator_oferta(True, fatores_do_usuario["oferta"])
        
        # Fator Aproveitamento
        f_sample_aprov = fator_aproveitamento(linha.get("APROVEITAMENTO", "URBANO"))
        if fatores_do_usuario["aproveitamento"] and f_sample_aprov != 0:
            fator_aproveitamento_calculado = limitar_fator(f_avaliado_aprov / f_sample_aprov)
        else:
            fator_aproveitamento_calculado = 1.0

        # Fator Topografia
        f_sample_topog = fator_topografia(linha.get("BOA TOPOGRAFIA?", "NÃO"))
        if fatores_do_usuario["topografia"] and f_sample_topog != 0:
            fator_topografia_calculado = limitar_fator(f_avaliado_topog / f_sample_topog)
        else:
            fator_topografia_calculado = 1.0

        # Fator Pedologia
        f_sample_pedol = fator_pedologia(linha.get("PEDOLOGIA ALAGÁVEL? ", "NÃO"))
        if fatores_do_usuario["pedologia"] and f_sample_pedol != 0:
            fator_pedologia_calculado = limitar_fator(f_avaliado_pedol / f_sample_pedol)
        else:
            fator_pedologia_calculado = 1.0

        # -------- Fator Pavimentação --------
        amostra_pav = get_multi(linha, "PAVIMENTACAO?", "PAVIMENTAÇÃO?", "PAVIMENTACAO ?")
        aval_pav    = get_multi(dados_avaliando, "PAVIMENTACAO?", "PAVIMENTAÇÃO?")
        f_sample_pavim   = fator_pavimentacao(amostra_pav)
        f_avaliado_pavim = fator_pavimentacao(aval_pav)

        # Se amostra vier vazia/None ⇒ usa o mesmo do avaliado (garante 1.00 em SIM/SIM e NÃO/NÃO)
        if amostra_pav is None:
            f_sample_pavim = f_avaliado_pavim

        if fatores_do_usuario["pavimentacao"] and f_sample_pavim != 0:
            fator_pavimentacao_calculado = limitar_fator(f_avaliado_pavim / f_sample_pavim)
        else:
            fator_pavimentacao_calculado = 1.0

        # Fator Esquina
        f_sample_esq = fator_esquina(linha.get(" ESQUINA?", "NÃO"))
        if fatores_do_usuario["esquina"] and f_sample_esq != 0:
            fator_esquina_calculado = limitar_fator(f_avaliado_esq / f_sample_esq)
        else:
            fator_esquina_calculado = 1.0

        logger.info(f"[FPA] AM={identificador_amostra} | Avaliado={aval_pav} ({f_avaliado_pavim}) "
            f"| Amostra={amostra_pav} ({f_sample_pavim}) => FPA={fator_pavimentacao_calculado:.2f}")

        # -------- Fator Acessibilidade --------
        amostra_ace = get_multi(linha, "ACESSIBILIDADE?", "ACESSIBILIDADE ?")
        aval_ace    = get_multi(dados_avaliando, "ACESSIBILIDADE?")
        f_sample_acess   = fator_acessibilidade(amostra_ace)
        f_avaliado_acess = fator_acessibilidade(aval_ace)

        if amostra_ace is None:
            f_sample_acess = f_avaliado_acess

        if fatores_do_usuario["acessibilidade"] and f_sample_acess != 0:
            fator_acessibilidade_calculado = limitar_fator(f_avaliado_acess / f_sample_acess)
        else:
            fator_acessibilidade_calculado = 1.0 
                
        logger.info(f"[FAC] AM={identificador_amostra} | Avaliado={aval_ace} ({f_avaliado_acess}) "
            f"| Amostra={amostra_ace} ({f_sample_acess}) => FAC={fator_acessibilidade_calculado:.2f}")     
                  
        # Fator Localização
        if fatores_do_usuario.get("localizacao_mesma_regiao", False):
            fator_localizacao_calculado = 1.0
        else:
            try:
                distancia_amostra = float(linha.get("DISTANCIA CENTRO", 0))
                distancia_avaliando = float(dados_avaliando.get("DISTANCIA CENTRO", 0))
                if distancia_amostra > 0 and distancia_avaliando > 0:
                    fator_item_comparativo = 1 / (distancia_amostra ** 0.1)
                    fator_bem_avaliando   = 1 / (distancia_avaliando ** 0.1)
                    fator_localizacao_calculado = limitar_localizacao(fator_bem_avaliando / fator_item_comparativo)
                else:
                    fator_localizacao_calculado = 1.0
            except:
                fator_localizacao_calculado = 1.0

        # 2) Cálculo do Valor Total Homogeneizado
        valor_total_homogeneizado = (
            valor_total *
            fator_area *
            fator_oferta *
            fator_localizacao_calculado *
            fator_aproveitamento_calculado *
            fator_topografia_calculado *
            fator_pedologia_calculado *
            fator_pavimentacao_calculado *
            fator_esquina_calculado *
            fator_acessibilidade_calculado
        )

        # 3) Monta o texto de memória de cálculo
        bloco_texto = []
        bloco_texto.append(f"AM {identificador_amostra}")
        bloco_texto.append("")
        # Mantém o texto original para o valor ofertado:
        bloco_texto.append(f"- VALOR TOTAL OFERTADO: {formatar_moeda_brasil(valor_total)}")
        bloco_texto.append(f"- ÁREA DA AMOSTRA (m²): {formatar_numero_brasileiro(area_da_amostra)}")
        bloco_texto.append("")

        bloco_texto.append("- Fator Área:")
        bloco_texto.append(f"   Avaliado: {formatar_numero_brasileiro(area_do_avaliando)}")
        bloco_texto.append(f"   Amostra: {formatar_numero_brasileiro(area_da_amostra)} - Cálculo => {fator_area:.2f}\n")

        bloco_texto.append("- Fator Oferta:")
        bloco_texto.append(f"   (fixo 0.90 se habilitado) => {fator_oferta:.2f}\n")

        bloco_texto.append("- Fator Aproveitamento (f_avaliado / f_amostra):")
        bloco_texto.append(f"   Avaliado: {f_avaliado_aprov:.2f}")
        bloco_texto.append(f"   Amostra: {f_sample_aprov:.2f}")
        bloco_texto.append(f"   => {fator_aproveitamento_calculado:.2f}\n")

        bloco_texto.append("- Fator Topografia (f_avaliado / f_amostra):")
        bloco_texto.append(f"   Avaliado: {f_avaliado_topog:.2f}")
        bloco_texto.append(f"   Amostra: {f_sample_topog:.2f}")
        bloco_texto.append(f"   => {fator_topografia_calculado:.2f}\n")

        bloco_texto.append("- Fator Pedologia (f_avaliado / f_amostra):")
        bloco_texto.append(f"   Avaliado: {f_avaliado_pedol:.2f}")
        bloco_texto.append(f"   Amostra: {f_sample_pedol:.2f}")
        bloco_texto.append(f"   => {fator_pedologia_calculado:.2f}\n")

        bloco_texto.append("- Fator Pavimentação (f_avaliado / f_amostra):")
        bloco_texto.append(f"   Avaliado: {f_avaliado_pavim:.2f}")
        bloco_texto.append(f"   Amostra: {f_sample_pavim:.2f}")
        bloco_texto.append(f"   => {fator_pavimentacao_calculado:.2f}\n")

        bloco_texto.append("- Fator Esquina (f_avaliado / f_amostra):")
        bloco_texto.append(f"   Avaliado: {f_avaliado_esq:.2f}")
        bloco_texto.append(f"   Amostra: {f_sample_esq:.2f}")
        bloco_texto.append(f"   => {fator_esquina_calculado:.2f}\n")

        bloco_texto.append("- Fator Acessibilidade (f_avaliado / f_amostra):")
        bloco_texto.append(f"   Avaliado: {f_avaliado_acess:.2f}")
        bloco_texto.append(f"   Amostra: {f_sample_acess:.2f}")
        bloco_texto.append(f"   => {fator_acessibilidade_calculado:.2f}\n")

        bloco_texto.append("- Fator Localização:")
        bloco_texto.append(f"   => {fator_localizacao_calculado:.2f}\n")

        # 4) Em vez de exibir o Valor Total Homogeneizado, agora exibe o Valor Unitário Homogeneizado (VUH)
        if area_da_amostra > 0:
            valor_unit_homog = valor_total_homogeneizado / area_da_amostra
        else:
            valor_unit_homog = 0.0

        bloco_texto.append(
            f"=> VUH (Valor Unitário Homogeneizado): {formatar_moeda_brasil(valor_unit_homog)}"
        )

        lista_memoria_completa.append("\n".join(bloco_texto))

    return lista_memoria_completa



def inserir_texto_memoria_calculo_no_placeholder(documento, marcador_placeholder, lista_memorias):
    for paragrafo in documento.paragraphs:
        if marcador_placeholder in paragrafo.text:
            paragrafo.text = paragrafo.text.replace(marcador_placeholder, "")
            paragrafo_atual = paragrafo
            for indice_bloco, bloco in enumerate(lista_memorias):
                if indice_bloco >= 1:
                    paragrafo_branco = inserir_paragrafo_apos(paragrafo_atual, "")
                    execucao_branco = paragrafo_branco.add_run("\n")
                    execucao_branco.font.size = Pt(10)
                    execucao_branco.font.name = "Arial"
                    paragrafo_atual = paragrafo_branco

                novo_paragrafo = inserir_paragrafo_apos(paragrafo_atual, "")
                linhas_texto = bloco.split("\n")

                for indice_linha, conteudo_linha in enumerate(linhas_texto):
                    execucao_texto = novo_paragrafo.add_run(conteudo_linha + "\n")
                    execucao_texto.font.name = "Arial"

                    if conteudo_linha.strip().startswith("=> VUH"):
                        execucao_texto.font.size = Pt(13)
                        execucao_texto.font.bold = True
                   
                   
                    elif indice_linha == 0 and conteudo_linha.strip().startswith("AM "):
                        execucao_texto.font.size = Pt(13)
                        execucao_texto.font.bold = True
                    else:
                        execucao_texto.font.size = Pt(10)
                        execucao_texto.font.bold = False

                novo_paragrafo.paragraph_format.line_spacing = 1.15
                paragrafo_atual = novo_paragrafo
            break


###############################################################################
# SUBSTITUIR PLACEHOLDER POR TEXTO OU IMAGEM
###############################################################################
def substituir_placeholder_por_texto_formatado(documento, marcador, texto, tamanho_fonte=Pt(12), negrito=False):
    """
    Substitui o placeholder por texto com fonte e tamanho definidos.
    """
    for paragrafo in documento.paragraphs:
        if marcador in paragrafo.text:
            paragrafo.text = paragrafo.text.replace(marcador, "")
            execucao = paragrafo.add_run(texto)
            execucao.font.name = "Arial"
            execucao.font.size = tamanho_fonte
            execucao.bold = negrito

    for tabela in documento.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for parag in celula.paragraphs:
                    if marcador in parag.text:
                        parag.text = parag.text.replace(marcador, "")
                        execucao = parag.add_run(texto)
                        execucao.font.name = "Arial"
                        execucao.font.size = tamanho_fonte
                        execucao.bold = negrito

def substituir_placeholder_por_imagem(documento, marcador, caminho_imagem, largura=Inches(5)):
    """
    Substitui o placeholder por uma imagem alinhada ao centro.
    """
    for paragrafo in documento.paragraphs:
        if marcador in paragrafo.text:
            paragrafo.text = paragrafo.text.replace(marcador, "")
            runn = paragrafo.add_run()
            runn.add_picture(caminho_imagem, width=largura)
            paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            runn.font.name = "Arial"
            runn.font.size = Pt(12)

    for tabela in documento.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for parag in celula.paragraphs:
                    if marcador in parag.text:
                        parag.text = parag.text.replace(marcador, "")
                        runn = parag.add_run()
                        runn.add_picture(caminho_imagem, width=largura)
                        parag.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        runn.font.name = "Arial"
                        runn.font.size = Pt(12)

def remover_paragrafo_por_marcador(documento, marcador):
    """
    Remove o parágrafo que contém o marcador especificado.
    """
    for paragrafo in documento.paragraphs:
        if marcador in paragrafo.text:
            p = paragrafo._element
            p.getparent().remove(p)
            break
            
def substituir_placeholder_por_imagem_em_todo_documento(documento, marcador, img_path, largura):
    # corpo
    substituir_placeholder_por_imagem(documento, marcador, img_path, largura)
    # cabeçalhos e rodapés
    for sec in documento.sections:
        for part in (sec.header, sec.footer):
            for par in part.paragraphs:
                if marcador in par.text:
                    par.text = par.text.replace(marcador, "")
                    run = par.add_run()
                    run.add_picture(img_path, width=largura)


###############################################################################
# TABELA DE GRAU DE PRECISÃO ([texto_grau_precisao])
###############################################################################
def inserir_tabela_classificacao_de_precisao(documento, marcador, amplitude_ic80):
    """
    Insere a tabela padrão da ABNT NBR 14653 e destaca a classificação conforme o IC.
    """
    grau_obtido = classificar_grau_de_precisao(amplitude_ic80)
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    from docx.oxml.shared import OxmlElement
    from lxml import etree

    from docx.oxml.ns import nsdecls, qn

    for paragrafo in documento.paragraphs:
        if marcador in paragrafo.text:
            paragrafo.text = paragrafo.text.replace(marcador, "")
            titulo_paragrafo = inserir_paragrafo_apos(paragrafo, "GRAU DE PRECISÃO")
            for execucao in titulo_paragrafo.runs:
                execucao.font.name = "Arial"
                execucao.font.size = Pt(12)
                execucao.font.bold = True
            titulo_paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

            tabela = documento.add_table(rows=3, cols=4)
            tabela.style = "Table Grid"

            tabela.cell(0,0).text = "Descrição"
            tabela.cell(0,1).text = "GRAU III"
            tabela.cell(0,2).text = "GRAU II"
            tabela.cell(0,3).text = "GRAU I"

            tabela.cell(1,0).text = (
                "Amplitude do intervalo de confiança de 80%\n"
                "em torno da estimativa de tendência central"
            )
            tabela.cell(1,1).text = "≤ 30%"
            tabela.cell(1,2).text = "≤ 40%"
            tabela.cell(1,3).text = "≤ 50%"

            celula_nota = tabela.cell(2, 0).merge(tabela.cell(2, 1)).merge(tabela.cell(2, 2)).merge(tabela.cell(2, 3))
            celula_nota.text = (
                "NOTA: Quando a amplitude do intervalo de confiança ultrapassar 50%, "
                "não há classificação do resultado quanto à precisão e é necessária justificativa "
                "com base no diagnóstico do mercado."
            )

            for linha in tabela.rows:
                for celula in linha.cells:
                    props = celula._tc.get_or_add_tcPr()
                    vAlign = OxmlElement('w:vAlign')
                    vAlign.set(qn('w:val'), "center")
                    props.append(vAlign)
                    for par_cel in celula.paragraphs:
                        par_cel.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for runn in par_cel.runs:
                            runn.font.name = "Arial"
                            runn.font.size = Pt(10)

            shading_azul = etree.fromstring(r'<w:shd {} w:fill="BDD7EE" w:val="clear"/>'.format(nsdecls('w')))
            if grau_obtido == "GRAU III":
                tabela.cell(0,1)._tc.get_or_add_tcPr().append(shading_azul)
                tabela.cell(1,1)._tc.get_or_add_tcPr().append(shading_azul)
            elif grau_obtido == "GRAU II":
                tabela.cell(0,2)._tc.get_or_add_tcPr().append(shading_azul)
                tabela.cell(1,2)._tc.get_or_add_tcPr().append(shading_azul)
            elif grau_obtido == "GRAU I":
                tabela.cell(0,3)._tc.get_or_add_tcPr().append(shading_azul)
                tabela.cell(1,3)._tc.get_or_add_tcPr().append(shading_azul)
            else:
                celula_nota._tc.get_or_add_tcPr().append(shading_azul)

            titulo_paragrafo._p.addnext(tabela._element)

            novo_paragrafo = inserir_paragrafo_apos(
                titulo_paragrafo,
                f"Amplitude IC 80% calculada: {amplitude_ic80:.2f}% — Classificação: {grau_obtido}"
            )
            novo_paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for exec_novo in novo_paragrafo.runs:
                exec_novo.font.name = "Arial"
                exec_novo.font.size = Pt(10)
            break


###############################################################################
# FUNÇÕES DE CLASSIFICAÇÃO (FUNDAMENTAÇÃO / ENQUADRAMENTO)
###############################################################################
def inserir_fundamentacao_e_enquadramento(
    documento,
    placeholder_fundamentacao="[FUNDAMENTACAO]",
    placeholder_enquadramento="[enquadramento final]",
    tipo_imovel="Gleba Rural",
    quantidade_amostras_validadas=12,
    lista_todos_os_fatores=None
):
    """
    Substitui no documento as tabelas de Grau de Fundamentação (ABNT NBR 14653)
    para imóveis urbanos (4 itens) ou rurais (5 itens). Aplica as regras:
    
    1) Se quantidade_amostras_validadas < 4 => Mensagem "Laudo sem grau de fundamentação..."
    2) Se qualquer fator < 0.20 ou > 2.00 => Mensagem "Laudo sem enquadramento..."
    3) Do contrário, cada item recebe um grau (III, II, I ou nada), soma pontos
       e no final aplica as obrigatoriedades e pontuação mínima.
       
    OBS: `lista_todos_os_fatores` deve conter TODOS os valores de fator utilizados
         (FA, FO, FAP, FT, FP, FPA, FE, FAC, FL etc.) para verificar se estão
         dentro de [0.80..1.25], [0.50..1.40], [0.20..2.00] ou fora disso.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement
    from lxml import etree

    from docx.oxml.ns import qn, nsdecls
    from docx.shared import Pt

    # Se nenhum fator for informado, considere lista vazia
    if lista_todos_os_fatores is None:
        lista_todos_os_fatores = []

    tipo_lower = tipo_imovel.strip().lower()

    #
    # 1) Checagens que podem ANULAR imediatamente o grau de fundamentação
    #
    if quantidade_amostras_validadas < 4:
        # Imprime a mensagem e sai
        for par in documento.paragraphs:
            if placeholder_fundamentacao in par.text:
                par.text = par.text.replace(placeholder_fundamentacao, "")
                runx = par.add_run(
                    "LAUDO SEM GRAU DE FUNDAMENTAÇÃO POR POSSUIR MENOS DO QUE 4 AMOSTRAS DE MERCADO COLETADAS."
                )
                runx.font.name = "Arial"
                runx.font.size = Pt(11)
                runx.font.bold = True
                break
        # Remove o [enquadramento final] se existir
        for par in documento.paragraphs:
            if placeholder_enquadramento in par.text:
                par.text = par.text.replace(placeholder_enquadramento, "")
        return

    # Verificar se algum fator está <0.20 ou >2.00
    extrapolou_fatores = False
    for f in lista_todos_os_fatores:
        if f < 0.20 or f > 2.00:
            extrapolou_fatores = True
            break

    if extrapolou_fatores:
        # Imprime a mensagem e sai
        for par in documento.paragraphs:
            if placeholder_fundamentacao in par.text:
                par.text = par.text.replace(placeholder_fundamentacao, "")
                runx = par.add_run(
                    "LAUDO SEM ENQUADRAMENTO EM GRAU DE FUNDAMENTAÇÃO POR EXTRAPOLAÇÃO NOS FATORES DE HOMOGENEIZAÇÃO."
                )
                runx.font.name = "Arial"
                runx.font.size = Pt(11)
                runx.font.bold = True
                break
        # Remove o [enquadramento final] se existir
        for par in documento.paragraphs:
            if placeholder_enquadramento in par.text:
                par.text = par.text.replace(placeholder_enquadramento, "")
        return

    

    def pintar_celula_azul(cell):
        shading_azul = etree.fromstring(
            r'<w:shd {} w:fill="BDD7EE" w:val="clear"/>'.format(nsdecls('w'))
        )
        cell._tc.get_or_add_tcPr().append(shading_azul)

    # -----------------------------
    # FUNÇÕES AUXILIARES DE CLASSIFICAÇÃO
    # -----------------------------
    def classificar_item2_por_amostras(n_valid):
        """Retorna (grau, pontos) para o item 2, dada a quantidade de amostras."""
        if n_valid >= 10:
            return ("III", 3)
        elif n_valid >= 6:
            return ("II", 2)
        elif n_valid >= 4:
            return ("I", 1)
        else:
            return ("-", 0)

    def classificar_por_intervalo_fatores(lista_fat, tipo_imovel="urbano"):
        """
        - Para URBANO:
            III → 0,80-1,25
            II  → 0,50-1,40
            I   → 0,20-2,00
        - Para RURAL:
            III → 0,80-1,25
            II  → 0,70-1,40
            I   → 0,50-2,00
        """
        if all(0.80 <= x <= 1.25 for x in lista_fat):
            return "III", 3
    
        if tipo_imovel.lower().startswith("rur"):
            # Faixa rural
            if all(0.70 <= x <= 1.40 for x in lista_fat):
                return "II", 2
            elif all(0.50 <= x <= 2.00 for x in lista_fat):
                return "I", 1
        else:
            # Faixa urbana
            if all(0.50 <= x <= 1.40 for x in lista_fat):
                return "II", 2
            elif all(0.20 <= x <= 2.00 for x in lista_fat):
                return "I", 1
    
        return "-", 0       # não deveria acontecer – já haveria “extrapolação”


    # Montamos as tabelas e pintamos conforme o resultado.
    # Em paralelo, somamos pontos.
    # No final, usamos a pontuação + itens obrigatórios para definir GRAU final.

    # --------------------------------------------------------------------
    # Se for RURAL => 5 itens. Se for URBANO => 4 itens.
    # Você pode ajustar de acordo com a sua tabela. Abaixo, implemento
    # a lógica pedida especialmente para o URBANO e, se for RURAL,
    # exemplifico item 5 = intervalos de fatores.
    # --------------------------------------------------------------------

    # Verificar se é rural
    is_rural = ("rural" in tipo_lower)

    # Classificação dos itens
    # Para URBANO: item1=III, item2=?, item3=II, item4=ver fatores
    # Para RURAL: item1=III, item2=?, item3=?, item4=?, item5=ver fatores
    # Ajuste seu item3/4 para rural conforme a sua tabela.

    pontos_item1 = 3
    grau_item1 = "III"  # "Completa quanto a todos os atributos"

    grau_item2, pontos_item2 = classificar_item2_por_amostras(quantidade_amostras_validadas)

    # item3
    # URBANO => "Informações essenciais" => GRAU II => 2pts
    if not is_rural:
        grau_item3 = "II"
        pontos_item3 = 2
    else:
        # Exemplo para RURAL: item3 => "Apresentação dos dados"
        # Digamos que fixamos GRAU II => 2 pts (ou você ajusta a seu critério).
        grau_item3 = "II"
        pontos_item3 = 2

    # item4 => se for URBANO, é o Intervalo admissível. Se for RURAL, pode ser "Origem dos fatores".
    # No RURAL, item5 será o Intervalo.
    if not is_rural:
        # URBANO => item4 = intervalos
        grau_item4, pontos_item4 = classificar_por_intervalo_fatores(lista_todos_os_fatores)
        # item5 não existe, setamos 0
        grau_item5, pontos_item5 = ("-", 0)
    else:
        # RURAL => item4 => "Origem dos fatores". Aqui vou supor GRAU II => 2pts fixo,
        # mas você pode adequar a lógica real.
        grau_item4 = "II"
        pontos_item4 = 2
        # item5 => intervalos
        grau_item5, pontos_item5 = classificar_por_intervalo_fatores(
            lista_todos_os_fatores, tipo_imovel="rural"
        )             
      
                
    # ----------------------------------------------------------------------------------
    # Montar a tabela (difere entre rural e urbano).
    # E pintar as células conforme o grau de cada item.
    # ----------------------------------------------------------------------------------

    shading_azul = etree.fromstring(r'<w:shd {} w:fill="BDD7EE" w:val="clear"/>'.format(nsdecls('w')))

    # Função que pinta a célula de acordo com "III", "II", "I"
    def pintar_grau_urbano(tabela, row, grau_txt):
        if grau_txt == "III":
            pintar_celula_azul(tabela.cell(row, 2))
        elif grau_txt == "II":
            pintar_celula_azul(tabela.cell(row, 3))
        elif grau_txt == "I":
            pintar_celula_azul(tabela.cell(row, 4))

    def pintar_grau_rural(tabela, row, grau_txt):
        if grau_txt == "III":
            pintar_celula_azul(tabela.cell(row, 2))
        elif grau_txt == "II":
            pintar_celula_azul(tabela.cell(row, 3))
        elif grau_txt == "I":
            pintar_celula_azul(tabela.cell(row, 4))

    if not is_rural:
        # Tabela URBANA (4 itens)
        for paragrafo in documento.paragraphs:
            if placeholder_fundamentacao in paragrafo.text:
                paragrafo.text = paragrafo.text.replace(placeholder_fundamentacao, "")

                tabela_urb = documento.add_table(rows=5, cols=5, style="Table Grid")
                tabela_urb.cell(0,0).text = "Item"
                tabela_urb.cell(0,1).text = "Descrição"
                tabela_urb.cell(0,2).text = "III"
                tabela_urb.cell(0,3).text = "II"
                tabela_urb.cell(0,4).text = "I"

                tabela_urb.cell(1,0).text = "1"
                tabela_urb.cell(1,1).text = "Caracterização do imóvel avaliando"
                tabela_urb.cell(1,2).text = "Completa quanto a todos os fatores analisados"
                tabela_urb.cell(1,3).text = "Completa quanto aos fatores utilizados no tratamento"
                tabela_urb.cell(1,4).text = "Adoção de situação paradigma"

                tabela_urb.cell(2,0).text = "2"
                tabela_urb.cell(2,1).text = "Quantidade mínima de dados de mercado, efetivamente utilizados"
                tabela_urb.cell(2,2).text = "≥ 12"
                tabela_urb.cell(2,3).text = "≥ 5"
                tabela_urb.cell(2,4).text = "≥ 3"

                tabela_urb.cell(3,0).text = "3"
                tabela_urb.cell(3,1).text = "Identificação dos dados de mercado"
                tabela_urb.cell(3,2).text = "Apresentação de informações relativas a todas as características dos dados analisadas, com foto e características observadas pelo autor do laudo"
                tabela_urb.cell(3,3).text = "Apresentação de informações relativas a todas as características dos dados analisadas"
                tabela_urb.cell(3,4).text = "Apresentação de informações relativas a todas as características dos dados correspondentes aos fatores utilizados"

                tabela_urb.cell(4,0).text = "4"
                tabela_urb.cell(4,1).text = "Intervalo admissível de ajuste para o conjunto de fatores"
                tabela_urb.cell(4,2).text = "0,80 a 1,25"
                tabela_urb.cell(4,3).text = "0,50 a 2,00"
                tabela_urb.cell(4,4).text = "0,40 a 2,50"

                # Pintar item1 => "III"
                pintar_grau_urbano(tabela_urb, 1, grau_item1)
                # Pintar item2
                pintar_grau_urbano(tabela_urb, 2, grau_item2)
                # Pintar item3 => "II"
                pintar_grau_urbano(tabela_urb, 3, grau_item3)
                # Pintar item4 => intervalos
                pintar_grau_urbano(tabela_urb, 4, grau_item4)

                # Ajustar a formatação das células
                for rr in range(len(tabela_urb.rows)):
                    for cc in range(5):
                        c_ = tabela_urb.cell(rr, cc)
                        # Centralizar verticalmente
                        props = c_._tc.get_or_add_tcPr()
                        vAlign = OxmlElement('w:vAlign')
                        vAlign.set(qn('w:val'), "center")
                        props.append(vAlign)
                        # Centralizar horizontal
                        for parx in c_.paragraphs:
                            parx.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for runn in parx.runs:
                                runn.font.name = "Arial"
                                runn.font.size = Pt(9)

                paragrafo._p.addnext(tabela_urb._element)
                break

    else:
        # Tabela RURAL (5 itens)
        for paragrafo in documento.paragraphs:
            if placeholder_fundamentacao in paragrafo.text:
                paragrafo.text = paragrafo.text.replace(placeholder_fundamentacao, "")

                tabela_rur = documento.add_table(rows=6, cols=5, style="Table Grid")
                tabela_rur.cell(0,0).text = "Item"
                tabela_rur.cell(0,1).text = "Descrição"
                tabela_rur.cell(0,2).text = "III"
                tabela_rur.cell(0,3).text = "II"
                tabela_rur.cell(0,4).text = "I"

                tabela_rur.cell(1,0).text = "1"
                tabela_rur.cell(1,1).text = "Caracterização do bem avaliado"
                tabela_rur.cell(1,2).text = "Completa quanto a todos os atributos analisados"
                tabela_rur.cell(1,3).text = "Completa quanto aos atributos utilizados no tratamento"
                tabela_rur.cell(1,4).text = "Adoção de situação paradigma"

                tabela_rur.cell(2,0).text = "2"
                tabela_rur.cell(2,1).text = "Quantidade mínima de dados efetivamente utilizados"
                tabela_rur.cell(2,2).text = "≥ 12"
                tabela_rur.cell(2,3).text = "≥ 5"
                tabela_rur.cell(2,4).text = "≥ 3"

                tabela_rur.cell(3,0).text = "3"
                tabela_rur.cell(3,1).text = "Apresentação dos dados"
                tabela_rur.cell(3,2).text = "Atributos relativos a todos os dados e variáveis analisados na modelagem, com foto"
                tabela_rur.cell(3,3).text = "Atributos relativos a todos os dados e variáveis analisados na modelagem"
                tabela_rur.cell(3,4).text = "Atributos relativos aos dados e variáveis efetivamente utilizados no modelo"

                tabela_rur.cell(4,0).text = "4"
                tabela_rur.cell(4,1).text = "Origem dos fatores de homogeneização (conforme 7.7.2.1)"
                tabela_rur.cell(4,2).text = "Estudos embasados em metodologia científica"
                tabela_rur.cell(4,3).text = "Publicações"
                tabela_rur.cell(4,4).text = "Análise do avaliador"

                tabela_rur.cell(5,0).text = "5"
                tabela_rur.cell(5,1).text = "Intervalo admissível de ajuste para o conjunto de fatores"
                tabela_rur.cell(5,2).text = "0,80 a 1,25"
                tabela_rur.cell(5,3).text = "0,70 a 1,40"
                tabela_rur.cell(5,4).text = "0,50 a 2,00"

                # Pintar item1 => "III"
                pintar_grau_rural(tabela_rur, 1, grau_item1)
                # Pintar item2
                pintar_grau_rural(tabela_rur, 2, grau_item2)
                # Pintar item3
                pintar_grau_rural(tabela_rur, 3, grau_item3)
                # Pintar item4
                pintar_grau_rural(tabela_rur, 4, grau_item4)
                # Pintar item5 => intervalos
                pintar_grau_rural(tabela_rur, 5, grau_item5)

                # Ajustar formatações
                for rr in range(len(tabela_rur.rows)):
                    for cc in range(5):
                        c_ = tabela_rur.cell(rr, cc)
                        props = c_._tc.get_or_add_tcPr()
                        vAlign = OxmlElement('w:vAlign')
                        vAlign.set(qn('w:val'), "center")
                        props.append(vAlign)
                        for parx in c_.paragraphs:
                            parx.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for runn in parx.runs:
                                runn.font.name = "Arial"
                                runn.font.size = Pt(9)

                paragrafo._p.addnext(tabela_rur._element)
                break

    #
    # 3) Calcular a pontuação total e aplicar as regras para o ENQUADRAMENTO FINAL
    #
    if not is_rural:
        soma_pontos = pontos_item1 + pontos_item2 + pontos_item3 + pontos_item4
        # Itens obrigatórios = 2 e 4
        # Se (2 e 4) não atingirem certo grau, rebaixa.

        # item2 >= "III"? => grau_item2=="III"
        # item4 >= "III"? => grau_item4=="III"

        # Definir enquadramento final:
        if soma_pontos >= 10:
            # Tentar GRAU III, mas só se item2 e item4 = III
            if (grau_item2 == "III") and (grau_item4 == "III"):
                classificacao_final = "III"
            else:
                # Rebaixa para II
                classificacao_final = "II"
        elif soma_pontos >= 6:
            # Tentar GRAU II => mas item2 >= II e item4 >= II
            # Se não satisfizer, rebaixa p/ I
            cond_ii = (grau_item2 in ["II","III"]) and (grau_item4 in ["II","III"])
            if cond_ii:
                classificacao_final = "II"
            else:
                classificacao_final = "I"
        elif soma_pontos >= 4:
            # Tentar GRAU I => mas todos >= I
            # item1= III => ok, item2 >=I => item3=II => item4 >=I
            # Se item2 ou item4 for "-" => sem enquadramento
            cond_i = (grau_item1 in ["I","II","III"]) and (grau_item2 in ["I","II","III"]) \
                     and (grau_item3 in ["I","II","III"]) and (grau_item4 in ["I","II","III"])
            if cond_i:
                classificacao_final = "I"
            else:
                classificacao_final = "SEM ENQUADRAMENTO"
        else:
            classificacao_final = "SEM ENQUADRAMENTO"
    else:
        # RURAL => 5 itens
        soma_pontos = pontos_item1 + pontos_item2 + pontos_item3 + pontos_item4 + pontos_item5
        # Itens obrigatórios = 2 e 5
        if soma_pontos >= 13:
            # Tentar III => item2=III e item5=III
            if (grau_item2 == "III") and (grau_item5 == "III"):
                classificacao_final = "III"
            else:
                classificacao_final = "II"
        elif soma_pontos >= 8:
            # Tentar II => item2>=II e item5>=II
            cond_ii = (grau_item2 in ["II","III"]) and (grau_item5 in ["II","III"])
            if cond_ii:
                classificacao_final = "II"
            else:
                classificacao_final = "I"
        elif soma_pontos >= 5:
            # Tentar I => todos >= I
            cond_i = (grau_item1 in ["I","II","III"]) and (grau_item2 in ["I","II","III"]) \
                     and (grau_item3 in ["I","II","III"]) and (grau_item4 in ["I","II","III"]) \
                     and (grau_item5 in ["I","II","III"])
            if cond_i:
                classificacao_final = "I"
            else:
                classificacao_final = "SEM ENQUADRAMENTO"
        else:
            classificacao_final = "SEM ENQUADRAMENTO"

    # Se no passo final a string ficou "SEM ENQUADRAMENTO", mas percebemos
    # que foi por rebaixar algo, reclassifique. (Acima já fizemos a lógica.)

    if classificacao_final not in ["III","II","I"]:
        classificacao_final = "SEM ENQUADRAMENTO"

    # --------------  BLOCO [TABELA DE ENQUADRAMENTO FINAL] --------------
    # -- Localize-o rapidamente com Ctrl + F nesta marca acima.
    for paragrafo in documento.paragraphs:
        if placeholder_enquadramento in paragrafo.text:
            # ------------------------------------------------------------------
            # 1. Limpa o placeholder e cria a Tabela-Resumo (3 linhas × 4 colunas)
            # ------------------------------------------------------------------
            paragrafo.text = paragrafo.text.replace(placeholder_enquadramento, "")
            tabela2 = documento.add_table(rows=3, cols=4, style="Table Grid")
    
            # ------------------------ CONFIGURAÇÃO DINÂMICA -------------------
            if is_rural:                                               # Imóvel rural
                pontos_minimos = {"III": "13", "II": "8", "I": "5"}
                itens_obrig = {
                    "III": "Itens 2 e 5 no Grau III",
                    "II": "Itens 2 e 5 no Grau II",
                    "I":  "Todos, no mínimo no Grau I",
                }
            else:                                                      # Imóvel urbano
                pontos_minimos = {"III": "10", "II": "6", "I": "4"}
                itens_obrig = {
                    "III": "Itens 2 e 4 no Grau III",
                    "II": "Itens 2 e 4 no Grau II",
                    "I":  "Todos, no mínimo no Grau I",
                }
    
            # Cabeçalhos da primeira linha
            tabela2.cell(0, 0).text = "Graus"
            tabela2.cell(0, 1).text = "III"
            tabela2.cell(0, 2).text = "II"
            tabela2.cell(0, 3).text = "I"
    
            # Linha de Pontos mínimos
            tabela2.cell(1, 0).text = "Pontos mínimos"
            tabela2.cell(1, 1).text = pontos_minimos["III"]
            tabela2.cell(1, 2).text = pontos_minimos["II"]
            tabela2.cell(1, 3).text = pontos_minimos["I"]
    
            # Linha de Itens obrigatórios
            tabela2.cell(2, 0).text = "Itens obrigatórios"
            tabela2.cell(2, 1).text  = itens_obrig["III"]
            tabela2.cell(2, 2).text  = itens_obrig["II"]
            tabela2.cell(2, 3).text  = itens_obrig["I"]
    
            # --------------------------------------------------------------
            # 2. Função auxiliar para pintar uma célula em azul-claro (BDD7EE)
            # --------------------------------------------------------------
            def _pinta(cel):
                cel._tc.get_or_add_tcPr().append(
                    etree.fromstring(
                        r'<w:shd {} w:fill="BDD7EE" w:val="clear"/>'
                        .format(nsdecls("w"))
                    )
                )
    
            # --------------------------------------------------------------
            # 3. Pinta TODAS as células da coluna correspondente ao
            #    grau final calculado (classificacao_final)
            # --------------------------------------------------------------
            col = {"III": 1, "II": 2, "I": 3}.get(classificacao_final)
            if col is not None:                       # Só pinta se houver enquadramento
                for linha in range(3):                # linhas 0,1,2
                    _pinta(tabela2.cell(linha, col))
    
            # --------------------------------------------------------------
            # 4. Ajusta alinhamentos verticais/horizontais e fonte da tabela
            # --------------------------------------------------------------
            from docx.enum.table import WD_ALIGN_VERTICAL
            from docx.enum.text  import WD_ALIGN_PARAGRAPH
    
            for r in range(3):
                for c in range(4):
                    cel = tabela2.cell(r, c)
                    cel.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    for p in cel.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for rn in p.runs:
                            rn.font.name = "Arial"
                            rn.font.size = Pt(9)
    
            # Posiciona a tabela logo após o parágrafo do placeholder
            paragrafo._p.addnext(tabela2._element)
    
            # ------------------------------------------------------------------
            # 5. Insere o parágrafo final de resumo da pontuação
            # ------------------------------------------------------------------
            from docx.text.paragraph import Paragraph
            novo_p = OxmlElement("w:p")
            paragrafo._p.addnext(novo_p)
            parag_fim = Paragraph(novo_p, paragrafo._parent)
    
            if classificacao_final == "SEM ENQUADRAMENTO":
                texto_final = (
                    f"Pontuação total obtida: {soma_pontos} ponto(s). "
                    "SEM ENQUADRAMENTO FINAL."
                )
            else:
                texto_final = (
                    f"Pontuação total obtida: {soma_pontos} ponto(s). "
                    f"Grau de Fundamentação final: GRAU {classificacao_final}"
                )
    
            run_fim = parag_fim.add_run(texto_final)
            run_fim.bold = True
            run_fim.font.name = "Arial"
            run_fim.font.size = Pt(11)
            parag_fim.alignment = WD_ALIGN_PARAGRAPH.CENTER
            break  # Sai do loop: placeholder encontrado e tratado
    # -------------------------------------------------------------------
    # --------------------------  FIM DO BLOCO  -------------------------

###############################################################################
# INSERIR FOTOS
###############################################################################
# def inserir_fotos_no_placeholder(documento, placeholder, caminhos_fotos):
#     """
#     Insere as fotos no local do placeholder [FOTOS] organizadas em blocos de até 4 (2x2).
#     """
#     from docx.enum.text import WD_ALIGN_PARAGRAPH
#     bloco_fotos = []
#     largura_imagem = Inches(3)

#     paragrafo_alvo = None
#     for paragrafo in documento.paragraphs:
#         if placeholder in paragrafo.text:
#             paragrafo_alvo = paragrafo
#             break

#     if not paragrafo_alvo:
#         return

#     paragrafo_alvo.text = paragrafo_alvo.text.replace(placeholder, "")



def inserir_fotos_no_placeholder(documento, placeholder, caminhos_fotos, largura_imagem=Inches(3), um_por_pagina=False):
    from docx.oxml.ns import qn
    import os

    # Localizar o parágrafo que contém o marcador
    paragrafo_alvo = None
    for paragrafo in documento.paragraphs:
        if placeholder in paragrafo.text:
            paragrafo_alvo = paragrafo
            break

    if not paragrafo_alvo:
        logger.warning(f"⚠️ Placeholder {placeholder} não encontrado.")
        return

    # Remove o texto do marcador
    paragrafo_alvo.text = ""

    if um_por_pagina:
        for i, caminho in enumerate(caminhos_fotos):
            if os.path.exists(caminho):
                novo_par = adicionar_paragrafo_apos(paragrafo_alvo)
                run = novo_par.add_run()
                run.add_picture(caminho, width=largura_imagem)
                novo_par.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # ✅ Evita página em branco após a última imagem
                if i < len(caminhos_fotos) - 1:
                    run.add_break(WD_BREAK.PAGE)

                logger.info(f"✅ Imagem inserida em página separada: {caminho}")
                # Atualiza o parágrafo de referência para inserir o próximo após ele
                paragrafo_alvo = novo_par
            else:
                logger.warning(f"⚠️ Imagem não encontrada: {caminho}")
    else:
        blocos_fotos = [caminhos_fotos[i:i + 4] for i in range(0, len(caminhos_fotos), 4)]

        for bloco in blocos_fotos:
            tabela_fotos = documento.add_table(rows=2, cols=2)
            tabela_fotos.style = "Table Grid"

            idx = 0
            for linha in range(2):
                for coluna in range(2):
                    if idx < len(bloco):
                        caminho_img = bloco[idx]
                        if os.path.exists(caminho_img):
                            celula = tabela_fotos.cell(linha, coluna)
                            run = celula.paragraphs[0].add_run()
                            run.add_picture(caminho_img, width=largura_imagem)
                            celula.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            logger.info(f"✅ Imagem inserida em tabela: {caminho_img}")
                        else:
                            logger.warning(f"⚠️ Imagem não encontrada: {caminho_img}")
                    idx += 1

            paragrafo_alvo._p.addnext(tabela_fotos._element)





    # Função interna claramente isolada
    # def inserir_quatro_fotos(documento, paragrafo_referencia, fotos, largura_imagem):
    #     qtd_fotos = len(fotos)
    #     tabela_fotos = documento.add_table(rows=2, cols=2)
    #     tabela_fotos.style = "Table Grid"

    #     indice_foto = 0
    #     for linha_idx in range(2):
    #         for col_idx in range(2):
    #             if indice_foto < qtd_fotos:
    #                 caminho = fotos[indice_foto]
    #                 par = tabela_fotos.rows[linha_idx].cells[col_idx].paragraphs[0]
    #                 run_image = par.add_run()
    #                 try:
    #                     run_image.add_picture(caminho, width=largura_imagem)
    #                     logger.info(f"✅ Imagem inserida: {caminho}")
    #                 except Exception as e:
    #                     logger.error(f"Erro ao inserir imagem: {caminho}, erro: {e}")
    #                 par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #                 indice_foto += 1

    #     paragrafo_referencia._p.addnext(tabela_fotos._element)
    #     inserir_paragrafo_apos(paragrafo_referencia, "")

    # # Loop claramente isolado com variável única (sem redefinição)
    # for idx, caminho_foto in enumerate(caminhos_fotos, start=1):
    #     fotos_para_inserir.append(caminho_foto)
    #     if (idx % 4) == 0:
    #         inserir_quatro_fotos(documento, paragrafo_alvo, fotos_para_inserir, largura_imagem)
    #         fotos_para_inserir = []

    # if fotos_para_inserir:
    #     inserir_quatro_fotos(documento, paragrafo_alvo, fotos_para_inserir, largura_imagem)









###############################################################################
# INSERIR LOGO (OPCIONAL)
###############################################################################
def inserir_logo_no_placeholder(documento, placeholder, caminho_logo):
    """
    Substitui [logo] pela imagem do logotipo, alinhado à direita.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    for paragrafo in documento.paragraphs:
        if placeholder in paragrafo.text:
            paragrafo.text = paragrafo.text.replace(placeholder, "")
            runn = paragrafo.add_run()
            runn.add_picture(caminho_logo, width=Inches(3))
            paragrafo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            return
# ###############################################################################
# # TABELA DE RESUMO DE VALORES ([RESUMO VALORES])
# # AGORA MODIFICADA PARA EXIBIR MÚLTIPLAS RESTRIÇÕES
# ###############################################################################
# def inserir_tabela_resumo_de_valores(documento, marcador, informacoes_de_resumo, area_utilizada):
#     """
#     Cria a tabela de resumo de valores, compatível com versões antigas do python-docx,
#     sem usar get_or_add_tblPr(), e forçando que a primeira letra do valor por extenso 
#     seja maiúscula, ex.: "Trinta e um mil, cento e setenta e dois reais e seis centavos".
    
#     Parâmetros em `informacoes_de_resumo`:
#       - valor_unitario (str) => ex: "R$ 35,37/m²"
#       - area_total_considerada (str) => ex: "1.000,00 m²"
#       - texto_descritivo_restricoes (str) => ex: "Múltiplas restrições aplicadas"
#       - restricoes (list[dict]) => cada item: {
#             "area": 345.0,
#             "percentualDepreciacao": 34,
#             "fator": 0.66,
#             "tipo": "APP",
#             "subtotal": "R$ 8.053,23"
#         }
#       - valor_total_indenizatorio (str) => ex: "R$ 30.979,30"
#       - valor_por_extenso (str) => se vier vazio, será calculado via num2words; 
#         em seguida, a inicial é forçada para maiúsculo.
#     """
#     import re
#     from lxml import etree
#     from docx.oxml.ns import nsdecls, qn
#     from docx.enum.text import WD_ALIGN_PARAGRAPH
#     from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL
#     from docx.shared import Pt

#     # Se tiver num2words, usamos para converter valor em texto extenso.
#     try:
#         from num2words import num2words
#     except ImportError:
#         num2words = None

#     def extrair_valor_numerico(texto_monetario):
#         """
#         Ex: "R$ 30.979,30" => 30979.30 (float).
#         Remove caracteres que não sejam dígitos ou vírgula, então substitui ',' por '.'.
#         """
#         somente_num_virg = re.sub(r"[^\d,]", "", texto_monetario)
#         somente_num_ponto = somente_num_virg.replace(",", ".")
#         try:
#             return float(somente_num_ponto)
#         except:
#             return 0.0

#     def gerar_extenso_por_num2words(texto_valor):
#         """
#         Converte "R$ 30.979,30" em algo como 
#         "Trinta e um mil, cento e setenta e nove reais e trinta centavos",
#         usando a biblioteca num2words(lang='pt_BR'). 
#         Em seguida, forçamos a primeira letra para maiúscula.
#         """
#         if not num2words:
#             return "(num2words não instalado)"

#         val = extrair_valor_numerico(texto_valor)
#         inteiro = int(val)
#         centavos = round((val - inteiro) * 100)
#         if inteiro == 0 and centavos == 0:
#             return "Zero real"

#         extenso_inteiro = num2words(inteiro, lang='pt_BR')
#         if centavos > 0:
#             extenso_centavos = num2words(centavos, lang='pt_BR')
#             texto_final = f"{extenso_inteiro} reais e {extenso_centavos} centavos"
#         else:
#             texto_final = f"{extenso_inteiro} reais"

#         # Forçar a primeira letra para maiúsculo, se não estiver vazio:
#         if texto_final:
#             texto_final = texto_final[0].upper() + texto_final[1:]
#         return texto_final

#     # -------------------------------------------------------------------------
#     # Localiza o placeholder no documento
#     for paragrafo in documento.paragraphs:
#         if marcador in paragrafo.text:
#             # Remove o texto do placeholder
#             paragrafo.text = paragrafo.text.replace(marcador, "")

#             # Carrega dados
#             valor_unit = informacoes_de_resumo.get("valor_unitario", "N/D")
#             area_total = informacoes_de_resumo.get("area_total_considerada", "N/D")
#             sit_rest = informacoes_de_resumo.get("texto_descritivo_restricoes", "N/D")
#             restricoes = informacoes_de_resumo.get("restricoes", [])
#             valor_total = informacoes_de_resumo.get("valor_total_indenizatorio", "R$ 0,00")
#             valor_extenso = informacoes_de_resumo.get("valor_por_extenso", "").strip()

#             # Se valor_por_extenso for vazio, gerar automaticamente
#             if not valor_extenso:
#                 valor_extenso = gerar_extenso_por_num2words(valor_total)

#             # Cria a tabela principal (7 linhas, 2 colunas)
#             tabela_principal = documento.add_table(rows=7, cols=2)
#             tabela_principal.style = "Table Grid"
#             tabela_principal.alignment = WD_TABLE_ALIGNMENT.CENTER

#             # (0) Cabeçalho mesclado
#             cel_titulo = tabela_principal.cell(0, 0).merge(tabela_principal.cell(0, 1))
#             cel_titulo.text = "RESUMO DOS VALORES TOTAIS"
#             shading_cab = etree.fromstring(r'<w:shd {} w:fill="D9D9D9" w:val="clear"/>'.format(nsdecls('w')))
#             cel_titulo._tc.get_or_add_tcPr().append(shading_cab)
#             for p_ in cel_titulo.paragraphs:
#                 p_.alignment = WD_ALIGN_PARAGRAPH.CENTER
#                 for run_ in p_.runs:
#                     run_.font.name = "Arial"
#                     run_.font.size = Pt(11)
#                     run_.font.bold = True

#             # (1) Valor Unitário Calculado
#             tabela_principal.cell(1,0).text = "Valor Unitário Calculado:"
#             tabela_principal.cell(1,1).text = valor_unit

#            # (2) Área Total de Interesse
#             tabela_principal.cell(2, 0).text = "Área Total de Interesse:"
#             tabela_principal.cell(2, 1).text = informacoes_de_resumo["area_total_considerada"]



#             # (3) Situação das Restrições
#             tabela_principal.cell(3,0).text = "Situação das Restrições:"
#             tabela_principal.cell(3,1).text = sit_rest

#             # (4) Sub-tabela => célula mesclada
#             cel_sub = tabela_principal.cell(4,0).merge(tabela_principal.cell(4,1))
#             shading_light_blue = etree.fromstring(r'<w:shd {} w:fill="E0ECF8" w:val="clear"/>'.format(nsdecls('w')))
#             cel_sub._tc.get_or_add_tcPr().append(shading_light_blue)

#             # Remove margens internas da célula mesclada
#             tc_pr_sub = cel_sub._tc.get_or_add_tcPr()
#             tc_margins_sub = tc_pr_sub.xpath('./w:tcMar')
#             if not tc_margins_sub:
#                 tcMar = OxmlElement('w:tcMar')
#                 tcMar.set(qn('w:top'), "0")
#                 tcMar.set(qn('w:left'), "0")
#                 tcMar.set(qn('w:right'), "0")
#                 tcMar.set(qn('w:bottom'), "0")
#                 tc_pr_sub.append(tcMar)
#             else:
#                 for m_ in tc_margins_sub:
#                     m_.set(qn('w:top'), "0")
#                     m_.set(qn('w:left'), "0")
#                     m_.set(qn('w:right'), "0")
#                     m_.set(qn('w:bottom'), "0")

#             # Se não tiver restrições, mostra texto simples
#             if not restricoes:
#                 cel_sub.text = "Nenhuma restrição aplicada."
#                 for r_ in cel_sub.paragraphs[0].runs:
#                     r_.font.name = "Arial"
#                     r_.font.size = Pt(10)
#             else:
#                 # Cria sub-tabela sem bordas
#                 subtab = documento.add_table(rows=len(restricoes)+1, cols=5)
#                 borders = subtab._element.xpath(".//w:tblBorders")
#                 for b_ in borders:
#                     b_.getparent().remove(b_)

#                 # Adicionar manualmente <w:tblPr>, se não existir
#                 tblPr = subtab._element.tblPr
#                 if tblPr is None:
#                     tblPr = OxmlElement('w:tblPr')
#                     subtab._element.insert(0, tblPr)

#                 # <w:tblCellMar> p/ zerar margens
#                 tblCellMar = OxmlElement('w:tblCellMar')
#                 tblCellMar.set(qn('w:top'), "0")
#                 tblCellMar.set(qn('w:left'), "0")
#                 tblCellMar.set(qn('w:right'), "0")
#                 tblCellMar.set(qn('w:bottom'), "0")
#                 tblPr.append(tblCellMar)

#                 # Cabeçalhos
#                 cabecalhos = ["Área (m²)", "% Depreciação", "Fator aplicado", "Tipo Restrição", "Subtotal (R$)"]
#                 for cidx, hh in enumerate(cabecalhos):
#                     subtab.cell(0,cidx).text = hh
#                     for run_ in subtab.cell(0,cidx).paragraphs[0].runs:
#                         run_.font.name = "Arial"
#                         run_.font.size = Pt(9)
#                         run_.font.bold = True
#                     subtab.cell(0,cidx).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

#                 # Linhas de dados
#                 for i, restr in enumerate(restricoes, start=1):
#                     area_ = formatar_area_brasil(restr.get("area", ""))
#                     perc_ = restr.get("percentualDepreciacao", 0)
#                     fat_ = restr.get("fator", 1.0)
#                     tipo_ = restr.get("tipo", "")
#                     sub_ = restr.get("subtotal", "R$ 0,00")

#                     subtab.cell(i,0).text = area_
#                     subtab.cell(i,1).text = f"{perc_:.0f}%"
#                     subtab.cell(i,2).text = f"{fat_:.2f}"
#                     subtab.cell(i,3).text = tipo_
#                     subtab.cell(i,4).text = sub_

#                     for cc_ in range(5):
#                         p_run = subtab.cell(i, cc_).paragraphs[0]
#                         p_run.alignment = WD_ALIGN_PARAGRAPH.CENTER
#                         for r_ in p_run.runs:
#                             r_.font.name = "Arial"
#                             r_.font.size = Pt(9)

#                 # Fundo azul e remover margens em todas as células
#                 for row_ in subtab.rows:
#                     for cell_ in row_.cells:
#                         shade_cell = etree.fromstring(r'<w:shd {} w:fill="E0ECF8" w:val="clear"/>'.format(nsdecls('w')))
#                         cell_._tc.get_or_add_tcPr().append(shade_cell)
#                         cell_.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
#                         tcpr = cell_._tc.get_or_add_tcPr()
#                         tc_marg = tcpr.xpath('./w:tcMar')
#                         if not tc_marg:
#                             newMar = OxmlElement('w:tcMar')
#                             newMar.set(qn('w:top'), "0")
#                             newMar.set(qn('w:left'), "0")
#                             newMar.set(qn('w:right'), "0")
#                             newMar.set(qn('w:bottom'), "0")
#                             tcpr.append(newMar)
#                         else:
#                             for mm in tc_marg:
#                                 mm.set(qn('w:top'), "0")
#                                 mm.set(qn('w:left'), "0")
#                                 mm.set(qn('w:right'), "0")
#                                 mm.set(qn('w:bottom'), "0")

#                 # Anexa a sub-tabela à célula
#                 cel_sub._tc.clear_content()
#                 cel_sub._tc.append(subtab._element)

#             # (5) Valor Total Indenizatório
#             tabela_principal.cell(5,0).text = "Valor Total Indenizatório:"
#             tabela_principal.cell(5,1).text = valor_total

#             # (6) Valor por Extenso
#             cel_ext = tabela_principal.cell(6,0).merge(tabela_principal.cell(6,1))
#             cel_ext.text = valor_extenso

#             # Ajustes de layout da Tabela Principal
#             for row_idx in range(7):
#                 row_ = tabela_principal.rows[row_idx]
#                 row_.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
#                 row_.height = Pt(18)
#                 for col_idx in range(2):
#                     c_ = row_.cells[col_idx]
#                     c_.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
#                     for pp_ in c_.paragraphs:
#                         pp_.alignment = WD_ALIGN_PARAGRAPH.CENTER
#                         for rr_ in pp_.runs:
#                             rr_.font.name = "Arial"
#                             rr_.font.size = Pt(10)
#                             rr_.font.bold = False

#             # Valor Unitário (linha 1 => col 1) e Valor Total (linha 5 => col 1) em negrito
#             for run_ in tabela_principal.rows[1].cells[1].paragraphs[0].runs:
#                 run_.font.bold = True
#             for run_ in tabela_principal.rows[5].cells[1].paragraphs[0].runs:
#                 run_.font.bold = True
#                 run_.font.size = Pt(11)

#             # Valor por Extenso (linha 6) => central e em negrito
#             for p_ in tabela_principal.rows[6].cells[0].paragraphs:
#                 p_.alignment = WD_ALIGN_PARAGRAPH.CENTER
#                 for rn_ in p_.runs:
#                     rn_.font.size = Pt(10)
#                     rn_.font.bold = True

#             # Insere a tabela após o parágrafo do placeholder
#             paragrafo._p.addnext(tabela_principal._element)
#             break

###############################################################################
# DIAGNÓSTICO DE MERCADO
###############################################################################
def inserir_tabela_diagnostico_de_mercado(documento, marcador, escolha_estrutura, escolha_conduta, escolha_desempenho):
    """
    Monta a tabela de diagnóstico de mercado (Estrutura, Conduta, Desempenho),
    destacando (sombreando) a opção escolhida.
    """
    from docx.enum.table import WD_ROW_HEIGHT_RULE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    from lxml import etree

    from docx.oxml.ns import nsdecls, qn

    dados_tabela = [
        ("Estrutura",  "BOA",         "Facilidade em se encontrar imóveis na região do avaliando"),
        ("Estrutura",  "LIMITADA",    "Dificuldade em se encontrar imóveis na região do avaliando."),
        ("Conduta",    "DESESTAGNADO","Boa movimentação do mercado imobiliário."),
        ("Conduta",    "ESTAGNADA",   "Pouca movimentação do mercado imobiliário."),
        ("Desempenho", "ALTO",        "Ótima atratividade comercial para negócios imobiliários."),
        ("Desempenho", "MÉDIO",       "Atratividade moderada para negócios imobiliários."),
        ("Desempenho", "BAIXO",       "Baixa atratividade da região para negócios imobiliários.")
    ]

    def verificar_se_destacar(dim, opc):
        if dim.lower() == "estrutura":
            return opc.upper() == escolha_estrutura.upper()
        elif dim.lower() == "conduta":
            return opc.upper() == escolha_conduta.upper()
        elif dim.lower() == "desempenho":
            return opc.upper() == escolha_desempenho.upper()
        return False

    for paragrafo in documento.paragraphs:
        if marcador in paragrafo.text:
            paragrafo.text = paragrafo.text.replace(marcador, "")
            tabela = documento.add_table(rows=len(dados_tabela), cols=3)
            tabela.style = "Table Grid"
            tabela.allow_autofit = False

            for i, (dimensao, opcao, descricao) in enumerate(dados_tabela):
                c_dim = tabela.rows[i].cells[0]
                c_opc = tabela.rows[i].cells[1]
                c_desc = tabela.rows[i].cells[2]

                # Exibir a dimensão apenas na "parte do meio" de cada bloco
                if i == 1:
                    c_dim.text = "Estrutura"
                elif i == 3:
                    c_dim.text = "Conduta"
                elif i == 5:
                    c_dim.text = "Desempenho"
                else:
                    c_dim.text = ""

                c_opc.text = opcao
                c_desc.text = descricao

                if verificar_se_destacar(dimensao, opcao):
                    shading_azul = etree.fromstring(
                        r'<w:shd {} w:fill="BDD7EE" w:val="clear"/>'.format(nsdecls('w'))
                    )
                    c_opc._tc.get_or_add_tcPr().append(shading_azul)
                    c_desc._tc.get_or_add_tcPr().append(shading_azul)

            # Mesclar primeira coluna
            tabela.cell(0, 0).merge(tabela.cell(1, 0))
            tabela.cell(2, 0).merge(tabela.cell(3, 0))
            tabela.cell(4, 0).merge(tabela.cell(5, 0))
            tabela.cell(4, 0).merge(tabela.cell(6, 0))

            for row_index, row in enumerate(tabela.rows):
                row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                row.height = Pt(28)
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        paragraph.paragraph_format.space_before = Pt(3)
                        paragraph.paragraph_format.space_after = Pt(3)
                        for run in paragraph.runs:
                            run.font.name = "Arial"
                            run.font.size = Pt(10)

            paragrafo._p.addnext(tabela._element)
            break

###############################################################################
# INSERIR TABELA DE RESUMO GERAL (EXEMPLO)
###############################################################################
def inserir_tabela_resumo_geral_completo(documento, placeholder, info_resumo_geral):
    """
    Exemplo de inserção de uma tabela extra de 'Resumo Geral da Avaliação',
    demonstrando como manipular dados caso queira algo mais completo.
    Substitui [RESUMO GERAL] por uma tabela no documento.
    """
    for paragrafo in documento.paragraphs:
        if placeholder in paragrafo.text:
            paragrafo.text = paragrafo.text.replace(placeholder, "")
            # Aqui apenas um exemplo (poderia adaptar)
            from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Pt
            from docx.oxml.shared import OxmlElement
            from lxml import etree

            from docx.oxml.ns import nsdecls, qn

            tabela_resumo = documento.add_table(rows=5, cols=2, style="Table Grid")

            tabela_resumo.cell(0,0).text = "Proprietário"
            tabela_resumo.cell(0,1).text = info_resumo_geral.get("proprietario","S/N")

            tabela_resumo.cell(1,0).text = "Documento"
            tabela_resumo.cell(1,1).text = info_resumo_geral.get("documento_imovel","(N/D)")

            tabela_resumo.cell(2,0).text = "Cartório"
            tabela_resumo.cell(2,1).text = info_resumo_geral.get("cartorio","(N/D)")

            tabela_resumo.cell(3,0).text = "Comarca"
            tabela_resumo.cell(3,1).text = info_resumo_geral.get("comarca","(N/D)")

            tabela_resumo.cell(4,0).text = "Endereço"
            tabela_resumo.cell(4,1).text = info_resumo_geral.get("endereco_imovel","(N/D)")

            for rr in range(len(tabela_resumo.rows)):
                for cc in range(2):
                    cell_ = tabela_resumo.cell(rr, cc)
                    for pp_ in cell_.paragraphs:
                        pp_.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run_ in pp_.runs:
                            run_.font.name = "Arial"
                            run_.font.size = Pt(9)

            paragrafo._p.addnext(tabela_resumo._element)
            break
###############################################################################
# >>>>>> RELATÓRIO PRINCIPAL – TRECHO REESCRITO (placeholders + finalidade) <<<<
###############################################################################
from docx.shared import Pt                      # já havia sido importado antes
from datetime import datetime
from docx import Document



    
###############################################################################
# >>>>>> FUNÇÃO PRINCIPAL (MAIN) COM A NOVA LÓGICA DAS RESTRIÇÕES ― 03-mai-2025
###############################################################################
def main():
    """
    PASSOS EXECUTADOS PELO SCRIPT
    ────────────────────────────
    1) Seleciona a planilha Excel de amostras de mercado.
    2) Pergunta todos os dados necessários ao usuário (proprietário, fatores,
       finalidade, etc.).
    3) Permite cadastrar qualquer quantidade de restrições, aplicando TRAVAS
       para impedir que a soma ultrapasse a “área de interesse”:
          • Se a finalidade é DESAPROPRIAÇÃO ou SERVIDÃO  → área digitada
            pelo usuário.
          • Caso contrário                              → área total lida da
            planilha.
    4) Valida a consistência; se tudo OK, processa estatísticas,
       gera gráficos e monta o relatório DOCX.
    """

    # =================================================================== ▒ SETUP
    barra_progresso = tqdm(total=6, desc="Processando", ncols=80)

    caminho_planilha = os.path.join(pasta_temp, "planilha.xlsx")
    request.files["planilha_excel"].save(caminho_planilha)

    root_plan.destroy()
    if not caminho_planilha:
        logger.info("Nenhuma planilha selecionada. Encerrando."); return

    # ================================================================= ▒ ENTRADAS
    nome_proprietario = input("Informe o nome completo do proprietário do imóvel: ").strip()

    def perguntar_sim_ou_nao(msg: str) -> bool:
        """Retorna True para S/s e False para N/n; repete até resposta válida."""
        while True:
            r = input(msg).strip().lower()
            if r in ("s", "n"):
                return r == "s"
            print("Opção inválida. Digite S ou N.")

    if perguntar_sim_ou_nao("Existem outros proprietários? (S/N): "):
        nome_proprietario += " e Outros"

    telefone_proprietario = ""
    if perguntar_sim_ou_nao("Deseja incluir telefone do proprietário? (S/N): "):
        telefone_proprietario = "Contato do Proprietário: " + input("Digite o telefone: ").strip()

    email_proprietario = ""
    if perguntar_sim_ou_nao("Deseja incluir email do proprietário? (S/N): "):
        email_proprietario = "E-mail do Proprietário: " + input("Digite o email: ").strip()

    nome_solicitante   = input("Informe o nome do solicitante do laudo: ").strip()
    nome_avaliador     = input("Informe o nome do avaliador responsável: ").strip()
    registro_avaliador = input("Informe o registro profissional do avaliador: ").strip()

    # ----------------------------------------------------- ▒ TIPO DE IMÓVEL
    print("\nQual tipo de imóvel está sendo avaliado?")
    opcoes_de_imovel = [
        "Apartamento residencial", "Casa residencial", "Terreno urbano",
        "Gleba urbana", "Terreno rural", "Gleba rural"
    ]
    for i, opc in enumerate(opcoes_de_imovel, 1):
        print(f"{i}) {opc}")
    while True:
        esc = input("Escolha o número: ").strip()
        if esc.isdigit() and 1 <= int(esc) <= len(opcoes_de_imovel):
            tipo_imovel_escolhido = opcoes_de_imovel[int(esc) - 1]
            break
        print("Opção inválida.")

    # =================================================== ▒ FINALIDADE + ÁREA
    print("\nQual a finalidade do laudo?")
    opcoes_finalidade = [
        "Desapropriação do Imóvel",
        "Avaliação para Garantia Bancária",
        "Avaliação para processos de Inventário e Partilha",
        "Avaliação de mercado para Compra e venda do Imóvel",
        "Avaliação para obtenção do valor de Locação do Imóvel",
        "Avaliação para Regularização Fiscal do Imóvel",
        "Avaliação para a obtenção de Seguro patrimonial do Imóvel",
        "Avaliação para Implantação de Servidão Administrativa",
        "Avaliação para Regularização do Imóvel"
    ]
    for i, fin in enumerate(opcoes_finalidade, 1):
        print(f"{i}) {fin}")
    while True:
        esc_fin = input("Escolha o número da finalidade: ").strip()
        if esc_fin.isdigit() and 1 <= int(esc_fin) <= len(opcoes_finalidade):
            idx_fin          = int(esc_fin)
            texto_finalidade = opcoes_finalidade[idx_fin - 1]
            break
        print("Opção inválida.")

    finalidade_lida = ("desapropriacao" if idx_fin == 1 else
                       "servidao"       if idx_fin == 8 else
                       "mercado")

    area_disponivel: float = 0.0  # sempre será definida antes da 1ª validação
    area_parcial   : float = 0.0  # nome usado no template

    def solicitar_area(msg: str) -> float:
        """Pergunta área numérica positiva; aceita vírgula ou ponto como separador."""
        while True:
            txt = input(msg).strip().replace(".", "").replace(",", ".")
            try:
                val = float(txt)
                if val > 0:
                    return val
                print("O valor deve ser maior que zero.")
            except ValueError:
                print("Valor inválido, tente novamente.")

    if finalidade_lida == "desapropriacao":
        area_disponivel = solicitar_area("Área desapropriada (m²): ")
    elif finalidade_lida == "servidao":
        area_disponivel = solicitar_area("Área para servidão (m²): ")
    else:
        print("A área total será lida automaticamente da planilha após o carregamento.\n")

    area_parcial = area_disponivel  # se ainda for zero, será atualizado depois

    # =================================================== ▒ FATORES BÁSICOS
    def perguntar_fator(msg): return perguntar_sim_ou_nao(msg)
    fatores_do_usuario = {
        "area"  : perguntar_fator("Usar fator Área? (S/N): "),
        "oferta": perguntar_fator("Usar fator Oferta? (S/N): ")
    }

    # =================================================== ▒ RESTRIÇÕES
    restricoes_lista      = []
    soma_areas_restricoes = 0.0
    print("\n--- Cadastro de Restrições ---")
    while perguntar_sim_ou_nao("Deseja cadastrar alguma restrição? (S/N): "):
        tipo_rest = input("Tipo de restrição (ex: APP, Servidão, Outro): ").strip()
        area_rest = solicitar_area("Área (m²) com essa restrição: ")

        # TRAVA IMEDIATA: se já sabemos area_disponivel (>0), nunca deixa exceder
        if area_disponivel > 0 and soma_areas_restricoes + area_rest > area_disponivel:
            exced = soma_areas_restricoes + area_rest - area_disponivel
            logger.warning(f"A soma das áreas de restrição excede a área disponível "
                  f"em {exced:.2f} m². Digite um valor menor.\n")
            continue

        deprec = solicitar_area("Porcentagem de depreciação (ex: 40 → 40%): ")
        deprec = min(deprec, 100.0)
        fator_restr = (100.0 - deprec) / 100.0

        restricoes_lista.append({
            "tipo"                 : tipo_rest,
            "area"                 : area_rest,
            "percentualDepreciacao": deprec,
            "fator"                : fator_restr
        })
        soma_areas_restricoes += area_rest
        print(f"Restrição '{tipo_rest}' adicionada. "
              f"Total de áreas restritas: {soma_areas_restricoes:.2f} m².\n")

    fatores_do_usuario["restricoes"] = restricoes_lista

    # =================================================== ▒ OUTROS FATORES
    fatores_do_usuario["aproveitamento"]           = perguntar_fator("Usar fator Aproveitamento? (S/N): ")
    print("\nO imóvel avaliando está na mesma região (~1 km) das amostras? (S/N)")
    fatores_do_usuario["localizacao_mesma_regiao"] = perguntar_fator("Escolha (S/N): ")
    fatores_do_usuario["topografia"]               = perguntar_fator("Usar fator Topografia? (S/N): ")
    fatores_do_usuario["pedologia"]                = perguntar_fator("Usar fator Pedologia? (S/N): ")
    fatores_do_usuario["pavimentacao"]             = perguntar_fator("Usar fator Pavimentação? (S/N): ")
    fatores_do_usuario["esquina"]                  = perguntar_fator("Usar fator Esquina? (S/N): ")
    fatores_do_usuario["acessibilidade"]           = perguntar_fator("Usar fator Acessibilidade? (S/N): ")

    num_doc = input("Número da matrícula (ex: 12345): ").strip()
    fatores_do_usuario["documentacaoImovel"] = f"Matrícula n° {num_doc}" if num_doc else "Documentação não informada"
    fatores_do_usuario["nomeCartorio"]       = input("Nome do cartório: ").strip()
    fatores_do_usuario["nomeComarca"]        = input("Nome da comarca: ").strip()
    fatores_do_usuario["enderecoCompleto"]   = input("Endereço completo do imóvel: ").strip()

    fatores_do_usuario.update({
        "nomeSolicitante"     : nome_solicitante,
        "avaliadorNome"       : nome_avaliador,
        "avaliadorRegistro"   : registro_avaliador,
        "tipoImovel"          : tipo_imovel_escolhido,
        "finalidadeTexto"     : texto_finalidade,
        "nomeProprietario"    : nome_proprietario,
        "telefoneProprietario": telefone_proprietario,
        "emailProprietario"   : email_proprietario
    })

    # =================================================== ▒ DIAGNÓSTICO DE MERCADO
    logger.info("\n=== DIAGNÓSTICO DE MERCADO ===")
    print("Estrutura:\n 1) BOA\n 2) LIMITADA")
    while True:
        e = input("Escolha (1 ou 2): ").strip()
        if e in ("1", "2"):
            fatores_do_usuario["estrutura_escolha"] = "BOA" if e == "1" else "LIMITADA"
            break
        print("Opção inválida.")
    print("\nConduta:\n 1) DESESTAGNADO\n 2) ESTAGNADA")
    while True:
        c = input("Escolha (1 ou 2): ").strip()
        if c in ("1", "2"):
            fatores_do_usuario["conduta_escolha"] = "DESESTAGNADO" if c == "1" else "ESTAGNADA"
            break
        print("Opção inválida.")
    print("\nDesempenho:\n 1) ALTO\n 2) MÉDIO\n 3) BAIXO")
    while True:
        d = input("Escolha (1, 2 ou 3): ").strip()
        if d in ("1", "2", "3"):
            fatores_do_usuario["desempenho_escolha"] = ("ALTO" if d == "1"
                                                        else "MÉDIO" if d == "2"
                                                        else "BAIXO")
            break
        print("Opção inválida.")

    # =================================================== ▒ LER PLANILHA
    barra_progresso.update(1)
    dataframe_amostras, dados_avaliando = ler_planilha_excel(caminho_planilha)
    df_amostras, dados_imovel = ler_planilha_excel(caminho_planilha)
    print(df_amostras.head())
    print(dados_imovel)
    area_total_planilha = float(dados_avaliando.get("AREA TOTAL", 0))

    # Se ainda não definimos area_disponivel (caso “mercado”), usamos a da planilha
    if area_disponivel == 0.0:
        area_disponivel = area_total_planilha
    if area_parcial == 0.0:
        area_parcial = area_disponivel

    # =================================================== ▒ VALIDAÇÃO FINAL
    
  
    if finalidade_lida in ("desapropriacao", "servidao"):
        if area_parcial > area_total_planilha:
            print(f"\nATENÇÃO: A área digitada ({area_parcial:,.2f} m²) "
                  f"é MAIOR que a área total do imóvel ({area_total_planilha:,.2f} m²).")
            print("Verifique os dados e tente novamente.")
            return    
    
    
    if soma_areas_restricoes > area_disponivel:
        logger.warning(f"\nATENÇÃO: A soma das áreas restritas ({soma_areas_restricoes:.2f} m²) "
              f"ultrapassa a área de interesse ({area_disponivel:.2f} m²).")
        logger.info("Encerrando o script, pois os dados estão inconsistentes.")
        return

    barra_progresso.update(1)

    # =================================================== ▒ FILTRAGEM / CÁLCULOS
    (dataframe_amostras_filtrado, indices_excluidos, amostras_excluidas,
     media_chauvenet, desvio_chauvenet, menor_valor_chauvenet,
     maior_valor_chauvenet, mediana_chauvenet) = aplicar_chauvenet_e_filtrar(dataframe_amostras)
    barra_progresso.update(1)
    print("Filtrado:", dataframe_amostras_filtrado.head())
    print("Média:", media_chauvenetia, "Mediana:", mediana_chauvenet)


    valores_homogeneizados_validos = homogeneizar_amostras(
        dataframe_amostras_filtrado, dados_avaliando,
        fatores_do_usuario, finalidade_lida
    )
    print("Homogeneizadas:", valores_homogeneizados_validos)
    lista_valores_originais_iniciais = dataframe_amostras_filtrado["VALOR TOTAL"].tolist()

    arquivo_aderencia = "grafico_aderencia_totais.png"
    gerar_grafico_aderencia_totais(
        dataframe_amostras_filtrado, valores_homogeneizados_validos,
        arquivo_aderencia
    )
    barra_progresso.update(1)

    arquivo_dispersao = "grafico_dispersao_mediana.png"
    gerar_grafico_dispersao_mediana(valores_homogeneizados_validos, arquivo_dispersao)
    barra_progresso.update(1)

    # =================================================== ▒ SELECIONAR FOTOS / LOGO
    
    fatores_do_usuario["caminhoLogo"] = caminho_logo  # já definido a partir de request.files["arquivo_logo"]
    

###############################################################################
# LEITURA DA PLANILHA EXCEL  —  distância sempre em relação ao CENTRO-URBANO
###############################################################################
import re, pandas as pd, numpy as np
from geopy.geocoders import Nominatim
from time import sleep
from math import radians, sin, cos, sqrt, atan2   # usado pelo haversine_km

# --------------------------------------------------------------------------
# Helpers internos (os dois já estavam no nosso “arsenal”)
# --------------------------------------------------------------------------
def _parse_coord(val: Union[str, float, int]):
    """
    Converte qualquer string de coordenada (-29.08°, 53,842 etc.) em float.
    Retorna None se não conseguir.
    """
    if pd.isna(val):                                # NaN do pandas
        return None
    if isinstance(val, str):
        val = val.replace(",", ".")
        m = re.search(r"[-+]?\d*\.?\d+", val)
        if m:
            try:
                return float(m.group())
            except ValueError:
                return None
        return None
    try:
        return float(val)
    except Exception:
        return None


def haversine_km(lat1, lon1, lat2, lon2):
    """Distância grande-círculo (km). Entradas em graus decimais."""
    for v in (lat1, lon1, lat2, lon2):
        if v is None or np.isnan(v):
            return np.nan
    R = 6371.0088                                # raio médio da Terra (km)
    φ1, λ1, φ2, λ2 = map(radians, (lat1, lon1, lat2, lon2))
    dφ, dλ = φ2 - φ1, λ2 - λ1
    a = sin(dφ/2)**2 + cos(φ1)*cos(φ2)*sin(dλ/2)**2
    c = 2*atan2(sqrt(a), sqrt(1-a))
    return R * c


# ==========================================================================
# =============================================================================
# LEITURA DA PLANILHA EXCEL  +  LIMPEZA  +  FILTRO POR RAIO
# =============================================================================
import re, math
import pandas as pd
from geopy.geocoders import Nominatim

# --------------------------------------------------------------
# AUXILIARES • caso não existam ainda no seu script
# --------------------------------------------------------------
def _parse_coord(txt):
    """
    Recebe algo como '-29.040298°' ou '-29,040298' e devolve float ou NaN.
    """
    if pd.isna(txt):
        return float('nan')
    txt = str(txt).strip()
    txt = txt.replace(",", ".").replace("°", "")
    try:
        return float(txt)
    except ValueError:
        return float('nan')

def haversine_km(lat1, lon1, lat2, lon2):
    """
    Distância esférica aproximada entre dois pontos (km).
    Retorna NaN se alguma coordenada faltar.
    """
    if any(math.isnan(x) for x in (lat1, lon1, lat2, lon2)):
        return float('nan')
    R = 6371.0088  # raio médio da Terra (km)
    phi1, phi2 = math.radians(lat1), math.radians(lat2)
    dphi       = math.radians(lat2 - lat1)
    dlambda    = math.radians(lon2 - lon1)
    a = (math.sin(dphi/2)**2 +
         math.cos(phi1) * math.cos(phi2) * math.sin(dlambda/2)**2)
    return 2 * R * math.asin(math.sqrt(a))

# ==============================================================
#  UTILIDADES DE CONVERSÃO NUMÉRICA E LEITURA DA PLANILHA
# ==============================================================

# ▸ Dependências mínimas
import re
import numbers
import pandas as pd
import numpy   as np
# geopy, haversine_km e _parse_coord devem estar importados/
# definidos no seu script principal ou em módulo auxiliar.

# --------------------------------------------------------------
# 1. HIGIENIZADOR NUMÉRICO
# --------------------------------------------------------------
def _to_float(valor):
    """
    Converte entradas brasileiras/financeiras em `float`.

    Aceita:
      • strings com R$ / pontos / vírgulas / espaços
      • ints, floats, Decimals, numpy numbers
      • valores ausentes (NaN, None, "", "—", etc.)

    Retorna:
      • `float`
      • `pd.NA` se não puder converter
    """
    # ── 1) Nulos continuam nulos ───────────────────────────────
    if pd.isna(valor):
        return pd.NA

    # ── 2) Já é número?  (int, float, numpy, Decimal…) ─────────
    if isinstance(valor, (numbers.Number, np.number)):
        try:
            return float(valor)
        except Exception:       # Decimal ou afins que falharem
            return pd.NA

    # ── 3) Caso seja string: limpeza br-pt ─────────────────────
    txt = str(valor).strip()

    # remove tudo que não seja dígito, vírgula, ponto ou sinal
    txt = re.sub(r"[^\d,.\-]", "", txt)

    # decide qual é separador decimal
    if "," in txt and "." in txt:
        # padrão "36.841,00" → "36841.00"
        txt = txt.replace(".", "").replace(",", ".")
    elif "," in txt:
        # padrão "36841,00"  → "36841.00"
        txt = txt.replace(",", ".")
    else:
        # padrão "36.841"    → "36841"
        txt = txt.replace(".", "")

    try:
        return float(txt)
    except ValueError:
        return pd.NA


#




###############################################################################
# HOMOGENEIZAR AMOSTRAS (DATAFRAME FILTRADO)
###############################################################################
def homogeneizar_amostras(dataframe_amostras_validas, dados_avaliando, fatores_do_usuario, finalidade_do_laudo):
    """
    Aplica os fatores de homogeneização às amostras e retorna uma lista com o valor unitário homogeneizado
    de cada amostra (em R$/m²). 
    """
    import math

    # Área do imóvel avaliado
    area_do_avaliando = float(dados_avaliando.get("AREA TOTAL", 0))

    # Fatores do imóvel avaliado
    f_avaliado_aprov = fator_aproveitamento(dados_avaliando.get("APROVEITAMENTO", "URBANO"))
    f_avaliado_topog = fator_topografia(dados_avaliando.get("BOA TOPOGRAFIA?", "NÃO"))
    f_avaliado_pedol = fator_pedologia(dados_avaliando.get("PEDOLOGIA ALAGÁVEL? ", "NÃO"))
    f_avaliado_pavim = fator_pavimentacao(dados_avaliando.get("PAVIMENTACAO?", "NÃO"))
    f_avaliado_esq   = fator_esquina(dados_avaliando.get(" ESQUINA?", "NÃO"))
    f_avaliado_acess = fator_acessibilidade(dados_avaliando.get("ACESSIBILIDADE?", "NÃO"))

    lista_valores_unitarios = []

    for _, linha in dataframe_amostras_validas.iterrows():
        valor_total_amostra = linha["VALOR TOTAL"]
        area_da_amostra = float(linha.get("AREA TOTAL", 0))

        # Cálculo dos fatores conforme a lógica original:
        fator_area = calcular_fator_area(area_do_avaliando, area_da_amostra, fatores_do_usuario["area"])
        fator_oferta = calcular_fator_oferta(True, fatores_do_usuario["oferta"])

        # Fator localização se "localizacao_mesma_regiao" for falso,
        # faz a comparação, senão = 1.0
        if fatores_do_usuario.get("localizacao_mesma_regiao", False):
            fator_localiz_calc = 1.0
        else:
            try:
                dist_amostra = float(linha.get("DISTANCIA CENTRO", 0))
                dist_avalia = float(dados_avaliando.get("DISTANCIA CENTRO", 0))
                if dist_amostra > 0 and dist_avalia > 0:
                    fa_item = 1.0 / (dist_amostra ** 0.1)
                    fa_avaliado = 1.0 / (dist_avalia ** 0.1)
                    fator_localiz_calc = limitar_fator(fa_avaliado / fa_item)
                else:
                    fator_localiz_calc = 1.0
            except:
                fator_localiz_calc = 1.0
            fator_localiz_calc = limitar_fator(fator_localiz_calc)

        # Fator aproveitamento (f_avaliado / f_amostra)
        f_sample_aprov = fator_aproveitamento(linha.get("APROVEITAMENTO", "URBANO"))
        if fatores_do_usuario["aproveitamento"] and f_sample_aprov != 0:
            fator_aprov_calc = limitar_fator(f_avaliado_aprov / f_sample_aprov)
        else:
            fator_aprov_calc = 1.0

        # Fator topografia
        f_sample_topog = fator_topografia(linha.get("BOA TOPOGRAFIA?", "NÃO"))
        if fatores_do_usuario["topografia"] and f_sample_topog != 0:
            fator_topog_calc = limitar_fator(f_avaliado_topog / f_sample_topog)
        else:
            fator_topog_calc = 1.0

        # Fator pedologia
        f_sample_pedol = fator_pedologia(linha.get("PEDOLOGIA ALAGÁVEL? ", "NÃO"))
        if fatores_do_usuario["pedologia"] and f_sample_pedol != 0:
            fator_pedol_calc = limitar_fator(f_avaliado_pedol / f_sample_pedol)
        else:
            fator_pedol_calc = 1.0

        # Fator pavimentação
        f_sample_pavim = fator_pavimentacao(linha.get("PAVIMENTACAO?", "NÃO"))
        if fatores_do_usuario["pavimentacao"] and f_sample_pavim != 0:
            fator_pavim_calc = limitar_fator(f_avaliado_pavim / f_sample_pavim)
        else:
            fator_pavim_calc = 1.0

        # Fator esquina
        f_sample_esq = fator_esquina(linha.get(" ESQUINA?", "NÃO"))
        if fatores_do_usuario["esquina"] and f_sample_esq != 0:
            fator_esq_calc = limitar_fator(f_avaliado_esq / f_sample_esq)
        else:
            fator_esq_calc = 1.0

        # Fator acessibilidade
        f_sample_acess = fator_acessibilidade(linha.get("ACESSIBILIDADE?", "NÃO"))
        if fatores_do_usuario["acessibilidade"] and f_sample_acess != 0:
            fator_acess_calc = limitar_fator(f_avaliado_acess / f_sample_acess)
        else:
            fator_acess_calc = 1.0

        # Valor homogeneizado
        valor_homog = (
            valor_total_amostra *
            fator_area *
            fator_oferta *
            fator_localiz_calc *
            fator_aprov_calc *
            fator_topog_calc *
            fator_pedol_calc *
            fator_pavim_calc *
            fator_esq_calc *
            fator_acess_calc
        )

        # Converte o valor total homogeneizado em valor unitário (R$/m²)
        if area_da_amostra > 0:
            valor_unitario = valor_homog / area_da_amostra
        else:
            valor_unitario = 0.0

        lista_valores_unitarios.append(valor_unitario)

    return lista_valores_unitarios


import os
import math
import unicodedata
from datetime import datetime

# --- Ciência de dados / Estatística ---
import numpy as np
import pandas as pd
import scipy.stats
from scipy.stats import gaussian_kde
from tqdm import tqdm

# --- Plotagem ---
import matplotlib.pyplot as plt
import matplotlib.ticker as ticker

# --- Manipulação de Word (python‑docx) ---
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement
from lxml import etree

from docx.oxml.ns import qn, nsdecls
from docx.text.paragraph import Paragraph

# --- Interface gráfica / Diálogos de arquivo ---

###############################################################################
# FUNÇÕES DE SUPORTE GERAIS
###############################################################################
def remover_acentos(texto):
    """
    Remove acentos de uma string.
    """
    return ''.join(
        caractere 
        for caractere in unicodedata.normalize('NFD', texto)
        if unicodedata.category(caractere) != 'Mn'
    )

def formatar_moeda_brasil(valor):
    """
    Formata o valor em Real brasileiro, trocando ponto por vírgula
    e inserindo ponto a cada milhar.
    Exemplo: 12345.6 => 'R$ 12.345,60'
    """
    formato_texto = f"{valor:,.2f}"
    formato_texto = formato_texto.replace(",", "X").replace(".", ",").replace("X", ".")
    return f"R$ {formato_texto}"

def formatar_numero_brasileiro(valor):
    """
    Formata um número em padrão brasileiro, com separador de milhar e decimal invertidos.
    Exemplo: 12345.6 => '12.345,60'
    """
    formato_texto = f"{valor:,.2f}"
    formato_texto = formato_texto.replace(",", "X").replace(".", ",").replace("X", ".")
    return formato_texto


###############################################################################
# FATORES ESPECÍFICOS (para homogeneização)
###############################################################################
def fator_aproveitamento(aproveitamento_texto):
    """
    Se aproveitamento for 'urbano' ou 'urbana', retorna 1.00
    Se for 'rural', retorna 0.80
    Caso contrário, retorna 1.00
    """
    if isinstance(aproveitamento_texto, str):
        valor = aproveitamento_texto.strip().lower()
        if valor in ["urbano", "urbana"]:
            return 1.00
        elif valor == "rural":
            return 0.80
    return 1.00

def fator_topografia(texto_topografia):
    """
    Se 'sim' em 'BOA TOPOGRAFIA?', retorna 1.10; senão, 1.00
    """
    if isinstance(texto_topografia, str):
        if "sim" in remover_acentos(texto_topografia.lower()):
            return 1.10
    return 1.00

def fator_pedologia(texto_pedologia):
    """
    Se 'sim' em 'PEDOLOGIA ALAGÁVEL?', retorna 0.70; senão, 1.00
    """
    if isinstance(texto_pedologia, str):
        if "sim" in remover_acentos(texto_pedologia.lower()):
            return 0.70
    return 1.00

def fator_pavimentacao(texto_pavimentacao):
    """
    Se 'sim' em 'PAVIMENTACAO?', retorna 1.00; senão, 0.90
    """
    if isinstance(texto_pavimentacao, str):
        if "sim" in remover_acentos(texto_pavimentacao.lower()):
            return 1.00
        else:
            return 0.90
    return 1.00

def fator_esquina(texto_esquina):
    """
    Se 'sim' em 'ESQUINA?', retorna 1.15; senão, 1.00
    """
    if isinstance(texto_esquina, str):
        if "sim" in remover_acentos(texto_esquina.lower()):
            return 1.15
    return 1.00

def fator_acessibilidade(texto_acessibilidade):
    """
    Se 'sim' em 'ACESSIBILIDADE?', retorna 1.00; senão, 0.90
    """
    if isinstance(texto_acessibilidade, str):
        if "sim" in remover_acentos(texto_acessibilidade.lower()):
            return 1.00
        else:
            return 0.90
    return 1.00


# ###############################################################################
# # INTERVALO DE CONFIANÇA (IC) VIA BOOTSTRAP DA MEDIANA - 80%
# ###############################################################################
# def intervalo_confianca_bootstrap_mediana(valores_numericos, numero_amostras=1000, nivel_confianca=0.80):
#     """
#     Calcula o intervalo de confiança (IC) para a mediana via bootstrap.
#     Retorna (limite_inferior, limite_superior).
#     """
#     array_valores = numpy.array(valores_numericos)
#     quantidade = len(array_valores)
#     lista_medianas = []
#     for _ in range(numero_amostras):
#         amostra_sorteada = numpy.random.choice(array_valores, size=quantidade, replace=True)
#         lista_medianas.append(numpy.median(amostra_sorteada))
#     array_medianas = numpy.array(lista_medianas)
#     limite_inferior = numpy.percentile(array_medianas, (1 - nivel_confianca) / 2 * 100)
#     limite_superior = numpy.percentile(array_medianas, (1 + nivel_confianca) / 2 * 100)
#     return limite_inferior, limite_superior

###############################################################################
# INTERVALO DE CONFIANÇA (IC) VIA BOOTSTRAP DA MEDIANA - 80% (robusto)
###############################################################################
def intervalo_confianca_bootstrap_mediana(valores_numericos, numero_amostras=1000, nivel_confianca=0.80):
    """
    Aceita lista/iterável de floats, numpy.ndarray, ou lista de dicts
    (usa 'valor_unitario' como padrão; fallback: 'valor_estimado'/'VUH'/'vuh').
    Retorna (LI, LS).
    """
    import math
    import numpy

    # 1) Normaliza entrada -> lista de floats finitos
    def _to_floats(seq):
        # Nunca use "seq or []" com numpy arrays
        if seq is None:
            iterable = []
        elif isinstance(seq, numpy.ndarray):
            iterable = seq.ravel().tolist()
        else:
            iterable = list(seq)

        out = []
        for v in iterable:
            if isinstance(v, dict):
                v = v.get("valor_unitario", v.get("valor_estimado", v.get("VUH", v.get("vuh", 0))))
            try:
                x = float(v)
            except (TypeError, ValueError):
                continue
            if math.isfinite(x):
                out.append(x)
        return out

    vals = _to_floats(valores_numericos)
    if not vals:
        return (0.0, 0.0)

    arr = numpy.asarray(vals, dtype=float)
    n = arr.size

    # Casos triviais
    if n == 1:
        m = float(arr[0])
        return (m, m)

    # 2) Bootstrap vetorizado
    rng = numpy.random.default_rng()
    samples = rng.choice(arr, size=(numero_amostras, n), replace=True)
    meds = numpy.median(samples, axis=1)

    alpha = 1.0 - float(nivel_confianca)
    li = float(numpy.percentile(meds, (alpha/2)*100))
    ls = float(numpy.percentile(meds, (1 - alpha/2)*100))
    return (li, ls)



###############################################################################
# CLASSIFICAÇÃO DO GRAU DE PRECISÃO (ABNT NBR 14653)
###############################################################################
def classificar_grau_de_precisao(amplitude_ic80):
    """
    Classifica o resultado quanto à precisão, de acordo com a amplitude do IC 80%
    e a tabela da ABNT NBR 14653.
    """
    if amplitude_ic80 <= 30:
        return "GRAU III"
    elif amplitude_ic80 <= 40:
        return "GRAU II"
    elif amplitude_ic80 <= 50:
        return "GRAU I"
    else:
        return "NÃO CLASSIFICADO"


# ###############################################################################
# # GRÁFICO DE DENSIDADE (KDE)
# ###############################################################################
# def gerar_grafico_densidade_kernel(valores_homogeneizados, nome_arquivo):
#     """
#     Gera o gráfico de densidade (KDE) e salva em 'nome_arquivo'.
#     Se o conjunto de dados tiver menos de 2 elementos, salva um gráfico com uma mensagem de aviso.
#     """
#     from scipy.stats import gaussian_kde
#     import numpy as np
#     import matplotlib.pyplot as plt

#     array_valores = np.array(valores_homogeneizados, dtype=float)
#     if array_valores.size < 2:
#         # Dados insuficientes para calcular o KDE.
#         plt.figure(figsize=(8, 6))
#         plt.text(0.5, 0.5, "Dados insuficientes para calcular KDE", 
#                  horizontalalignment='center', verticalalignment='center', 
#                  transform=plt.gca().transAxes, fontsize=12)
#         plt.title("Histograma de Densidade de Kernel (KDE)")
#         plt.xlabel("Valores Homogeneizados")
#         plt.ylabel("Densidade")
#         plt.tight_layout()
#         plt.savefig(nome_arquivo, bbox_inches='tight')
#         plt.close()
#         return

#     media_valores = np.mean(array_valores)
#     mediana_valores = np.median(array_valores)

#     eixo_x = np.linspace(array_valores.min(), array_valores.max(), 300)
#     funcao_densidade = gaussian_kde(array_valores)
#     valores_densidade = funcao_densidade(eixo_x)

#     plt.figure(figsize=(8, 6))
#     plt.fill_between(eixo_x, valores_densidade, alpha=0.6)
#     plt.title("Histograma de Densidade de Kernel (KDE)")
#     plt.xlabel("Valores Homogeneizados")
#     plt.ylabel("Densidade")
#     plt.axvline(x=media_valores, color='red', linestyle='--', label=f"Média: {media_valores:,.2f}")
#     plt.axvline(x=mediana_valores, color='green', linestyle='-', label=f"Mediana: {mediana_valores:,.2f}")
#     plt.legend()
#     plt.tight_layout()
#     plt.savefig(nome_arquivo, bbox_inches='tight')
#     plt.close()

###############################################################################
# GRÁFICO DE DENSIDADE (KDE) — robusto a lista de dicts ou floats
###############################################################################
def gerar_grafico_densidade_kernel(valores_homogeneizados, nome_arquivo):
    """
    Aceita:
      - lista de floats (valores unitários), ou
      - lista de dicts com 'valor_unitario' ou 'valor_estimado'.
    Gera a curva de densidade (KDE) e salva em nome_arquivo.
    """
    import numpy as np
    import matplotlib.pyplot as plt
    from scipy.stats import gaussian_kde

    # 1) Normaliza entrada → lista de floats
    if not valores_homogeneizados:
        plt.figure(figsize=(8, 6))
        plt.text(0.5, 0.5, "Sem dados para KDE",
                 ha='center', va='center', transform=plt.gca().transAxes, fontsize=12)
        plt.title("Densidade (KDE)")
        plt.xlabel("Valor unitário (homogeneizado)")
        plt.ylabel("Densidade")
        plt.tight_layout()
        plt.savefig(nome_arquivo, bbox_inches='tight', dpi=150)
        plt.close()
        return

    first = valores_homogeneizados[0]
    if isinstance(first, dict):
        vals = [
            float(d.get("valor_unitario", d.get("valor_estimado", 0)) or 0)
            for d in valores_homogeneizados
        ]
    else:
        vals = [float(v) for v in valores_homogeneizados]

    # 2) Limpeza: remove NaN/inf e valores <= 0 (KDE pode ficar ruim)
    vals = [v for v in vals if np.isfinite(v) and v > 0]

    # 3) Casos de borda
    if len(vals) < 2:
        plt.figure(figsize=(8, 6))
        if len(vals) == 1:
            v = vals[0]
            plt.axvline(v)
            plt.title("Densidade (KDE) — dado único")
            plt.xlabel("Valor unitário (homogeneizado)")
            plt.ylabel("Densidade")
            plt.tight_layout()
        else:
            plt.text(0.5, 0.5, "Dados insuficientes para calcular KDE",
                     ha='center', va='center', transform=plt.gca().transAxes, fontsize=12)
            plt.title("Densidade (KDE)")
            plt.xlabel("Valor unitário (homogeneizado)")
            plt.ylabel("Densidade")
            plt.tight_layout()
        plt.savefig(nome_arquivo, bbox_inches='tight', dpi=150)
        plt.close()
        return

    arr = np.array(vals, dtype=float)
    media = float(np.mean(arr))
    mediana = float(np.median(arr))

    xs = np.linspace(arr.min(), arr.max(), 300)
    kde = gaussian_kde(arr)
    ys = kde(xs)

    plt.figure(figsize=(8, 6))
    plt.fill_between(xs, ys, alpha=0.6)
    plt.title("Densidade (KDE)")
    plt.xlabel("Valor unitário (homogeneizado)")
    plt.ylabel("Densidade")
    plt.axvline(x=media, linestyle='--', label=f"Média: {media:,.2f}")
    plt.axvline(x=mediana, linestyle='-', label=f"Mediana: {mediana:,.2f}")
    plt.legend()
    plt.tight_layout()
    plt.savefig(nome_arquivo, bbox_inches='tight', dpi=150)
    plt.close()

###############################################################################
# APLICAÇÃO DO CRITÉRIO DE CHAUVENET (NOVO MÉTODO DE SANEAMENTO)
###############################################################################
def aplicar_chauvenet_e_filtrar(dataframe_amostras, limiar=1.0, usar_log=True, fator_mad=2.5):
    """
    Aplica (1) Log (opcional), (2) Critério de Chauvenet e (3) Filtro robusto via MAD.
    Retorna os dados filtrados e estatísticas básicas.

    Retorno:
    --------
    (
        dataframe_valido,           # DataFrame sem outliers (Chauvenet + MAD)
        indices_excluidos,          # Lista final de índices excluídos
        amostras_excluidas,         # Lista de identificações das amostras excluídas
        media_final,
        desvio_padrao_final,
        menor_valor_final,
        maior_valor_final,
        mediana_valor_final
    )
    """
    import math
    from math import erf
    import numpy as np
    
    # Se a coluna "VALOR TOTAL" não existir ou estiver vazia, retorne vazio
    if "VALOR TOTAL" not in dataframe_amostras.columns or len(dataframe_amostras) == 0:
        return (
            dataframe_amostras.copy(),
            [],
            [],
            0.0,
            0.0,
            0.0,
            0.0,
            0.0
        )
    
    # Cria uma cópia para não alterar o DataFrame original
    df_copy = dataframe_amostras.copy().reset_index(drop=True)
    
    # 1) Verifica se é possível usar log
    valores = df_copy["VALOR TOTAL"].values.astype(float)
    pode_usar_log = usar_log and np.all(valores > 0)
    
    if pode_usar_log:
        # Transformação log
        valores_transformados = np.log(valores)
    else:
        valores_transformados = valores
    
    # 2) Critério de Chauvenet no domínio transformado
    media_ch = np.mean(valores_transformados)
    desvio_ch = np.std(valores_transformados, ddof=1)
    n = len(valores_transformados)
    
    indices_validos_chauvenet = []
    indices_excluidos_chauvenet = []
    
    for idx, vt in enumerate(valores_transformados):
        if desvio_ch > 0:
            z = abs(vt - media_ch) / desvio_ch
        else:
            z = 0.0
        prob_in = 0.5 * (1 + erf(z / math.sqrt(2)))
        crit_chauvenet = n * prob_in

        if crit_chauvenet > limiar:
            indices_validos_chauvenet.append(idx)
        else:
            indices_excluidos_chauvenet.append(idx)
    
    df_chauvenet = df_copy.iloc[indices_validos_chauvenet].copy().reset_index(drop=True)
    
    # Lista de amostras excluídas pelo Chauvenet
    amostras_excl_chauvenet = []
    if "AM" in df_copy.columns:
        for idx_exc in indices_excluidos_chauvenet:
            amostras_excl_chauvenet.append(str(df_copy.iloc[idx_exc]["AM"]))
    else:
        for idx_exc in indices_excluidos_chauvenet:
            amostras_excl_chauvenet.append(f"Linha#{idx_exc+1}")
    
    # 3) Filtro robusto via MAD (Median Absolute Deviation)
    val_chauv = df_chauvenet["VALOR TOTAL"].values.astype(float)
    if pode_usar_log:
        val_chauv_transf = np.log(val_chauv)
    else:
        val_chauv_transf = val_chauv
    
    if len(val_chauv_transf) > 0:
        mediana_tf = np.median(val_chauv_transf)
        mad_raw = np.median(np.abs(val_chauv_transf - mediana_tf))
        if mad_raw == 0:
            # Se der zero (pouca variação), evita divisão por zero
            mad_raw = 1e-9
    else:
        mediana_tf = 0.0
        mad_raw = 1e-9
    
    indices_validos_mad = []
    indices_excluidos_mad = []
    
    for idx_m, vtf in enumerate(val_chauv_transf):
        z_rob = (vtf - mediana_tf) / (mad_raw * 1.4826)
        if abs(z_rob) > fator_mad:
            indices_excluidos_mad.append(idx_m)
        else:
            indices_validos_mad.append(idx_m)
    
    df_mad = df_chauvenet.iloc[indices_validos_mad].copy().reset_index(drop=True)

    # Identificações excluídas pelo MAD 
    amostras_excl_mad = []
    for idx_m2 in indices_excluidos_mad:
        idx_original_mad = df_chauvenet.index[idx_m2]
        if "AM" in df_copy.columns:
            amostras_excl_mad.append(str(df_copy.loc[idx_original_mad, "AM"]))
        else:
            amostras_excl_mad.append(f"Linha#{idx_original_mad+1}")
    
    # 4) Combinar exclusões: Chauvenet OU MAD
    set_chauv = set(indices_excluidos_chauvenet)
    indices_excl_mad_original = [df_chauvenet.index[idxk] for idxk in indices_excluidos_mad]
    set_mad = set(indices_excl_mad_original)
    set_excl_total = set_chauv.union(set_mad)
    indices_excluidos_final = sorted(list(set_excl_total))
    set_amostras_final = set(amostras_excl_chauvenet).union(set(amostras_excl_mad))
    list_amostras_excluidas_total = sorted(set_amostras_final)
    
    df_valido_final = df_mad.reset_index(drop=True)
    
    # 6) Estatísticas finais (no domínio original: "VALOR TOTAL")
    array_final = df_valido_final["VALOR TOTAL"].values.astype(float)
    
    if len(array_final) > 0:
        media_final = np.mean(array_final)
        desvio_padrao_final = np.std(array_final, ddof=1)
        menor_valor_final = array_final.min()
        maior_valor_final = array_final.max()
        mediana_valor_final = np.median(array_final)
    else:
        media_final = 0.0
        desvio_padrao_final = 0.0
        menor_valor_final = 0.0
        maior_valor_final = 0.0
        mediana_valor_final = 0.0
    
    return (
        df_valido_final,
        indices_excluidos_final,
        list_amostras_excluidas_total,
        media_final,
        desvio_padrao_final,
        menor_valor_final,
        maior_valor_final,
        mediana_valor_final
    )


###############################################################################
# CÁLCULO DE FATORES BÁSICOS
###############################################################################
def calcular_fator_area(area_do_avaliando, area_da_amostra, usar_fator_area):
    """
    Calcula o fator área se usar_fator_area for True.
    Mantém a fórmula padrão: (Área da Amostra / Área do Avaliado)^(1/4),
    limitado a [0.5, 1.4].
    """
    if not usar_fator_area or area_do_avaliando <= 0 or area_da_amostra <= 0:
        return 1.0
    fator_calculado = (area_da_amostra / area_do_avaliando) ** 0.25
    return max(0.5, min(1.4, fator_calculado))

def limitar_fator(x):
    """
    Limita o valor do fator x ao intervalo [0.50, 2.0].
    """
    return max(0.50, min(1.4, x))

def calcular_fator_oferta(oferta_aplicada, usar_fator_oferta):
    """
    Retorna 0.9 se usar_fator_oferta e oferta_aplicada forem True; senão, 1.0.
    (Fator fixo)
    """
    return 0.9 if (usar_fator_oferta and oferta_aplicada) else 1.0



###############################################################################
# GRÁFICOS DE ADERÊNCIA E DISPERSÃO
###############################################################################
# def gerar_grafico_aderencia_totais(dataframe, valores_homogeneizados_unitarios, nome_arquivo_imagem):
#     """
#     Gera um gráfico comparando os VALORES TOTAIS ORIGINAIS de cada amostra 
#     com os VALORES TOTAIS ESTIMADOS, calculados a partir do valor unitário homogeneizado (R$/m²)
#     multiplicado pela área de cada amostra.
#     """
#     import numpy as np
#     import matplotlib.pyplot as plt
#     from scipy.stats import linregress
#     import matplotlib.ticker as ticker

#     # 1) Obter os valores totais originais
#     valores_originais_totais = dataframe["VALOR TOTAL"].tolist()

#     # 2) Calcular os valores estimados
#     valores_estimados_totais = []
#     for i, valor_unit in enumerate(valores_homogeneizados_unitarios):
#         area = dataframe.iloc[i]["AREA TOTAL"]
#         if area > 0:
#             valor_total_estimado = valor_unit * area
#         else:
#             valor_total_estimado = 0.0
#         valores_estimados_totais.append(valor_total_estimado)

#     x = np.array(valores_originais_totais, dtype=float)
#     y = np.array(valores_estimados_totais, dtype=float)

#     fig, ax = plt.subplots(figsize=(8, 6))
#     ax.scatter(x, y, color='blue', label='Amostras')

#     if x.size > 0 and y.size > 0:
#         limite_min = min(np.min(x), np.min(y))
#         limite_max = max(np.max(x), np.max(y))
#     else:
#         limite_min, limite_max = 0, 1

#     if len(x) >= 2 and len(y) >= 2:
#         slope, intercept, r_value, p_value, std_err = linregress(x, y)
#         x_fit = np.linspace(limite_min, limite_max, 100)
#         y_fit = slope * x_fit + intercept
#         ax.plot(x_fit, y_fit, 'r-', label=f'Reta Ajustada (R² = {r_value**2:.2f})')
#     else:
#         ax.text(0.5, 0.5, "Dados insuficientes para regressão", 
#                 horizontalalignment='center', verticalalignment='center', 
#                 transform=ax.transAxes, fontsize=12, color='red')

#     ax.set_title("Gráfico de Aderência - Valores Totais")
#     ax.set_xlabel("Valor Total Original (R$)")
#     ax.set_ylabel("Valor Total Estimado (R$)")
#     ax.legend()
#     ax.grid(True)
#     ax.tick_params(axis='x', rotation=45)
   
#     def formatar_valor_em_reais(valor, pos):
#         return formatar_moeda_brasil(valor)

#     formatador = ticker.FuncFormatter(formatar_valor_em_reais)
#     ax.xaxis.set_major_formatter(formatador)
#     ax.yaxis.set_major_formatter(formatador)

#     fig.tight_layout()
#     fig.savefig(nome_arquivo_imagem, bbox_inches='tight')
#     plt.close(fig)
def gerar_grafico_aderencia_totais(dataframe, valores_homogeneizados_unitarios, nome_arquivo_imagem):
    """
    Gera um gráfico comparando os VALORES TOTAIS ORIGINAIS de cada amostra
    com os VALORES TOTAIS ESTIMADOS, calculados a partir do valor unitário homogeneizado (R$/m²)
    multiplicado pela área de cada amostra.
    """
    import numpy as np
    import matplotlib.pyplot as plt
    from scipy.stats import linregress
    import matplotlib.ticker as ticker

    # 1) Obter os valores totais originais
    valores_originais_totais = dataframe["VALOR TOTAL"].tolist()

    # 2) Calcular os valores estimados a partir do VUH e da área
    valores_estimados_totais = []
    for i, valor_unit in enumerate(valores_homogeneizados_unitarios):
        area = float(dataframe.iloc[i]["AREA TOTAL"])
        valor_total_estimado = (float(valor_unit) * area) if area > 0 else 0.0
        valores_estimados_totais.append(valor_total_estimado)

    x = np.asarray(valores_originais_totais, dtype=float)
    y = np.asarray(valores_estimados_totais, dtype=float)

    # --- Deriva os identificadores para rotular os pontos (1, 2, 3...) ---
    if "idx" in dataframe.columns:
        try:
            idxs = dataframe["idx"].astype(int).tolist()
        except Exception:
            idxs = list(range(1, len(x) + 1))
    elif "AM" in dataframe.columns:
        try:
            idxs = dataframe["AM"].astype(int).tolist()
        except Exception:
            idxs = list(range(1, len(x) + 1))
    else:
        idxs = list(range(1, len(x) + 1))

    # Garante mesmo comprimento (evita "x and y must be the same size")
    n = min(len(x), len(y), len(idxs))
    x, y, idxs = x[:n], y[:n], idxs[:n]

    fig, ax = plt.subplots(figsize=(8, 6))
    ax.scatter(x, y, label="Amostras")

    # --- Anota cada ponto com o número da amostra ---
    for xi, yi, i_amostra in zip(x, y, idxs):
        ax.annotate(
            f"{int(i_amostra)}", (float(xi), float(yi)),
            xytext=(3, 3), textcoords="offset points",
            fontsize=8, color="black",
            bbox=dict(boxstyle="round,pad=0.1", fc="white", alpha=0.7)
        )

    # Limites e regressão
    if x.size > 0 and y.size > 0:
        limite_min = min(np.min(x), np.min(y))
        limite_max = max(np.max(x), np.max(y))
    else:
        limite_min, limite_max = 0.0, 1.0

    if n >= 2:
        slope, intercept, r_value, p_value, std_err = linregress(x, y)
        x_fit = np.linspace(limite_min, limite_max, 100)
        y_fit = slope * x_fit + intercept
        ax.plot(x_fit, y_fit, "r-", label=f"Reta Ajustada (R² = {r_value**2:.2f})")
    else:
        ax.text(0.5, 0.5, "Dados insuficientes para regressão",
                ha="center", va="center", transform=ax.transAxes, fontsize=12, color="red")

    ax.set_title("Gráfico de Aderência - Valores Totais")
    ax.set_xlabel("Valor Total Original (R$)")
    ax.set_ylabel("Valor Total Estimado (R$)")
    ax.legend()
    ax.grid(True)
    ax.tick_params(axis="x", rotation=45)

    def formatar_valor_em_reais(valor, pos):
        return formatar_moeda_brasil(valor)

    formatador = ticker.FuncFormatter(formatar_valor_em_reais)
    ax.xaxis.set_major_formatter(formatador)
    ax.yaxis.set_major_formatter(formatador)

    fig.tight_layout()
    fig.savefig(nome_arquivo_imagem, bbox_inches="tight")
    plt.close(fig)


### essa é a original do PAULO
# def gerar_grafico_dispersao_mediana(valores_homogeneizados, nome_arquivo):
#     """
#     Gera um gráfico de dispersão simples (index vs. valores homogeneizados)
#     e destaca a mediana com uma linha horizontal.
#     """
#     import numpy as np
#     import matplotlib.pyplot as plt

#     arr = np.array(valores_homogeneizados, dtype=float)
#     if arr.size < 1:
#         plt.figure()
#         plt.text(0.5, 0.5, "Sem valores para exibir", 
#                  ha='center', va='center', 
#                  transform=plt.gca().transAxes, fontsize=12)
#         plt.title("Dispersão dos Valores Homogeneizados")
#         plt.savefig(nome_arquivo, bbox_inches='tight')
#         plt.close()
#         return

#     indices = np.arange(1, len(arr) + 1)

#     plt.figure(figsize=(8, 6))
#     plt.scatter(indices, arr, marker='o', label="Valores Homogeneizados")
#     mediana = np.median(arr)
#     plt.axhline(y=mediana, color='r', linestyle='--', label=f"Mediana: {mediana:,.2f}")

#     plt.xlabel("Índice da Amostra")
#     plt.ylabel("Valor Unitário Homogeneizado (R$/m²)")
#     plt.title("Gráfico de Dispersão dos Valores Homogeneizados")
#     plt.legend()
#     plt.tight_layout()
#     plt.savefig(nome_arquivo, bbox_inches='tight')
#     plt.close()

# ACRESCIMO PARA VISUALIZR OS GRAFICOS COM CORES DIFERENTES


##### essa funcionou razoavelmente bem
# def gerar_grafico_dispersao_mediana(
#     homog,
#     caminho_saida,
#     idx_amostras_ativas,
#     idx_amostras_usuario_retirou,
#     idx_amostras_chauvenet_retirou
# ):
#     plt.figure(figsize=(8, 6))

#     # Ativos válidos (ativos - chauvenet)
#     ativos_validos_idx = [
#         idx for idx in idx_amostras_ativas if idx not in idx_amostras_chauvenet_retirou
#     ]
#     ativos_validos_valores = []
#     for idx in ativos_validos_idx:
#         if idx in idx_amostras_ativas:
#             pos = idx_amostras_ativas.index(idx)
#             if pos < len(homog):  # segurança extra
#                 ativos_validos_valores.append(homog[pos])

#     plt.scatter(ativos_validos_idx, ativos_validos_valores, color='blue', label='Amostras Ativas')

#     # Usuário retirou (ativos - usuario)
#     usuario_retirou_idx_filtrados = [
#         idx for idx in idx_amostras_usuario_retirou if idx in idx_amostras_ativas
#     ]
#     usuario_retirou_valores = []
#     for idx in usuario_retirou_idx_filtrados:
#         if idx in idx_amostras_ativas:
#             pos = idx_amostras_ativas.index(idx)
#             if pos < len(homog):  # segurança extra
#                 usuario_retirou_valores.append(homog[pos])

#     if usuario_retirou_idx_filtrados and usuario_retirou_valores:
#         plt.scatter(usuario_retirou_idx_filtrados, usuario_retirou_valores, color='gray', label='Retiradas pelo Usuário')

#     # Chauvenet retirou
#     chauvenet_idx_filtrados = [
#         idx for idx in idx_amostras_chauvenet_retirou if idx in idx_amostras_ativas
#     ]
#     chauvenet_valores = []
#     for idx in chauvenet_idx_filtrados:
#         if idx in idx_amostras_ativas:
#             pos = idx_amostras_ativas.index(idx)
#             if pos < len(homog):  # segurança extra
#                 chauvenet_valores.append(homog[pos])

#     if chauvenet_idx_filtrados and chauvenet_valores:
#         plt.scatter(chauvenet_idx_filtrados, chauvenet_valores, color='red', label='Retiradas Chauvenet')

#     # Mediana linha (seguro)
#     if ativos_validos_valores:
#         plt.axhline(np.median(ativos_validos_valores), color='green', linestyle='--',
#                     label=f'Mediana: {np.median(ativos_validos_valores):.2f}')

#     plt.xlabel('Índice da Amostra')
#     plt.ylabel('Valor Unitário Homogeneizado (R$/m²)')
#     plt.title('Gráfico de Dispersão das Amostras Selecionadas')
#     plt.legend()
#     plt.grid(True)
#     plt.tight_layout()

#     plt.savefig(caminho_saida)
#     plt.close()


#     import matplotlib.pyplot as plt
# import numpy as np



# def gerar_grafico_dispersao_mediana(
#     df_filtrado,
#     homog,
#     caminho_saida,
#     idx_amostras_ativas,
#     idx_amostras_usuario_retirou,
#     idx_amostras_chauvenet_retirou
# ):
#     plt.figure(figsize=(8, 6))

#     # Mapeamento claro e explícito (idx_amostra: valor homog)
#     mapa_homog = dict(zip(idx_amostras_ativas, homog))

#     # Ativos válidos (ativos - chauvenet)
#     # Ativos válidos (ativos - chauvenet) - versão segura e robusta definitiva
#     ativos_validos_idx = []
#     ativos_validos_valores = []

#     for idx in idx_amostras_ativas:
#         if idx not in idx_amostras_chauvenet_retirou and idx in mapa_homog:
#             ativos_validos_idx.append(idx)
#             ativos_validos_valores.append(mapa_homog[idx])

#     plt.scatter(ativos_validos_idx, ativos_validos_valores, color='blue', label='Amostras Ativas')

#     # Retiradas pelo usuário (em cinza)
#     usuario_retirou_idx_filtrados = [
#         idx for idx in idx_amostras_usuario_retirou if idx in mapa_homog
#     ]
#     usuario_retirou_valores = [mapa_homog[idx] for idx in usuario_retirou_idx_filtrados]

#     if usuario_retirou_idx_filtrados:
#         plt.scatter(usuario_retirou_idx_filtrados, usuario_retirou_valores, color='gray', label='Retiradas pelo Usuário')

#     # Retiradas por Chauvenet (em vermelho)
#     chauvenet_idx_filtrados = [
#         idx for idx in idx_amostras_chauvenet_retirou if idx in mapa_homog
#     ]
#     chauvenet_valores = [mapa_homog[idx] for idx in chauvenet_idx_filtrados]

#     if chauvenet_idx_filtrados:
#         plt.scatter(chauvenet_idx_filtrados, chauvenet_valores, color='red', label='Retiradas Chauvenet')

#     # Linha mediana (seguros)
#     if ativos_validos_valores:
#         plt.axhline(np.median(ativos_validos_valores), color='green', linestyle='--',
#                     label=f'Mediana: {np.median(ativos_validos_valores):.2f}')

#     plt.xlabel('Índice da Amostra')
#     plt.ylabel('Valor Unitário Homogeneizado (R$/m²)')
#     plt.title('Gráfico de Dispersão das Amostras Selecionadas')
#     plt.legend()
#     plt.grid(True)
#     plt.tight_layout()

#     plt.savefig(caminho_saida)
#     plt.close()

# def gerar_grafico_dispersao_mediana(
#     df_filtrado,
#     homog,
#     caminho_saida,
#     ativos_frontend,
#     amostras_usuario_retirou,
#     amostras_chauvenet_retirou  # ← adicione explicitamente este argumento faltante
# ):
#     import matplotlib.pyplot as plt
#     import numpy as np

#     plt.figure(figsize=(8, 6))

#     mapa_homog = dict(zip(df_filtrado["idx"], homog))

#     ativos_validos_idx = [
#         idx for idx in ativos_frontend if idx not in amostras_chauvenet_retirou
#     ]

#     ativos_validos_valores = [mapa_homog[idx] for idx in ativos_validos_idx if idx in mapa_homog]

#     plt.scatter(ativos_validos_idx, ativos_validos_valores, color='blue', label='Amostras Ativas')

#     # Amostras excluídas explicitamente pelo usuário
#     usuario_retirou_valores = [mapa_homog[idx] for idx in amostras_usuario_retirou if idx in mapa_homog]
#     if amostras_usuario_retirou and usuario_retirou_valores:
#         plt.scatter(amostras_usuario_retirou, usuario_retirou_valores, color='gray', label='Retiradas pelo Usuário')

#     # Amostras excluídas por Chauvenet
#     chauvenet_valores = [mapa_homog[idx] for idx in amostras_chauvenet_retirou if idx in mapa_homog]
#     if amostras_chauvenet_retirou and chauvenet_valores:
#         plt.scatter(amostras_chauvenet_retirou, chauvenet_valores, color='red', label='Retiradas Chauvenet')

#     # Linha da mediana
#     if ativos_validos_valores:
#         plt.axhline(np.median(ativos_validos_valores), color='green', linestyle='--',
#                     label=f'Mediana: {np.median(ativos_validos_valores):.2f}')

#     plt.xlabel('Índice da Amostra')
#     plt.ylabel('Valor Unitário Homogeneizado (R$/m²)')
#     plt.title('Gráfico de Dispersão das Amostras Selecionadas')
#     plt.legend()
#     plt.grid(True)
#     plt.tight_layout()

#     plt.savefig(caminho_saida)
#     plt.close()

def gerar_grafico_dispersao_mediana(
    df_filtrado,
    homog,
    caminho_saida,
    ativos_frontend,
    amostras_usuario_retirou,
    amostras_chauvenet_retirou  # ← permanece explícito
):
    import matplotlib.pyplot as plt
    import numpy as np

    # --- helper: extrai float de item que pode ser dict ou número ---
    def _to_float(v):
        if isinstance(v, dict):
            v = v.get("valor_unitario", v.get("valor_estimado", v.get("VUH", v.get("vuh", 0.0))))
        try:
            return float(v)
        except (TypeError, ValueError):
            return None

    # Mapa idx -> valor (já em float), respeitando a ordem de df_filtrado
    idxs = df_filtrado["idx"].tolist()
    valores = [ _to_float(h) for h in homog ]
    mapa_homog = {i: v for i, v in zip(idxs, valores) if v is not None}

    # Separações com consistência (só o que existe no mapa)
    ativos_validos_idx = [
        i for i in ativos_frontend
        if (i in mapa_homog) and (i not in set(amostras_chauvenet_retirou))
    ]
    ativos_validos_valores = [ mapa_homog[i] for i in ativos_validos_idx ]

    user_out_idx = [ i for i in amostras_usuario_retirou if i in mapa_homog ]
    user_out_vals = [ mapa_homog[i] for i in user_out_idx ]

    chauv_out_idx = [ i for i in amostras_chauvenet_retirou if i in mapa_homog ]
    chauv_out_vals = [ mapa_homog[i] for i in chauv_out_idx ]

    plt.figure(figsize=(8, 6))

    # Plot: ativas
    if ativos_validos_idx and ativos_validos_valores:
        plt.scatter(ativos_validos_idx, ativos_validos_valores, label='Amostras Ativas')
        # rótulos numéricos ao lado dos pontos ativos
        for x, y in zip(ativos_validos_idx, ativos_validos_valores):
            plt.annotate(
                f"{int(x)}", (x, y),
                xytext=(3, 3), textcoords="offset points",
                fontsize=8
            )

    # Plot: retiradas pelo usuário
    if user_out_idx and user_out_vals:
        plt.scatter(user_out_idx, user_out_vals, label='Retiradas pelo Usuário')
        for x, y in zip(user_out_idx, user_out_vals):
            plt.annotate(
                f"{int(x)}", (x, y),
                xytext=(3, 3), textcoords="offset points",
                fontsize=8
            )

    # Plot: retiradas por Chauvenet
    if chauv_out_idx and chauv_out_vals:
        plt.scatter(chauv_out_idx, chauv_out_vals, label='Retiradas Chauvenet')
        for x, y in zip(chauv_out_idx, chauv_out_vals):
            plt.annotate(
                f"{int(x)}", (x, y),
                xytext=(3, 3), textcoords="offset points",
                fontsize=8
            )

    # Linha da mediana (apenas das ativas válidas)
    if ativos_validos_valores:
        med = float(np.median(ativos_validos_valores))
        plt.axhline(med, linestyle='--', label=f'Mediana: {med:.2f}')

    plt.xlabel('Índice da Amostra')
    plt.ylabel('Valor Unitário Homogeneizado (R$/m²)')
    plt.title('Gráfico de Dispersão das Amostras Selecionadas')
    plt.legend()
    plt.grid(True)
    plt.tight_layout()
    plt.savefig(caminho_saida)
    plt.close()





    
#########################################################################################################################
# TABELA DE AMOSTRAS HOMOGENEIZADAS
##########################################################################################################################
def calcular_detalhes_amostras(dataframe_amostras_validas, dados_avaliando, fatores_do_usuario, finalidade_do_laudo):
    """
    Monta uma lista de dicionários com os detalhes e o 'Valor Total Homogeneizado'
    para gerar a tabela final.
    
    As colunas "VU" e "VUH" são calculadas, respectivamente, com os valores unitários originais
    e os valores unitários homogenizados.
    """
    import math

    lista_detalhes = []
    area_do_avaliando = float(dados_avaliando.get("AREA TOTAL", 0))

    # Fatores do Avaliado
    f_avaliado_aprov = fator_aproveitamento(dados_avaliando.get("APROVEITAMENTO", "URBANO"))
    f_avaliado_topog = fator_topografia(dados_avaliando.get("BOA TOPOGRAFIA?", "NÃO"))
    f_avaliado_pedol = fator_pedologia(dados_avaliando.get("PEDOLOGIA ALAGÁVEL? ", "NÃO"))
    f_avaliado_pavim = fator_pavimentacao(dados_avaliando.get("PAVIMENTACAO?", "NÃO"))
    f_avaliado_esq   = fator_esquina(dados_avaliando.get(" ESQUINA?", "NÃO"))
    f_avaliado_acess = fator_acessibilidade(dados_avaliando.get("ACESSIBILIDADE?", "NÃO"))

    for indice, linha in dataframe_amostras_validas.iterrows():
        identificador_amostra = str(linha.get("AM", indice + 1))
        valor_total_amostra = linha["VALOR TOTAL"]
        area_amostra = float(linha.get("AREA TOTAL", 0))

        # Fatores básicos
        fator_area = calcular_fator_area(area_do_avaliando, area_amostra, fatores_do_usuario["area"])
        fator_oferta = calcular_fator_oferta(True, fatores_do_usuario["oferta"])

        # Cálculo dos fatores individuais (f_avaliado / f_amostra)
        f_sample_aprov = fator_aproveitamento(linha.get("APROVEITAMENTO", "URBANO"))
        if fatores_do_usuario["aproveitamento"] and f_sample_aprov != 0:
            fator_aproveitamento_calculado = f_avaliado_aprov / f_sample_aprov
        else:
            fator_aproveitamento_calculado = 1.0

        f_sample_topog = fator_topografia(linha.get("BOA TOPOGRAFIA?", "NÃO"))
        if fatores_do_usuario["topografia"] and f_sample_topog != 0:
            fator_topografia_calculado = f_avaliado_topog / f_sample_topog
        else:
            fator_topografia_calculado = 1.0

        f_sample_pedol = fator_pedologia(linha.get("PEDOLOGIA ALAGÁVEL? ", "NÃO"))
        if fatores_do_usuario["pedologia"] and f_sample_pedol != 0:
            fator_pedologia_calculado = f_avaliado_pedol / f_sample_pedol
        else:
            fator_pedologia_calculado = 1.0

        f_sample_pavim = fator_pavimentacao(linha.get("PAVIMENTACAO?", "NÃO"))
        if fatores_do_usuario["pavimentacao"] and f_sample_pavim != 0:
            fator_pavimentacao_calculado = f_avaliado_pavim / f_sample_pavim
        else:
            fator_pavimentacao_calculado = 1.0

        f_sample_esq = fator_esquina(linha.get(" ESQUINA?", "NÃO"))
        if fatores_do_usuario["esquina"] and f_sample_esq != 0:
            fator_esquina_calculado = f_avaliado_esq / f_sample_esq
        else:
            fator_esquina_calculado = 1.0

        f_sample_acess = fator_acessibilidade(linha.get("ACESSIBILIDADE?", "NÃO"))
        if fatores_do_usuario["acessibilidade"] and f_sample_acess != 0:
            fator_acessibilidade_calculado = f_avaliado_acess / f_sample_acess
        else:
            fator_acessibilidade_calculado = 1.0

        # Fator localização
        if fatores_do_usuario.get("localizacao_mesma_regiao", False):
            fator_localizacao_calculado = 1.0
        else:
            try:
                distancia_amostra = float(linha.get("DISTANCIA CENTRO", 0))
                distancia_avaliando = float(dados_avaliando.get("DISTANCIA CENTRO", 0))
                if distancia_amostra > 0 and distancia_avaliando > 0:
                    fator_item_comparativo = 1 / math.pow(distancia_amostra, 1/10)
                    fator_bem_avaliando = 1 / math.pow(distancia_avaliando, 1/10)
                    fator_localizacao_calculado = fator_bem_avaliando / fator_item_comparativo
                    if fator_localizacao_calculado > 1.40:
                        fator_localizacao_calculado = 1.40
                    elif fator_localizacao_calculado < 0.50:
                        fator_localizacao_calculado = 0.50
                else:
                    fator_localizacao_calculado = 1.0
            except:
                fator_localizacao_calculado = 1.0

        # Cálculo do valor total homogenizado aplicando todos os fatores
        valor_total_homogeneizado = (
            valor_total_amostra *
            fator_area *
            fator_oferta *
            fator_localizacao_calculado *
            fator_aproveitamento_calculado *
            fator_topografia_calculado *
            fator_pedologia_calculado *
            fator_pavimentacao_calculado *
            fator_esquina_calculado *
            fator_acessibilidade_calculado
        )

        # Cálculo dos valores unitários:
        # VU  => Valor unitário original (da planilha)
        # VUH => Valor unitário homogenizado
        if area_amostra > 0:
            vu = valor_total_amostra / area_amostra
            vuh = valor_total_homogeneizado / area_amostra
        else:
            vu = 0.0
            vuh = 0.0

        linha_detalhes = {
            "AM": identificador_amostra,
            "AREA": formatar_numero_brasileiro(area_amostra),
            "VU": formatar_moeda_brasil(vu),
            "FA": f"{fator_area:.2f}",
            "FO": f"{fator_oferta:.2f}",
            "FAP": f"{fator_aproveitamento_calculado:.2f}",
            "FT": f"{fator_topografia_calculado:.2f}",
            "FP": f"{fator_pedologia_calculado:.2f}",
            "FPA": f"{fator_pavimentacao_calculado:.2f}",
            "FE": f"{fator_esquina_calculado:.2f}",
            "FAC": f"{fator_acessibilidade_calculado:.2f}",
            "FL": f"{fator_localizacao_calculado:.2f}",
            "VUH": formatar_moeda_brasil(vuh)
        }
        lista_detalhes.append(linha_detalhes)

    return lista_detalhes


def inserir_tabela_amostras_calculadas(documento, lista_detalhes, col_widths=None):
    """
    Insere a tabela de amostras homogeneizadas no local do placeholder [tabelaSimilares].
    
    As colunas da tabela são:
      "AM", "AREA", "VU",
      "FA", "FO", "FAP", "FT", "FP", "FPA",
      "FE", "FAC", "FL", "VUH"
      
    Nesta versão, para as colunas de fatores (FA, FO, FAP, FT, FP, FPA, FE, FAC, FL),
    o valor exibido é convertido para float, limitado ao intervalo [0.50, 2.0] pela função
    limitar_fator() e formatado com duas casas decimais.
    """
    from docx.shared import Pt, Inches
    from lxml import etree

    from docx.oxml.ns import nsdecls
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.table import WD_TABLE_ALIGNMENT

    if not lista_detalhes:
        return

    # Definição padrão de larguras: VU e VUH um pouco maiores
    if col_widths is None:
        col_widths = [
            0.6,  # AM
            1.2,  # AREA
            1.5,  # VU
            0.6,  # FA
            0.6,  # FO
            0.6,  # FAP
            0.6,  # FT
            0.6,  # FP
            0.6,  # FPA
            0.6,  # FE
            0.6,  # FAC
            0.6,  # FL
            1.5   # VUH
        ]

    # Lista das colunas que representam fatores
    colunas_fator = ["FA", "FO", "FAP", "FT", "FP", "FPA", "FE", "FAC", "FL"]

    # Achar o placeholder
    for paragrafo in documento.paragraphs:
        if "[tabelaSimilares]" in paragrafo.text:
            # Remove o placeholder
            paragrafo.text = paragrafo.text.replace("[tabelaSimilares]", "")

            nomes_colunas = [
                "AM", "AREA", "VU",
                "FA", "FO", "FAP", "FT", "FP", "FPA",
                "FE", "FAC", "FL", "VUH"
            ]
            qtd_colunas = len(nomes_colunas)
            qtd_linhas = len(lista_detalhes) + 1

            # Cria a tabela
            tabela = documento.add_table(rows=qtd_linhas, cols=qtd_colunas)
            tabela.allow_autofit = False
            tabela.style = "Table Grid"
            tabela.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Ajusta as larguras de cada coluna
            for i, w_inches in enumerate(col_widths):
                for row in tabela.rows:
                    row.cells[i].width = Inches(w_inches)

            # Cabeçalho com sombreamento azul-claro
            shading_azul_claro = etree.fromstring(
                r'<w:shd {} w:val="clear" w:fill="BDD7EE"/>'.format(nsdecls('w'))
            )
            # Preenche o cabeçalho
            for col_idx, nome in enumerate(nomes_colunas):
                cell_head = tabela.rows[0].cells[col_idx]
                cell_head.text = nome
                cell_head._tc.get_or_add_tcPr().append(shading_azul_claro)
                for run in cell_head.paragraphs[0].runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)
                    run.font.bold = True
                cell_head.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Preenche as linhas de dados
            for i, dic_amostra in enumerate(lista_detalhes, start=1):
                for col_idx, nome_col in enumerate(nomes_colunas):
                    cell_data = tabela.rows[i].cells[col_idx]

                    
                    
                    # Se a coluna representa um fator, converte, aplica *clamp específico* e formata
                    if nome_col in colunas_fator:
                        try:
                            valor_num = float(dic_amostra.get(nome_col, 0))
                            # Aplica a limitação ao intervalo [0.50, 2.0]
                            valor_cel = f"{limitar_fator(valor_num):.2f}"
                        except Exception:
                            valor_cel = str(dic_amostra.get(nome_col, ""))
                    else:
                        valor_cel = str(dic_amostra.get(nome_col, ""))



                    cell_data.text = valor_cel

                    for run in cell_data.paragraphs[0].runs:
                        run.font.name = "Arial"
                        run.font.bold = False
                        # Para as colunas "VU" e "VUH", utiliza fonte menor
                        if nome_col in ["VU", "VUH"]:
                            run.font.size = Pt(8)
                        else:
                            run.font.size = Pt(9)
                    cell_data.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Insere a tabela após o parágrafo
            paragrafo._p.addnext(tabela._element)

            # Insere legenda abaixo da tabela
            paragrafo_legenda = inserir_paragrafo_apos(paragrafo, "")
            texto_legenda = (
                "Legendas:\n"
                "- AM = Amostra\n"
                "- AREA = Área do Imóvel (m²)\n"
                "- VU = Valor Unitário Ofertado\n"
                "- FA = Fator Área\n"
                "- FO = Fator Oferta\n"
                "- FAP = Fator Aproveitamento\n"
                "- FT = Fator Topografia\n"
                "- FP = Fator Pedologia\n"
                "- FPA = Fator Pavimentação\n"
                "- FE = Fator Esquina\n"
                "- FAC = Fator Acessibilidade\n"
                "- FL = Fator Localização\n"
                "- VUH = Valor Unitário Homogeneizado\n"
            )
            paragrafo_legenda2 = inserir_paragrafo_apos(paragrafo_legenda, texto_legenda)
            for run in paragrafo_legenda2.runs:
                run.font.name = "Arial"
                run.font.size = Pt(9)
            break


#######################################################################
# FUNÇÕES DE FORMATAÇÃO
#######################################################################
def inserir_tabela_amostras_originais(documento, dataframe):
    """
    Substitui o placeholder [amostras original] pela tabela de amostras originais,
    com as colunas: AM, VALOR TOTAL, ÁREA TOTAL (m²), VALOR UNITÁRIO (R$/m²), CIDADE, FONTE.
    Agora, deixamos um espaço um pouco maior entre as linhas.
    """
    from docx.shared import Pt, Inches
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement
    from lxml import etree

    from docx.oxml.ns import nsdecls, qn

    # Ajuste conforme as larguras desejadas (em polegadas) para cada coluna
    col_widths = [0.2, 1.3, 1.1, 0.8, 2.0, 2.9]

    # Títulos visíveis no cabeçalho
    colunas_visiveis = [
        "AM",
        "VALOR TOTAL",
        "ÁREA TOTAL (m²)",
        "VALOR UNITÁRIO (R$/m²)",
        "CIDADE",
        "FONTE"
    ]

    # Colunas correspondentes do DataFrame (caso precise filtrar ou renomear)
    colunas_df = [
        "idx",
        "VALOR TOTAL",
        "AREA TOTAL",
        "valor_unitario",
        "cidade",
        "fonte"
    ]

    # Localiza o parágrafo onde o placeholder [amostras original] está
    paragrafo_alvo = None
    for paragrafo in documento.paragraphs:
        if "[amostras original]" in paragrafo.text:
            paragrafo_alvo = paragrafo
            break

    # Se não encontrou o placeholder, não faz nada
    if not paragrafo_alvo:
        return

    # Remove o texto do placeholder
    paragrafo_alvo.text = paragrafo_alvo.text.replace("[amostras original]", "")

    # Número de linhas = registros do dataframe + 1 (para o cabeçalho)
    num_linhas = len(dataframe) + 1
    # Número de colunas = quantidade de títulos visíveis
    num_colunas = len(colunas_visiveis)

    # Cria a tabela
    tabela = documento.add_table(rows=num_linhas, cols=num_colunas, style="Table Grid")
    tabela.allow_autofit = False
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Função para centralizar verticalmente a célula
    def set_vertical_alignment(cell):
        tcPr = cell._tc.get_or_add_tcPr()
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), "center")
        tcPr.append(vAlign)

    # --- Cabeçalho ---
    for c, titulo_exib in enumerate(colunas_visiveis):
        cell_header = tabela.rows[0].cells[c]
        cell_header.text = titulo_exib

        # Fundo azul claro no cabeçalho
        shading_xml = etree.fromstring(
            f'<w:shd {nsdecls("w")} w:fill="BDD7EE" w:val="clear"/>'
        )
        cell_header._tc.get_or_add_tcPr().append(shading_xml)

        # Formatação da fonte do cabeçalho
        for run in cell_header.paragraphs[0].runs:
            run.font.name = "Arial"
            run.font.size = Pt(10)
            run.font.bold = True

        # Alinhamento horizontal e vertical do cabeçalho
        cell_header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_vertical_alignment(cell_header)

    # --- Linhas de dados ---
    for i, (_, row) in enumerate(dataframe.iterrows(), start=1):
        # Monta a lista de valores (na mesma ordem das colunas do cabeçalho)
        valores_linha = []

        # AM
        am_str = str(row.get("AM", ""))
        valores_linha.append(am_str)

        # VALOR TOTAL (exemplo de formatação de moeda)
        try:
            vt_str = f"R$ {row['VALOR TOTAL']:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        except:
            vt_str = str(row.get("VALOR TOTAL", ""))
        valores_linha.append(vt_str)

        # ÁREA TOTAL
        try:
            area_str = f"{row['AREA TOTAL']:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        except:
            area_str = str(row.get("AREA TOTAL", ""))
        valores_linha.append(area_str)

        # VALOR UNITÁRIO
        try:
           vu_str = f"R$ {row['valor_unitario']:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        except:
            vu_str = str(row.get("VALOR UNITARIO", ""))
        valores_linha.append(vu_str)

        # CIDADE
        cidade_str = str(row.get("cidade", ""))
        valores_linha.append(cidade_str)

        # FONTE
        fonte_str = str(row.get("fonte", ""))
        valores_linha.append(fonte_str)

        # Preenche as células
        for col_index, valor_cel in enumerate(valores_linha):
            cell_data = tabela.rows[i].cells[col_index]
            cell_data.text = valor_cel

            # Formatação da fonte das células de dados
            for run in cell_data.paragraphs[0].runs:
                run.font.name = "Arial"
                run.font.size = Pt(8)
                run.font.bold = False

            # Alinhamento horizontal
            cell_data.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Espaçamento vertical dentro da célula
            paragraph_format = cell_data.paragraphs[0].paragraph_format
            paragraph_format.space_before = Pt(2)
            paragraph_format.space_after = Pt(2)

            # Alinhamento vertical
            set_vertical_alignment(cell_data)

    # --- Ajuste de altura das linhas e largura das colunas ---
    for row_index in range(num_linhas):
        if row_index == 0:
            # Aumenta a altura da linha do cabeçalho
            tabela.rows[row_index].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            tabela.rows[row_index].height = Pt(40)
        else:
            # Aumenta a altura das linhas de dados
            tabela.rows[row_index].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            tabela.rows[row_index].height = Pt(26)

        # Ajusta a largura de cada coluna
        for col_index, w_inch in enumerate(col_widths):
            tabela.rows[row_index].cells[col_index].width = Inches(w_inch)

    # Insere a tabela logo depois do parágrafo alvo
    paragrafo_alvo._p.addnext(tabela._element)



###############################################################################
# >>>>>>> TEXTO DETALHADO DE SANEAMENTO (CHAUVENET) <<<<<<
###############################################################################
def gerar_texto_saneamento_chauvenet_extremamente_detalhado(
    dataframe_inicial,
    dataframe_filtrado,
    indices_excluidos,
    amostras_excluidas,
    media,
    desvio_padrao,
    menor_valor,
    maior_valor,
    mediana_valor
):
    """
    Gera um texto completo sobre o critério de Chauvenet e as estatísticas.
    """
    n_inicial = len(dataframe_inicial)
    n_filtrado = len(dataframe_filtrado)
    n_eliminadas = len(indices_excluidos)

    valores_filtrados_para_IC = dataframe_filtrado["VALOR TOTAL"].values
    if len(valores_filtrados_para_IC) > 1:
        limite_inf_ic, limite_sup_ic = intervalo_confianca_bootstrap_mediana(valores_filtrados_para_IC, 1000, 0.80)
        amplitude_ic = ((limite_sup_ic - limite_inf_ic) / numpy.median(valores_filtrados_para_IC)) * 100.0
    else:
        amplitude_ic = 0.0

    if desvio_padrao > 0:
        ds_menor = abs(media - menor_valor) / desvio_padrao
        ds_maior = abs(maior_valor - mediana_valor) / desvio_padrao
    else:
        ds_menor = 0.0
        ds_maior = 0.0

    if n_eliminadas > 0:
        identificacoes_excluidas = ", ".join(amostras_excluidas)
    else:
        identificacoes_excluidas = "Nenhuma"

    texto_exemplo = (
        f"-SANEAMENTO DOS DADOS AMOSTRAIS (CRITÉRIO DE CHAUVENET)\n"
        f"Quantidade de Amostras Válidas: {n_filtrado} unid.\n\n"
        f"-TESTANDO A AMOSTRA DE VALOR MAIS REDUZIDO-\n"
        f"D/S calc. = (Média - Menor Valor) / Desvio Padrão\n"
        f"D/S calc. para o MENOR valor = {ds_menor:.4f}\n\n"
        f"-TESTANDO A AMOSTRA DE VALOR MAIS ELEVADO-:\n"
        f"D/S calc. = (Maior Valor - Mediana) / Desvio Padrão\n"
        f"D/S calc. para o MAIOR valor = {ds_maior:.4f}\n\n"
        f"-CONCLUSÃO-:\n"
        f"* ALGUMAS AMOSTRAS PODEM NÃO SER VÁLIDAS (caso não satisfaçam o critério de Chauvenet) *\n\n"
        f"-RESUMO ESTATÍSTICO DA AVALIAÇÃO-\n"
        f"Quantidade de Amostras Válidas (Utilizadas no cálculo): {n_filtrado}\n"
        f"Quantidade de Amostras Eliminadas pelo critério de Chauvenet: {n_eliminadas}\n"
        f"Identificação das Amostras Eliminadas: {identificacoes_excluidas}\n"
    )

    return texto_exemplo

def inserir_texto_saneamento_no_placeholder(documento, marcador_placeholder, texto_saneamento):
    """
    Substitui o placeholder [texto_relatorio_resumo_saneamento] por um texto explicativo.
    """
    for paragrafo in documento.paragraphs:
        if marcador_placeholder in paragrafo.text:
            paragrafo.text = paragrafo.text.replace(marcador_placeholder, "")
            paragrafo_titulo = inserir_paragrafo_apos(paragrafo, "• SANEAMENTO DOS DADOS AMOSTRAIS\n")
            for execucao in paragrafo_titulo.runs:
                execucao.font.name = "Arial"
                execucao.font.size = Pt(12)
                execucao.font.bold = True
            paragrafo_titulo.paragraph_format.line_spacing = 1.15

            paragrafo_saneamento = inserir_paragrafo_apos(paragrafo_titulo, "")
            linhas = texto_saneamento.split("\n")
            for linha_texto in linhas:
                execucao_linha = paragrafo_saneamento.add_run(linha_texto + "\n")
                execucao_linha.font.name = "Arial"
                execucao_linha.font.size = Pt(12)
                execucao_linha.font.bold = False

            paragrafo_saneamento.paragraph_format.line_spacing = 1.15
            break


###############################################################################
# MEMÓRIA DE CÁLCULO DETALHADA
###############################################################################
def gerar_lista_memoria_calculo(dataframe_amostras, dados_avaliando, fatores_do_usuario, finalidade_do_laudo):
    import math
    
    lista_memoria_completa = []
    area_do_avaliando = float(dados_avaliando.get("AREA TOTAL", 0))

    # Fatores do Avaliado (utilizando as funções auxiliares já definidas)
    f_avaliado_aprov = fator_aproveitamento(dados_avaliando.get("APROVEITAMENTO", "URBANO"))
    f_avaliado_topog = fator_topografia(dados_avaliando.get("BOA TOPOGRAFIA?", "NÃO"))
    f_avaliado_pedol = fator_pedologia(dados_avaliando.get("PEDOLOGIA ALAGÁVEL? ", "NÃO"))
    f_avaliado_pavim = fator_pavimentacao(dados_avaliando.get("PAVIMENTACAO?", "NÃO"))
    f_avaliado_esq   = fator_esquina(dados_avaliando.get(" ESQUINA?", "NÃO"))
    f_avaliado_acess = fator_acessibilidade(dados_avaliando.get("ACESSIBILIDADE?", "NÃO"))

    for indice, linha in dataframe_amostras.iterrows():
        identificador_amostra = str(linha.get("AM", indice+1))
        valor_total = linha["VALOR TOTAL"]
        area_da_amostra = float(linha.get("AREA TOTAL", 0))

        # 1) Cálculo dos fatores básicos
        fator_area = calcular_fator_area(area_do_avaliando, area_da_amostra, fatores_do_usuario["area"])
        fator_oferta = calcular_fator_oferta(True, fatores_do_usuario["oferta"])
        
        # Fator Aproveitamento
        f_sample_aprov = fator_aproveitamento(linha.get("APROVEITAMENTO", "URBANO"))
        if fatores_do_usuario["aproveitamento"] and f_sample_aprov != 0:
            fator_aproveitamento_calculado = limitar_fator(f_avaliado_aprov / f_sample_aprov)
        else:
            fator_aproveitamento_calculado = 1.0

        # Fator Topografia
        f_sample_topog = fator_topografia(linha.get("BOA TOPOGRAFIA?", "NÃO"))
        if fatores_do_usuario["topografia"] and f_sample_topog != 0:
            fator_topografia_calculado = limitar_fator(f_avaliado_topog / f_sample_topog)
        else:
            fator_topografia_calculado = 1.0

        # Fator Pedologia
        f_sample_pedol = fator_pedologia(linha.get("PEDOLOGIA ALAGÁVEL? ", "NÃO"))
        if fatores_do_usuario["pedologia"] and f_sample_pedol != 0:
            fator_pedologia_calculado = limitar_fator(f_avaliado_pedol / f_sample_pedol)
        else:
            fator_pedologia_calculado = 1.0

        # Fator Pavimentação
        f_sample_pavim = fator_pavimentacao(linha.get("PAVIMENTACAO?", "NÃO"))
        if fatores_do_usuario["pavimentacao"] and f_sample_pavim != 0:
            fator_pavimentacao_calculado = limitar_fator(f_avaliado_pavim / f_sample_pavim)
        else:
            fator_pavimentacao_calculado = 1.0

        # Fator Esquina
        f_sample_esq = fator_esquina(linha.get(" ESQUINA?", "NÃO"))
        if fatores_do_usuario["esquina"] and f_sample_esq != 0:
            fator_esquina_calculado = limitar_fator(f_avaliado_esq / f_sample_esq)
        else:
            fator_esquina_calculado = 1.0

         # Fator Acessibilidade
        f_sample_acess = fator_acessibilidade(linha.get("ACESSIBILIDADE?", "NÃO"))
        if fatores_do_usuario["acessibilidade"] and f_sample_acess != 0:
            fator_acessibilidade_calculado = limitar_fator(f_avaliado_acess / f_sample_acess)
        else:
            fator_acessibilidade_calculado = 1.0      
              
                  
        # Fator Localização
        if fatores_do_usuario.get("localizacao_mesma_regiao", False):
            fator_localizacao_calculado = 1.0
        else:
            try:
                distancia_amostra = float(linha.get("DISTANCIA CENTRO", 0))
                distancia_avaliando = float(dados_avaliando.get("DISTANCIA CENTRO", 0))
                if distancia_amostra > 0 and distancia_avaliando > 0:
                    fator_item_comparativo = 1 / (distancia_amostra ** 0.1)
                    fator_bem_avaliando   = 1 / (distancia_avaliando ** 0.1)
                    fator_localizacao_calculado = limitar_fator(fator_bem_avaliando / fator_item_comparativo)
                else:
                    fator_localizacao_calculado = 1.0
            except:
                fator_localizacao_calculado = 1.0

        # 2) Cálculo do Valor Total Homogeneizado
        valor_total_homogeneizado = (
            valor_total *
            fator_area *
            fator_oferta *
            fator_localizacao_calculado *
            fator_aproveitamento_calculado *
            fator_topografia_calculado *
            fator_pedologia_calculado *
            fator_pavimentacao_calculado *
            fator_esquina_calculado *
            fator_acessibilidade_calculado
        )

        # 3) Monta o texto de memória de cálculo
        bloco_texto = []
        bloco_texto.append(f"AM {identificador_amostra}")
        bloco_texto.append("")
        # Mantém o texto original para o valor ofertado:
        bloco_texto.append(f"- VALOR TOTAL OFERTADO: {formatar_moeda_brasil(valor_total)}")
        bloco_texto.append(f"- ÁREA DA AMOSTRA (m²): {formatar_numero_brasileiro(area_da_amostra)}")
        bloco_texto.append("")

        bloco_texto.append("- Fator Área:")
        bloco_texto.append(f"   Avaliado: {formatar_numero_brasileiro(area_do_avaliando)}")
        bloco_texto.append(f"   Amostra: {formatar_numero_brasileiro(area_da_amostra)} - Cálculo => {fator_area:.2f}\n")

        bloco_texto.append("- Fator Oferta:")
        bloco_texto.append(f"   (fixo 0.90 se habilitado) => {fator_oferta:.2f}\n")

        bloco_texto.append("- Fator Aproveitamento (f_avaliado / f_amostra):")
        bloco_texto.append(f"   Avaliado: {f_avaliado_aprov:.2f}")
        bloco_texto.append(f"   Amostra: {f_sample_aprov:.2f}")
        bloco_texto.append(f"   => {fator_aproveitamento_calculado:.2f}\n")

        bloco_texto.append("- Fator Topografia (f_avaliado / f_amostra):")
        bloco_texto.append(f"   Avaliado: {f_avaliado_topog:.2f}")
        bloco_texto.append(f"   Amostra: {f_sample_topog:.2f}")
        bloco_texto.append(f"   => {fator_topografia_calculado:.2f}\n")

        bloco_texto.append("- Fator Pedologia (f_avaliado / f_amostra):")
        bloco_texto.append(f"   Avaliado: {f_avaliado_pedol:.2f}")
        bloco_texto.append(f"   Amostra: {f_sample_pedol:.2f}")
        bloco_texto.append(f"   => {fator_pedologia_calculado:.2f}\n")

        bloco_texto.append("- Fator Pavimentação (f_avaliado / f_amostra):")
        bloco_texto.append(f"   Avaliado: {f_avaliado_pavim:.2f}")
        bloco_texto.append(f"   Amostra: {f_sample_pavim:.2f}")
        bloco_texto.append(f"   => {fator_pavimentacao_calculado:.2f}\n")

        bloco_texto.append("- Fator Esquina (f_avaliado / f_amostra):")
        bloco_texto.append(f"   Avaliado: {f_avaliado_esq:.2f}")
        bloco_texto.append(f"   Amostra: {f_sample_esq:.2f}")
        bloco_texto.append(f"   => {fator_esquina_calculado:.2f}\n")

        bloco_texto.append("- Fator Acessibilidade (f_avaliado / f_amostra):")
        bloco_texto.append(f"   Avaliado: {f_avaliado_acess:.2f}")
        bloco_texto.append(f"   Amostra: {f_sample_acess:.2f}")
        bloco_texto.append(f"   => {fator_acessibilidade_calculado:.2f}\n")

        bloco_texto.append("- Fator Localização:")
        bloco_texto.append(f"   => {fator_localizacao_calculado:.2f}\n")

        # 4) Em vez de exibir o Valor Total Homogeneizado, agora exibe o Valor Unitário Homogeneizado (VUH)
        if area_da_amostra > 0:
            valor_unit_homog = valor_total_homogeneizado / area_da_amostra
        else:
            valor_unit_homog = 0.0

        bloco_texto.append(
            f"=> VUH (Valor Unitário Homogeneizado): {formatar_moeda_brasil(valor_unit_homog)}"
        )

        lista_memoria_completa.append("\n".join(bloco_texto))

    return lista_memoria_completa



def inserir_texto_memoria_calculo_no_placeholder(documento, marcador_placeholder, lista_memorias):
    for paragrafo in documento.paragraphs:
        if marcador_placeholder in paragrafo.text:
            paragrafo.text = paragrafo.text.replace(marcador_placeholder, "")
            paragrafo_atual = paragrafo
            for indice_bloco, bloco in enumerate(lista_memorias):
                if indice_bloco >= 1:
                    paragrafo_branco = inserir_paragrafo_apos(paragrafo_atual, "")
                    execucao_branco = paragrafo_branco.add_run("\n")
                    execucao_branco.font.size = Pt(10)
                    execucao_branco.font.name = "Arial"
                    paragrafo_atual = paragrafo_branco

                novo_paragrafo = inserir_paragrafo_apos(paragrafo_atual, "")
                linhas_texto = bloco.split("\n")

                for indice_linha, conteudo_linha in enumerate(linhas_texto):
                    execucao_texto = novo_paragrafo.add_run(conteudo_linha + "\n")
                    execucao_texto.font.name = "Arial"

                    if conteudo_linha.strip().startswith("=> VUH"):
                        execucao_texto.font.size = Pt(13)
                        execucao_texto.font.bold = True
                   
                   
                    elif indice_linha == 0 and conteudo_linha.strip().startswith("AM "):
                        execucao_texto.font.size = Pt(13)
                        execucao_texto.font.bold = True
                    else:
                        execucao_texto.font.size = Pt(10)
                        execucao_texto.font.bold = False

                novo_paragrafo.paragraph_format.line_spacing = 1.15
                paragrafo_atual = novo_paragrafo
            break


###############################################################################
# SUBSTITUIR PLACEHOLDER POR TEXTO OU IMAGEM
###############################################################################
def substituir_placeholder_por_texto_formatado(documento, marcador, texto, tamanho_fonte=Pt(12), negrito=False):
    """
    Substitui o placeholder por texto com fonte e tamanho definidos.
    """
    for paragrafo in documento.paragraphs:
        if marcador in paragrafo.text:
            paragrafo.text = paragrafo.text.replace(marcador, "")
            execucao = paragrafo.add_run(texto)
            execucao.font.name = "Arial"
            execucao.font.size = tamanho_fonte
            execucao.bold = negrito

    for tabela in documento.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for parag in celula.paragraphs:
                    if marcador in parag.text:
                        parag.text = parag.text.replace(marcador, "")
                        execucao = parag.add_run(texto)
                        execucao.font.name = "Arial"
                        execucao.font.size = tamanho_fonte
                        execucao.bold = negrito

def substituir_placeholder_por_imagem(documento, marcador, caminho_imagem, largura=Inches(5)):
    """
    Substitui o placeholder por uma imagem alinhada ao centro.
    """
    for paragrafo in documento.paragraphs:
        if marcador in paragrafo.text:
            paragrafo.text = paragrafo.text.replace(marcador, "")
            runn = paragrafo.add_run()
            runn.add_picture(caminho_imagem, width=largura)
            paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            runn.font.name = "Arial"
            runn.font.size = Pt(12)

    for tabela in documento.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for parag in celula.paragraphs:
                    if marcador in parag.text:
                        parag.text = parag.text.replace(marcador, "")
                        runn = parag.add_run()
                        runn.add_picture(caminho_imagem, width=largura)
                        parag.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        runn.font.name = "Arial"
                        runn.font.size = Pt(12)

def remover_paragrafo_por_marcador(documento, marcador):
    """
    Remove o parágrafo que contém o marcador especificado.
    """
    for paragrafo in documento.paragraphs:
        if marcador in paragrafo.text:
            p = paragrafo._element
            p.getparent().remove(p)
            break


###############################################################################
# TABELA DE GRAU DE PRECISÃO ([texto_grau_precisao])
###############################################################################
def inserir_tabela_classificacao_de_precisao(documento, marcador, amplitude_ic80):
    """
    Insere a tabela padrão da ABNT NBR 14653 e destaca a classificação conforme o IC.
    """
    grau_obtido = classificar_grau_de_precisao(amplitude_ic80)
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    from docx.oxml.shared import OxmlElement
    from lxml import etree

    from docx.oxml.ns import nsdecls, qn

    for paragrafo in documento.paragraphs:
        if marcador in paragrafo.text:
            paragrafo.text = paragrafo.text.replace(marcador, "")
            titulo_paragrafo = inserir_paragrafo_apos(paragrafo, "GRAU DE PRECISÃO")
            for execucao in titulo_paragrafo.runs:
                execucao.font.name = "Arial"
                execucao.font.size = Pt(12)
                execucao.font.bold = True
            titulo_paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

            tabela = documento.add_table(rows=3, cols=4)
            tabela.style = "Table Grid"

            tabela.cell(0,0).text = "Descrição"
            tabela.cell(0,1).text = "GRAU III"
            tabela.cell(0,2).text = "GRAU II"
            tabela.cell(0,3).text = "GRAU I"

            tabela.cell(1,0).text = (
                "Amplitude do intervalo de confiança de 80%\n"
                "em torno da estimativa de tendência central"
            )
            tabela.cell(1,1).text = "≤ 30%"
            tabela.cell(1,2).text = "≤ 40%"
            tabela.cell(1,3).text = "≤ 50%"

            celula_nota = tabela.cell(2, 0).merge(tabela.cell(2, 1)).merge(tabela.cell(2, 2)).merge(tabela.cell(2, 3))
            celula_nota.text = (
                "NOTA: Quando a amplitude do intervalo de confiança ultrapassar 50%, "
                "não há classificação do resultado quanto à precisão e é necessária justificativa "
                "com base no diagnóstico do mercado."
            )

            for linha in tabela.rows:
                for celula in linha.cells:
                    props = celula._tc.get_or_add_tcPr()
                    vAlign = OxmlElement('w:vAlign')
                    vAlign.set(qn('w:val'), "center")
                    props.append(vAlign)
                    for par_cel in celula.paragraphs:
                        par_cel.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for runn in par_cel.runs:
                            runn.font.name = "Arial"
                            runn.font.size = Pt(10)

            shading_azul = etree.fromstring(r'<w:shd {} w:fill="BDD7EE" w:val="clear"/>'.format(nsdecls('w')))
            if grau_obtido == "GRAU III":
                tabela.cell(0,1)._tc.get_or_add_tcPr().append(shading_azul)
                tabela.cell(1,1)._tc.get_or_add_tcPr().append(shading_azul)
            elif grau_obtido == "GRAU II":
                tabela.cell(0,2)._tc.get_or_add_tcPr().append(shading_azul)
                tabela.cell(1,2)._tc.get_or_add_tcPr().append(shading_azul)
            elif grau_obtido == "GRAU I":
                tabela.cell(0,3)._tc.get_or_add_tcPr().append(shading_azul)
                tabela.cell(1,3)._tc.get_or_add_tcPr().append(shading_azul)
            else:
                celula_nota._tc.get_or_add_tcPr().append(shading_azul)

            titulo_paragrafo._p.addnext(tabela._element)

            novo_paragrafo = inserir_paragrafo_apos(
                titulo_paragrafo,
                f"Amplitude IC 80% calculada: {amplitude_ic80:.2f}% — Classificação: {grau_obtido}"
            )
            novo_paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for exec_novo in novo_paragrafo.runs:
                exec_novo.font.name = "Arial"
                exec_novo.font.size = Pt(10)
            break



###############################################################################
# INSERIR LOGO (OPCIONAL)
###############################################################################
def inserir_logo_no_placeholder(documento, placeholder, caminho_logo):
    """
    Substitui [logo] pela imagem do logotipo, alinhado à direita.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    for paragrafo in documento.paragraphs:
        if placeholder in paragrafo.text:
            paragrafo.text = paragrafo.text.replace(placeholder, "")
            runn = paragrafo.add_run()
            runn.add_picture(caminho_logo, width=Inches(3))
            paragrafo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            return


###############################################################################
# TABELA DE RESUMO DE VALORES ([RESUMO VALORES])
# AGORA MODIFICADA PARA EXIBIR MÚLTIPLAS RESTRIÇÕES
###############################################################################
def inserir_tabela_resumo_de_valores(documento, marcador, informacoes_de_resumo, area_utilizada):
    """
    Cria a tabela de resumo de valores, compatível com versões antigas do python-docx,
    sem usar get_or_add_tblPr(), e forçando que a primeira letra do valor por extenso 
    seja maiúscula, ex.: "Trinta e um mil, cento e setenta e dois reais e seis centavos".
    
    Parâmetros em `informacoes_de_resumo`:
      - valor_unitario (str) => ex: "R$ 35,37/m²"
      - area_total_considerada (str) => ex: "1.000,00 m²"
      - texto_descritivo_restricoes (str) => ex: "Múltiplas restrições aplicadas"
      - restricoes (list[dict]) => cada item: {
            "area": 345.0,
            "percentualDepreciacao": 34,
            "fator": 0.66,
            "tipo": "APP",
            "subtotal": "R$ 8.053,23"
        }
      - valor_total_indenizatorio (str) => ex: "R$ 30.979,30"
      - valor_por_extenso (str) => se vier vazio, será calculado via num2words; 
        em seguida, a inicial é forçada para maiúsculo.
    """
    import re
    from lxml import etree
    from docx.oxml.ns import nsdecls, qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL
    from docx.shared import Pt

    # Se tiver num2words, usamos para converter valor em texto extenso.
    try:
        from num2words import num2words
    except ImportError:
        num2words = None

    def extrair_valor_numerico(texto_monetario):
        """
        Ex: "R$ 30.979,30" => 30979.30 (float).
        Remove caracteres que não sejam dígitos ou vírgula, então substitui ',' por '.'.
        """
        somente_num_virg = re.sub(r"[^\d,]", "", texto_monetario)
        somente_num_ponto = somente_num_virg.replace(",", ".")
        try:
            return float(somente_num_ponto)
        except:
            return 0.0

    def gerar_extenso_por_num2words(texto_valor):
        """
        Converte "R$ 30.979,30" em algo como 
        "Trinta e um mil, cento e setenta e nove reais e trinta centavos",
        usando a biblioteca num2words(lang='pt_BR'). 
        Em seguida, forçamos a primeira letra para maiúscula.
        """
        if not num2words:
            return "(num2words não instalado)"

        val = extrair_valor_numerico(texto_valor)
        inteiro = int(val)
        centavos = round((val - inteiro) * 100)
        if inteiro == 0 and centavos == 0:
            return "Zero real"

        extenso_inteiro = num2words(inteiro, lang='pt_BR')
        if centavos > 0:
            extenso_centavos = num2words(centavos, lang='pt_BR')
            texto_final = f"{extenso_inteiro} reais e {extenso_centavos} centavos"
        else:
            texto_final = f"{extenso_inteiro} reais"

        # Forçar a primeira letra para maiúsculo, se não estiver vazio:
        if texto_final:
            texto_final = texto_final[0].upper() + texto_final[1:]
        return texto_final


    # -----------------------------------------------------------------
    # ►►  FORMATADOR DE ÁREA  ◄◄
    #     2.000      →  "2.000"
    #     2_000.50   →  "2.000,50"
    #     80_234.89  →  "80.234,89"
    # -----------------------------------------------------------------
    def formatar_area_brasil(valor):
        try:
            v = float(valor)
        except Exception:
            return str(valor)                   # cai fora se não for número

        # Inteiro?  -> sem decimais
        if abs(v - round(v)) < 1e-6:
            return f"{int(v):,}".replace(",", ".")
        # Tem centavos -> 2 casas decimais
        texto = f"{v:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        # elimina “,00” caso o arredondamento zere
        return texto.rstrip("0").rstrip(",")
 



    
    # -------------------------------------------------------------------------
    # Localiza o placeholder no documento
    for paragrafo in documento.paragraphs:
        if marcador in paragrafo.text:
            # Remove o texto do placeholder
            paragrafo.text = paragrafo.text.replace(marcador, "")

            # Carrega dados
            valor_unit = informacoes_de_resumo.get("valor_unitario", "N/D")
            area_total = informacoes_de_resumo.get("area_total_considerada", "N/D")
            sit_rest = informacoes_de_resumo.get("texto_descritivo_restricoes", "N/D")
            restricoes = informacoes_de_resumo.get("restricoes", [])
            valor_total = informacoes_de_resumo.get("valor_total_indenizatorio", "R$ 0,00")
            valor_extenso = informacoes_de_resumo.get("valor_por_extenso", "").strip()

            # Se valor_por_extenso for vazio, gerar automaticamente
            if not valor_extenso:
                valor_extenso = gerar_extenso_por_num2words(valor_total)

            # Cria a tabela principal (7 linhas, 2 colunas)
            tabela_principal = documento.add_table(rows=7, cols=2)
            tabela_principal.style = "Table Grid"
            tabela_principal.alignment = WD_TABLE_ALIGNMENT.CENTER

            # (0) Cabeçalho mesclado
            cel_titulo = tabela_principal.cell(0, 0).merge(tabela_principal.cell(0, 1))
            cel_titulo.text = "RESUMO DOS VALORES TOTAIS"
            shading_cab = etree.fromstring(r'<w:shd {} w:fill="D9D9D9" w:val="clear"/>'.format(nsdecls('w')))
            cel_titulo._tc.get_or_add_tcPr().append(shading_cab)
            for p_ in cel_titulo.paragraphs:
                p_.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run_ in p_.runs:
                    run_.font.name = "Arial"
                    run_.font.size = Pt(11)
                    run_.font.bold = True

            # (1) Valor Unitário Calculado
            tabela_principal.cell(1,0).text = "Valor Unitário Calculado:"
            tabela_principal.cell(1,1).text = valor_unit

           # (2) Área Total de Interesse
            tabela_principal.cell(2, 0).text = "Área Total de Interesse:"
            tabela_principal.cell(2, 1).text = f"{area_utilizada:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")



            # (3) Situação das Restrições
            tabela_principal.cell(3,0).text = "Situação das Restrições:"
            tabela_principal.cell(3,1).text = sit_rest

            # (4) Sub-tabela => célula mesclada
            cel_sub = tabela_principal.cell(4,0).merge(tabela_principal.cell(4,1))
            shading_light_blue = etree.fromstring(r'<w:shd {} w:fill="E0ECF8" w:val="clear"/>'.format(nsdecls('w')))
            cel_sub._tc.get_or_add_tcPr().append(shading_light_blue)

            # Remove margens internas da célula mesclada
            tc_pr_sub = cel_sub._tc.get_or_add_tcPr()
            tc_margins_sub = tc_pr_sub.xpath('./w:tcMar')
            if not tc_margins_sub:
                tcMar = OxmlElement('w:tcMar')
                tcMar.set(qn('w:top'), "0")
                tcMar.set(qn('w:left'), "0")
                tcMar.set(qn('w:right'), "0")
                tcMar.set(qn('w:bottom'), "0")
                tc_pr_sub.append(tcMar)
            else:
                for m_ in tc_margins_sub:
                    m_.set(qn('w:top'), "0")
                    m_.set(qn('w:left'), "0")
                    m_.set(qn('w:right'), "0")
                    m_.set(qn('w:bottom'), "0")

            # Se não tiver restrições, mostra texto simples
            if not restricoes:
                cel_sub.text = "Nenhuma restrição aplicada."
                for r_ in cel_sub.paragraphs[0].runs:
                    r_.font.name = "Arial"
                    r_.font.size = Pt(10)
            else:
                # Cria sub-tabela sem bordas
                subtab = documento.add_table(rows=len(restricoes)+1, cols=5)
                borders = subtab._element.xpath(".//w:tblBorders")
                for b_ in borders:
                    b_.getparent().remove(b_)

                # Adicionar manualmente <w:tblPr>, se não existir
                tblPr = subtab._element.tblPr
                if tblPr is None:
                    tblPr = OxmlElement('w:tblPr')
                    subtab._element.insert(0, tblPr)

                # <w:tblCellMar> p/ zerar margens
                tblCellMar = OxmlElement('w:tblCellMar')
                tblCellMar.set(qn('w:top'), "0")
                tblCellMar.set(qn('w:left'), "0")
                tblCellMar.set(qn('w:right'), "0")
                tblCellMar.set(qn('w:bottom'), "0")
                tblPr.append(tblCellMar)

                # Cabeçalhos
                cabecalhos = ["Área (m²)", "% Depreciação", "Fator aplicado", "Tipo Restrição", "Subtotal (R$)"]
                for cidx, hh in enumerate(cabecalhos):
                    subtab.cell(0,cidx).text = hh
                    for run_ in subtab.cell(0,cidx).paragraphs[0].runs:
                        run_.font.name = "Arial"
                        run_.font.size = Pt(9)
                        run_.font.bold = True
                    subtab.cell(0,cidx).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Linhas de dados
                for i, restr in enumerate(restricoes, start=1):
                    area_ = str(restr.get("area", ""))
                    perc_ = restr.get("percentualDepreciacao", 0)
                    fat_ = restr.get("fator", 1.0)
                    tipo_ = restr.get("tipo", "")
                    sub_ = restr.get("subtotal", "R$ 0,00")

                    subtab.cell(i,0).text = area_
                    subtab.cell(i,1).text = f"{perc_:.0f}%"
                    subtab.cell(i,2).text = f"{fat_:.2f}"
                    subtab.cell(i,3).text = tipo_
                    subtab.cell(i,4).text = sub_

                    for cc_ in range(5):
                        p_run = subtab.cell(i, cc_).paragraphs[0]
                        p_run.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for r_ in p_run.runs:
                            r_.font.name = "Arial"
                            r_.font.size = Pt(9)

                # Fundo azul e remover margens em todas as células
                for row_ in subtab.rows:
                    for cell_ in row_.cells:
                        shade_cell = etree.fromstring(r'<w:shd {} w:fill="E0ECF8" w:val="clear"/>'.format(nsdecls('w')))
                        cell_._tc.get_or_add_tcPr().append(shade_cell)
                        cell_.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        tcpr = cell_._tc.get_or_add_tcPr()
                        tc_marg = tcpr.xpath('./w:tcMar')
                        if not tc_marg:
                            newMar = OxmlElement('w:tcMar')
                            newMar.set(qn('w:top'), "0")
                            newMar.set(qn('w:left'), "0")
                            newMar.set(qn('w:right'), "0")
                            newMar.set(qn('w:bottom'), "0")
                            tcpr.append(newMar)
                        else:
                            for mm in tc_marg:
                                mm.set(qn('w:top'), "0")
                                mm.set(qn('w:left'), "0")
                                mm.set(qn('w:right'), "0")
                                mm.set(qn('w:bottom'), "0")

                # Anexa a sub-tabela à célula
                cel_sub._tc.clear_content()
                cel_sub._tc.append(subtab._element)

            # (5) Valor Total Indenizatório
            tabela_principal.cell(5,0).text = "Valor Total Indenizatório:"
            tabela_principal.cell(5,1).text = valor_total

            # (6) Valor por Extenso
            cel_ext = tabela_principal.cell(6,0).merge(tabela_principal.cell(6,1))
            cel_ext.text = valor_extenso

            # Ajustes de layout da Tabela Principal
            for row_idx in range(7):
                row_ = tabela_principal.rows[row_idx]
                row_.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
                row_.height = Pt(18)
                for col_idx in range(2):
                    c_ = row_.cells[col_idx]
                    c_.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    for pp_ in c_.paragraphs:
                        pp_.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for rr_ in pp_.runs:
                            rr_.font.name = "Arial"
                            rr_.font.size = Pt(10)
                            rr_.font.bold = False

            # Valor Unitário (linha 1 => col 1) e Valor Total (linha 5 => col 1) em negrito
            for run_ in tabela_principal.rows[1].cells[1].paragraphs[0].runs:
                run_.font.bold = True
            for run_ in tabela_principal.rows[5].cells[1].paragraphs[0].runs:
                run_.font.bold = True
                run_.font.size = Pt(11)

            # Valor por Extenso (linha 6) => central e em negrito
            for p_ in tabela_principal.rows[6].cells[0].paragraphs:
                p_.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for rn_ in p_.runs:
                    rn_.font.size = Pt(10)
                    rn_.font.bold = True

            # Insere a tabela após o parágrafo do placeholder
            paragrafo._p.addnext(tabela_principal._element)
            break




###############################################################################
# DIAGNÓSTICO DE MERCADO
###############################################################################
def inserir_tabela_diagnostico_de_mercado(documento, marcador, escolha_estrutura, escolha_conduta, escolha_desempenho):
    """
    Monta a tabela de diagnóstico de mercado (Estrutura, Conduta, Desempenho),
    destacando (sombreando) a opção escolhida.
    """
    from docx.enum.table import WD_ROW_HEIGHT_RULE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    from lxml import etree

    from docx.oxml.ns import nsdecls, qn

    dados_tabela = [
        ("Estrutura",  "BOA",         "Facilidade em se encontrar imóveis na região do avaliando"),
        ("Estrutura",  "LIMITADA",    "Dificuldade em se encontrar imóveis na região do avaliando."),
        ("Conduta",    "DESESTAGNADO","Boa movimentação do mercado imobiliário."),
        ("Conduta",    "ESTAGNADA",   "Pouca movimentação do mercado imobiliário."),
        ("Desempenho", "ALTO",        "Ótima atratividade comercial para negócios imobiliários."),
        ("Desempenho", "MÉDIO",       "Atratividade moderada para negócios imobiliários."),
        ("Desempenho", "BAIXO",       "Baixa atratividade da região para negócios imobiliários.")
    ]

    def verificar_se_destacar(dim, opc):
        if dim.lower() == "estrutura":
            return opc.upper() == escolha_estrutura.upper()
        elif dim.lower() == "conduta":
            return opc.upper() == escolha_conduta.upper()
        elif dim.lower() == "desempenho":
            return opc.upper() == escolha_desempenho.upper()
        return False

    for paragrafo in documento.paragraphs:
        if marcador in paragrafo.text:
            paragrafo.text = paragrafo.text.replace(marcador, "")
            tabela = documento.add_table(rows=len(dados_tabela), cols=3)
            tabela.style = "Table Grid"
            tabela.allow_autofit = False

            for i, (dimensao, opcao, descricao) in enumerate(dados_tabela):
                c_dim = tabela.rows[i].cells[0]
                c_opc = tabela.rows[i].cells[1]
                c_desc = tabela.rows[i].cells[2]

                # Exibir a dimensão apenas na "parte do meio" de cada bloco
                if i == 1:
                    c_dim.text = "Estrutura"
                elif i == 3:
                    c_dim.text = "Conduta"
                elif i == 5:
                    c_dim.text = "Desempenho"
                else:
                    c_dim.text = ""

                c_opc.text = opcao
                c_desc.text = descricao

                if verificar_se_destacar(dimensao, opcao):
                    shading_azul = etree.fromstring(
                        r'<w:shd {} w:fill="BDD7EE" w:val="clear"/>'.format(nsdecls('w'))
                    )
                    c_opc._tc.get_or_add_tcPr().append(shading_azul)
                    c_desc._tc.get_or_add_tcPr().append(shading_azul)

            # Mesclar primeira coluna
            tabela.cell(0, 0).merge(tabela.cell(1, 0))
            tabela.cell(2, 0).merge(tabela.cell(3, 0))
            tabela.cell(4, 0).merge(tabela.cell(5, 0))
            tabela.cell(4, 0).merge(tabela.cell(6, 0))

            for row_index, row in enumerate(tabela.rows):
                row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                row.height = Pt(28)
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        paragraph.paragraph_format.space_before = Pt(3)
                        paragraph.paragraph_format.space_after = Pt(3)
                        for run in paragraph.runs:
                            run.font.name = "Arial"
                            run.font.size = Pt(10)

            paragrafo._p.addnext(tabela._element)
            break


###############################################################################
# INSERIR TABELA DE RESUMO GERAL (EXEMPLO)
###############################################################################
def inserir_tabela_resumo_geral_completo(documento, placeholder, info_resumo_geral):
    """
    Exemplo de inserção de uma tabela extra de 'Resumo Geral da Avaliação',
    demonstrando como manipular dados caso queira algo mais completo.
    Substitui [RESUMO GERAL] por uma tabela no documento.
    """
    for paragrafo in documento.paragraphs:
        if placeholder in paragrafo.text:
            paragrafo.text = paragrafo.text.replace(placeholder, "")
            # Aqui apenas um exemplo (poderia adaptar)
            from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Pt
            from docx.oxml.shared import OxmlElement
            from lxml import etree

            from docx.oxml.ns import nsdecls, qn

            tabela_resumo = documento.add_table(rows=5, cols=2, style="Table Grid")

            tabela_resumo.cell(0,0).text = "Proprietário"
            tabela_resumo.cell(0,1).text = info_resumo_geral.get("proprietario","S/N")

            tabela_resumo.cell(1,0).text = "Documento"
            tabela_resumo.cell(1,1).text = info_resumo_geral.get("documento_imovel","(N/D)")

            tabela_resumo.cell(2,0).text = "Cartório"
            tabela_resumo.cell(2,1).text = info_resumo_geral.get("cartorio","(N/D)")

            tabela_resumo.cell(3,0).text = "Comarca"
            tabela_resumo.cell(3,1).text = info_resumo_geral.get("comarca","(N/D)")

            tabela_resumo.cell(4,0).text = "Endereço"
            tabela_resumo.cell(4,1).text = info_resumo_geral.get("endereco_imovel","(N/D)")

            for rr in range(len(tabela_resumo.rows)):
                for cc in range(2):
                    cell_ = tabela_resumo.cell(rr, cc)
                    for pp_ in cell_.paragraphs:
                        pp_.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run_ in pp_.runs:
                            run_.font.name = "Arial"
                            run_.font.size = Pt(9)

            paragrafo._p.addnext(tabela_resumo._element)
            break


###############################################################################
# >>>>>> RELATÓRIO PRINCIPAL - GERAÇÃO DO WORD <<<<<<
###############################################################################
def gerar_relatorio_avaliacao_com_template(
    dados_avaliando,
    dataframe_amostras_inicial,
    dataframe_amostras_filtrado,
    indices_excluidos,
    amostras_excluidas,
    media,
    desvio_padrao,
    menor_valor,
    maior_valor,
    mediana_valor,
    valores_originais_iniciais,
    valores_homogeneizados_validos,
    caminho_imagem_aderencia,
    caminho_imagem_dispersao,
    uuid_atual,                          # obrigatório, sem valor padrão
    finalidade_do_laudo,                # agora obrigatório também
    area_parcial_afetada,              # idem — valor digitado deve ser usado sempre
    fatores_do_usuario=None,
    caminhos_fotos_avaliando=None,
    caminhos_fotos_adicionais=None,
    caminhos_fotos_proprietario=None,
    caminhos_fotos_planta=None,
    caminho_template="template.docx",
    nome_arquivo_word="relatorio.docx"
):
    # DEFINIÇÃO DEFINITIVA DA ÁREA UTILIZADA (corrigido!)
    area_utilizada = area_parcial_afetada if finalidade_do_laudo in ["desapropriacao", "servidao"] else float(dados_avaliando.get("AREA TOTAL", 0))
    logger.info(f"✅ Área utilizada definitiva definida imediatamente após parâmetros: {area_utilizada}")
    # Insira logs aqui para depuração detalhada:
    logger.info(f"Valores originais recebidos: {valores_originais_iniciais}")
    logger.info(f"Valores homogeneizados válidos recebidos: {valores_homogeneizados_validos}")
    logger.info(f"Área Parcial Afetada recebida: {area_parcial_afetada}")

    # ──────────────────────────────────────────────────────
    # Alias para compatibilizar o novo nome:
    # logger.info(f"🔴 Área Parcial Afetada recebida no main.py: {area_parcial_afetada}")
    # area_utilizada = area_parcial_afetada
    # logger.info(f"🟢 Área utilizada atribuída no main.py: {area_utilizada}")

    # ──────────────────────────────────────────────────────
    """
    Gera o relatório Word completo, exibindo todos os itens e incluindo
    o tratamento de múltiplas restrições.
    """
    from docx import Document
    from datetime import datetime

    data_atual = datetime.now().strftime("%d/%m/%Y")

    # Carregar template
    documento = Document(caminho_template)

    cidade_nome = fatores_do_usuario.get("cidade", "CIDADE NÃO INFORMADA").strip().upper()
    data_formatada = datetime.now().strftime("%d-%m-%Y")

    substituir_placeholder_por_texto_formatado(
        documento,
        "[cidade]",
        f"{cidade_nome}, {data_formatada}",
        Pt(12),
        False
    )


    
    # ------------------------------------------------------------------
    # MAPA DE AMOSTRAS - LOCALIZAÇÃO DOS DADOS DE MERCADO E AVALIANDO
    # ------------------------------------------------------------------
    pasta_saida = os.path.join("static", "arquivos", f"avaliacao_{uuid_atual}")
    os.makedirs(pasta_saida, exist_ok=True)


    caminho_mapa = os.path.join(pasta_saida, "mapa_amostras.png")

    gerar_mapa_amostras(dataframe_amostras_filtrado, dados_avaliando, nome_png=caminho_mapa)
    # INSIRA ESSA VERIFICAÇÃO LOG AQUI:
    if os.path.exists(caminho_mapa):
        logger.info(f"✅ MAPA AMOSTRAS encontrado: {caminho_mapa}")
    else:
        logger.warning(f"❌ MAPA AMOSTRAS NÃO encontrado: {caminho_mapa}")
    
    if caminho_mapa and os.path.exists(caminho_mapa):  # <- ESSA É A LINHA CORRIGIDA
        substituir_placeholder_por_imagem(
            documento, "[MAPAAMOSTRAS]", caminho_mapa, largura=Inches(6)
        )
    else:
        substituir_placeholder_por_texto_formatado(
            documento, "[MAPAAMOSTRAS]",
            "Mapa de localização não disponível — coordenadas incompletas",
            Pt(12), False
        )
    
    # Inserir a tabela de amostras originais
    logger.info("🔎 DataFrame que será enviado para inserir_tabela_amostras_originais:")
    logger.info(dataframe_amostras_inicial.head())
    logger.info(f"🔎 Colunas disponíveis: {list(dataframe_amostras_inicial.columns)}")
    inserir_tabela_amostras_originais(documento, dataframe_amostras_inicial)

    # Preencher alguns placeholders básicos
    def substituir_placeholder_por_titulo_e_valor(documento, marcador, titulo, valor, tamanho_fonte):
        for paragrafo in documento.paragraphs:
            if marcador in paragrafo.text:
                paragrafo.text = paragrafo.text.replace(marcador, "")
                run_titulo = paragrafo.add_run(titulo)
                run_titulo.font.name = "Arial"
                run_titulo.font.size = tamanho_fonte
                run_titulo.bold = True
                run_valor = paragrafo.add_run(valor)
                run_valor.font.name = "Arial"
                run_valor.font.size = tamanho_fonte
                run_valor.bold = False
        for tabela in documento.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    for parag in celula.paragraphs:
                        if marcador in parag.text:
                            parag.text = parag.text.replace(marcador, "")
                            run_titulo = parag.add_run(titulo)
                            run_titulo.font.name = "Arial"
                            run_titulo.font.size = tamanho_fonte
                            run_titulo.bold = True
                            run_valor = parag.add_run(valor)
                            run_valor.font.name = "Arial"
                            run_valor.font.size = tamanho_fonte
                            run_valor.bold = False

    substituir_placeholder_por_texto_formatado(documento, "[created]", data_atual, Pt(13), False)
    substituir_placeholder_por_texto_formatado(documento, "[avaliadorNome]", fatores_do_usuario.get("avaliadorNome", ""), Pt(13), False)
    substituir_placeholder_por_texto_formatado(documento, "[avaliadorRegistro]", fatores_do_usuario.get("avaliadorRegistro", ""), Pt(13), False)

    substituir_placeholder_por_titulo_e_valor(
        documento,
        "[nomeSolicitante]",
        "• Solicitante: ",
        fatores_do_usuario.get("nomeSolicitante", ""),
        Pt(11)
    )
    substituir_placeholder_por_titulo_e_valor(
        documento,
        "[nome proprietário]",
        "• Nome do Proprietário: ",
        fatores_do_usuario.get("nomeProprietario", ""),
        Pt(11)
    )
    substituir_placeholder_por_titulo_e_valor(
        documento,
        "[telefone proprietario]",
        "• Contato do Proprietário: ",
        fatores_do_usuario.get("telefoneProprietario", ""),
        Pt(11)
    )
    substituir_placeholder_por_titulo_e_valor(
        documento,
        "[email]",
        "• E-mail do Proprietário: ",
        fatores_do_usuario.get("emailProprietario", ""),
        Pt(11)
    )
    substituir_placeholder_por_titulo_e_valor(
        documento,
        "[tipoImovel]",
        "• Tipo do Imóvel: ",
        fatores_do_usuario.get("tipoImovel", ""),
        Pt(11)
    )
    substituir_placeholder_por_titulo_e_valor(
        documento,
        "[enderecoCompleto]",
        "• Endereço do Imóvel: ",
        fatores_do_usuario.get("enderecoCompleto", ""),
        Pt(11)
    )
    substituir_placeholder_por_titulo_e_valor(
        documento,
        "[finalidade]",
        "",  # título vazio
        "",  # valor vazio
        Pt(11)
    )


    area_total_lida = float(dados_avaliando.get("AREA TOTAL", 0))
    area_total_str = f"{formatar_numero_brasileiro(area_total_lida)} m²"

    substituir_placeholder_por_titulo_e_valor(
        documento,
        "[areaTotal]",
        "• Área Total: ",
        area_total_str,
        Pt(11)
    )

    substituir_placeholder_por_titulo_e_valor(
        documento,
        "[documentacaoImovel]",
        "• Documentação do Imóvel: ",
        fatores_do_usuario.get("documentacaoImovel", ""),
        Pt(11)
    )
    substituir_placeholder_por_titulo_e_valor(
        documento,
        "[cartorio]",
        "• Cartório: ",
        fatores_do_usuario.get("nomeCartorio", ""),
        Pt(11)
    )
    substituir_placeholder_por_titulo_e_valor(
        documento,
        "[comarca]",
        "• Comarca: ",
        fatores_do_usuario.get("nomeComarca", ""),
        Pt(11)
    )

    texto_topo = str(dados_avaliando.get("BOA TOPOGRAFIA?", "")).strip().lower()
    topo_formatado = "Sim" if "sim" in remover_acentos(texto_topo) else "Não"
    substituir_placeholder_por_titulo_e_valor(
        documento,
        "[topografia]",
        "• Possui Boa Topografia? ",
        topo_formatado,
        Pt(11)
    )

    texto_pav = str(dados_avaliando.get("PAVIMENTACAO?", "")).strip().lower()
    pav_formatado = "Sim" if "sim" in remover_acentos(texto_pav) else "Não"
    substituir_placeholder_por_titulo_e_valor(
        documento,
        "[pavimentacao]",
        "• Imóvel Pavimentado? ",
        pav_formatado,
        Pt(11)
    )

    texto_esq = str(dados_avaliando.get(" ESQUINA?", "")).strip().lower()
    esq_formatado = "Sim" if "sim" in remover_acentos(texto_esq) else "Não"
    substituir_placeholder_por_titulo_e_valor(
        documento,
        "[terreno_de_esquina]",
        "• Imóvel de Esquina? ",
        esq_formatado,
        Pt(11)
    )

    valor_distancia = dados_avaliando.get("DISTANCIA CENTRO", 0)
    distancia_km = float(valor_distancia) if pd.notna(valor_distancia) else 0.0
    distancia_str = f"{distancia_km:.2f} km"
    substituir_placeholder_por_titulo_e_valor(
        documento,
        "[distanciaAvaliando]",
        "• Distância Avaliando ao Centro (Km): ",
        distancia_str,
        Pt(11)
    )

    # Gráfico KDE
    nome_arquivo_grafico_kernel = "grafico_kernel.png"
    gerar_grafico_densidade_kernel(valores_homogeneizados_validos, nome_arquivo_grafico_kernel)
    # INSIRA ESSA VERIFICAÇÃO LOG AQUI:
    if os.path.exists(nome_arquivo_grafico_kernel):
        logger.info(f"✅ Gráfico Kernel encontrado: {nome_arquivo_grafico_kernel}")
    else:
        logger.warning(f"❌ Gráfico Kernel NÃO encontrado: {nome_arquivo_grafico_kernel}")
    substituir_placeholder_por_imagem(documento, "[graficoKernel]", nome_arquivo_grafico_kernel, largura=Inches(5))

    # Tabela de amostras homogeneizadas
    lista_detalhes_calc = calcular_detalhes_amostras(
        dataframe_amostras_filtrado,
        dados_avaliando,
        fatores_do_usuario,
        finalidade_do_laudo
    )
    inserir_tabela_amostras_calculadas(documento, lista_detalhes_calc)
  
    
    # ------------------------------------------------------------------
    # COLESTE TODOS OS FATORES DAS AMOSTRAS PARA O ENQUADRAMENTO
    # ------------------------------------------------------------------
    lista_todos_os_fatores = []
    for det in lista_detalhes_calc:
        for chave in ("FA", "FO", "FAP", "FT", "FP", "FPA", "FE", "FAC", "FL"):
            try:
                lista_todos_os_fatores.append(float(det[chave]))
            except Exception:
                pass  # ignora caso não converta

    

    # Memória de cálculo
    lista_memoria_calculo = gerar_lista_memoria_calculo(
        dataframe_amostras_filtrado,
        dados_avaliando,
        fatores_do_usuario,
        finalidade_do_laudo
    )
    inserir_texto_memoria_calculo_no_placeholder(documento, "[texto_tabela_fatores]", lista_memoria_calculo)

    # Texto de saneamento (Chauvenet)
    texto_saneamento_chauvenet = gerar_texto_saneamento_chauvenet_extremamente_detalhado(
        dataframe_amostras_inicial, dataframe_amostras_filtrado,
        indices_excluidos, amostras_excluidas,
        media, desvio_padrao, menor_valor, maior_valor, mediana_valor
    )
    inserir_texto_saneamento_no_placeholder(documento, "[texto_relatorio_resumo_saneamento]", texto_saneamento_chauvenet)

    # Diagnóstico de mercado
    inserir_tabela_diagnostico_de_mercado(
        documento,
        "[DIAGNOSTICO]",
        fatores_do_usuario.get("estrutura_escolha", "BOA"),
        fatores_do_usuario.get("conduta_escolha", "ESTAGNADA"),
        fatores_do_usuario.get("desempenho_escolha", "ALTO")
    )

    # IC 80% e valores
    # Normaliza a entrada: aceita lista de floats ou lista de dicts com 'valor_unitario'/'valor_estimado'
    def _extrair_valores_float(seq):
        import math
        out = []
        for v in seq or []:
            if isinstance(v, dict):
                val = v.get("valor_unitario", v.get("valor_estimado", v.get("VUH", v.get("vuh", 0))))
            else:
                val = v
            try:
                val = float(val)
            except (TypeError, ValueError):
                continue
            if math.isfinite(val) and val > 0:
                out.append(val)
        return out

    valores_h_num = _extrair_valores_float(valores_homogeneizados_validos)
    array_validados = numpy.array(valores_h_num, dtype=float)

    if len(array_validados) > 0:
        limite_inferior_ic, limite_superior_ic = intervalo_confianca_bootstrap_mediana(array_validados, 1000, 0.80)
        valor_minimo = limite_inferior_ic
        valor_maximo = limite_superior_ic
        valor_mediano = numpy.median(array_validados)
    else:
        valor_minimo = 0.0
        valor_maximo = 0.0
        valor_mediano = 0.0

    # =========================================================================
    # DETERMINAÇÃO DA ÁREA DE CÁLCULO
    # =========================================================================
    # Para finalidades regulares: usa a área total da planilha
    # Para desapropriação/servidão: usa a área digitada pelo usuário no formulário
    # Esta área será utilizada para todos os cálculos de valor e restrições
    # =========================================================================
    # DEFINIÇÃO CRÍTICA: Qual área usar para cálculos
    # (desapropriação/servidão → área digitada // outros → área da planilha)
    if finalidade_do_laudo in ["desapropriacao", "servidao"]:
        area_utilizada = area_parcial_afetada
        logger.info(f"DEBUG: Usando área parcial afetada (usuário): {area_utilizada} m²")
    else:
        area_utilizada = area_total_lida
        logger.info(f"DEBUG: Usando área total da planilha: {area_utilizada} m²")

  

    restricoes_usuario = fatores_do_usuario.get("restricoes", [])

    def calcular_valor_total_com_restricoes(valor_unit):
        soma_area_restricoes = 0.0
        valor_acumulado = 0.0
        lista_subtotais = []
        for r_ in restricoes_usuario:
            a_ = r_["area"]
            f_ = r_["fator"]
            if a_ > 0:
                subtotal = valor_unit * a_ * f_
            else:
                subtotal = 0.0
            lista_subtotais.append(subtotal)
            valor_acumulado += subtotal
            soma_area_restricoes += a_
        sobra = area_utilizada - soma_area_restricoes
        if sobra > 0:
            valor_acumulado += (valor_unit * sobra)
        return valor_acumulado, lista_subtotais, sobra

    valor_total_minimo, _, _ = calcular_valor_total_com_restricoes(valor_minimo)
    valor_total_mediano, subtotais_medianos, sobra_median = calcular_valor_total_com_restricoes(valor_mediano)
    valor_total_maximo, _, _ = calcular_valor_total_com_restricoes(valor_maximo)

    substituir_placeholder_por_texto_formatado(
        documento, "[avaliacaoValorTotalMinimo]",
        formatar_moeda_brasil(valor_total_minimo), Pt(18), False
    )
    substituir_placeholder_por_texto_formatado(
        documento, "[avaliacaoValorTotal]",
        formatar_moeda_brasil(valor_total_mediano), Pt(23), True
    )
    substituir_placeholder_por_texto_formatado(
        documento, "[avaliacaoValorTotalMaximo]",
        formatar_moeda_brasil(valor_total_maximo), Pt(18), False
    )

    substituir_placeholder_por_texto_formatado(
        documento, "[avaliacaoValorTotalMinimoUnitario]",
        f"{formatar_moeda_brasil(valor_minimo)}/m²", Pt(12), False
    )
    substituir_placeholder_por_texto_formatado(
        documento, "[avaliacaoValorTotalUnitario]",
        f"{formatar_moeda_brasil(valor_mediano)}/m²", Pt(12), True
    )
    substituir_placeholder_por_texto_formatado(
        documento, "[avaliacaoValorTotalMaximoUnitario]",
        f"{formatar_moeda_brasil(valor_maximo)}/m²", Pt(12), False
    )

    # Montar lista de restrições "oficiais" para exibir em [RESUMO VALORES]
    soma_atual = 0.0
    restricoes_detalhadas_final = []
    for i, r_ in enumerate(restricoes_usuario):
        a_ = r_["area"]
        f_ = r_["fator"]
        perc_dep = r_["percentualDepreciacao"]
        tipo_ = r_["tipo"]
        subt = valor_mediano * a_ * f_ if a_ > 0 else 0.0
        restricoes_detalhadas_final.append({
            "tipo": tipo_,
            "area": a_,
            "percentualDepreciacao": perc_dep,
            "fator": f_,
            "subtotal": formatar_moeda_brasil(subt)
        })
        soma_atual += a_
    sobra_of = area_utilizada - soma_atual
    if sobra_of > 0:
        valor_sobra = valor_mediano * sobra_of
        restricoes_detalhadas_final.append({
            "tipo": "Área Livre",
            "area": sobra_of,
            "percentualDepreciacao": 0.0,
            "fator": 1.0,
            "subtotal": formatar_moeda_brasil(valor_sobra)
        })

    if len(restricoes_usuario) == 0:
        texto_rest = "Não aplicada"
    elif len(restricoes_usuario) == 1:
        if abs(restricoes_usuario[0]["area"] - area_utilizada) < 1e-3:
            texto_rest = "Aplicada a toda a área"
        else:
            texto_rest = "Aplicada parcialmente"
    else:
        texto_rest = "Múltiplas restrições aplicadas"

    if finalidade_do_laudo in ["desapropriacao", "servidao"]:
        area_final = area_parcial_afetada
    else:
        area_final = float(dados_avaliando.get("AREA TOTAL", 0))

    info_resumo = {
        "valor_unitario": f"{formatar_moeda_brasil(valor_mediano)}/m²",
        "area_total_considerada": f"{formatar_numero_brasileiro(area_final)} m²",
        "texto_descritivo_restricoes": texto_rest,
        "restricoes": restricoes_detalhadas_final,
        "valor_total_indenizatorio": formatar_moeda_brasil(valor_total_mediano),
        "valor_por_extenso": ""
    }

    inserir_tabela_resumo_de_valores(documento, "[RESUMO VALORES]", info_resumo, area_utilizada)

    # Gráficos de aderência e dispersão
    substituir_placeholder_por_imagem(documento, "[graficoAderencia2]", caminho_imagem_aderencia, largura=Inches(5))
    substituir_placeholder_por_imagem(documento, "[graficoDispersao]", caminho_imagem_dispersao, largura=Inches(5))

    # Grau de precisão
    if len(valores_homogeneizados_validos) > 0:
        # Converte lista (floats ou dicts) → lista de floats
        def _extrair_valores_float(seq, chave_preferida="valor_unitario"):
            import math
            out = []
            for v in seq or []:
                if isinstance(v, dict):
                    # <<< corrigido: sem aspas após a variável >>>
                    val = v.get(chave_preferida, v.get("valor_estimado", v.get("VUH", v.get("vuh", 0))))
                else:
                    val = v
                try:
                    val = float(val)
                except (TypeError, ValueError):
                    continue
                if math.isfinite(val):
                    out.append(val)
            return out

        vals = _extrair_valores_float(valores_homogeneizados_validos, "valor_unitario")

        if vals:
            arr = numpy.array(vals, dtype=float)
            mediana_hom = float(numpy.median(arr))
            valor_minimo = float(numpy.min(arr))
            valor_maximo = float(numpy.max(arr))
            amplitude_ic80 = ((valor_maximo - valor_minimo) / mediana_hom) * 100 if mediana_hom > 0 else 0.0
        else:
            mediana_hom = 0.0
            amplitude_ic80 = 0.0
    else:
        mediana_hom = 0.0
        amplitude_ic80 = 0.0

    inserir_tabela_classificacao_de_precisao(documento, "[texto_grau_precisao]", amplitude_ic80)

    # Fundamentação e enquadramento
    quantidade_amostras_validadas = len(dataframe_amostras_filtrado)

    inserir_fundamentacao_e_enquadramento(
            documento,
            "[FUNDAMENTACAO]",
            "[enquadramento final]",
            fatores_do_usuario.get("tipoImovel", "Gleba Rural"),
            quantidade_amostras_validadas,
            lista_todos_os_fatores          # << novo argumento
    )  
    
    # Fotos do avaliando (agora é uma lista de listas)
    for grupo in caminhos_fotos_avaliando:
        for caminho in grupo:
            if os.path.exists(caminho):
                logger.info(f"✅ Foto do avaliando encontrada: {caminho}")
            else:
                logger.warning(f"❌ Foto do avaliando NÃO encontrada: {caminho}")

    # Documentos adicionais (matrícula)
    for grupo in caminhos_fotos_adicionais:
        for caminho in grupo:
            if os.path.exists(caminho):
                logger.info(f"✅ Documento adicional (matrícula) encontrado: {caminho}")
            else:
                logger.warning(f"❌ Documento adicional (matrícula) NÃO encontrado: {caminho}")

    # Documentos do proprietário
    for grupo in caminhos_fotos_proprietario:
        for caminho in grupo:
            if os.path.exists(caminho):
                logger.info(f"✅ Documento do proprietário encontrado: {caminho}")
            else:
                logger.warning(f"❌ Documento do proprietário NÃO encontrado: {caminho}")

    # Documentos da planta
    for grupo in caminhos_fotos_planta:
        for caminho in grupo:
            if os.path.exists(caminho):
                logger.info(f"✅ Documento de planta encontrado: {caminho}")
            else:
                logger.warning(f"❌ Documento de planta NÃO encontrado: {caminho}")


    # Verificar se o logo existe
    caminho_logo = fatores_do_usuario.get("caminhoLogo", "")
    if caminho_logo:
        if os.path.exists(caminho_logo):
            logger.info(f"✅ Logo encontrado: {caminho_logo}")
        else:
            logger.warning(f"❌ Logo NÃO encontrado: {caminho_logo}")

    from itertools import chain

    # Inserir fotos do imóvel
    if caminhos_fotos_avaliando:
        todas_as_fotos = list(chain.from_iterable(caminhos_fotos_avaliando))
        inserir_fotos_no_placeholder(documento, "[FOTOS]", todas_as_fotos, largura_imagem=Inches(3.2), um_por_pagina=False)
    else:
        substituir_placeholder_por_texto_formatado(
            documento,
            "[FOTOS]",
            "FOTOS DO IMÓVEL AVALIADO NÃO FORNECIDAS",
            Pt(12),
            True
        )


    
    # Inserir documentos adicionais (matrícula)
    if caminhos_fotos_adicionais:
        todas_matriculas = list(chain.from_iterable(caminhos_fotos_adicionais))
        inserir_fotos_no_placeholder(documento, "[MATRICULA]", todas_matriculas, largura_imagem=Inches(5), um_por_pagina=True)
    else:
        substituir_placeholder_por_texto_formatado(
            documento,
            "[MATRICULA]",
            "DOCUMENTAÇÃO ADICIONAL NÃO FORNECIDA",
            Pt(12),
            True
        )


    
    # Inserir documentação do proprietário
    if caminhos_fotos_proprietario:
        todos_proprietarios = list(chain.from_iterable(caminhos_fotos_proprietario))
        inserir_fotos_no_placeholder(documento, "[PROPRIETARIO]", todos_proprietarios, largura_imagem=Inches(5), um_por_pagina=True)
    else:
        substituir_placeholder_por_texto_formatado(
            documento,
            "[PROPRIETARIO]",
            "DOCUMENTAÇÃO DO PROPRIETÁRIO NÃO FORNECIDA",
            Pt(12),
            True
        )

    # Inserir documentação da planta
    if caminhos_fotos_planta:
        todas_plantas = list(chain.from_iterable(caminhos_fotos_planta))
        inserir_fotos_no_placeholder(documento, "[PLANTA]", todas_plantas, largura_imagem=Inches(5), um_por_pagina=True)
    else:
        substituir_placeholder_por_texto_formatado(
            documento,
            "[PLANTA]",
            "PLANTA DO IMÓVEL NÃO FORNECIDA",
            Pt(12),
            True
        )
  
    # Logo
    caminho_logo = fatores_do_usuario.get("caminhoLogo", "")
    if caminho_logo and os.path.exists(caminho_logo):
        inserir_logo_no_placeholder(documento, "[logo]", caminho_logo)

    # (Exemplo) Inserir tabela [RESUMO GERAL] se existir placeholder
    # A função 'inserir_tabela_resumo_geral_completo' foi meramente ilustrativa
    # no código anterior. Você pode chamar se quiser:
    # inserir_tabela_resumo_geral_completo(documento, "[RESUMO GERAL]", {...})

    # Salvar
    documento.save(nome_arquivo_word)
    # Limpar arquivos PNG temporários gerados a partir de PDFs
    def limpar_arquivos_temp_png(lista_de_caminhos):
        for caminho in lista_de_caminhos:
            if isinstance(caminho, str) and caminho.endswith(".png") and os.path.exists(caminho):
                try:
                    os.remove(caminho)
                    logger.info(f"🗑️ PNG temporário removido: {caminho}")
                except Exception as e:
                    logger.warning(f"⚠️ Falha ao remover {caminho}: {e}")

    # Apagar apenas os arquivos gerados a partir de PDFs
    # Flatten antes de filtrar PNGs
    limpar_arquivos_temp_png([c for grupo in caminhos_fotos_adicionais for c in grupo if c.endswith(".png")])
    limpar_arquivos_temp_png([c for grupo in caminhos_fotos_proprietario for c in grupo if c.endswith(".png")])
    limpar_arquivos_temp_png([c for grupo in caminhos_fotos_planta for c in grupo if c.endswith(".png")])


    try:
        os.startfile(nome_arquivo_word)
    except:
        pass




###############################################################################
# LEITURA DA PLANILHA EXCEL
###############################################################################
def ler_planilha_excel(caminho_arquivo_excel: str, raio_limite_km: float = 150.0):
    import pandas as pd
    from math import radians, sin, cos, sqrt, atan2
    from geopy.geocoders import Nominatim

    def _to_float(v):
        if isinstance(v, str):
            v = v.replace(".", "").replace(",", ".").strip()
        try:
            return float(v)
        except:
            return pd.NA

    def _parse_coord(coord):
        try:
            if isinstance(coord, str):
                coord = coord.replace("°", "").replace(",", ".").strip()
            return float(coord)
        except:
            return pd.NA


    def haversine_km(lat1, lon1, lat2, lon2):
        if pd.isna(lat1) or pd.isna(lon1) or pd.isna(lat2) or pd.isna(lon2):
            return pd.NA
        R = 6371.0
        dlat = radians(lat2 - lat1)
        dlon = radians(lon2 - lon1)
        a = sin(dlat/2)**2 + cos(radians(lat1)) * cos(radians(lat2)) * sin(dlon/2)**2
        c = 2 * atan2(sqrt(a), sqrt(1 - a))
        return R * c

    df = pd.read_excel(caminho_arquivo_excel)
    print(df.head()) 
    df.dropna(how="all", inplace=True)
    df.reset_index(drop=True, inplace=True)

    for col in ("VALOR TOTAL", "AREA TOTAL", "VALOR UNITARIO"):
        if col in df.columns:
            df[col] = df[col].apply(_to_float)

    dados_avaliando = df.iloc[-1].to_dict()
    dataframe_amostras = df.iloc[:-1].copy()

    if {"VALOR TOTAL", "AREA TOTAL"}.issubset(dataframe_amostras.columns):
        dataframe_amostras["VALOR UNITARIO"] = (
            dataframe_amostras["VALOR TOTAL"] / dataframe_amostras["AREA TOTAL"].replace({0: pd.NA})
        )

    lat_av = _parse_coord(dados_avaliando.get("LATITUDE"))
    lon_av = _parse_coord(dados_avaliando.get("LONGITUDE"))

    nome_cidade = str(dados_avaliando.get("CIDADE", "")).strip()
    if nome_cidade:
        try:
            geoloc = Nominatim(user_agent="aval-geo")
            loc = geoloc.geocode(f"{nome_cidade}, Brazil", timeout=10)
            lat_ctr, lon_ctr = loc.latitude, loc.longitude if loc else (lat_av, lon_av)
        except:
            lat_ctr, lon_ctr = lat_av, lon_av
    else:
        lat_ctr, lon_ctr = lat_av, lon_av

    dados_avaliando["DISTANCIA CENTRO"] = haversine_km(lat_av, lon_av, lat_ctr, lon_ctr)

    dataframe_amostras["LAT_PARS"] = dataframe_amostras["LATITUDE"].apply(_parse_coord)
    dataframe_amostras["LON_PARS"] = dataframe_amostras["LONGITUDE"].apply(_parse_coord)
    dataframe_amostras["DISTANCIA CENTRO"] = dataframe_amostras.apply(
        lambda r: haversine_km(r["LAT_PARS"], r["LON_PARS"], lat_ctr, lon_ctr), axis=1
    )

    logger.info(f"✅ Linhas antes do filtro crítico: {len(dataframe_amostras)}")
    logger.info(f"Valores nulos em 'VALOR TOTAL': {dataframe_amostras['VALOR TOTAL'].isna().sum()}")
    logger.info(f"Valores nulos em 'AREA TOTAL': {dataframe_amostras['AREA TOTAL'].isna().sum()}")
    logger.info(f"Valores nulos em 'DISTANCIA CENTRO': {dataframe_amostras['DISTANCIA CENTRO'].isna().sum()}")

    logger.info(f"Antes da exclusão, dataframe_amostras:\n{dataframe_amostras.head()}")


    mask_excluir = (
        (dataframe_amostras["DISTANCIA CENTRO"] > raio_limite_km) |
        (dataframe_amostras["DISTANCIA CENTRO"].isna()) |
        (dataframe_amostras["VALOR TOTAL"].isna()) |
        (dataframe_amostras["AREA TOTAL"].isna()) |
        (dataframe_amostras["AREA TOTAL"] == 0)
    )
    logger.info(f"Máscara de exclusão:\n{ mask_excluir.head()}")
    logger.info(f"Depois da exclusão, dataframe_amostras:\n{ dataframe_amostras.loc[~mask_excluir].head() }")
    dataframe_amostras = dataframe_amostras.loc[~mask_excluir].reset_index(drop=True)
    logger.info(f"✅ Linhas após o filtro crítico: {len(dataframe_amostras)}")
    dataframe_amostras.drop(columns=["LAT_PARS", "LON_PARS"], inplace=True)

    logger.info(f"Antes da exclusão, dataframe_amostras:\n{ dataframe_amostras }")
    logger.info(f"Mascara de exclusão:\n{ mask_excluir }")
    logger.info(f"Depois da exclusão, dataframe_amostras:\n{ dataframe_amostras.loc[~mask_excluir] }")

    return dataframe_amostras, dados_avaliando



###############################################################################
# HOMOGENEIZAR AMOSTRAS (DATAFRAME FILTRADO)
###############################################################################
def homogeneizar_amostras(dataframe_amostras_validas, dados_avaliando, fatores_do_usuario, finalidade_do_laudo):
    import math
    import numpy as np

    area_do_avaliando = float(dados_avaliando.get("AREA TOTAL", 0))

    f_avaliado_aprov = fator_aproveitamento(dados_avaliando.get("APROVEITAMENTO", "URBANO"))
    f_avaliado_topog = fator_topografia(dados_avaliando.get("BOA TOPOGRAFIA?", "NÃO"))
    f_avaliado_pedol = fator_pedologia(dados_avaliando.get("PEDOLOGIA ALAGÁVEL? ", "NÃO"))
    f_avaliado_pavim = fator_pavimentacao(dados_avaliando.get("PAVIMENTACAO?", "NÃO"))
    f_avaliado_esq   = fator_esquina(dados_avaliando.get(" ESQUINA?", "NÃO"))
    f_avaliado_acess = fator_acessibilidade(dados_avaliando.get("ACESSIBILIDADE?", "NÃO"))

    lista_valores_unitarios = []
    lista_residuos_relativos = []
    lista_valores_estimados = []
    lista_valores_originais = []

    for i, (_, linha) in enumerate(dataframe_amostras_validas.iterrows()):
        valor_total_amostra = linha["VALOR TOTAL"]
        area_da_amostra = float(linha.get("AREA TOTAL", 0))
        valor_unitario_original = linha.get("VALOR UNITARIO")

        
        if valor_unitario_original is None:
            valor_unitario_original = valor_total_amostra / area_da_amostra if area_da_amostra > 0 else 0.0
        else:
            # Garante conversão correta mesmo se vier string:
            try:
                valor_unitario_original = float(str(valor_unitario_original).replace("R$", "").replace(".", "").replace(",", "."))
            except Exception:
                valor_unitario_original = valor_total_amostra / area_da_amostra if area_da_amostra > 0 else 0.0


        fator_area = calcular_fator_area(area_do_avaliando, area_da_amostra, fatores_do_usuario["area"])
        fator_oferta = calcular_fator_oferta(True, fatores_do_usuario["oferta"])

        if fatores_do_usuario.get("localizacao_mesma_regiao", False):
            fator_localiz_calc = 1.0
        else:
            try:
                dist_amostra = float(linha.get("DISTANCIA CENTRO", 0))
                dist_avalia = float(dados_avaliando.get("DISTANCIA CENTRO", 0))
                if dist_amostra > 0 and dist_avalia > 0:
                    fa_item = 1.0 / (dist_amostra ** 0.1)
                    fa_avaliado = 1.0 / (dist_avalia ** 0.1)
                    fator_localiz_calc = limitar_fator(fa_avaliado / fa_item)
                else:
                    fator_localiz_calc = 1.0
            except:
                fator_localiz_calc = 1.0
            fator_localiz_calc = limitar_fator(fator_localiz_calc)

        f_sample_aprov = fator_aproveitamento(linha.get("APROVEITAMENTO", "URBANO"))
        if fatores_do_usuario["aproveitamento"] and f_sample_aprov != 0:
            fator_aprov_calc = limitar_fator(f_avaliado_aprov / f_sample_aprov)
        else:
            fator_aprov_calc = 1.0

        f_sample_topog = fator_topografia(linha.get("BOA TOPOGRAFIA?", "NÃO"))
        if fatores_do_usuario["topografia"] and f_sample_topog != 0:
            fator_topog_calc = limitar_fator(f_avaliado_topog / f_sample_topog)
        else:
            fator_topog_calc = 1.0

        f_sample_pedol = fator_pedologia(linha.get("PEDOLOGIA ALAGÁVEL? ", "NÃO"))
        if fatores_do_usuario["pedologia"] and f_sample_pedol != 0:
            fator_pedol_calc = limitar_fator(f_avaliado_pedol / f_sample_pedol)
        else:
            fator_pedol_calc = 1.0

        f_sample_pavim = fator_pavimentacao(linha.get("PAVIMENTACAO?", "NÃO"))
        if fatores_do_usuario["pavimentacao"] and f_sample_pavim != 0:
            fator_pavim_calc = limitar_fator(f_avaliado_pavim / f_sample_pavim)
        else:
            fator_pavim_calc = 1.0

        f_sample_esq = fator_esquina(linha.get(" ESQUINA?", "NÃO"))
        if fatores_do_usuario["esquina"] and f_sample_esq != 0:
            fator_esq_calc = limitar_fator(f_avaliado_esq / f_sample_esq)
        else:
            fator_esq_calc = 1.0

        f_sample_acess = fator_acessibilidade(linha.get("ACESSIBILIDADE?", "NÃO"))
        if fatores_do_usuario["acessibilidade"] and f_sample_acess != 0:
            fator_acess_calc = limitar_fator(f_avaliado_acess / f_sample_acess)
        else:
            fator_acess_calc = 1.0

        valor_homog = (
            valor_total_amostra *
            fator_area *
            fator_oferta *
            fator_localiz_calc *
            fator_aprov_calc *
            fator_topog_calc *
            fator_pedol_calc *
            fator_pavim_calc *
            fator_esq_calc *
            fator_acess_calc
        )

        if area_da_amostra > 0:
            valor_unitario = valor_homog / area_da_amostra
        else:
            valor_unitario = 0.0

        valor_unitario_avaliando = dados_avaliando.get("valor_unitario_medio", 0)
        if valor_unitario_avaliando:
            residuo_rel = 100 * (valor_unitario - valor_unitario_avaliando) / valor_unitario_avaliando
        else:
            residuo_rel = 0.0
        
        lista_valores_unitarios.append(valor_unitario)
        lista_valores_estimados.append(valor_unitario)
        lista_valores_originais.append(valor_unitario_original)
        lista_residuos_relativos.append(residuo_rel)

        print(f"""
        Amostra: {linha.get('IDENTIFICADOR', f'Amostra {i+1}')}
        VALOR TOTAL: {valor_total_amostra}
        AREA: {area_da_amostra}
        FA: {fator_area}
        FO: {fator_oferta}
        FAP: {fator_aprov_calc}
        FT: {fator_topog_calc}
        FP: {fator_pedol_calc}
        FPA: {fator_pavim_calc}
        FE: {fator_esq_calc}
        FAC: {fator_acess_calc}
        FL: {fator_localiz_calc}
        VALOR_HOMOG (Numerador): {valor_homog}
        VUH CALCULADO: {valor_unitario}
        VUH ESPERADO (planilha): [coloque o valor manual aqui para comparação]
        """)


    # FIM DO LOOP PRINCIPAL

    desvio_padrao_residuos = np.std(lista_residuos_relativos) if lista_residuos_relativos else 1
    lista_residuos_dp = [
        (residuo / desvio_padrao_residuos) if desvio_padrao_residuos > 0 else 0.0
        for residuo in lista_residuos_relativos
    ]

    amostras_resultantes = []
    for i, (_, linha) in enumerate(dataframe_amostras_validas.iterrows()):
        # Valor unitário original: tenta pegar da coluna, senão calcula
        # valor_unitario_original = linha.get("VALOR UNITARIO")
        # if valor_unitario_original is None:
        #     valor_unitario_original = linha["VALOR TOTAL"] / linha["AREA TOTAL"] if linha["AREA TOTAL"] > 0 else 0.0

        amostras_resultantes.append({
            "identificador": linha.get("IDENTIFICADOR", f"Amostra {i+1}"),
            "valor_total": linha["VALOR TOTAL"],
            "area": linha["AREA TOTAL"],
            "valor_unitario_original": lista_valores_originais[i],
            "valor_unitario": lista_valores_unitarios[i],
            "valor_estimado": lista_valores_estimados[i],
            "residuo_rel": lista_residuos_relativos[i],
            "residuo_dp": lista_residuos_dp[i],
            "idx": linha.get("idx") or linha.get("AM"),  # ← ESSENCIAL!
        })

    return amostras_resultantes



